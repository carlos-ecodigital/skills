#!/usr/bin/env python3
"""
LOI/NCNDA v3.0 — Document Generation Script (v3.1)

Usage:
    python generate_loi.py intake.yaml [--output path/to/output.docx]

Takes a YAML intake file and produces a complete, branded .docx ready to sign.
All clauses, conditionals, and schedules are generated from the intake data.
"""

import sys
import os
import yaml
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT


# ---------------------------------------------------------------------------
# Configuration -- DE brand standards (aligned with document-templates skill)
# ---------------------------------------------------------------------------

FONT_NAME = "Inter"
FONT_BODY = Pt(10)
FONT_HEADING1 = Pt(13)
FONT_HEADING2 = Pt(11)
FONT_SMALL = Pt(8)
LINE_SPACING = 1.15

NAVY = RGBColor(0x14, 0x29, 0x45)    # DE brand heading color
BLACK = RGBColor(0x1E, 0x29, 0x3B)   # DE brand body text
GREY = RGBColor(0x64, 0x74, 0x8B)    # DE brand slate/secondary
COBALT = RGBColor(0x00, 0x34, 0xAF)  # DE brand primary accent

# Logo path -- shared with document-factory skill
SCRIPT_DIR = Path(__file__).parent
LOGO_PATH = SCRIPT_DIR.parent / "document-factory" / "assets" / "DE_Logo_Black.png"
if not LOGO_PATH.exists():
    LOGO_PATH = SCRIPT_DIR / "DE_Logo_Black.png"

# DE entity footer text
DE_FOOTER = "Digital Energy Group AG  |  Baarerstrasse 43, 6300 Zug  |  CHE-408.639.320  |  digital-energy.group"


# ---------------------------------------------------------------------------
# YAML Loading + Validation
# ---------------------------------------------------------------------------

def load_intake(path: str) -> dict:
    with open(path, "r") as f:
        data = yaml.safe_load(f)
    validate(data)
    return data


def validate(d: dict):
    errors = []
    t = d.get("type", "")
    if t not in ("EndUser", "Distributor", "Wholesale"):
        errors.append(f"type must be EndUser, Distributor, or Wholesale (got: {d.get('type')})")
    for s in ("provider", "counterparty", "programme", "dates"):
        if s not in d:
            errors.append(f"Missing section: {s}")
    cp = d.get("counterparty", {})
    for f in ("name", "short", "description"):
        if not cp.get(f):
            errors.append(f"counterparty.{f} required")
    if not d.get("dates", {}).get("loi_date"):
        errors.append("dates.loi_date required")
    if t == "Distributor" and d.get("partnership_mode", "combined") == "combined":
        for f in ("partner_core_capability", "partner_contribution",
                   "combined_offering", "target_end_users", "partner_service_scope"):
            if not d.get("commercial", {}).get(f):
                errors.append(f"commercial.{f} required for Distributor Mode A")
    if t == "Wholesale":
        for f in ("dec_block_count", "indicative_mw"):
            if not d.get("commercial", {}).get(f):
                errors.append(f"commercial.{f} required for Wholesale")
    if t == "EndUser":
        if not d.get("commercial", {}).get("service_type"):
            errors.append("commercial.service_type required for EndUser")
    if errors:
        print("VALIDATION ERRORS:")
        for e in errors:
            print(f"  - {e}")
        sys.exit(1)


# ---------------------------------------------------------------------------
# Document Builder
# ---------------------------------------------------------------------------

class LOI:
    def __init__(self, data: dict):
        self.d = data
        self.t = data["type"]
        self.doc = Document()
        self._setup()
        self.party = "Partner" if self.t == "Distributor" else "Customer"

    def _setup(self):
        style = self.doc.styles["Normal"]
        style.font.name = FONT_NAME
        style.font.size = FONT_BODY
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = LINE_SPACING
        for s in self.doc.sections:
            s.page_width = Mm(210)
            s.page_height = Mm(297)
            s.top_margin = Mm(20)
            s.bottom_margin = Mm(35)
            s.left_margin = Mm(25)
            s.right_margin = Mm(20)
            s.different_first_page_header_footer = True

            # First page header: DE logo
            fph = s.first_page_header
            fph.is_linked_to_previous = False
            if LOGO_PATH.exists():
                fph.paragraphs[0].add_run().add_picture(str(LOGO_PATH), height=Mm(12))

            # Continuation header: small logo + document title
            h = s.header
            h.is_linked_to_previous = False
            hp = h.paragraphs[0]
            hp.paragraph_format.tab_stops.add_tab_stop(Mm(165), WD_TAB_ALIGNMENT.RIGHT)
            if LOGO_PATH.exists():
                hp.add_run().add_picture(str(LOGO_PATH), height=Mm(8))
            hp.add_run("\t")
            tr = hp.add_run("Letter of Intent")
            tr.font.size = Pt(9)
            tr.font.name = FONT_NAME
            tr.font.color.rgb = GREY

            # First page footer
            fpf = s.first_page_footer
            fpf.is_linked_to_previous = False
            fp = fpf.paragraphs[0]
            fp.paragraph_format.tab_stops.add_tab_stop(Mm(165), WD_TAB_ALIGNMENT.RIGHT)
            r = fp.add_run("Confidential    " + DE_FOOTER)
            r.font.size = Pt(8)
            r.font.name = FONT_NAME
            r.font.color.rgb = GREY

            # Continuation footer
            cf = s.footer
            cf.is_linked_to_previous = False
            cfp = cf.paragraphs[0]
            cfp.paragraph_format.tab_stops.add_tab_stop(Mm(165), WD_TAB_ALIGNMENT.RIGHT)
            r2 = cfp.add_run("Confidential    " + DE_FOOTER)
            r2.font.size = Pt(8)
            r2.font.name = FONT_NAME
            r2.font.color.rgb = GREY

    # --- Helpers ---

    def g(self, *keys, default=""):
        v = self.d
        for k in keys:
            v = v.get(k, default) if isinstance(v, dict) else default
        return v or default

    def choice(self, k):
        return bool(self.g("choices", k))

    def p(self, text, bold=False, italic=False, size=None, color=None, align=None, space_after=None):
        para = self.doc.add_paragraph()
        if align:
            para.alignment = align
        if space_after is not None:
            para.paragraph_format.space_after = Pt(space_after)
        run = para.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = size or FONT_BODY
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        return para

    def bp(self, label, text):
        """Bold label + normal text in one paragraph."""
        para = self.doc.add_paragraph()
        r1 = para.add_run(label)
        r1.bold = True
        r1.font.name = FONT_NAME
        r1.font.size = FONT_BODY
        r2 = para.add_run(text)
        r2.font.name = FONT_NAME
        r2.font.size = FONT_BODY
        return para

    def h(self, text, level=2):
        heading = self.doc.add_heading(text, level=level)
        for r in heading.runs:
            r.font.name = FONT_NAME
            r.font.color.rgb = NAVY

    def table(self, headers, rows):
        t = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.name = FONT_NAME
                    r.font.size = FONT_BODY
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = t.rows[ri + 1].cells[ci]
                cell.text = str(val)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = FONT_NAME
                        r.font.size = FONT_BODY
        return t

    def line(self):
        self.p("—" * 50, color=GREY, size=Pt(6), space_after=0)

    # --- Sections ---

    def letterhead(self):
        # Cover page: Title, then both parties listed together
        subjects = {
            "Distributor": "Letter of Intent and Non-Circumvention Non-Disclosure Agreement \u2014 Strategic Infrastructure Partnership",
            "Wholesale": "Letter of Intent and Non-Circumvention Non-Disclosure Agreement \u2014 Purpose-Built AI Colocation Capacity",
            "EndUser": "Letter of Intent \u2014 AI Compute Infrastructure Services",
        }
        self.p(subjects[self.t], bold=True, size=FONT_HEADING1, color=NAVY)
        self.p("")

        # Provider (Between:)
        prov = self.d.get("provider", {})
        self.p("Between:", color=BLACK)
        self.p(prov.get("legal_name", ""), bold=True)
        self.p(prov.get("address", ""), size=FONT_SMALL, color=GREY)
        kvk = prov.get("kvk", "")
        if kvk:
            self.p(f"KvK: {kvk}", size=FONT_SMALL, color=GREY)
        self.p(f"Date: {self.g('dates', 'loi_date')}", size=FONT_SMALL, color=GREY)
        self.p("")

        # Counterparty (And:)
        cp = self.d.get("counterparty", {})
        self.p("And:", color=BLACK)
        self.p(cp.get("name", ""), bold=True)
        self.p(cp.get("address", ""), size=FONT_SMALL, color=GREY)
        rt = cp.get("reg_type", "")
        rn = cp.get("reg_number", "")
        if rt and rn:
            self.p(f"{rt}: {rn}", size=FONT_SMALL, color=GREY)

    def addressee(self):
        # No longer used -- both parties are on cover page via letterhead()
        pass

    def subject(self):
        # No longer used -- title is in letterhead()
        pass

    def recitals(self):
        self.h("Recitals")
        prov = self.g("provider", "short_name")
        cp = self.g("counterparty", "short")
        desc = self.g("counterparty", "description")
        sites = self.g("programme", "site_count")

        if self.t == "Distributor":
            self.p(
                f'(A) {prov} (the "Provider") develops and operates Digital Energy Centers ("DECs"): '
                f"purpose-built, liquid-cooled colocation facilities designed for high-density accelerated "
                f"compute workloads. DECs integrate AI compute with energy recycling and generation. "
                f"The Provider's programme spans {sites} identified sites with secured energy and "
                f"grid access, positioning its platform as one of the leading sovereign AI infrastructure "
                f"programmes capable of delivering capacity at scale within 12 months of commercial "
                f"commitment. The Provider seeks qualified channel and integration partners to extend "
                f"its platform reach to end-user segments where the {self.party} holds established customer "
                f"relationships and domain expertise."
            )
        elif self.t == "Wholesale":
            self.p(
                f'(A) {prov} (the "Provider") develops and operates Digital Energy Centers ("DECs"): '
                f"purpose-built, liquid-cooled colocation facilities designed for high-density accelerated "
                f"compute workloads. DECs integrate AI compute with energy recycling and generation. "
                f"The Provider's programme spans {sites} identified sites with secured energy and "
                f"grid access, positioning its platform as one of the leading sovereign AI infrastructure "
                f"programmes capable of delivering capacity at scale within 12 months of commercial commitment."
            )
        else:
            self.p(
                f'(A) {prov} (the "Provider") operates Digital Energy Centers ("DECs"): '
                f"purpose-built, liquid-cooled colocation facilities designed for high-density AI and "
                f"accelerated compute workloads. DECs integrate AI compute with energy recycling and "
                f"generation. The Provider's programme spans {sites} identified sites with secured "
                f"energy and grid access, offering European data sovereignty and capacity across multiple "
                f"service models \u2014 including dedicated bare metal colocation, shared cloud environments, "
                f"and token-based GPU access."
            )

        self.p(f'(B) {cp} (the "{self.party}") {desc}.')

        if self.t == "Distributor":
            self.p(
                f"(C) The Parties wish to record their mutual interest in establishing a strategic "
                f"partnership under which the {self.party} would collaborate with the Provider to deliver "
                f"AI colocation services to end-user customers, on the indicative terms set out below. "
                f'This letter of intent and non-circumvention non-disclosure agreement (the "LOI") '
                f"reflects both Parties' strategic intent and is intended to form the basis for further "
                f"commercial negotiation toward a definitive partnership agreement or master services "
                f'agreement (the "MSA").'
            )
            self.p(
                "(D) The Parties recognise that the commercial discussions contemplated by this LOI "
                "will require the exchange of commercially sensitive and competitively valuable "
                "information, and wish to establish binding confidentiality and non-circumvention "
                "protections to facilitate those discussions."
            )
        elif self.t == "Wholesale":
            self.p(
                f"(C) The Parties wish to record their mutual interest in the Provider making available, "
                f"and the {self.party} procuring, dedicated AI colocation capacity at the Provider's DEC "
                f"facilities, on the indicative terms set out below. This letter of intent and "
                f'non-circumvention non-disclosure agreement (the "LOI") reflects both Parties\' strategic '
                f"intent to establish a long-term infrastructure partnership and is intended to form the "
                f"basis for further technical scoping and commercial negotiation toward a Master Services "
                f'Agreement (the "MSA").'
            )
            self.p(
                "(D) The Parties recognise that the commercial discussions contemplated by this LOI "
                "will require the exchange of commercially sensitive information, including site-specific "
                "data, pricing models, and infrastructure specifications, and wish to establish binding "
                "confidentiality and non-circumvention protections to facilitate those discussions."
            )
        else:
            self.p(
                f"(C) The Parties wish to record their mutual interest in the {self.party} procuring AI "
                f"compute infrastructure services at the Provider's DEC facilities, on the indicative "
                f'terms set out below. This letter of intent (the "LOI") reflects both Parties\' intent '
                f"and is intended to form the basis for further technical scoping and commercial "
                f'negotiation toward a service agreement (the "MSA").'
            )

    def definitions(self):
        self.h("1. Definitions")
        self.p("1.1 In this LOI, unless the context requires otherwise:")

        defs = []
        defs.append(('"Affiliate"', 'means, in relation to a Party, any entity that directly or indirectly controls, is controlled by, or is under common control with that Party, where "control" means the ownership of more than 50% of the voting rights or equivalent ownership interest;'))

        if self.t in ("Distributor", "Wholesale"):
            ac_text = (
                '"Associated Counterparty" means, in relation to a Site Identifier, any of the following '
                "with whom the Provider has a contractual, commercial, or active non-public engagement "
                "in connection with that site: (a) landowners and land lessors; (b) greenhouse operators "
                "and agricultural partners; (c) heat offtake counterparties; (d) energy procurement "
                "counterparties; (e) engineering, procurement, and construction contractors engaged or "
                "under negotiation; and (f) grid connection holders and distribution system operators "
                "(to the extent of non-public engagement)."
            )
            if self.t == "Distributor":
                ac_text += (
                    " (g) government agencies, municipalities, and regulatory bodies with whom the "
                    "Provider has active non-public engagement (excluding general regulatory contacts "
                    "available to any market participant);"
                )
            else:
                ac_text += (
                    " For the avoidance of doubt, government agencies and regulatory bodies are "
                    "excluded unless the Provider has made a specific named introduction;"
                )
            defs.append(("", ac_text))

        defs.append(('"Business Day"', "means a day other than a Saturday, Sunday, or public holiday in the Netherlands;"))

        if self.t == "Distributor":
            defs.append(('"Competitor"', 'means any entity whose primary business includes the development, ownership, or operation of colocation facilities for high-density compute workloads, or the provision of AI infrastructure services, in the European Economic Area;'))

        ci_text = (
            '"Confidential Information" means all information (whether technical, commercial, financial, '
            "or otherwise) disclosed by a Party to the other in connection with the Purpose, whether in "
            "writing, orally, electronically, or by inspection, including any information that by its "
            "nature or the circumstances of disclosure would reasonably be understood to be confidential"
        )
        if self.t in ("Distributor", "Wholesale"):
            ci_text += ", and including metadata, EXIF data, digital artifacts, file names, folder names, and any derivative information"
        ci_text += ". The existence and contents of this LOI are Confidential Information;"
        defs.append(("", ci_text))

        defs.append(('"DEC"', "means a Digital Energy Center: a purpose-built colocation facility developed and operated by the Provider;"))
        defs.append(('"DEC Block"', "means a standardised compute unit of approximately 1.2 MW IT capacity" + (", factory-manufactured and modular;" if self.t != "EndUser" else ";")))
        defs.append(('"Financing Party"', "means any bank, financial institution, fund, security trustee, or other entity providing or arranging " + ("debt, mezzanine, or structured " if self.t != "EndUser" else "") + "finance to the Provider or any of its Affiliates in connection with the development or operation of any DEC;"))

        if self.t == "Distributor":
            defs.append(('"MSA"', "means the definitive Master Services Agreement or partnership agreement to be negotiated between the Parties;"))
        elif self.t == "Wholesale":
            defs.append(('"MSA"', "means the definitive Master Services Agreement to be negotiated between the Parties, incorporating the Sales Order Form, SLA Schedule, and Pricing Framework;"))
        else:
            defs.append(('"MSA"', "means the definitive service agreement to be negotiated between the Parties;"))

        if self.t == "Distributor":
            defs.append(('"Protected Business Information" or "PBI"', "means the Provider's proprietary methodologies, strategies, and frameworks including: (a) site-sourcing and land-acquisition methodology; (b) energy procurement and grid connection strategy; (c) regulatory and permitting playbook; (d) heat offtake and energy recycling commercial model; (e) financial and operational modelling frameworks; (f) DEC design specifications and engineering standards; and (g) colocation pricing methodology and deal economics;"))

        if self.t == "Distributor":
            defs.append(('"Purpose"', "means evaluating, negotiating, and progressing toward the execution of an MSA for the Transaction, including all activities under any companion agreements between the Parties (such as a Referral Agreement);"))
        elif self.t == "Wholesale":
            defs.append(('"Purpose"', "means evaluating, negotiating, and progressing toward the execution of an MSA for the provision of AI colocation services;"))
        else:
            defs.append(('"Purpose"', "means evaluating, negotiating, and progressing toward the execution of an MSA;"))

        if self.t in ("Distributor", "Wholesale"):
            defs.append(('"Ready-for-Service" or "RFS"', "means the date on which a DEC or DEC Block has been commissioned and is available for the provision of colocation services;"))

        defs.append(('"Representatives"', "means, in relation to a Party, its Affiliates, and its and their respective directors, officers, employees, agents, and professional advisers;"))

        if self.t in ("Distributor", "Wholesale"):
            defs.append(('"Sales Order Form"', "means the document that will form the basis for each binding service order under the MSA" + (", setting out confirmed capacity, pricing, site allocation, and RFS date;" if self.t == "Wholesale" else ";")))

        if self.t == "Wholesale":
            defs.append(('"Services"', "means the AI colocation services to be provided by the Provider, comprising power supply and distribution, liquid cooling infrastructure, physical security, building management, and facility operations;"))
        elif self.t == "EndUser":
            defs.append(('"Services"', "means the AI compute infrastructure services to be provided by the Provider, the scope of which will depend on the service model selected under Clause 3.1."))

        if self.t in ("Distributor", "Wholesale"):
            defs.append(('"Site Identifier"', "means any information that identifies or could reasonably be used to identify a specific DEC location, including: address, GPS coordinates, cadastral reference, project codename, photographs, aerial imagery, file names, folder names, metadata, and any information derived from the foregoing;"))

        if self.t == "Distributor":
            defs.append(('"Transaction"', "means the proposed strategic partnership between the Parties as described in Clause 3;"))

        if self.t in ("Distributor", "Wholesale"):
            defs.append(('"Whitespace"', "means designated floor area within a DEC available for the deployment of IT hardware and supporting infrastructure."))

        for label, text in defs:
            if label:
                self.bp(label + " ", text)
            else:
                self.p(text)

    def clause2(self):
        self.h("2. Purpose and Scope")
        if self.t == "Distributor":
            self.p(f"2.1 This LOI records the Parties' mutual interest in establishing a strategic partnership under which the {self.party} would participate in the delivery of AI colocation services to end-user customers through the Provider's DEC platform, on the indicative terms set out in Clauses 3 and 4.")
        elif self.t == "Wholesale":
            self.p(f"2.1 This LOI records the Parties' mutual interest in the {self.party} procuring dedicated AI colocation capacity at the Provider's DEC facilities, on the indicative terms set out in Clauses 3 and 4.")
        else:
            self.p(f"2.1 This LOI records the Parties' mutual interest in the {self.party} procuring AI compute infrastructure services at the Provider's DEC facilities, on the indicative terms set out in Clauses 3 and 4.")

        if self.t in ("Distributor", "Wholesale"):
            self.p("2.2 The Parties intend this LOI to provide " + ("a framework for further commercial discussion and negotiation" if self.t == "Distributor" else "the basis for technical scoping and commercial negotiation") + " toward a definitive MSA. The commercial terms in Clauses 3 and 4 are non-binding expressions of intent.")
            self.p("2.3 The confidentiality, non-circumvention, and general provisions in Clauses 5 through 8 are legally binding and enforceable from the date of execution.")
        else:
            self.p("2.2 The commercial terms in Clauses 3 and 4 are non-binding expressions of intent. The confidentiality and general provisions in Clauses 5 through 7 are legally binding and enforceable from the date of execution.")

    def clause3_ds(self):
        mode = self.d.get("partnership_mode", "combined")
        heading = "3. Partnership and Combined Offering (NON-BINDING)" if mode == "combined" else "3. Partnership and Referral Arrangement (NON-BINDING)"
        self.h(heading)
        comm = self.d.get("commercial", {})
        mode = self.d.get("partnership_mode", "combined")

        # 3.1
        self.bp("3.1 Partnership Overview. ",
                f"{self.g('counterparty', 'short')} provides {comm.get('partner_core_capability', '')}. "
                f"The Provider develops and operates purpose-built AI colocation facilities with secured "
                f"energy and grid access, liquid cooling, and full energy recovery infrastructure.")
        if mode == "combined":
            self.p(
                f"The Parties intend to combine the {self.party}'s {comm.get('partner_contribution', '')} "
                f"with the Provider's colocation platform to deliver {comm.get('combined_offering', '')} "
                f"to {comm.get('target_end_users', '')}."
            )
        else:
            self.p(
                f"The Parties intend to establish a referral arrangement under which the {self.party} "
                f"would introduce qualified end-user customers to the Provider's DEC platform, leveraging "
                f"the {self.party}'s established client relationships and market presence."
            )

        # 3.2
        if mode == "combined":
            self.bp("3.2 Combined Offering. ", "Under the envisaged partnership:")
            self.p("(a) the Provider would supply dedicated AI colocation capacity at its DEC facilities, including power distribution, liquid cooling, physical security, and facility operations;")
            self.p(f"(b) the {self.party} would contribute {comm.get('partner_service_scope', '')}; and")
            self.p(f"(c) the end-user customer would procure a single, integrated solution through the {self.party}, with the Provider's DEC platform as the infrastructure foundation.")
            self.p("The precise service boundaries, responsibility matrix, SLA allocation, and commercial terms of the combined offering will be defined in the MSA.")
        else:
            self.bp("3.2 Referral Arrangement. ",
                     f"The {self.party} would identify and introduce qualified end-user customers to the "
                     f"Provider's DEC platform, leveraging the {self.party}'s market presence, customer "
                     f"relationships, and domain credibility. The Provider would manage all commercial, "
                     f"technical, and contractual relationships with introduced customers directly. The "
                     f"{self.party}'s ongoing role following an introduction \u2014 including relationship "
                     f"support, account management, or technical liaison \u2014 will be agreed on a "
                     f"case-by-case basis and set out in the MSA or a separate Referral Agreement.")
            self.p("The economic terms of the referral arrangement, including qualifying introduction criteria, fee structure, payment terms, and duration, will be set out in a separate Referral Agreement between the Parties.")

        # 3.3
        territory = comm.get("territory", "")
        segments = comm.get("target_segments", "")
        mw = comm.get("indicative_mw", "")
        self.bp("3.3 Commercial Scope. ", "The Parties intend the partnership to address the following market:")
        self.p(f"(a) Territory: {territory};")
        self.p(f"(b) Target segments: {segments}; and")
        self.p(f"(c) Estimated capacity: The {self.party} has indicated an estimated annual capacity requirement of approximately {mw} MW IT across its end-user customer base. This estimate is non-binding and is provided to inform capacity planning.")

        # 3.4
        if self.choice("pricing"):
            self.bp("3.4 Indicative Economics. ", "Based on preliminary discussions, the Parties envisage the following indicative economic framework:")
            ch = self.d.get("choices", {})
            self.table(
                ["Parameter", "Indicative Terms"],
                [
                    ["Partner capacity rate", f"EUR {ch.get('partner_rate', '')} per kW per month"],
                    ["Commercial structure", ch.get("economics_description", "")],
                    ["Minimum annual commitment", f"{ch.get('min_commitment', '')} MW IT"],
                ]
            )
            self.p("All economic terms are indicative and subject to the terms of the MSA.")
        else:
            self.bp("3.4 Indicative Economics. ",
                     "The economic framework for the partnership will be determined during commercial scoping and set out in the MSA. The Provider will issue indicative terms upon agreement of partnership scope, capacity estimates, and contract structure.")

        # 3.5
        if self.choice("exclusivity"):
            scope = self.d.get("choices", {}).get("exclusivity_scope", "")
            self.bp("3.5 Preferred Partner. ",
                     f"Subject to execution of the MSA and achievement of the minimum capacity commitments set out therein, the Provider intends to designate the {self.party} as a preferred partner within the following scope: {scope}. This designation means the Provider will offer the {self.party} priority participation in opportunities within the defined scope before engaging other channel partners for the same opportunity. This LOI does not create any exclusivity obligation; the terms of any such designation will be agreed in the MSA.")
        else:
            self.bp("3.5 Non-Exclusivity. ", "The partnership contemplated by this LOI is non-exclusive. Both Parties retain the right to enter into similar arrangements with third parties.")

    def clause3_ws(self):
        self.h("3. Indicative Capacity and Commercial Terms (NON-BINDING)")
        comm = self.d.get("commercial", {})
        ch = self.d.get("choices", {})
        blocks = comm.get("dec_block_count", "")
        mw = comm.get("indicative_mw", "")

        self.bp("3.1 Capacity Requirement. ", f"The {self.party} has indicated interest in approximately {blocks} DEC Blocks ({mw} MW IT) of purpose-built AI colocation capacity. Capacity is provisioned in whole DEC Block increments of 1.2 MW IT each; partial requirements are rounded up to the nearest full DEC Block.")

        self.bp("3.2 Technical Specification. ", "All DEC facilities are designed for high-density AI compute workloads and are expected to include, at minimum: facility power supply and distribution, direct liquid cooling infrastructure supporting rack densities of 40 kW and above, building management, physical security, and 24/7 facility operations. The exact capacity, rack configuration, power density, and cooling requirements will be determined during the technical scoping phase following this LOI.")

        if self.choice("phasing"):
            self.bp("3.3 Deployment Phasing. ", f"The {self.party} has indicated the following high-level phasing interest:")
            phases = ch.get("phases", [])
            rows = [[f"Phase {i+1}", f"{p.get('mw', '')} MW IT", p.get("timeline", "")] for i, p in enumerate(phases)]
            self.table(["Phase", "Approximate Capacity", "Indicative Timeline"], rows)

        exp = comm.get("expansion_mw", "")
        self.bp("3.4 Expansion. ", f"The {self.party} has expressed interest in future expansion to approximately {exp} MW IT, subject to availability, commercial agreement, and the terms of the MSA. The Provider will use reasonable endeavours to accommodate expansion requirements within its DEC programme.")

        if self.choice("pricing"):
            self.bp("3.5 Indicative Pricing. ", f"Based on the {self.party}'s indicated capacity and term preferences, the Provider's indicative pricing is as follows:")
            self.table(
                ["Parameter", "Indicative Terms"],
                [
                    ["Base colocation fee", f"EUR {ch.get('base_price', '')} per kW per month"],
                    ["Contract term", f"{ch.get('term_years', '')} years"],
                    ["Power", "Metered IT Load (kWh) x PUE x Energy Rate (EUR/kWh) \u2014 billed additionally"],
                    ["Price escalation", "Subject to annual escalation, the mechanism for which will be agreed in the MSA"],
                    ["Indicative RFS", ch.get("target_rfs_date", "")],
                    ["Ramp schedule", ch.get("ramp_schedule", "")],
                ]
            )
            self.p(f"All commercial terms are indicative and subject to the outcome of the technical scoping phase, the {self.party}'s final capacity requirements, contract term, and credit profile.")
        else:
            self.bp("3.5 Indicative Pricing. ", f"Pricing for the Services will be determined following the technical scoping phase and set out in the Sales Order Form. The Provider will provide indicative pricing upon completion of the {self.party}'s capacity, density, and term requirements.")

        term = comm.get("min_term", "")
        self.bp("3.6 Term and Commitment. ", f"The {self.party} has indicated willingness to enter into a minimum commitment term of {term} under the MSA. The {self.party} acknowledges that the Provider intends the MSA to include take-or-pay provisions commensurate with the committed capacity, the terms of which will be negotiated in good faith.")

        self.bp("3.7 Credit Assessment. ", f"The Provider will complete a credit assessment of the {self.party} (or the entity that will execute the MSA) as part of the commercial process. The {self.party} agrees to cooperate with such assessment, which may include provision of audited financial statements, credit references, evidence of parent company support, or other financial information as reasonably requested. Where the {self.party} does not hold an investment-grade credit rating (or equivalent), the Parties will discuss appropriate credit support mechanisms, which may include a parent company guarantee, security deposit, or letter of credit.")

        self.bp("3.8 Site Allocation. ", f"The Provider is developing DEC facilities across multiple locations. The Provider will allocate capacity to the {self.party} based on the DEC(s) best suited to the {self.party}'s requirements, taking into account development readiness, grid availability, RFS timeline, and the {self.party}'s latency and connectivity needs. Site allocation will be confirmed during the technical scoping phase and formalised in the Sales Order Form.")

    def clause3_eu(self):
        self.h("3. Service Requirements (NON-BINDING)")
        comm = self.d.get("commercial", {})
        types = comm.get("service_type", [])

        self.bp("3.1 Service Model. ", f"The {self.party} has indicated interest in the following service model:")

        descs = {
            "bare_metal": (
                "Bare Metal Colocation \u2014 Dedicated, liquid-cooled rack space within a DEC, supporting "
                "densities of 40 kW per rack and above. The Provider supplies power distribution, direct "
                "liquid cooling, physical security, and 24/7 facility operations. The Customer deploys and "
                "manages its own hardware and software. The demarcation point is at the rack: everything "
                "below (facility infrastructure) is the Provider's responsibility; everything above (compute "
                "hardware, operating system, workloads) is the Customer's. Billed monthly on a per-kW basis "
                "with a committed term."
            ),
            "shared_cloud": (
                "Shared Cloud \u2014 Managed GPU compute capacity hosted on sovereign European infrastructure "
                "at the Provider's DEC facilities. The Customer accesses reserved or on-demand compute "
                "resources without procuring, deploying, or managing hardware. The Provider or a designated "
                "delivery partner operates the compute platform, including hardware lifecycle, scheduling, "
                "and platform software. The Customer manages workloads, data, and application layers. Billed "
                "on reserved capacity or metered usage, with a committed term or minimum spend."
            ),
            "tokens": (
                "Token-Based GPU Access \u2014 Flexible, on-demand access to GPU compute measured in GPU-hours "
                "or equivalent units. The Customer purchases compute tokens \u2014 pre-paid or pay-as-you-go "
                "\u2014 redeemable against available capacity across the Provider's DEC network. No dedicated "
                "hardware allocation or minimum infrastructure commitment. The Provider or a designated "
                "delivery partner manages all infrastructure and scheduling. The Customer submits workloads "
                "and consumes capacity as needed. Lowest entry threshold; designed for variable or exploratory "
                "workloads."
            ),
        }
        for st in types:
            if st in descs:
                self.p(descs[st])

        if any(st in ("shared_cloud", "tokens") for st in types):
            self.p("For Shared Cloud and Token-Based models: the Provider may deliver these services in partnership with a qualified delivery partner. The specific delivery partner, platform, and commercial terms will be confirmed in the MSA.", italic=True)

        cap = comm.get("indicative_capacity", "")
        self.bp("3.2 Indicative Capacity. ", f"The {self.party} has indicated interest in approximately {cap} of compute capacity. The exact capacity, configuration, and technical requirements will be determined during the technical scoping phase.")

        self.bp("3.3 Technical Requirements. ", f"The Provider's DEC facilities support rack densities of 40 kW and above with direct liquid cooling, designed for GPU-intensive training and inference workloads. The {self.party}'s specific requirements \u2014 including GPU type, rack density, network connectivity, and data sovereignty needs \u2014 will be confirmed during technical scoping and formalised in the MSA.")

        if self.choice("pricing"):
            ch = self.d.get("choices", {})
            self.bp("3.4 Indicative Pricing. ", "Based on preliminary discussions, the Provider's indicative pricing is as follows:")
            self.table(
                ["Parameter", "Indicative Terms"],
                [
                    ["Service rate", f"EUR {ch.get('base_price', '')} {ch.get('pricing_unit', 'per kW per month')}"],
                    ["Contract term", ch.get("term_years", "")],
                    ["Indicative RFS", ch.get("target_rfs_date", "")],
                ]
            )
            self.p(f"All terms are indicative and subject to the outcome of technical scoping and the {self.party}'s final requirements.")
        else:
            self.bp("3.4 Indicative Pricing. ", f"Pricing will be determined following the technical scoping phase and set out in the MSA. The Provider will provide indicative pricing upon agreement of the {self.party}'s capacity, density, and service model requirements.")

        term = comm.get("min_term", "")
        self.bp("3.5 Term. ", f"The {self.party} has indicated willingness to enter into a minimum commitment of {term} under the MSA.")

    def clause4(self):
        if self.t == "Distributor":
            self.h("4. Relationship Structure and Protection (NON-BINDING)")
            pbi = self.g("protection", "pbi_survival", default="10 years")
            self.bp("4.1 Protected Business Information. ",
                     f"The {self.party} acknowledges that the Provider's competitive position depends on the confidentiality of its Protected Business Information as defined in Clause 1. The {self.party} agrees that, for a period of {pbi} from the date of this LOI, it shall not directly or indirectly use any PBI to replicate, reverse-engineer, or independently develop any material element of the Provider's business model, site-sourcing methodology, energy procurement strategy, or commercial framework, whether for itself or for any third party. This obligation survives termination or expiry of this LOI and of any MSA.")
            self.bp("4.2 Change of Control. ",
                     f"If, during the term of this LOI, a Competitor acquires Control (as defined in Clause 1.1, \u201cAffiliate\u201d) of the {self.party}, the Provider may terminate this LOI by written notice with immediate effect. Upon such termination, the confidentiality and non-circumvention obligations in Clauses 6 and 7 shall continue for their stated survival periods.")
            self.bp("4.3 Associated Counterparties. ",
                     f"The {self.party} acknowledges that the Provider maintains relationships with Associated Counterparties at each DEC site. The sharing of any Site Identifier by the Provider shall be deemed an introduction of all Associated Counterparties for that site for the purposes of Clause 7.")
            self.bp("4.4 Governance. ", "The Parties intend to establish a joint steering committee or equivalent governance mechanism, the terms of which will be set out in the MSA.")
            self.bp("4.5 Implementation Roadmap. ", "Following execution of this LOI, the Parties intend to proceed as follows:")
            self.p("(a) Commercial scoping (target: 30 days post-LOI) \u2014 The Parties will define the partnership scope, target customer segments, capacity estimates, and economic framework.", bold=False)
            self.p("(b) Draft Partnership Agreement (target: 60 days post-LOI) \u2014 The Provider will issue a draft MSA or partnership agreement incorporating the agreed commercial terms.")
            self.p("(c) MSA negotiation and execution (target: 90 days post-LOI) \u2014 The Parties will negotiate and execute the MSA.")
            self.p("These timelines are indicative and non-binding.")

        elif self.t == "Wholesale":
            self.h("4. Relationship Structure and Next Steps (NON-BINDING)")
            self.bp("4.1 MSA Structure. ", "The Parties intend to execute a Master Services Agreement incorporating: (a) a Pricing Framework setting out the commercial terms for the Services; (b) Sales Order Forms for each capacity commitment; and (c) an SLA Schedule defining availability, performance, and remediation commitments. The MSA will be the sole binding commercial agreement between the Parties.")
            self.bp("4.2 Revenue Chain. ", "The Parties acknowledge the intended contractual chain: this LOI (non-binding commercial intent) \u2192 Sales Order Form (binding capacity commitment and pricing) \u2192 MSA (definitive commercial terms). Each stage is designed to provide increasing commercial certainty and to support the Provider's project finance activities.")
            self.bp("4.3 Direct Agreement Willingness. ", f"The {self.party} confirms its willingness, subject to commercially reasonable terms, to enter into a direct agreement with the Provider's Financing Parties if requested under Clause 5.3. The {self.party} acknowledges that its cooperation in this regard materially supports the Provider's ability to deliver the committed capacity on the indicative timeline.")
            self.bp("4.4 Expansion and Priority. ", f"The Provider will offer the {self.party} priority access to additional capacity within the DEC(s) allocated to the {self.party}, subject to availability. Expansion terms will be governed by the MSA.")
            self.bp("4.5 Implementation Roadmap. ", "Following execution of this LOI, the Parties intend to proceed as follows:")
            self.p("(a) Technical scoping (target: 30 days post-LOI) \u2014 Detailed technical discovery to determine exact capacity, rack layout, power distribution, cooling specifications, and connectivity requirements.")
            self.p(f"(b) Credit assessment (target: 30 days post-LOI, concurrent with technical scoping) \u2014 The Provider will complete a credit assessment of the {self.party} or the entity that will execute the MSA.")
            self.p("(c) Sales Order Form (target: 60 days post-LOI) \u2014 Upon completion of technical scoping, the Provider will issue a Sales Order Form setting out confirmed commercial terms, facility specifications, and service level framework.")
            self.p("(d) MSA negotiation and execution (target: 90 days post-LOI) \u2014 The Parties will negotiate and execute the MSA.")
            self.p("These timelines are indicative and non-binding.")
            self.bp("4.6 No Capacity Reservation. ", "This LOI does not reserve capacity at any DEC. Capacity allocation is on a first-come, first-served basis and will only be confirmed upon execution of the MSA.")

        else:  # EndUser
            self.h("4. Next Steps (NON-BINDING)")
            self.p("4.1 Following execution of this LOI, the Parties intend to proceed as follows:")
            self.p("(a) Technical scoping (target: 30 days post-LOI) \u2014 Detailed discovery to determine capacity, configuration, connectivity, and service model requirements.")
            self.p("(b) Commercial proposal (target: 60 days post-LOI) \u2014 The Provider will issue a commercial proposal setting out confirmed pricing, service scope, and SLA framework.")
            self.p("(c) MSA negotiation and execution (target: 90 days post-LOI) \u2014 The Parties will negotiate and execute the MSA.")
            self.p("These timelines are indicative and non-binding.")
            self.p("4.2 This LOI does not reserve capacity at any DEC. Capacity allocation will be confirmed upon execution of the MSA.")

    def clause5(self):
        self.h("5. Project Finance and Assignment (BINDING)")
        self.bp("5.1 Revenue Bankability. ", "The Parties acknowledge that the Provider intends to finance the development and operation of its DEC programme through a combination of equity investment and non-recourse project finance. The Provider's ability to secure favourable financing terms depends in part on demonstrating contracted or committed revenue streams. This LOI, while non-binding in its commercial terms, is intended to evidence the Parties' genuine commercial intent and to support the Provider's financing activities.")
        self.bp("5.2 Assignment. ", f"Neither Party may assign its rights or obligations under this LOI without the prior written consent of the other Party, except that:")
        self.p("(a) the Provider may assign this LOI, or any rights under it, to any Financing Party or security trustee as security for project finance obligations; and")
        self.p(f"(b) the Provider may assign this LOI to any Affiliate or special-purpose vehicle within its corporate group without the {self.party}'s consent, provided the Provider remains liable for the performance of the assignee's obligations.")
        self.bp("5.3 Lender Acknowledgment. ", f"The {self.party} acknowledges and agrees that, upon the Provider's written request, the {self.party} shall negotiate in good faith and execute a direct agreement (or lender acknowledgment letter) with the Provider's Financing Party within 30 Business Days of such request. Such direct agreement may include, as is customary in project finance transactions: (a) step-in rights for the Financing Party upon a Provider default; (b) cure periods in favour of the Financing Party; and (c) information rights enabling the Financing Party to monitor the commercial relationship. The terms of any such direct agreement shall be commercially reasonable and consistent with market practice for project finance transactions.")

    def clause6(self):
        self.h(f"6. Confidentiality {'and Non-Disclosure ' if self.t != 'EndUser' else ''}(BINDING)")
        existing = self.choice("existing_nda")

        if existing:
            nda_date = self.g("choices", "nda_date")
            transaction_suffix = " and the Transaction." if self.t == "Distributor" else (" and the proposed transaction." if self.t == "Wholesale" else ".")
            self.p(f'6.1 The Non-Disclosure Agreement between the Parties dated {nda_date} (the "NDA") remains in full force and effect and applies to all information exchanged in connection with this LOI{transaction_suffix}')
            self.p("6.2 The existence and contents of this LOI shall be treated as Confidential Information under the NDA.")
            if self.t != "EndUser":
                self.p("6.3 To the extent of any conflict between the NDA and this LOI, the provisions of this LOI shall prevail.")
                transaction_ref = "the Transaction" if self.t == "Distributor" else "the proposed transaction"
                self.p(f"6.4 Neither Party shall make any public announcement regarding this LOI or {transaction_ref} without the prior written consent of the other Party, except as required by applicable law.")
            else:
                self.p("6.3 Neither Party shall make any public announcement regarding this LOI without the prior written consent of the other Party, except as required by applicable law.")
        else:
            surv = self.g("protection", "confidentiality_survival", default="3 years")
            if self.t == "EndUser":
                # Tier A — 8 clauses
                self.p(f"6.1 Each Party shall keep confidential all Confidential Information received from the other Party and shall use such information solely for the Purpose.")
                self.p("6.2 Each Party may disclose Confidential Information only to its Representatives who have a genuine need to know for the Purpose and who are bound by confidentiality obligations no less restrictive than this Clause 6.")
                self.p("6.3 The obligations in Clauses 6.1 and 6.2 do not apply to information that the receiving Party can demonstrate: (a) is or becomes publicly available through no fault of the receiving Party; (b) was already in the receiving Party's possession without restriction; (c) was independently developed without use of the Confidential Information; or (d) was received from a third party without breach of any confidentiality obligation.")
                self.p("6.4 A Party may disclose Confidential Information to the extent required by applicable law, regulation, or court order, provided that (where legally permitted) the disclosing Party gives the other Party prior written notice and discloses only the minimum required.")
                self.p("6.5 Upon written request or expiry of this LOI, the receiving Party shall promptly return or destroy all Confidential Information and certify such return or destruction in writing. The receiving Party may retain copies required by applicable law or internal compliance policies, subject to continued confidentiality.")
                self.p("6.6 Neither Party shall make any public announcement regarding this LOI without the prior written consent of the other Party, except as required by applicable law.")
                self.p(f"6.7 The obligations in this Clause 6 shall survive for {surv} from the date of termination or expiry of this LOI. Obligations in respect of trade secrets shall continue indefinitely.")
                self.p("6.8 Each Party acknowledges that a breach of this Clause 6 may cause irreparable harm for which damages would not be an adequate remedy. The disclosing Party shall be entitled to seek injunctive or other equitable relief without the need to prove actual loss.")
            else:
                # Tier B — 16 clauses
                self.bp("6.1 Purpose Limitation. ", "Each Party shall use the other Party's Confidential Information solely for the Purpose and for no other purpose.")
                self.bp("6.2 Non-Disclosure. ", "Each Party shall keep confidential all Confidential Information received from the other Party and shall not disclose such information to any person except as permitted under this Clause 6.")
                self.bp("6.3 Standard of Care. ", "Each Party shall apply no less than reasonable care to protect the other Party's Confidential Information, and no less than the care it applies to its own confidential information of a similar nature.")
                self.bp("6.4 Permitted Disclosures. ", "A Party may disclose Confidential Information to:")
                self.p("(a) its Representatives who have a genuine need to know for the Purpose and who are bound by confidentiality obligations no less restrictive than this Clause 6 (whether by professional duty or written undertaking);")
                self.p("(b) its bona fide Financing Parties and potential co-investors, provided they are bound by confidentiality obligations no less restrictive than this Clause 6; and")
                self.p("(c) to the extent required by applicable law, regulation, court order, or the rules of any relevant regulatory authority or stock exchange, provided that (where legally permitted) the disclosing Party: (i) gives the other Party prior written notice as soon as reasonably practicable; (ii) consults with the other Party regarding the scope and manner of disclosure; and (iii) discloses only the minimum information required to comply.")
                self.bp("6.5 Liability for Representatives. ", "Each Party shall be responsible for any breach of this Clause 6 by its Representatives.")
                self.bp("6.6 Exclusions. ", "The obligations in Clauses 6.1 through 6.3 do not apply to information that the receiving Party can demonstrate:")
                self.p("(a) is or becomes publicly available through no fault of the receiving Party or its Representatives;")
                self.p("(b) was already in the lawful possession of the receiving Party before disclosure, without restriction as to use or disclosure;")
                self.p("(c) was independently developed by the receiving Party without use of or reference to the Confidential Information; or")
                self.p("(d) was received from a third party who was not, to the receiving Party's knowledge, under any obligation of confidentiality in respect of that information.")
                self.bp("6.7 No Implied Rights. ", "No licence or right is granted under this LOI to the receiving Party in respect of any intellectual property rights of the disclosing Party. All Confidential Information remains the property of the disclosing Party.")
                self.bp("6.8 Return and Destruction. ", "Upon the earlier of: (a) the disclosing Party's written request, or (b) the expiry or termination of this LOI, the receiving Party shall promptly return or destroy all documents, materials, and tangible items containing Confidential Information and certify such return or destruction in writing within 15 Business Days. The receiving Party may retain copies to the extent required by applicable law or its internal compliance policies, provided such retained copies remain subject to this Clause 6.")
                self.bp("6.9 Onward-Sharing Controls. ", "If the receiving Party receives an inquiry from any third party regarding the disclosing Party's Confidential Information, the receiving Party shall: (a) not respond to such inquiry without the disclosing Party's prior written consent; and (b) promptly notify the disclosing Party of such inquiry. The receiving Party shall not further distribute or re-disclose Confidential Information beyond the persons authorised under Clause 6.4 without the disclosing Party's prior written consent.")
                self.bp("6.10 Compliance Confirmation. ", "Upon the disclosing Party's reasonable written request (not more than once per calendar year), the receiving Party shall confirm in writing its compliance with the obligations in this Clause 6.")
                self.bp("6.11 Breach Notification. ", "Each Party shall notify the other Party in writing within 72 hours of becoming aware of any actual or suspected breach of this Clause 6, and shall take all reasonable steps to mitigate the effects of such breach.")
                self.bp("6.12 Disclaimer. ", 'All Confidential Information is disclosed "as is." The disclosing Party makes no representation or warranty, express or implied, as to the accuracy, completeness, or reliability of any Confidential Information. The receiving Party shall be solely responsible for its own assessment and due diligence.')
                self.bp("6.13 Metadata Protection. ", "Confidential Information includes metadata, EXIF data, geolocation data, timestamps, file names, folder names, and any digital artifacts associated with or derived from disclosed materials. The receiving Party shall not extract, analyse, or use such metadata except as necessary for the Purpose.")
                surv_text = f"6.14 Survival. The obligations in this Clause 6 shall survive termination or expiry of this LOI for a period of {surv} from the date of termination or expiry. Obligations in respect of information that constitutes a trade secret under applicable law shall continue indefinitely."
                if self.t == "Distributor":
                    surv_text += " Obligations in respect of Protected Business Information shall survive for the period specified in Clause 4.1."
                self.bp("6.14 Survival. ", surv_text.replace("6.14 Survival. ", ""))
                transaction_ref = "the Transaction" if self.t == "Distributor" else "the proposed transaction"
                self.bp("6.15 Announcements. ", f"Neither Party shall make any public announcement regarding this LOI or {transaction_ref} without the prior written consent of the other Party, except as required by applicable law.")
                self.bp("6.16 Remedies. ", "Each Party acknowledges that a breach of this Clause 6 may cause the disclosing Party irreparable harm for which damages would not be an adequate remedy. The disclosing Party shall be entitled to seek injunctive or other equitable relief from any court of competent jurisdiction, without the need to prove actual loss and without prejudice to any other rights or remedies.")

    def clause7_nc(self):
        if self.t == "EndUser":
            return  # No NC for End Users

        self.h("7. Non-Circumvention (BINDING)")
        nc_dur = self.g("protection", "nc_duration", default="24 months")

        self.p(f"7.1 The {self.party} shall not, directly or indirectly, without the prior written consent of the Provider:")

        if self.t == "Distributor":
            self.p(f"(a) contact, solicit, deal with, or enter into any business relationship with any Associated Counterparty introduced by the Provider in connection with the Purpose or the Transaction;")
            self.p(f"(b) circumvent, avoid, or bypass the Provider in order to deal directly or indirectly with any Associated Counterparty; or")
            self.p(f"(c) attempt to divert or appropriate any business opportunity disclosed by the Provider in connection with the Purpose or the Transaction.")
        else:  # Wholesale — lighter scope
            self.p(f"(a) contact, solicit, deal with, or enter into any business relationship with any Associated Counterparty introduced by the Provider in connection with this LOI or the Services; or")
            self.p(f"(b) circumvent, avoid, or bypass the Provider in order to deal directly or indirectly with any Associated Counterparty in connection with the development, ownership, or operation of colocation or energy infrastructure on or adjacent to a site identified by the Provider.")

        self.bp("7.2 Duration. ", f"The obligations in this Clause 7 shall continue for {nc_dur} after the earlier of: (a) the expiry or termination of this LOI; or (b) the expiry or termination of the MSA, if one is executed.")
        self.bp("7.3 Deemed Introduction. ", f"The Provider's sharing of any Site Identifier with the {self.party} shall constitute a deemed introduction of all Associated Counterparties for that site. The Provider is not required to separately name each Associated Counterparty.")

        if self.t == "Distributor":
            self.bp("7.4 Scope \u2014 Private and Public Bodies. ", "The non-circumvention obligations in Clause 7.1 apply to:")
            self.p("(a) Private Associated Counterparties (landowners, greenhouse operators, heat offtakers, energy counterparties, EPCs): in all cases where the Provider has introduced or disclosed their identity or involvement; and")
            self.p("(b) Public Bodies (government agencies, municipalities, regulatory bodies, distribution system operators): only where the Provider has made a specific, named introduction to an individual or department within that body. General knowledge that the Provider engages with a public body does not trigger non-circumvention protection.")
        else:
            self.bp("7.4 Scope Limitation. ", f"The non-circumvention obligations in this Clause 7 are limited to the Provider's supply-side relationships (site partners, energy counterparties, and infrastructure providers). Nothing in this Clause 7 restricts the {self.party} from conducting its own business with its existing or future end-user customers, cloud service customers, or compute buyers.")

        self.bp("7.5 Independent Knowledge Exception. ", f"The obligations in Clause 7.1 do not apply to any Associated Counterparty with whom the {self.party} can demonstrate, by contemporaneous written evidence, that it had an existing business relationship or substantive commercial contact before the date of this LOI or before the Provider's disclosure.")
        self.bp("7.6 MSA Supersession. ", "If the Parties execute an MSA, the non-circumvention provisions of the MSA shall replace this Clause 7 upon execution of the MSA, except that the survival period in Clause 7.2 shall apply to any introduction made before the MSA effective date that is not separately covered by the MSA.")

    def clause_general(self):
        cl = "8" if self.t != "EndUser" else "7"
        self.h(f"{cl}. General Provisions (BINDING)")
        val_date = self.g("dates", "validity_date")
        if not val_date:
            val_date = "[12 months from LOI date]"

        self.bp(f"{cl}.1 Non-Binding Status. ", "")
        if self.t != "EndUser":
            self.p(f"(a) Non-binding provisions. Clauses 2 through 4 and Schedule 1 of this LOI are non-binding expressions of the Parties' current intentions. They do not create legally enforceable obligations and are subject to the negotiation and execution of the MSA.")
            self.p(f"(b) Binding provisions. Clauses 5 (Project Finance and Assignment), 6 (Confidentiality), 7 (Non-Circumvention), and {cl} (General Provisions) are legally binding and enforceable obligations.")
        else:
            self.p("(a) Non-binding provisions. Clauses 2 through 4 of this LOI are non-binding expressions of the Parties' current intentions. They do not create legally enforceable obligations and are subject to the negotiation and execution of the MSA.")
            self.p(f"(b) Binding provisions. Clauses 5 (Project Finance and Assignment), 6 (Confidentiality), and {cl} (General Provisions) are legally binding and enforceable obligations.")

        scoping_phrase = "commercial scoping and negotiation" if self.t == "Distributor" else "technical scoping and commercial negotiation"
        self.bp(f"{cl}.2 Good Faith. ", f"The Parties agree to engage in the {scoping_phrase} process in good faith (te goeder trouw) and in accordance with the principles of reasonableness and fairness (redelijkheid en billijkheid) as contemplated by Article 6:248 of the Dutch Civil Code (Burgerlijk Wetboek). " + ("" if self.t == "EndUser" else "For the avoidance of doubt, t") + ("T" if self.t == "EndUser" else "t") + f"he good faith obligation does not oblige either Party to enter into the MSA. Either Party may discontinue negotiations at any time, provided it does so in good faith. Any liability arising from a breach of this good faith obligation shall be limited to verifiable reliance damages (negatief contractsbelang)" + ("." if self.t == "EndUser" else " and shall not extend to loss of profit or expectation damages (positief contractsbelang)."))

        survive = "Clauses 5.2, 5.3, 6, and 7" if self.t != "EndUser" else "Clauses 5.2, 5.3, and 6"
        self.bp(f"{cl}.3 Validity. ", f"This LOI shall remain valid until {val_date}, after which it shall lapse automatically unless extended by mutual written agreement. Upon lapse, {survive} shall survive for their respective stated periods.")
        self.bp(f"{cl}.4 Costs. ", "Each Party shall bear its own costs in connection with this LOI" + (" and the negotiation of the MSA." if self.t != "EndUser" else "."))
        self.bp(f"{cl}.5 Counterparts. ", "This LOI may be executed in counterparts, including by electronic signature within the meaning of the eIDAS Regulation (EU) No 910/2014." + (" Each counterpart constitutes an original." if self.t != "EndUser" else ""))
        self.bp(f"{cl}.6 Notices. ", "All notices under this LOI shall be in writing (including email) and addressed to the contact details in the preamble. A notice is effective upon receipt." + (" Each Party shall promptly notify the other of any change to its contact details." if self.t != "EndUser" else ""))
        self.bp(f"{cl}.7 Governing Law. ", "This LOI shall be governed by and construed in accordance with the laws of the Netherlands. The United Nations Convention on Contracts for the International Sale of Goods (CISG) is expressly excluded.")
        self.bp(f"{cl}.8 Jurisdiction. ", "The courts of Amsterdam (Rechtbank Amsterdam) shall have exclusive jurisdiction over any dispute arising out of or in connection with the binding provisions of this LOI.")

        if self.t != "EndUser":
            transaction_ref = "the Transaction" if self.t == "Distributor" else "the proposed transaction"
            self.bp(f"{cl}.9 Entire Agreement. ", f"This LOI, together with any NDA referenced in Clause 6 (ALT-A), constitutes the entire agreement between the Parties in relation to its subject matter and supersedes all prior negotiations, representations, and agreements relating to {transaction_ref}. Nothing in this Clause limits liability for fraud.")
            self.bp(f"{cl}.10 Partnership Disclaimer. ", "Nothing in this LOI shall be construed as creating a partnership, joint venture, agency, or employment relationship between the Parties. Neither Party has authority to bind the other or to incur any obligation on the other's behalf.")
        else:
            self.bp(f"{cl}.9 Entire Agreement. ", "This LOI constitutes the entire agreement between the Parties in relation to its subject matter and supersedes all prior negotiations, representations, and agreements. Nothing in this Clause limits liability for fraud.")

    def signature(self):
        self.line()
        closing = "We look forward to working with you."
        bespoke = self.g("choices", "bespoke_closing")
        if bespoke:
            closing += f" {bespoke}"
        self.p(closing)
        self.p("")
        self.p("Yours faithfully,")
        self.p("")

        prov = self.d.get("provider", {})
        self.p(f"For and on behalf of {prov.get('legal_name', '')}", bold=True)
        self.p(f"KvK: {prov.get('kvk', '')}")
        self.p("")
        self.p("Signature: ____________________________")
        self.p(f"Name: {prov.get('signatory_name', '')}")
        self.p(f"Title: {prov.get('signatory_title', '')}")
        self.p("Date: ____________________________")

        self.p("")
        self.line()
        self.p("")
        self.p("ACKNOWLEDGED AND AGREED:", bold=True)
        self.p("")

        cp = self.d.get("counterparty", {})
        self.p(f"For and on behalf of {cp.get('name', '')}", bold=True)
        rt = cp.get("reg_type", "")
        rn = cp.get("reg_number", "")
        if rt and rn:
            self.p(f"{rt}: {rn}")
        self.p("")
        self.p("Signature: ____________________________")
        self.p(f"Name: {cp.get('signatory_name', '')}")
        self.p(f"Title: {cp.get('signatory_title', '')}")
        self.p("Date: ____________________________")

    def schedule(self):
        # Schedule starts on its own page
        self.doc.add_page_break()
        if self.t == "Distributor":
            self.line()
            self.h("Schedule 1 \u2014 Partnership Details (NON-BINDING)")
            comm = self.d.get("commercial", {})
            self.table(
                ["Item", "Detail"],
                [
                    ["Partnership Mode", self.d.get("partnership_mode", "combined").title()],
                    ["Target Customer Segments", comm.get("target_segments", "")],
                    ["Geographic Scope", comm.get("territory", "")],
                    ["Estimated Annual Capacity", f"{comm.get('indicative_mw', '')} MW IT"],
                    ["Commercial scoping complete", f"{self.g('dates', 'loi_date')} + 30 days"],
                    ["Draft MSA issued", f"{self.g('dates', 'loi_date')} + 60 days"],
                    ["MSA executed", f"{self.g('dates', 'loi_date')} + 90 days"],
                ]
            )
        elif self.t == "Wholesale":
            self.line()
            self.h("Schedule 1 \u2014 Capacity and Technical Requirements (NON-BINDING)")
            comm = self.d.get("commercial", {})
            self.table(
                ["Item", "Detail"],
                [
                    ["Indicative Capacity", f"{comm.get('dec_block_count', '')} DEC Blocks ({comm.get('indicative_mw', '')} MW IT)"],
                    ["GPU / Accelerator Type", "[To be confirmed during technical scoping]"],
                    ["Target Rack Density", "[To be confirmed]"],
                    ["Cooling Requirement", "Direct liquid cooling"],
                    ["Technical scoping complete", f"{self.g('dates', 'loi_date')} + 30 days"],
                    ["Credit assessment complete", f"{self.g('dates', 'loi_date')} + 30 days"],
                    ["Sales Order Form issued", f"{self.g('dates', 'loi_date')} + 60 days"],
                    ["MSA executed", f"{self.g('dates', 'loi_date')} + 90 days"],
                    ["Expansion Target", f"{comm.get('expansion_mw', '')} MW IT"],
                ]
            )
        else:  # EndUser
            self.line()
            self.h("Schedule 1 \u2014 Service Requirements (NON-BINDING)")
            comm = self.d.get("commercial", {})
            types = comm.get("service_type", [])
            type_str = ", ".join(t.replace("_", " ").title() for t in types)
            self.table(
                ["Item", "Detail"],
                [
                    ["Service Model", type_str],
                    ["Indicative Capacity", comm.get("indicative_capacity", "")],
                    ["GPU / Accelerator Type", "[To be confirmed during technical scoping]"],
                    ["Data Sovereignty", "[To be confirmed]"],
                    ["Technical scoping complete", f"{self.g('dates', 'loi_date')} + 30 days"],
                    ["Commercial proposal issued", f"{self.g('dates', 'loi_date')} + 60 days"],
                    ["MSA executed", f"{self.g('dates', 'loi_date')} + 90 days"],
                ]
            )

    def footer(self):
        # Real Word footers are set up in _setup(). Add version reference at end of body.
        self.p("")
        self.p(
            f"DE-LOI-{self.t}-v3.0",
            italic=True, size=FONT_SMALL, color=GREY
        )

    def build(self) -> Document:
        self.letterhead()
        # Cover page ends here -- agreement body starts on next page
        self.doc.add_page_break()
        self.recitals()
        self.definitions()
        self.clause2()

        if self.t == "Distributor":
            self.clause3_ds()
        elif self.t == "Wholesale":
            self.clause3_ws()
        else:
            self.clause3_eu()

        self.clause4()
        self.clause5()
        self.clause6()
        self.clause7_nc()
        self.clause_general()
        self.signature()
        self.schedule()
        self.footer()
        return self.doc


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    if len(sys.argv) < 2:
        print("Usage: python generate_loi.py <intake.yaml> [--output path.docx]")
        sys.exit(1)

    data = load_intake(sys.argv[1])

    output = None
    if "--output" in sys.argv:
        idx = sys.argv.index("--output")
        if idx + 1 < len(sys.argv):
            output = sys.argv[idx + 1]

    if not output:
        loi_date = data.get("dates", {}).get("loi_date", "")
        try:
            dt = datetime.strptime(loi_date, "%d %B %Y")
            date_str = dt.strftime("%Y%m%d")
        except ValueError:
            date_str = datetime.now().strftime("%Y%m%d")
        cp = data.get("counterparty", {}).get("short", "Unknown")
        output = f"{date_str}_DEG_LOI-{data['type']}_{cp}_(DRAFT).docx"

    loi = LOI(data)
    doc = loi.build()
    doc.save(output)

    print(f"Generated: {output}")
    print(f"Type: DE-LOI-{data['type']}-v3.0")
    print(f"Counterparty: {data.get('counterparty', {}).get('name', 'Unknown')}")
    print(f"Clauses: ALL (full document)")


if __name__ == "__main__":
    main()
