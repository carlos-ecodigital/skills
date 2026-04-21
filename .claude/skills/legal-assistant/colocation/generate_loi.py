#!/usr/bin/env python3
"""
LOI/NCNDA v3.2 — Document Generation Script

Usage:
    python generate_loi.py intake.yaml [--output path/to/output.docx]
                                        [--override R-11,R-14]
                                        [--override-reason "..."]

Takes a YAML intake file and produces a complete, branded .docx ready to sign.
All clauses, conditionals, and schedules are generated from the intake data.

v3.2 changes (2026-04-16):
- Recital A is library-sourced (see _shared/loi-recital-a-library.md). Variants:
  default | sovereignty | integration | bespoke.
- Customer-facing "DEC Block" language removed. Capacity expressed in MW IT +
  Designated Sites. Deprecated YAML field: commercial.dec_block_count.
- "Minimum commitment term of 5 years" replaced with "approximately 5 years,
  indicative only".
- Cl. 4.2 "Revenue Chain" with Unicode arrows replaced with "Contractual
  Sequence" prose + numbered list.
- Closing line: hardcoded single-sentence "We look forward to working with you."
  Honors OPEN-1 (commit 2097f52) — bespoke_closing is not supported.
  choices.bespoke_closing in YAML is silently ignored for backward compat.
- Schedule 1 title no longer carries "(NON-BINDING)" suffix. Italic prefatory
  note on the schedule; authoritative binding/non-binding assignment in Cl. 5.1.
- Pre-save QA linter. Rules in _shared/loi-qa-gate.md. Severity: fail / warn /
  info. Overridable via --override / --override-reason.

Planned (not yet wired in this file): new types StrategicSupplier (DE-LOI-SS-v1.0)
and EcosystemPartnership (DE-LOI-EP-v1.0). See CHANGELOG.md.
"""

import sys
import os
import yaml
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

# Import shared cover page and branding from document-factory
# Script now lives under legal-assistant/colocation/, one level deeper than the
# original loi-generator/ layout, so resolve document-factory via parent.parent.
SCRIPT_DIR = Path(__file__).parent
sys.path.insert(0, str(SCRIPT_DIR.parent.parent / "document-factory"))
from common import (  # noqa: E402
    add_cover, Party,
    COBALT, SLATE, SLATE_800, SLATE_900, WHITE,
    setup_first_page_header, setup_cont_header,
    setup_first_footer, setup_cont_footer,
)


# ---------------------------------------------------------------------------
# Configuration -- LOI clause formatting (body text, not cover page)
# ---------------------------------------------------------------------------

FONT_NAME = "Inter"
FONT_BODY = Pt(10)
FONT_HEADING1 = Pt(13)
FONT_HEADING2 = Pt(11)
FONT_SMALL = Pt(8)
LINE_SPACING = 1.15

# Local aliases for clause body text (imported colors used for cover page)
NAVY = SLATE_900
BLACK = SLATE_800
GREY = SLATE


# ---------------------------------------------------------------------------
# YAML Loading + Validation
# ---------------------------------------------------------------------------

def load_entities_register() -> dict:
    """Load the entities register from `config/entities.yaml`.

    v3.5.2 scope Q: single source of truth for DE's legal entities. Returns
    empty dict if register file is absent (backward-compat — intake YAMLs
    can still use explicit `provider.legal_name`, `address`, `kvk` fields).
    """
    here = os.path.dirname(os.path.abspath(__file__))
    register_path = os.path.join(here, "..", "config", "entities.yaml")
    if not os.path.exists(register_path):
        return {}
    try:
        with open(register_path, "r") as f:
            return yaml.safe_load(f) or {}
    except Exception:
        return {}


def expand_provider_from_register(data: dict, register: dict) -> dict:
    """Expand `provider.entity: "de_nl"` references into full provider record.

    v3.5.2 scope Q. Behaviour:
      - If `provider.entity` is set and register has the key, expand it.
        Any fields explicitly set on the intake YAML itself (e.g. signatory
        overrides) take precedence over the register defaults.
      - If `provider.entity` is not set but the register has a
        `type_defaults` block with an entry for the intake's `type`,
        populate the entity + signatory_mode from that per-type default
        (v3.5.3-cont scope J12 wiring). A minimal intake YAML with only
        `type: Wholesale` now auto-expands to NL BV / Carlos / Director.
      - If neither is set, the provider dict is returned unchanged
        (backward-compat with v3.5.1 and earlier).
      - `signatory_mode` on the intake YAML selects the default signatory
        variant from the register (pre_msa / post_msa for NL BV; ceo for AG).

    Mutates and returns `data`.
    """
    prov = data.get("provider", {}) or {}
    entity_key = prov.get("entity")

    # v3.5.3-cont scope J12: if no explicit entity is set AND the provider
    # dict is "minimal" (contains at most entity/signatory_mode/short_name —
    # i.e. no explicit identity fields that the user clearly set themselves),
    # apply the per-LOI-type default from the register's `type_defaults`
    # block. This is a pure fallback — explicit `provider.entity` wins, and
    # explicit-fields intakes (legal_name/address/kvk/etc. set) are
    # preserved untouched for backward-compat with v3.5.1.
    if not entity_key:
        # "Identity fields" — presence of ANY of these means the user wrote
        # an explicit-fields intake and type_defaults must NOT override.
        _EXPLICIT_IDENTITY_FIELDS = (
            "legal_name", "address", "kvk", "reg_number",
            "reg_type", "jurisdiction", "legal_form",
        )
        is_minimal = not any(prov.get(f) for f in _EXPLICIT_IDENTITY_FIELDS)
        if is_minimal:
            type_defaults = register.get("type_defaults", {}) or {}
            type_key = data.get("type")
            td = type_defaults.get(type_key)
            if td and isinstance(td, dict):
                entity_key = td.get("entity")
                if entity_key and "signatory_mode" not in prov:
                    prov = dict(prov)
                    prov["signatory_mode"] = td.get("signatory_mode", "pre_msa")
                    data["provider"] = prov

    if not entity_key:
        return data
    entities = register.get("entities", {})
    entity = entities.get(entity_key)
    if not entity:
        return data

    # Build the expanded provider record. Explicit intake values WIN over
    # register defaults — this lets intake YAMLs override signatory per-deal
    # without touching the register.
    expanded = {
        "legal_name": entity.get("legal_name", ""),
        "short_name": entity.get("short_name", ""),
        "abbreviation": entity.get("abbreviation", ""),
        "legal_form": entity.get("legal_form", ""),
        "jurisdiction": entity.get("jurisdiction", ""),
        "address": entity.get("address", ""),
        "reg_type": entity.get("reg_type", ""),
        "reg_number": entity.get("reg_number", ""),
        "rsin": entity.get("rsin", ""),
        "parent": entity.get("parent"),
    }
    # Pick signatory based on signatory_mode
    mode = prov.get("signatory_mode", "pre_msa")
    sig = None
    if mode == "pre_msa" and "default_signatory_pre_msa" in entity:
        sig = entity["default_signatory_pre_msa"]
    elif mode == "post_msa" and "default_signatory_post_msa" in entity:
        sig = entity["default_signatory_post_msa"]
    elif "default_signatory" in entity:
        sig = entity["default_signatory"]
    if sig:
        expanded["signatory_name"] = sig.get("name", "")
        expanded["signatory_title"] = sig.get("title", "")

    # Apply explicit intake overrides (intake wins)
    for k, v in prov.items():
        if k in ("entity", "signatory_mode"):
            continue
        if v:
            expanded[k] = v

    data["provider"] = expanded
    return data


def load_intake(path: str) -> dict:
    with open(path, "r") as f:
        data = yaml.safe_load(f)
    # v3.5.2 scope Q: expand provider.entity references BEFORE validation so
    # downstream validators see the fully-populated provider record.
    register = load_entities_register()

    # v3.5 polish: fail fast on unknown `provider.entity` keys. Without this
    # check, typos like `entity: "de_nnl"` would silently fall through to
    # the backward-compat path (treating the YAML's explicit fields as
    # ground truth) and the user would only notice at document-render time.
    prov = data.get("provider", {}) or {}
    entity_key = prov.get("entity")
    if entity_key:
        entities = register.get("entities", {}) or {}
        if entity_key not in entities:
            available = sorted(entities.keys())
            print(
                f"ERROR: provider.entity '{entity_key}' not found in "
                f"config/entities.yaml. Available keys: "
                f"{', '.join(available) if available else '(none — register missing or empty)'}",
                file=sys.stderr,
            )
            sys.exit(1)

    data = expand_provider_from_register(data, register)
    validate(data)
    return data


def validate(d: dict):
    """Validate intake YAML. Raises SystemExit on failure.

    v3.2: rule R-18 (fail) — deprecated field commercial.dec_block_count.
    """
    errors = []
    t = d.get("type", "")
    if t not in ("EndUser", "Distributor", "Wholesale",
                  "StrategicSupplier", "EcosystemPartnership"):
        errors.append(
            "type must be one of: EndUser, Distributor, Wholesale, "
            f"StrategicSupplier, EcosystemPartnership (got: {d.get('type')})"
        )
    for s in ("provider", "counterparty", "programme", "dates"):
        if s not in d:
            errors.append(f"Missing section: {s}")
    cp = d.get("counterparty", {})
    for f in ("name", "short", "description"):
        if not cp.get(f):
            errors.append(f"counterparty.{f} required")
    if not d.get("dates", {}).get("loi_date"):
        errors.append("dates.loi_date required")

    # R-18 — deprecated field migration error
    if "dec_block_count" in d.get("commercial", {}):
        errors.append(
            "[R-18] commercial.dec_block_count is deprecated in v3.2. "
            "Use commercial.indicative_mw instead. See CHANGELOG.md § Migration."
        )

    if t == "Distributor" and d.get("partnership_mode", "combined") == "combined":
        for f in ("partner_core_capability", "partner_contribution",
                   "combined_offering", "target_end_users", "partner_service_scope"):
            if not d.get("commercial", {}).get(f):
                errors.append(f"commercial.{f} required for Distributor Mode A")
    if t == "Wholesale":
        if not d.get("commercial", {}).get("indicative_mw"):
            errors.append("commercial.indicative_mw required for Wholesale (MW IT)")
    if t == "EndUser":
        if not d.get("commercial", {}).get("service_type"):
            errors.append("commercial.service_type required for EndUser")
    if t == "StrategicSupplier":
        supplier = d.get("supplier", {})
        if not supplier.get("capability_category"):
            errors.append("supplier.capability_category required for StrategicSupplier")
        if not supplier.get("core_capability"):
            errors.append("supplier.core_capability required for StrategicSupplier")
        purposes = supplier.get("strategic_purposes", [])
        if not purposes or not (1 <= len(purposes) <= 2):
            errors.append("supplier.strategic_purposes required for StrategicSupplier (1-2 items)")
        valid_purposes = {"capacity_lock_in", "pricing_volume",
                           "supply_chain_de_risking", "engineering_integration",
                           "pipeline_visibility"}
        for p in purposes:
            if p not in valid_purposes:
                errors.append(
                    f"supplier.strategic_purposes: invalid '{p}'. "
                    f"Valid: {sorted(valid_purposes)}"
                )
        if "capacity_lock_in" in purposes and not supplier.get("lead_time_target"):
            errors.append(
                "supplier.lead_time_target required when strategic_purposes "
                "includes capacity_lock_in"
            )
        if "pricing_volume" in purposes and not supplier.get("volume_indicative"):
            errors.append(
                "supplier.volume_indicative required when strategic_purposes "
                "includes pricing_volume"
            )
        if ("engineering_integration" in purposes
                and not d.get("choices", {}).get("joint_ip")):
            errors.append(
                "choices.joint_ip required when strategic_purposes "
                "includes engineering_integration (none | background | foreground)"
            )
    if t == "EcosystemPartnership":
        eco = d.get("ecosystem", {})
        if not eco.get("relationship_type"):
            errors.append("ecosystem.relationship_type required for EcosystemPartnership")
        valid_rel = {"standards_body", "university", "research_consortium",
                      "co_marketing", "industry_association", "policy_partner",
                      "other"}
        if eco.get("relationship_type") and eco.get("relationship_type") not in valid_rel:
            errors.append(
                f"ecosystem.relationship_type must be one of {sorted(valid_rel)}"
            )
        if not eco.get("collaboration_themes"):
            errors.append("ecosystem.collaboration_themes required for EcosystemPartnership")
        if not eco.get("joint_activity_categories"):
            errors.append("ecosystem.joint_activity_categories required for EcosystemPartnership")

    # Recital A variant validation
    variant = d.get("programme", {}).get("recital_a_variant", "default")
    if variant not in ("default", "sovereignty", "integration", "bespoke"):
        errors.append(
            f"programme.recital_a_variant must be one of: default, sovereignty, "
            f"integration, bespoke (got: {variant})"
        )
    if variant == "bespoke" and not d.get("programme", {}).get("recital_a_bespoke"):
        errors.append(
            "programme.recital_a_bespoke required when recital_a_variant=bespoke"
        )

    if errors:
        print("VALIDATION ERRORS:")
        for e in errors:
            print(f"  - {e}")
        sys.exit(1)


# ---------------------------------------------------------------------------
# Recital A (v3.4 — single canonical body + per-type tails)
# ---------------------------------------------------------------------------
# Source of truth is _shared/loi-recital-a-library.md. v3.4 change: collapsed
# three variants (default / sovereignty / integration) into a single canonical
# body. User-approved wording 2026-04-17. Per-type tails now cover all 5 types
# (v3.3 only had tails for DS/SS/EP; EU/WS were empty strings).

RECITAL_A_BODY = (
    # v3.5.2 brand-name rename: defined term "Digital Energy" is established
    # in the Parties Preamble; Recital A uses it directly (no "(the "Provider")"
    # parenthetical). Any reference to "Digital Energy" in the body has been
    # replaced with "Digital Energy".
    '{prov} develops and operates Digital Energy Centers '
    '("DECs"), distributed energy hubs for liquid-cooled AI colocation, '
    "integrating accelerated compute with heat recycling and behind-the-meter "
    "(BTM) power production, engineered as one integrated system. Digital Energy "
    "is building an integrated sovereign AI infrastructure platform for "
    "enterprise and institutional customers, designed for edge inference."
)

# Per-type tail appended to the shared body. Customer-facing types (EU, WS)
# use "Digital Energy's integrated platform [verb]" subject (procurement pattern);
# relationship-facing types (DS, SS, EP) use "Digital Energy [verb]" subject.
RECITAL_A_TAIL_BY_TYPE = {
    "EndUser": (
        " Digital Energy's integrated platform supplies dedicated AI inference "
        "capacity to enterprises, AI labs, and research institutions requiring "
        "sovereign, high-density and low-latency infrastructure on European soil."
    ),
    "Distributor": (
        " Digital Energy seeks qualified channel and integration partners to "
        "extend its platform reach to end-user segments where the Partner "
        "holds established customer relationships and domain expertise."
    ),
    "Wholesale": (
        " Digital Energy's integrated platform contracts liquid-cooled AI "
        "colocation capacity at megawatt scale to NeoCloud operators and "
        "GPU cloud providers internationally."
    ),
    "StrategicSupplier": (
        " Digital Energy seeks qualified EPC contractors, modular infrastructure "
        "manufacturers, and OEM vendors to deliver the DEC platform and secure "
        "supply continuity across its active development pipeline."
    ),
    "EcosystemPartnership": (
        " Digital Energy engages with ecosystem partners on sovereign AI "
        "infrastructure, sustainable datacentre design, and European "
        "industrial policy alignment."
    ),
}


def resolve_recital_a(d: dict) -> str:
    """Return the Recital A body + type-specific tail (without the '(A) ' prefix).

    v3.4: single canonical body formatted with provider short name, plus a
    per-type tail. Bespoke override path (programme.recital_a_variant='bespoke'
    with programme.recital_a_bespoke) returns the raw bespoke string, with the
    QA linter enforcing forbidden-pattern rules.

    Legacy YAMLs with recital_a_variant set to 'default' / 'sovereignty' /
    'integration' still render correctly — those variant keys are ignored and
    the single canonical body is used. No migration break.
    """
    prov = d.get("provider", {}).get("short_name", "Digital Energy")
    t = d.get("type", "EndUser")
    variant = d.get("programme", {}).get("recital_a_variant", "default")
    if variant == "bespoke":
        bespoke = d.get("programme", {}).get("recital_a_bespoke", "")
        if bespoke:
            return bespoke
    body = RECITAL_A_BODY.format(prov=prov)
    tail = RECITAL_A_TAIL_BY_TYPE.get(t, "")
    return body + tail


# ---------------------------------------------------------------------------
# Document Builder
# ---------------------------------------------------------------------------

_AGREEMENT_TYPE_BY_LOI = {
    "Distributor": "Letter of Intent and NCNDA",
    "Wholesale": "Letter of Intent and NCNDA",
    "EndUser": "Letter of Intent",
    "StrategicSupplier": "Letter of Intent and NCNDA",
    "EcosystemPartnership": "Letter of Intent",
}

# v3.4: Wholesale subject dropped "Purpose-Built" marketing modifier.
_SUBJECT_BY_LOI = {
    "Distributor": "Strategic Infrastructure Partnership",
    "Wholesale": "AI Colocation Capacity",
    "EndUser": "AI Compute Infrastructure Services",
    "StrategicSupplier": "Strategic Supply and Infrastructure Partnership",
    "EcosystemPartnership": "Strategic Ecosystem Collaboration",
}


class LOI:
    def __init__(self, data: dict):
        self.d = data
        self.t = data["type"]
        self.agreement_type = _AGREEMENT_TYPE_BY_LOI[self.t]
        self.subject = _SUBJECT_BY_LOI[self.t]
        self.doc = Document()
        self._setup()
        party_by_type = {
            "Distributor": "Partner",
            "StrategicSupplier": "Supplier",
            "EcosystemPartnership": "Partner",
            "Wholesale": "Customer",
            "EndUser": "Customer",
        }
        self.party = party_by_type.get(self.t, "Customer")
        # v3.5.2 (brand-name defined term): provider is defined by its
        # short name throughout the body. The v3.5.2 brand rename replaced
        # every prior "the Provider" reference with "Digital Energy".
        self.provider_term = self._derive_provider_term(data)
        # QA linter accumulators (populated during build)
        self.overrides = set(self.d.get("_overrides", []))
        self.override_reason = self.d.get("_override_reason", "")
        self.qa_findings = []

    def _setup(self):
        style = self.doc.styles["Normal"]
        style.font.name = FONT_NAME
        style.font.size = FONT_BODY
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = LINE_SPACING
        # v3.5 scope A'''': derive footer entity from provider.legal_name so
        # BV-signed instruments render "Digital Energy Netherlands B.V."
        # footer and AG-signed instruments render "Digital Energy Group AG"
        # footer. Prior default of "ag" caused BV LOIs to ship with the
        # Swiss parent's entity in the footer — material misidentification.
        entity = self._derive_footer_entity(
            self.d.get("provider", {}).get("legal_name", "")
        )

        for s in self.doc.sections:
            s.page_width = Mm(210)
            s.page_height = Mm(297)
            s.top_margin = Mm(20)
            s.bottom_margin = Mm(35)
            s.left_margin = Mm(25)
            s.right_margin = Mm(20)
            s.different_first_page_header_footer = True

            # Shared header/footer from document-factory
            setup_first_page_header(s)
            setup_cont_header(s, title=self.agreement_type)
            setup_first_footer(s, classification="Confidential", entity=entity)
            setup_cont_footer(s, classification="Confidential", entity=entity)

    # --- Helpers ---

    def g(self, *keys, default=""):
        v = self.d
        for k in keys:
            v = v.get(k, default) if isinstance(v, dict) else default
        return v or default

    def choice(self, k):
        return bool(self.g("choices", k))

    @staticmethod
    def _is_tbc(value):
        """Detect placeholder / unresolved sentinel values."""
        if value is None:
            return True
        s = str(value).strip()
        if not s:
            return True
        return s.upper() in ("[TBC]", "[TO BE CONFIRMED]", "TBC", "TO BE CONFIRMED", "XXXXXXXX")

    @staticmethod
    def _derive_provider_term(data):
        """Return the defined-term short-name used throughout the body.

        v3.5 polish: derive from `provider.short_name` with "Digital Energy"
        fallback — preserves brand when AG signs (short_name still
        "Digital Energy") while supporting future subsidiary/JV instruments
        with different short names without another body-wide rename.

        Extracted as a static helper (matches `_derive_footer_entity`) so
        the derivation is unit-testable without constructing a full
        DocBuilder (which triggers `_setup()` and document-factory image
        loading — brittle in CI across different path layouts).
        """
        prov = (data or {}).get("provider") or {}
        return prov.get("short_name") or "Digital Energy"

    @staticmethod
    def _derive_footer_entity(legal_name):
        """Map provider legal name to document-factory footer entity key.

        v3.5 scope A'''': dedicated helper so the derivation is unit-testable.
        Returns "nl" for Digital Energy Netherlands B.V.; "ag" for Digital
        Energy Group AG; "ag" as safe default for anything unrecognized.
        """
        if not legal_name:
            return "ag"
        name = str(legal_name)
        if "Netherlands" in name or "B.V." in name or " BV" in name:
            return "nl"
        return "ag"

    def _render_placeholder(self, value, context):
        """Return render-appropriate string for a value that may be a placeholder.

        context values:
          - "sig_block_name"   → blank fillable line
          - "sig_block_title"  → blank fillable line
          - "body_clause"      → keep "[TBC]" as visible draft marker
          - default            → return empty string

        v3.5 scope J5 (Jonathan memo): `[TBC]` should never surface literally
        in the sig block of an external-facing draft — renders as a fillable
        blank line per standard legal instrument convention. The helper is
        also used for other render contexts; behaviour is context-scoped.
        """
        is_tbc = self._is_tbc(value)
        if context in ("sig_block_name", "sig_block_title"):
            return "____________________________" if is_tbc else str(value)
        if context == "body_clause":
            return "[TBC]" if is_tbc else str(value)
        return "" if is_tbc else str(value)

    def p(self, text, bold=False, italic=False, size=None, color=None, align=None, space_after=None):
        para = self.doc.add_paragraph()
        if align:
            para.alignment = align
        if space_after is not None:
            para.paragraph_format.space_after = Pt(space_after)
        run = para.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = size or FONT_BODY
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        return para

    def bp(self, label, text):
        """Bold label + normal text in one paragraph."""
        para = self.doc.add_paragraph()
        r1 = para.add_run(label)
        r1.bold = True
        r1.font.name = FONT_NAME
        r1.font.size = FONT_BODY
        r2 = para.add_run(text)
        r2.font.name = FONT_NAME
        r2.font.size = FONT_BODY
        return para

    def h(self, text, level=2):
        heading = self.doc.add_heading(text, level=level)
        for r in heading.runs:
            r.font.name = FONT_NAME
            r.font.color.rgb = NAVY

    def table(self, headers, rows):
        t = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.name = FONT_NAME
                    r.font.size = FONT_BODY
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = t.rows[ri + 1].cells[ci]
                cell.text = str(val)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = FONT_NAME
                        r.font.size = FONT_BODY
        return t

    def line(self):
        self.p("—" * 50, color=GREY, size=Pt(6), space_after=0)

    # --- Sections ---

    def letterhead(self):
        """Render IB-standard cover page via document-factory's add_cover()."""
        agreement_type = self.agreement_type
        subject = self.subject

        # Build Party objects from YAML data
        prov = self.d.get("provider", {})
        provider_party = Party(
            legal_name=prov.get("legal_name", ""),
            address=prov.get("address", ""),
            registration_type="KvK" if prov.get("kvk") else None,
            registration_number=prov.get("kvk"),
            parent=prov.get("parent"),
        )

        cp = self.d.get("counterparty", {})
        counterparty = Party(
            legal_name=cp.get("name", ""),
            address=cp.get("address", ""),
            registration_type=cp.get("reg_type"),
            registration_number=cp.get("reg_number"),
        )

        # Delegate to document-factory's shared cover page
        add_cover(
            self.doc,
            agreement_type=agreement_type,
            subject=subject,
            date_str=self.g("dates", "loi_date"),
            parties=[provider_party, counterparty],
            formality="non_binding",
            classification="Confidential",
        )

    def parties(self):
        """Render Parties Preamble — legal identification block in body.

        v3.5.2 scope A''' (Jonathan memo / user directive): prior documents
        identified parties only on the cover page. Cover page is a
        presentation layer, not a legal identification layer. Proper
        contract structure has parties named in the body of the instrument
        itself. This block appears between the cover and Recital A; its
        job is to make the legal identification of both parties verifiable
        from the body alone, independent of the cover page.

        Defined terms established here propagate to every clause:
          - Provider defined term: "Digital Energy" (brand-name — v3.5.2)
          - Counterparty defined term: "the Customer" / "the Partner" /
            "the Supplier" per type
        """
        prov = self.d.get("provider", {})
        cp = self.d.get("counterparty", {})
        loi_date = self.g("dates", "loi_date") or "____________________________"

        # v3.5.5 spacing fix: prior implementation emitted explicit blank
        # `self.p("")` paragraphs between each block, doubling the vertical
        # gap (blank paragraph height + default 3pt space_after × 2). The
        # preamble now uses `space_after=6` on the intro + party paragraphs
        # (single consistent gap) and a default-spaced closing line. This
        # matches the tighter cadence of the body clauses and keeps the
        # preamble block proportional to the rest of the page.
        self.p(
            f'THIS LETTER OF INTENT (the "LOI") is dated {loi_date} and entered into between:',
            space_after=6,
        )

        # Party 1 — Digital Energy (Provider)
        prov_legal = prov.get("legal_name", "")
        prov_form = prov.get("legal_form") or "a private limited liability company"
        # Ensure form starts with an article — grammatical Parties Preamble
        if prov_form and not prov_form.lower().startswith(("a ", "an ", "the ")):
            prov_form = f"a {prov_form}"
        prov_jur = prov.get("jurisdiction") or "the Netherlands"
        prov_addr = prov.get("address", "")
        prov_reg_type = prov.get("reg_type") or ("KvK" if prov.get("kvk") else "")
        prov_reg_number = prov.get("reg_number") or prov.get("kvk", "")

        party1_frag = f"(1) {prov_legal}, {prov_form} incorporated under the laws of {prov_jur}"
        if prov_addr and not self._is_tbc(prov_addr):
            party1_frag += f", with registered office at {prov_addr}"
        if prov_reg_type and prov_reg_number and not self._is_tbc(prov_reg_number):
            party1_frag += f" and registered with the {prov_reg_type} under number {prov_reg_number}"
        # Brand-name defined term — no "the" prefix; the brand name itself
        # is the defined short-form used throughout the body.
        party1_frag += f' ("{self.provider_term}"); and'
        self.p(party1_frag, space_after=6)

        # Party 2 — Counterparty
        cp_legal = cp.get("name", "")
        cp_jur = cp.get("jurisdiction") or self._render_placeholder(
            cp.get("jurisdiction"), "body_clause"
        )
        cp_addr = cp.get("address", "")
        cp_reg_type = cp.get("reg_type", "")
        cp_reg_number = cp.get("reg_number", "")

        party2_frag = f"(2) {cp_legal}, an entity incorporated under the laws of {cp_jur or '[TBC]'}"
        if cp_addr and not self._is_tbc(cp_addr):
            party2_frag += f", with registered office at {cp_addr}"
        # v3.6.0 bug 7: Party 1 has a KvK fallback; Party 2 previously
        # required BOTH cp_reg_type AND cp_reg_number, silently dropping
        # the registration number when reg_type was absent. Now emit the
        # number whenever it's present (using a generic label when
        # reg_type is unset — operator can override via YAML).
        if cp_reg_number and not self._is_tbc(cp_reg_number):
            if cp_reg_type:
                party2_frag += f" and registered with the {cp_reg_type} under number {cp_reg_number}"
            else:
                party2_frag += f" with company number {cp_reg_number}"
        party2_frag += f' (the "{self.party}").'
        self.p(party2_frag, space_after=6)

        self.p('(each a "Party" and together the "Parties")')

    def recitals(self):
        self.h("Recitals")
        cp = self.g("counterparty", "short")
        desc = self.g("counterparty", "description")
        # v3.6.0 bug 5: strip trailing period from description before the
        # engine appends its own. Avoids double-period when operator ends
        # the YAML value with '.'.
        if isinstance(desc, str):
            desc = desc.rstrip().rstrip(".")

        # Recital A — v3.4: resolve_recital_a() returns body + type-specific tail.
        self.p(f"(A) {resolve_recital_a(self.d)}")

        self.p(f'(B) {cp} (the "{self.party}") {desc}.')

        if self.t == "Distributor":
            self.p(
                f"(C) The Parties wish to record their mutual interest in establishing a strategic "
                f"partnership under which the {self.party} would collaborate with Digital Energy to deliver "
                f"AI colocation services to end-user customers, on the indicative terms set out below. "
                f'This letter of intent and non-circumvention non-disclosure agreement (the "LOI") '
                f"is subject to the negotiation and execution of a definitive partnership "
                f'agreement or master services agreement (the "MSA").'
            )
            self.p(
                "(D) The Parties will exchange commercially sensitive and competitively valuable "
                "information in connection with this LOI, and agree to the binding confidentiality "
                "and non-circumvention provisions set out in Clauses 6 and 7."
            )
        elif self.t == "Wholesale":
            self.p(
                f"(C) The Parties wish to record their mutual interest in Digital Energy making available, "
                f"and the {self.party} procuring, dedicated AI colocation capacity at Digital Energy's DEC "
                f"facilities, on the indicative terms set out below. This letter of intent and "
                f'non-circumvention non-disclosure agreement (the "LOI") is subject to the '
                f'negotiation and execution of a Master Services Agreement (the "MSA").'
            )
            self.p(
                "(D) The Parties will exchange commercially sensitive information, including "
                "site-specific data, pricing models, and infrastructure specifications, in connection "
                "with this LOI, and agree to the binding confidentiality and non-circumvention "
                "provisions set out in Clauses 6 and 7."
            )
        elif self.t == "StrategicSupplier":
            self.p(
                f"(C) The Parties wish to record their mutual interest in establishing a strategic "
                f"supply partnership under which the {self.party} would contribute to the delivery of "
                f"Digital Energy's DEC platform, on the indicative terms set out below. This letter of "
                f'intent and non-circumvention non-disclosure agreement (the "LOI") is subject '
                f"to the negotiation and execution of a Framework Agreement and statements of "
                f"work for named Provider projects."
            )
            self.p(
                "(D) The Parties will exchange commercially sensitive and competitively valuable "
                "information, including site-specific information about Provider projects and Partner "
                "design and manufacturing methodology, in connection with this LOI, and agree to the "
                "binding confidentiality and non-circumvention provisions set out in Clauses 6 and 7."
            )
        elif self.t == "EcosystemPartnership":
            themes_list = self.d.get("ecosystem", {}).get("collaboration_themes", [])
            themes_str = ", ".join(themes_list) if themes_list else "[COLLABORATION_THEMES]"
            self.p(
                f"(C) The Parties wish to record their mutual interest in a non-commercial ecosystem "
                f"collaboration focused on {themes_str}, on the framework set out in Clauses 3 through 5. "
                f'This letter of intent (the "LOI") is non-commercial in scope and does not '
                f"contemplate any payment, capacity purchase, or commercial commitment between "
                f"the Parties."
            )
            self.p(
                "(D) The Parties will exchange commercially sensitive or non-public information in "
                "connection with research, policy, or programme activities under this LOI, and agree "
                "to the mutual confidentiality provisions set out in Clause 6."
            )
        else:  # EndUser
            self.p(
                f"(C) The Parties wish to record their mutual interest in the {self.party} procuring "
                f"AI compute infrastructure services at Digital Energy's DEC facilities, on the indicative "
                f'terms set out below. This letter of intent (the "LOI") is subject to the '
                f'negotiation and execution of a service agreement (the "MSA").'
            )

    def definitions(self):
        self.h("1. Definitions")
        self.p("1.1 In this LOI, unless the context requires otherwise:")

        defs = []
        defs.append(('"Affiliate"', 'means, in relation to a Party, any entity that directly or indirectly controls, is controlled by, or is under common control with that Party, where "control" means the ownership of more than 50% of the voting rights or equivalent ownership interest;'))

        if self.t in ("Distributor", "Wholesale", "StrategicSupplier"):
            ac_text = (
                '"Associated Counterparty" means, in relation to a Site Identifier, any of the following '
                "with whom Digital Energy has a contractual, commercial, or active non-public engagement "
                "in connection with that site: (a) landowners and land lessors; (b) greenhouse operators "
                "and agricultural partners; (c) heat offtake counterparties; (d) energy procurement "
                "counterparties; (e) engineering, procurement, and construction contractors engaged or "
                "under negotiation; and (f) grid connection holders and distribution system operators "
                "(to the extent of non-public engagement)."
            )
            if self.t == "Distributor":
                ac_text += (
                    " (g) government agencies, municipalities, and regulatory bodies with whom the "
                    "Provider has active non-public engagement (excluding general regulatory contacts "
                    "available to any market participant);"
                )
            else:
                ac_text += (
                    " For the avoidance of doubt, government agencies and regulatory bodies are "
                    "excluded unless Digital Energy has made a specific named introduction;"
                )
            defs.append(("", ac_text))

        defs.append(('"Business Day"', "means a day other than a Saturday, Sunday, or public holiday in the Netherlands;"))

        if self.t == "Distributor":
            defs.append(('"Competitor"', 'means any entity whose primary business includes the development, ownership, or operation of colocation facilities for high-density compute workloads, or the provision of AI infrastructure services, in the European Economic Area;'))

        ci_text = (
            '"Confidential Information" means all information (whether technical, commercial, financial, '
            "or otherwise) disclosed by a Party to the other in connection with the Purpose, whether in "
            "writing, orally, electronically, or by inspection, including any information that by its "
            "nature or the circumstances of disclosure would reasonably be understood to be confidential"
        )
        if self.t in ("Distributor", "Wholesale", "StrategicSupplier"):
            ci_text += ", and including metadata, EXIF data, digital artifacts, file names, folder names, and any derivative information"
        ci_text += ". The existence and contents of this LOI are Confidential Information;"
        defs.append(("", ci_text))

        defs.append(('"DEC"', "means a Digital Energy Center: a liquid-cooled AI colocation facility developed and operated by Digital Energy;"))
        # v3.2: "DEC Block" removed from customer-facing definitions. Capacity expressed in MW IT + Designated Sites.
        defs.append(('"Designated Site"', "means any DEC or DEC site designated by Digital Energy for the delivery of capacity to the " + self.party + ", as confirmed in the MSA;"))
        defs.append(('"Financing Party"', "means any bank, financial institution, fund, security trustee, or other entity providing or arranging " + ("debt, mezzanine, or structured " if self.t != "EndUser" else "") + "finance to Digital Energy or any of its Affiliates in connection with the development or operation of any DEC;"))

        if self.t == "StrategicSupplier":
            defs.append((
                '"Framework Agreement"',
                "means the definitive Framework Agreement and accompanying statements of work to be negotiated between the Parties, setting out the commercial and operational framework for the supply relationship;"
            ))

        if self.t == "Distributor":
            defs.append(('"MSA"', "means the definitive Master Services Agreement or partnership agreement to be negotiated between the Parties;"))
        elif self.t == "Wholesale":
            defs.append(('"MSA"', "means the definitive Master Services Agreement to be negotiated between the Parties, incorporating the Sales Order Form, SLA Schedule, and Pricing Framework;"))
        elif self.t == "EndUser":
            defs.append(('"MSA"', "means the definitive service agreement to be negotiated between the Parties;"))
        # StrategicSupplier uses Framework Agreement (above) — no MSA definition.

        if self.t == "Distributor":
            defs.append(('"Protected Business Information" or "PBI"', "means Digital Energy's proprietary methodologies, strategies, and frameworks including: (a) site-sourcing and land-acquisition methodology; (b) energy procurement and grid connection strategy; (c) regulatory and permitting playbook; (d) heat offtake and energy recycling commercial model; (e) financial and operational modelling frameworks; (f) DEC design specifications and engineering standards; and (g) colocation pricing methodology and deal economics;"))

        # SS: PBI only if engineering_integration in strategic purposes
        if self.t == "StrategicSupplier":
            purposes = set(self.d.get("supplier", {}).get("strategic_purposes", []))
            if "engineering_integration" in purposes:
                defs.append(('"Protected Business Information" or "PBI"', "means Digital Energy's proprietary design, engineering, and infrastructure integration methodologies, interface specifications, and technical reference architectures disclosed in connection with the design integration collaboration under Clause 3.7;"))

        if self.t == "Distributor":
            defs.append(('"Purpose"', "means evaluating, negotiating, and progressing toward the execution of an MSA for the Transaction, including all activities under any companion agreements between the Parties (such as a Referral Agreement);"))
        elif self.t == "Wholesale":
            defs.append(('"Purpose"', "means evaluating, negotiating, and progressing toward the execution of an MSA for the provision of AI colocation services;"))
        elif self.t == "StrategicSupplier":
            defs.append(('"Purpose"', "means evaluating, negotiating, and progressing toward the execution of a Framework Agreement and accompanying statements of work for the supply relationship;"))
        else:
            defs.append(('"Purpose"', "means evaluating, negotiating, and progressing toward the execution of an MSA;"))

        if self.t in ("Distributor", "Wholesale", "StrategicSupplier"):
            defs.append(('"Ready-for-Service" or "RFS"', "means the date on which a DEC (or a discrete capacity allocation within a DEC) has been commissioned and is available for the provision of colocation services;"))

        defs.append(('"Representatives"', "means, in relation to a Party, its Affiliates, and its and their respective directors, officers, employees, agents, and professional advisers;"))

        if self.t in ("Distributor", "Wholesale"):
            defs.append(('"Sales Order Form"', "means the document that will form the basis for each binding service order under the MSA" + (", setting out confirmed capacity, pricing, site allocation, and RFS date;" if self.t == "Wholesale" else ";")))

        if self.t == "Wholesale":
            defs.append(('"Services"', "means the AI colocation services to be provided by Digital Energy, comprising power supply and distribution, liquid cooling infrastructure, physical security, building management, and facility operations;"))
        elif self.t == "EndUser":
            defs.append(('"Services"', "means the AI compute infrastructure services to be provided by Digital Energy, the scope of which will depend on the service model selected under Clause 3.1."))

        if self.t in ("Distributor", "Wholesale", "StrategicSupplier"):
            defs.append(('"Site Identifier"', "means any information that identifies or could reasonably be used to identify a specific DEC location, including: address, GPS coordinates, cadastral reference, project codename, photographs, aerial imagery, file names, folder names, metadata, and any information derived from the foregoing;"))

        if self.t == "Distributor":
            defs.append(('"Transaction"', "means the proposed strategic partnership between the Parties as described in Clause 3;"))

        if self.t in ("Distributor", "Wholesale"):
            defs.append(('"Whitespace"', "means designated floor area within a DEC available for the deployment of IT hardware and supporting infrastructure."))

        for label, text in defs:
            if label:
                self.bp(label + " ", text)
            else:
                self.p(text)

    def clause2(self):
        self.h("2. Purpose and Scope")
        if self.t == "Distributor":
            self.p(f"2.1 This LOI records the Parties' mutual interest in establishing a strategic partnership under which the {self.party} would participate in the delivery of AI colocation services to end-user customers through Digital Energy's DEC platform, on the indicative terms set out in Clauses 3 and 4.")
        elif self.t == "Wholesale":
            self.p(f"2.1 This LOI records the Parties' mutual interest in the {self.party} procuring dedicated AI colocation capacity at Digital Energy's DEC facilities, on the indicative terms set out in Clauses 3 and 4.")
        elif self.t == "StrategicSupplier":
            self.p(f"2.1 This LOI records the Parties' mutual interest in establishing a strategic supply partnership under which the {self.party} would contribute to the delivery of Digital Energy's DEC platform, on the indicative terms set out in Clauses 3 and 4.")
        else:
            self.p(f"2.1 This LOI records the Parties' mutual interest in the {self.party} procuring AI compute infrastructure services at Digital Energy's DEC facilities, on the indicative terms set out in Clauses 3 and 4.")

        if self.t in ("Distributor", "Wholesale", "StrategicSupplier"):
            if self.t == "Distributor":
                scoping_phrase = "a framework for further commercial discussion and negotiation toward a definitive MSA"
            elif self.t == "Wholesale":
                scoping_phrase = "the basis for technical scoping and commercial negotiation toward a definitive MSA"
            else:  # StrategicSupplier
                scoping_phrase = "a framework for further technical and commercial discussion toward a definitive Framework Agreement"
            self.p(f"2.2 The Parties intend this LOI to provide {scoping_phrase}. The commercial terms in Clauses 3 and 4 are non-binding expressions of intent.")
            self.p("2.3 The confidentiality, non-circumvention, and general provisions in Clauses 5 through 8 are legally binding and enforceable from the date of execution.")
        else:
            self.p("2.2 The commercial terms in Clauses 3 and 4 are non-binding expressions of intent. The confidentiality and general provisions in Clauses 5 through 7 are legally binding and enforceable from the date of execution.")

    def clause3_ds(self):
        mode = self.d.get("partnership_mode", "combined")
        heading = "3. Partnership and Combined Offering" if mode == "combined" else "3. Partnership and Referral Arrangement"
        self.h(heading)
        comm = self.d.get("commercial", {})
        mode = self.d.get("partnership_mode", "combined")

        # 3.1
        self.bp("3.1 Partnership Overview. ",
                f"{self.g('counterparty', 'short')} provides {comm.get('partner_core_capability', '')}. "
                f"Digital Energy develops and operates liquid-cooled AI colocation facilities with secured "
                f"energy and grid access, liquid cooling, and full energy recovery infrastructure.")
        if mode == "combined":
            self.p(
                f"The Parties intend to combine the {self.party}'s {comm.get('partner_contribution', '')} "
                f"with Digital Energy's AI colocation platform to deliver {comm.get('combined_offering', '')} "
                f"to {comm.get('target_end_users', '')}."
            )
        else:
            self.p(
                f"The Parties intend to establish a referral arrangement under which the {self.party} "
                f"would introduce qualified end-user customers to Digital Energy's DEC platform, leveraging "
                f"the {self.party}'s established client relationships and market presence."
            )

        # 3.2
        if mode == "combined":
            self.bp("3.2 Combined Offering. ", "Under the envisaged partnership:")
            self.p("(a) Digital Energy would supply dedicated AI colocation capacity at its DEC facilities, including power distribution, liquid cooling, physical security, and facility operations;")
            self.p(f"(b) the {self.party} would contribute {comm.get('partner_service_scope', '')}; and")
            self.p(f"(c) the end-user customer would procure a single, integrated solution through the {self.party}, with Digital Energy's DEC platform as the infrastructure foundation.")
            self.p("The precise service boundaries, responsibility matrix, SLA allocation, and commercial terms of the combined offering will be defined in the MSA.")
        else:
            self.bp("3.2 Referral Arrangement. ",
                     f"The {self.party} would identify and introduce qualified end-user customers to the "
                     f"Provider's DEC platform, leveraging the {self.party}'s market presence, customer "
                     f"relationships, and domain credibility. Digital Energy would manage all commercial, "
                     f"technical, and contractual relationships with introduced customers directly. The "
                     f"{self.party}'s ongoing role following an introduction \u2014 including relationship "
                     f"support, account management, or technical liaison \u2014 will be agreed on a "
                     f"case-by-case basis and set out in the MSA or a separate Referral Agreement.")
            self.p("The economic terms of the referral arrangement, including qualifying introduction criteria, fee structure, payment terms, and duration, will be set out in a separate Referral Agreement between the Parties.")

        # 3.3
        territory = comm.get("territory", "")
        segments = comm.get("target_segments", "")
        mw = comm.get("indicative_mw", "")
        self.bp("3.3 Commercial Scope. ", "The Parties intend the partnership to address the following market:")
        self.p(f"(a) Territory: {territory};")
        self.p(f"(b) Target segments: {segments}; and")
        self.p(f"(c) Estimated capacity: The {self.party} has indicated an estimated annual capacity requirement of approximately {mw} MW IT across its end-user customer base. This estimate is non-binding and is provided to inform capacity planning.")

        # 3.4
        if self.choice("pricing"):
            self.bp("3.4 Indicative Economics. ", "Based on preliminary discussions, the Parties envisage the following indicative economic framework:")
            ch = self.d.get("choices", {})
            self.table(
                ["Parameter", "Indicative Terms"],
                [
                    ["Partner capacity rate", f"EUR {ch.get('partner_rate', '')} per kW per month"],
                    ["Commercial structure", ch.get("economics_description", "")],
                    ["Minimum annual commitment", f"{ch.get('min_commitment', '')} MW IT"],
                ]
            )
            self.p("All economic terms are indicative and subject to the terms of the MSA.")
        else:
            self.bp("3.4 Indicative Economics. ",
                     "The economic framework for the partnership will be determined during commercial scoping and set out in the MSA. Digital Energy will issue indicative terms upon agreement of partnership scope, capacity estimates, and contract structure.")

        # 3.5
        if self.choice("exclusivity"):
            scope = self.d.get("choices", {}).get("exclusivity_scope", "")
            self.bp("3.5 Preferred Partner. ",
                     f"Subject to execution of the MSA and achievement of the minimum capacity commitments set out therein, Digital Energy intends to designate the {self.party} as a preferred partner within the following scope: {scope}. This designation means Digital Energy will offer the {self.party} priority participation in opportunities within the defined scope before engaging other channel partners for the same opportunity. This LOI does not create any exclusivity obligation; the terms of any such designation will be agreed in the MSA.")
        else:
            self.bp("3.5 Non-Exclusivity. ", "The partnership contemplated by this LOI is non-exclusive. Both Parties retain the right to enter into similar arrangements with third parties.")

    def clause3_ws(self):
        self.h("3. Indicative Capacity and Commercial Terms")
        comm = self.d.get("commercial", {})
        ch = self.d.get("choices", {})
        mw = comm.get("indicative_mw", "")

        # v3.2: capacity expressed in MW IT + Designated Sites; no "DEC Block" customer-facing.
        self.bp("3.1 Indicative Capacity. ", f"The {self.party} has indicated interest in approximately {mw} MW IT of liquid-cooled AI colocation capacity, to be delivered across one or more Designated Sites. Site configuration, phasing, and delivery milestones will be set out in the MSA.")

        # v3.5 scope J1 (Jonathan memo): Cl. 3.2 parametrized. Default density
        # bumped from "40 kW and above" to "approximately 130 kW and above"
        # matching NVIDIA GB200/GB300 NVL72 reference architectures. Cooling
        # topology explicit. Both overridable via YAML for non-AI workloads.
        rack_density = comm.get("rack_density_kw", "130")
        cooling_topology = comm.get(
            "cooling_topology",
            "direct-to-chip liquid cooling (consistent with NVIDIA GB200 NVL72 and GB300 NVL72 reference architectures, which target approximately 120\u2013140 kW per rack at full configuration)",
        )
        self.bp(
            "3.2 Technical Specification. ",
            f"All DEC facilities are designed for high-density AI compute workloads and are expected to include, at minimum: facility power supply and distribution, {cooling_topology} supporting rack densities of approximately {rack_density} kW and above, rear-door heat exchangers or equivalent where applicable, building management, physical security, and 24/7 facility operations. The exact capacity, rack configuration, power density, and cooling requirements will be determined during the technical scoping phase following this LOI.",
        )

        if self.choice("phasing"):
            self.bp("3.3 Deployment Phasing. ", f"The {self.party} has indicated the following high-level phasing interest:")
            phases = ch.get("phases", [])
            rows = [[f"Phase {i+1}", f"{p.get('mw', '')} MW IT", p.get("timeline", "")] for i, p in enumerate(phases)]
            self.table(["Phase", "Approximate Capacity", "Indicative Timeline"], rows)

        # v3.5 scope J3 (Jonathan memo): Cl. 3.4 Expansion branches on value
        # type. Numeric values render the approximate-MW sentence; empty,
        # [TBC], "to be discussed" and non-numeric strings render a generic
        # forward-expansion clause that doesn't insert broken grammar into
        # the instrument.
        exp = comm.get("expansion_mw", "")
        exp_str = str(exp).strip() if exp is not None else ""
        is_numeric_exp = bool(exp_str) and exp_str.lstrip("~<>= ").split()[0].split(".")[0].isdigit()
        if is_numeric_exp and not self._is_tbc(exp_str):
            self.bp(
                "3.4 Expansion. ",
                f"The {self.party} has expressed interest in future expansion to approximately {exp_str} MW IT, subject to availability, commercial agreement, and the terms of the MSA. Digital Energy will use reasonable endeavours to accommodate expansion requirements within its DEC programme.",
            )
        else:
            self.bp(
                "3.4 Expansion. ",
                f"The {self.party} has expressed interest in future expansion beyond the initial deployment, with scale to be determined following technical scoping and subject to availability, commercial agreement, and the terms of the MSA. Digital Energy will use reasonable endeavours to accommodate expansion requirements within its DEC programme.",
            )

        if self.choice("pricing"):
            self.bp("3.5 Indicative Pricing. ", f"Based on the {self.party}'s indicated capacity and term preferences, Digital Energy's indicative pricing is as follows:")
            self.table(
                ["Parameter", "Indicative Terms"],
                [
                    ["Base colocation fee", f"EUR {ch.get('base_price', '')} per kW per month"],
                    ["Contract term", f"{ch.get('term_years', '')} years"],
                    ["Power", "Metered IT Load (kWh) x PUE x Energy Rate (EUR/kWh) \u2014 billed additionally"],
                    ["Price escalation", "Subject to annual escalation, the mechanism for which will be agreed in the MSA"],
                    ["Indicative RFS", ch.get("target_rfs_date", "")],
                    ["Ramp schedule", ch.get("ramp_schedule", "")],
                ]
            )
            self.p(f"All commercial terms are indicative and subject to the outcome of the technical scoping phase, the {self.party}'s final capacity requirements, contract term, and credit profile.")
        else:
            self.bp("3.5 Indicative Pricing. ", f"Pricing for the Services will be determined following the technical scoping phase and set out in the Sales Order Form. Digital Energy will provide indicative pricing upon completion of the {self.party}'s capacity, density, and term requirements.")

        term = comm.get("min_term", "")
        # v3.2: "minimum commitment term" replaced with "approximately X, indicative only".
        self.bp("3.6 Indicative Term. ", f"The {self.party} anticipates a commitment term of approximately {term}, indicative only and subject to confirmation in the MSA. Actual term may be longer or shorter depending on final commercial terms. The {self.party} acknowledges that Digital Energy intends the MSA to include take-or-pay provisions commensurate with the committed capacity, the terms of which will be negotiated in good faith.")

        self.bp("3.7 Credit Assessment. ", f"Digital Energy will complete a credit assessment of the {self.party} (or the entity that will execute the MSA) as part of the commercial process. The {self.party} agrees to cooperate with such assessment, which may include provision of audited financial statements, credit references, evidence of parent company support, or other financial information as reasonably requested. Where the {self.party} does not hold an investment-grade credit rating (or equivalent), the Parties will discuss appropriate credit support mechanisms, which may include a parent company guarantee, security deposit, or letter of credit.")

        self.bp("3.8 Site Allocation. ", f"Digital Energy is developing DEC facilities across multiple locations. Digital Energy will allocate capacity to the {self.party} based on the DEC(s) best suited to the {self.party}'s requirements, taking into account development readiness, grid availability, RFS timeline, and the {self.party}'s latency and connectivity needs. Site allocation will be confirmed during the technical scoping phase and formalised in the Sales Order Form.")

    def clause3_eu(self):
        self.h("3. Service Requirements")
        comm = self.d.get("commercial", {})
        types = comm.get("service_type", [])

        self.bp("3.1 Service Model. ", f"The {self.party} has indicated interest in the following service model:")

        descs = {
            "bare_metal": (
                "Bare Metal Colocation \u2014 Dedicated, liquid-cooled rack space within a DEC, supporting "
                "densities of 40 kW per rack and above. Digital Energy supplies power distribution, direct "
                "liquid cooling, physical security, and 24/7 facility operations. The Customer deploys and "
                "manages its own hardware and software. The demarcation point is at the rack: everything "
                "below (facility infrastructure) is Digital Energy's responsibility; everything above (compute "
                "hardware, operating system, workloads) is the Customer's. Billed monthly on a per-kW basis "
                "with a committed term."
            ),
            "shared_cloud": (
                "Shared Cloud \u2014 Managed GPU compute capacity hosted on sovereign European infrastructure "
                "at Digital Energy's DEC facilities. The Customer accesses reserved or on-demand compute "
                "resources without procuring, deploying, or managing hardware. Digital Energy or a designated "
                "delivery partner operates the compute platform, including hardware lifecycle, scheduling, "
                "and platform software. The Customer manages workloads, data, and application layers. Billed "
                "on reserved capacity or metered usage, with a committed term or minimum spend."
            ),
            "tokens": (
                "Token-Based GPU Access \u2014 Flexible, on-demand access to GPU compute measured in GPU-hours "
                "or equivalent units. The Customer purchases compute tokens \u2014 pre-paid or pay-as-you-go "
                "\u2014 redeemable against available capacity across Digital Energy's DEC network. No dedicated "
                "hardware allocation or minimum infrastructure commitment. Digital Energy or a designated "
                "delivery partner manages all infrastructure and scheduling. The Customer submits workloads "
                "and consumes capacity as needed. Lowest entry threshold; designed for variable or exploratory "
                "workloads."
            ),
        }
        for st in types:
            if st in descs:
                self.p(descs[st])

        if any(st in ("shared_cloud", "tokens") for st in types):
            self.p("For Shared Cloud and Token-Based models: Digital Energy may deliver these services in partnership with a qualified delivery partner. The specific delivery partner, platform, and commercial terms will be confirmed in the MSA.", italic=True)

        cap = comm.get("indicative_capacity", "")
        self.bp("3.2 Indicative Capacity. ", f"The {self.party} has indicated interest in approximately {cap} of compute capacity. The exact capacity, configuration, and technical requirements will be determined during the technical scoping phase.")

        self.bp("3.3 Technical Requirements. ", f"Digital Energy's DEC facilities support rack densities of 40 kW and above with direct liquid cooling, designed for GPU-intensive training and inference workloads. The {self.party}'s specific requirements \u2014 including GPU type, rack density, network connectivity, and data sovereignty needs \u2014 will be confirmed during technical scoping and formalised in the MSA.")

        if self.choice("pricing"):
            ch = self.d.get("choices", {})
            self.bp("3.4 Indicative Pricing. ", "Based on preliminary discussions, Digital Energy's indicative pricing is as follows:")
            self.table(
                ["Parameter", "Indicative Terms"],
                [
                    ["Service rate", f"EUR {ch.get('base_price', '')} {ch.get('pricing_unit', 'per kW per month')}"],
                    ["Contract term", ch.get("term_years", "")],
                    ["Indicative RFS", ch.get("target_rfs_date", "")],
                ]
            )
            self.p(f"All terms are indicative and subject to the outcome of technical scoping and the {self.party}'s final requirements.")
        else:
            self.bp("3.4 Indicative Pricing. ", f"Pricing will be determined following the technical scoping phase and set out in the MSA. Digital Energy will provide indicative pricing upon agreement of the {self.party}'s capacity, density, and service model requirements.")

        term = comm.get("min_term", "")
        # v3.2: indicative, not "minimum commitment"
        self.bp("3.5 Indicative Term. ", f"The {self.party} anticipates a commitment term of approximately {term}, indicative only and subject to confirmation in the MSA.")

    def clause3_ss(self):
        """v3.3: Strategic Supplier Cl. 3 — Partnership Scope.

        3.1 Capability Contribution (ALWAYS). Then 3.2-3.8 are purpose-driven:
        - capacity_lock_in       -> 3.2 Capacity Reservation, 3.3 Lead-Time Targets
        - pricing_volume         -> 3.4 Pricing Framework, 3.5 Volume Tiers
        - supply_chain_de_risking-> 3.6 Dual-Source + Continuity
        - engineering_integration-> 3.7 Design Integration + IP Allocation
        - pipeline_visibility    -> 3.8 Preferred-Supplier / ROFR
        """
        self.h("3. Partnership Scope")
        supplier = self.d.get("supplier", {})
        purposes = set(supplier.get("strategic_purposes", []))

        # 3.1 Capability Contribution (ALWAYS)
        self.bp(
            "3.1 Capability Contribution. ",
            "Under the envisaged partnership:"
        )
        self.p(
            "(a) Digital Energy would contribute the site, energy procurement, grid "
            "connection, land and regulatory compliance, long-term operational "
            "responsibility, and integration into Digital Energy's broader DEC "
            "platform for each Designated Site; and"
        )
        # v3.6.0 bug 6: strip trailing period from core_capability to
        # avoid double-period when operator ends YAML value with '.'.
        _cc = supplier.get("core_capability", "[CAPABILITY DESCRIPTION TO BE CONFIRMED]")
        if isinstance(_cc, str):
            _cc = _cc.rstrip().rstrip(".")
        self.p(
            f"(b) the {self.party} would contribute {_cc}."
        )
        self.p(
            "The precise scope boundaries, responsibility matrix, and commercial "
            "terms will be defined in the Framework Agreement and accompanying "
            "statements of work."
        )

        # 3.2 Capacity Reservation — IF capacity_lock_in
        if "capacity_lock_in" in purposes:
            lead_time = supplier.get("lead_time_target", "[LEAD TIME TO BE CONFIRMED]")
            self.bp(
                "3.2 Capacity Reservation. ",
                f"Digital Energy has indicated a projected demand across its active "
                f"development pipeline. Subject to commercial agreement, the "
                f"{self.party} will reserve capacity in its manufacturing or "
                f"service-delivery plan to support Digital Energy's indicative "
                f"pipeline, with specific volumes and delivery windows set out in "
                f"the Framework Agreement."
            )
            # 3.3 Lead-Time Targets — same trigger
            self.bp(
                "3.3 Lead-Time Targets. ",
                f"The {self.party} targets {lead_time}. The Framework Agreement "
                f"will set the binding lead-time commitments, performance "
                f"measurement, and remedies for delay."
            )

        # 3.4 Pricing Framework — IF pricing_volume
        if "pricing_volume" in purposes:
            volume = supplier.get("volume_indicative", "[VOLUME TO BE CONFIRMED]")
            self.bp(
                "3.4 Pricing Framework. ",
                "The pricing framework will be structured as cost-plus, unit "
                "economics, or volume-tier model as appropriate to the supply "
                "category. Specific pricing, volume thresholds, and escalation "
                "mechanisms will be set out in the Framework Agreement."
            )
            # 3.5 Volume Tiers — same trigger
            self.bp(
                "3.5 Volume Tiers. ",
                f"Indicative volume contemplated: {volume}. The Framework "
                f"Agreement will set out final volume tiers and the commercial "
                f"principles applicable at each tier. All volumes and pricing "
                f"in this LOI are non-binding."
            )

        # 3.6 Dual-Source and Continuity — IF supply_chain_de_risking
        if "supply_chain_de_risking" in purposes:
            self.bp(
                "3.6 Dual-Source and Continuity Commitments. ",
                f"The {self.party} acknowledges Digital Energy's requirement for "
                f"supply-chain resilience. The Parties will discuss, and document "
                f"in the Framework Agreement: (a) component substitution and "
                f"second-source provisions; (b) continuity commitments, including "
                f"obligations to provide advance notice of discontinuation and to "
                f"support orderly transition; and (c) delivery service-level "
                f"commitments and remedies for failure to meet them."
            )

        # 3.7 Design Integration and IP Allocation — IF engineering_integration
        if "engineering_integration" in purposes:
            joint_ip = self.g("choices", "joint_ip", default="background").lower()
            if joint_ip == "none":
                ip_clause = (
                    "no foreground IP is contemplated as a deliverable of the "
                    "collaboration, and any IP developed will remain with the "
                    "Party that created it;"
                )
            elif joint_ip == "foreground":
                ip_clause = (
                    "foreground IP developed jointly in the course of the "
                    "collaboration will be jointly owned, with cross-licences on "
                    "terms to be agreed in the Framework Agreement;"
                )
            else:  # background (default)
                ip_clause = (
                    "each Party retains all right, title, and interest in its "
                    "background IP; any foreground IP developed jointly will be "
                    "allocated per the principles set out in the Framework Agreement;"
                )
            self.bp(
                "3.7 Design Integration and IP Allocation. ",
                f"The Parties will collaborate on design integration across the "
                f"specified interfaces between Digital Energy's DEC platform and "
                f"the {self.party}'s contribution. IP allocation: (a) {ip_clause} "
                f"(b) no Party assigns pre-existing IP by reason of this LOI or "
                f"any discussions under it; and (c) the full IP framework will be "
                f"set out in the Framework Agreement."
            )

        # 3.8 Preferred-Supplier / ROFR — IF pipeline_visibility
        if "pipeline_visibility" in purposes:
            self.bp(
                "3.8 Preferred-Supplier and Right of First Refusal. ",
                f"Subject to commercial agreement in the Framework Agreement, the "
                f"Provider intends to grant the {self.party} a right of first "
                f"refusal on the procurement of the supply scope set out in "
                f"Clause 3.1(b) across Digital Energy's active development pipeline. "
                f"The right of first refusal requires the {self.party} to submit "
                f"a compliant proposal within 20 Business Days of Digital Energy's "
                f"invitation. This LOI does not create any binding right of first "
                f"refusal; the terms will be set out in the Framework Agreement."
            )

    def clause4_ss(self):
        """v3.3: Strategic Supplier Cl. 4 — Pipeline Engagement.

        4.2 Contractual Sequence, 4.4 Change of Control, 4.6 Implementation
        Roadmap are ALWAYS. Others are purpose-driven.
        """
        self.h("4. Pipeline Engagement")
        purposes = set(self.d.get("supplier", {}).get("strategic_purposes", []))

        # 4.1 Project Introduction Process — IF pipeline_visibility
        if "pipeline_visibility" in purposes:
            self.bp(
                "4.1 Project Introduction Process. ",
                f"Following execution of this LOI, Digital Energy intends to "
                f"maintain a pipeline register of Designated Sites. The "
                f"{self.party} will be invited to participate in named project "
                f"opportunities on the cadence and criteria set out in the "
                f"Framework Agreement."
            )

        # 4.2 Contractual Sequence (ALWAYS) — mirrors WS Cl. 4.2 pattern
        self.bp(
            "4.2 Contractual Sequence. ",
            "The Parties acknowledge the intended progression of commercial instruments:"
        )
        self.p("(a) this LOI, setting out non-binding commercial intent and binding confidentiality and non-circumvention;")
        self.p("(b) a Framework Agreement, setting out the definitive commercial and operational framework for the supply relationship;")
        self.p("(c) one or more Statements of Work for named Provider projects, each executed under the Framework Agreement; and")
        self.p("(d) site-specific deliverables, schedules, or operational annexes executed under each Statement of Work.")
        # v3.6.0 bug 4 (SS): meta-commentary trailer removed — explains
        # the LOI rather than creating obligation; R-22 class.

        # 4.3 Joint-Development Governance — IF engineering_integration
        if "engineering_integration" in purposes:
            self.bp(
                "4.3 Joint-Development Governance. ",
                "The Parties intend to establish a joint design and engineering "
                "working group to coordinate interfaces, design reviews, and "
                "lifecycle decisions. The governance structure, decision rights, "
                "and escalation paths will be set out in the Framework Agreement."
            )

        # 4.4 Change of Control (ALWAYS)
        self.bp(
            "4.4 Change of Control. ",
            f"If, during the term of this LOI, a direct competitor of the "
            f"Provider (as determined by reference to its primary business being "
            f"the development, ownership, or operation of colocation facilities "
            f"for high-density compute workloads) acquires Control of the "
            f"{self.party}, Digital Energy may terminate this LOI by written notice "
            f"with immediate effect. Upon such termination, the confidentiality "
            f"and non-circumvention obligations in Clauses 6 and 7 shall "
            f"continue for their stated survival periods."
        )

        # 4.5 Exclusivity — IF capacity_lock_in AND choices.exclusivity
        if "capacity_lock_in" in purposes and self.choice("exclusivity"):
            self.bp(
                "4.5 Exclusivity. ",
                "Subject to execution of the Framework Agreement and achievement "
                "of the reservation commitments set out therein, the Parties "
                "intend to agree a scope of mutual exclusivity. This LOI does "
                "not create any exclusivity; the terms will be agreed in the "
                "Framework Agreement."
            )

        # 4.6 Implementation Roadmap (ALWAYS)
        self.bp(
            "4.6 Implementation Roadmap. ",
            "Following execution of this LOI, the Parties intend to proceed as follows:"
        )
        self.p("(a) Technical and commercial scoping (target: 30 days post-LOI) — detailed discovery on scope, capability, lead times, pricing framework, and IP allocation.")
        self.p("(b) Draft Framework Agreement (target: 60 days post-LOI) — Digital Energy will issue a draft Framework Agreement incorporating the agreed terms.")
        self.p("(c) Framework Agreement negotiation and execution (target: 90 days post-LOI) — the Parties will negotiate and execute the Framework Agreement.")
        self.p("These timelines are indicative and non-binding.")

    def clause4(self):
        if self.t == "Distributor":
            self.h("4. Relationship Structure and Protection")
            pbi = self.g("protection", "pbi_survival", default="10 years")
            self.bp("4.1 Protected Business Information. ",
                     f"The {self.party} acknowledges that Digital Energy's competitive position depends on the confidentiality of its Protected Business Information as defined in Clause 1. The {self.party} agrees that, for a period of {pbi} from the date of this LOI, it shall not directly or indirectly use any PBI to replicate, reverse-engineer, or independently develop any material element of Digital Energy's business model, site-sourcing methodology, energy procurement strategy, or commercial framework, whether for itself or for any third party. This obligation survives termination or expiry of this LOI and of any MSA.")
            self.bp("4.2 Change of Control. ",
                     f"If, during the term of this LOI, a Competitor acquires Control (as defined in Clause 1.1, \u201cAffiliate\u201d) of the {self.party}, Digital Energy may terminate this LOI by written notice with immediate effect. Upon such termination, the confidentiality and non-circumvention obligations in Clauses 6 and 7 shall continue for their stated survival periods.")
            self.bp("4.3 Associated Counterparties. ",
                     f"The {self.party} acknowledges that Digital Energy maintains relationships with Associated Counterparties at each DEC site. The sharing of any Site Identifier by Digital Energy shall be deemed an introduction of all Associated Counterparties for that site for the purposes of Clause 7.")
            self.bp("4.4 Governance. ", "The Parties intend to establish a joint steering committee or equivalent governance mechanism, the terms of which will be set out in the MSA.")
            self.bp("4.5 Implementation Roadmap. ", "Following execution of this LOI, the Parties intend to proceed as follows:")
            self.p("(a) Commercial scoping (target: 30 days post-LOI) \u2014 The Parties will define the partnership scope, target customer segments, capacity estimates, and economic framework.", bold=False)
            self.p("(b) Draft Partnership Agreement (target: 60 days post-LOI) \u2014 Digital Energy will issue a draft MSA or partnership agreement incorporating the agreed commercial terms.")
            self.p("(c) MSA negotiation and execution (target: 90 days post-LOI) \u2014 The Parties will negotiate and execute the MSA.")
            self.p("These timelines are indicative and non-binding.")

        elif self.t == "Wholesale":
            self.h("4. Relationship Structure and Next Steps")
            self.bp("4.1 MSA Structure. ", "The Parties intend to execute a Master Services Agreement incorporating: (a) a Pricing Framework setting out the commercial terms for the Services; (b) Sales Order Forms for each capacity commitment; and (c) an SLA Schedule defining availability, performance, and remediation commitments. The MSA will be the sole binding commercial agreement between the Parties.")
            # v3.2: Cl. 4.2 arrows replaced with institutional prose + numbered list.
            self.bp("4.2 Contractual Sequence. ", "The Parties acknowledge the intended progression of commercial instruments:")
            self.p("(a) this LOI, setting out non-binding commercial intent;")
            self.p("(b) a Sales Order Form or equivalent binding capacity commitment with indicative pricing;")
            self.p("(c) the Master Services Agreement (MSA), containing definitive commercial terms; and")
            self.p("(d) site-specific deliverables, schedules, or operational annexes executed under the MSA.")
            # v3.6.0 bug 4 (WS): meta-commentary trailer removed.
            self.bp("4.3 Direct Agreement Willingness. ", f"The {self.party} confirms its willingness, subject to commercially reasonable terms, to enter into a direct agreement with Digital Energy's Financing Parties if requested under Clause 5.3. The {self.party}'s cooperation in this regard materially supports delivery of the committed capacity on the indicative timeline.")
            self.bp("4.4 Expansion and Priority. ", f"Digital Energy will offer the {self.party} priority access to additional capacity within the DEC(s) allocated to the {self.party}, subject to availability. Expansion terms will be governed by the MSA.")
            self.bp("4.5 Implementation Roadmap. ", "Following execution of this LOI, the Parties intend to proceed as follows:")
            self.p("(a) Technical scoping (target: 30 days post-LOI) \u2014 Detailed technical discovery to determine exact capacity, rack layout, power distribution, cooling specifications, and connectivity requirements.")
            self.p(f"(b) Credit assessment (target: 30 days post-LOI, concurrent with technical scoping) \u2014 Digital Energy will complete a credit assessment of the {self.party} or the entity that will execute the MSA.")
            self.p("(c) Sales Order Form (target: 60 days post-LOI) \u2014 Upon completion of technical scoping, Digital Energy will issue a Sales Order Form setting out confirmed commercial terms, facility specifications, and service level framework.")
            self.p("(d) MSA negotiation and execution (target: 90 days post-LOI) \u2014 The Parties will negotiate and execute the MSA.")
            self.p("These timelines are indicative and non-binding.")
            self.bp("4.6 No Capacity Reservation. ", "This LOI does not reserve capacity at any DEC. Capacity allocation is on a first-come, first-served basis and will only be confirmed upon execution of the MSA.")

        else:  # EndUser
            self.h("4. Next Steps")
            self.p("4.1 Following execution of this LOI, the Parties intend to proceed as follows:")
            self.p("(a) Technical scoping (target: 30 days post-LOI) \u2014 Detailed discovery to determine capacity, configuration, connectivity, and service model requirements.")
            self.p("(b) Commercial proposal (target: 60 days post-LOI) \u2014 Digital Energy will issue a commercial proposal setting out confirmed pricing, service scope, and SLA framework.")
            self.p("(c) MSA negotiation and execution (target: 90 days post-LOI) \u2014 The Parties will negotiate and execute the MSA.")
            self.p("These timelines are indicative and non-binding.")
            self.p("4.2 This LOI does not reserve capacity at any DEC. Capacity allocation will be confirmed upon execution of the MSA.")

    def clause5(self):
        """Cl. 5 for EU / DS / WS — Project Finance and Assignment.
        v3.4: meta-commentary stripped from 5.1. Strategic Supplier uses
        clause5_ss() instead (supply-side bankability, not revenue-bankability).
        """
        self.h("5. Project Finance and Assignment (BINDING)")
        self.bp(
            "5.1 Project Finance Context. ",
            "Digital Energy is developing the DEC programme under a combination "
            "of equity investment and non-recourse project finance. This LOI "
            "is binding in Clauses 5, 6, 7, and 8 to support that financing "
            "structure."
        )
        self.bp("5.2 Assignment. ", f"Neither Party may assign its rights or obligations under this LOI without the prior written consent of the other Party, except that:")
        self.p("(a) Digital Energy may assign this LOI, or any rights under it, to any Financing Party or security trustee as security for project finance obligations; and")
        self.p(f"(b) Digital Energy may assign this LOI to any Affiliate or special-purpose vehicle within its corporate group without the {self.party}'s consent, provided Digital Energy remains liable for the performance of the assignee's obligations.")
        self.bp("5.3 Lender Acknowledgment. ", f"The {self.party} acknowledges and agrees that, upon Digital Energy's written request, the {self.party} shall negotiate in good faith and execute a direct agreement (or lender acknowledgment letter) with Digital Energy's Financing Party within 30 Business Days of such request. Such direct agreement may include, as is customary in project finance transactions: (a) step-in rights for the Financing Party upon a Provider default; (b) cure periods in favour of the Financing Party; and (c) information rights enabling the Financing Party to monitor the commercial relationship. The terms of any such direct agreement shall be commercially reasonable and consistent with market practice for project finance transactions.")

    def clause5_ss(self):
        """v3.4: Cl. 5 for Strategic Supplier — Supply Chain and Delivery Commitment.

        A supplier is not a revenue counterparty. The bankability signal DE
        cares about from a supplier is SUPPLY CERTAINTY and DELIVERY
        COMMITMENT, not contracted revenue. This clause reshapes Cl. 5 to
        that frame, keeping the Assignment and Financing Continuity
        Acknowledgment sub-clauses structurally parallel to the commercial
        Cl. 5 so a lender reading across the five LOI types sees consistent
        architecture.
        """
        self.h("5. Supply Chain and Delivery Commitment (BINDING)")
        self.bp(
            "5.1 Delivery Intent. ",
            f"The {self.party} confirms its intent to reserve manufacturing "
            f"and service-delivery capacity to support Digital Energy's active "
            f"development pipeline, on the terms to be agreed in the Framework "
            f"Agreement and accompanying statements of work. This LOI is "
            f"binding in Clauses 5, 6, 7, and 8 to support that supply "
            f"relationship."
        )
        self.bp(
            "5.2 Assignment. ",
            f"Neither Party may assign its rights or obligations under this "
            f"LOI without the prior written consent of the other Party (such "
            f"consent not to be unreasonably withheld), except that:"
        )
        self.p("(a) Digital Energy may assign this LOI, or any rights under it, to any Financing Party or security trustee as security for project finance obligations; and")
        self.p(f"(b) Digital Energy may assign this LOI to any Affiliate or special-purpose vehicle within its corporate group without the {self.party}'s consent, provided Digital Energy remains liable for the performance of the assignee's obligations.")
        self.bp(
            "5.3 Financing Continuity Acknowledgment. ",
            f"The {self.party} acknowledges that Digital Energy's DEC programme "
            f"is financed on a non-recourse, per-site basis, and that supply "
            f"continuity from qualified suppliers materially supports that "
            f"financing. Upon Digital Energy's written request, the {self.party} "
            f"shall negotiate in good faith with Digital Energy's Financing "
            f"Party to confirm supply arrangements on financed projects, "
            f"including reasonable cooperation with customary direct-agreement "
            f"or acknowledgment mechanics (step-in, cure periods, information "
            f"rights), within 30 Business Days of such request. The terms of "
            f"any such direct agreement shall be commercially reasonable and "
            f"consistent with market practice for project finance supply "
            f"arrangements."
        )

    def clause6(self):
        self.h(f"6. Confidentiality {'and Non-Disclosure ' if self.t != 'EndUser' else ''}(BINDING)")
        existing = self.choice("existing_nda")

        if existing:
            nda_date = self.g("choices", "nda_date")
            transaction_suffix = " and the Transaction." if self.t == "Distributor" else (" and the proposed transaction." if self.t == "Wholesale" else ".")
            self.p(f'6.1 The Non-Disclosure Agreement between the Parties dated {nda_date} (the "NDA") remains in full force and effect and applies to all information exchanged in connection with this LOI{transaction_suffix}')
            self.p("6.2 The existence and contents of this LOI shall be treated as Confidential Information under the NDA.")
            if self.t != "EndUser":
                self.p("6.3 To the extent of any conflict between the NDA and this LOI, the provisions of this LOI shall prevail.")
                transaction_ref = "the Transaction" if self.t == "Distributor" else "the proposed transaction"
                self.p(f"6.4 Neither Party shall make any public announcement regarding this LOI or {transaction_ref} without the prior written consent of the other Party, except as required by applicable law.")
            else:
                self.p("6.3 Neither Party shall make any public announcement regarding this LOI without the prior written consent of the other Party, except as required by applicable law.")
        else:
            surv = self.g("protection", "confidentiality_survival", default="3 years")
            if self.t == "EndUser":
                # Tier A — 8 clauses
                self.p(f"6.1 Each Party shall keep confidential all Confidential Information received from the other Party and shall use such information solely for the Purpose.")
                self.p("6.2 Each Party may disclose Confidential Information only to its Representatives who have a genuine need to know for the Purpose and who are bound by confidentiality obligations no less restrictive than this Clause 6.")
                self.p("6.3 The obligations in Clauses 6.1 and 6.2 do not apply to information that the receiving Party can demonstrate: (a) is or becomes publicly available through no fault of the receiving Party; (b) was already in the receiving Party's possession without restriction; (c) was independently developed without use of the Confidential Information; or (d) was received from a third party without breach of any confidentiality obligation.")
                self.p("6.4 A Party may disclose Confidential Information to the extent required by applicable law, regulation, or court order, provided that (where legally permitted) the disclosing Party gives the other Party prior written notice and discloses only the minimum required.")
                self.p("6.5 Upon written request or expiry of this LOI, the receiving Party shall promptly return or destroy all Confidential Information and certify such return or destruction in writing. The receiving Party may retain copies required by applicable law or internal compliance policies, subject to continued confidentiality.")
                self.p("6.6 Neither Party shall make any public announcement regarding this LOI without the prior written consent of the other Party, except as required by applicable law.")
                self.p(f"6.7 The obligations in this Clause 6 shall survive for {surv} from the date of termination or expiry of this LOI. Obligations in respect of trade secrets shall continue indefinitely.")
                self.p("6.8 Each Party acknowledges that a breach of this Clause 6 may cause irreparable harm for which damages would not be an adequate remedy. The disclosing Party shall be entitled to seek injunctive or other equitable relief without the need to prove actual loss.")
            else:
                # Tier B — 16 clauses
                self.bp("6.1 Purpose Limitation. ", "Each Party shall use the other Party's Confidential Information solely for the Purpose and for no other purpose.")
                self.bp("6.2 Non-Disclosure. ", "Each Party shall keep confidential all Confidential Information received from the other Party and shall not disclose such information to any person except as permitted under this Clause 6.")
                self.bp("6.3 Standard of Care. ", "Each Party shall apply no less than reasonable care to protect the other Party's Confidential Information, and no less than the care it applies to its own confidential information of a similar nature.")
                self.bp("6.4 Permitted Disclosures. ", "A Party may disclose Confidential Information to:")
                self.p("(a) its Representatives who have a genuine need to know for the Purpose and who are bound by confidentiality obligations no less restrictive than this Clause 6 (whether by professional duty or written undertaking);")
                self.p("(b) its bona fide Financing Parties and potential co-investors, provided they are bound by confidentiality obligations no less restrictive than this Clause 6; and")
                self.p("(c) to the extent required by applicable law, regulation, court order, or the rules of any relevant regulatory authority or stock exchange, provided that (where legally permitted) the disclosing Party: (i) gives the other Party prior written notice as soon as reasonably practicable; (ii) consults with the other Party regarding the scope and manner of disclosure; and (iii) discloses only the minimum information required to comply.")
                self.bp("6.5 Liability for Representatives. ", "Each Party shall be responsible for any breach of this Clause 6 by its Representatives.")
                self.bp("6.6 Exclusions. ", "The obligations in Clauses 6.1 through 6.3 do not apply to information that the receiving Party can demonstrate:")
                self.p("(a) is or becomes publicly available through no fault of the receiving Party or its Representatives;")
                self.p("(b) was already in the lawful possession of the receiving Party before disclosure, without restriction as to use or disclosure;")
                self.p("(c) was independently developed by the receiving Party without use of or reference to the Confidential Information; or")
                self.p("(d) was received from a third party who was not, to the receiving Party's knowledge, under any obligation of confidentiality in respect of that information.")
                self.bp("6.7 No Implied Rights. ", "No licence or right is granted under this LOI to the receiving Party in respect of any intellectual property rights of the disclosing Party. All Confidential Information remains the property of the disclosing Party.")
                self.bp("6.8 Return and Destruction. ", "Upon the earlier of: (a) the disclosing Party's written request, or (b) the expiry or termination of this LOI, the receiving Party shall promptly return or destroy all documents, materials, and tangible items containing Confidential Information and certify such return or destruction in writing within 15 Business Days. The receiving Party may retain copies to the extent required by applicable law or its internal compliance policies, provided such retained copies remain subject to this Clause 6.")
                self.bp("6.9 Onward-Sharing Controls. ", "If the receiving Party receives an inquiry from any third party regarding the disclosing Party's Confidential Information, the receiving Party shall: (a) not respond to such inquiry without the disclosing Party's prior written consent; and (b) promptly notify the disclosing Party of such inquiry. The receiving Party shall not further distribute or re-disclose Confidential Information beyond the persons authorised under Clause 6.4 without the disclosing Party's prior written consent.")
                self.bp("6.10 Compliance Confirmation. ", "Upon the disclosing Party's reasonable written request (not more than once per calendar year), the receiving Party shall confirm in writing its compliance with the obligations in this Clause 6.")
                self.bp("6.11 Breach Notification. ", "Each Party shall notify the other Party in writing within 72 hours of becoming aware of any actual or suspected breach of this Clause 6, and shall take all reasonable steps to mitigate the effects of such breach.")
                self.bp("6.12 Disclaimer. ", 'All Confidential Information is disclosed "as is." The disclosing Party makes no representation or warranty, express or implied, as to the accuracy, completeness, or reliability of any Confidential Information. The receiving Party shall be solely responsible for its own assessment and due diligence.')
                self.bp("6.13 Metadata Protection. ", "Confidential Information includes metadata, EXIF data, geolocation data, timestamps, file names, folder names, and any digital artifacts associated with or derived from disclosed materials. The receiving Party shall not extract, analyse, or use such metadata except as necessary for the Purpose.")
                surv_text = f"6.14 Survival. The obligations in this Clause 6 shall survive termination or expiry of this LOI for a period of {surv} from the date of termination or expiry. Obligations in respect of information that constitutes a trade secret under applicable law shall continue indefinitely."
                if self.t == "Distributor":
                    surv_text += " Obligations in respect of Protected Business Information shall survive for the period specified in Clause 4.1."
                self.bp("6.14 Survival. ", surv_text.replace("6.14 Survival. ", ""))
                transaction_ref = "the Transaction" if self.t == "Distributor" else "the proposed transaction"
                self.bp("6.15 Announcements. ", f"Neither Party shall make any public announcement regarding this LOI or {transaction_ref} without the prior written consent of the other Party, except as required by applicable law.")
                self.bp("6.16 Remedies. ", "Each Party acknowledges that a breach of this Clause 6 may cause the disclosing Party irreparable harm for which damages would not be an adequate remedy. The disclosing Party shall be entitled to seek injunctive or other equitable relief from any court of competent jurisdiction, without the need to prove actual loss and without prejudice to any other rights or remedies.")

    def clause7_nc(self):
        if self.t == "EndUser" or self.t == "EcosystemPartnership":
            return  # No NC for End Users or Ecosystem Partnerships

        self.h("7. Non-Circumvention (BINDING)")
        nc_dur = self.g("protection", "nc_duration", default="24 months")

        # v3.3: for SS, downstream agreement is "Framework Agreement"; otherwise "MSA".
        downstream = "Framework Agreement" if self.t == "StrategicSupplier" else "MSA"

        self.p(f"7.1 The {self.party} shall not, directly or indirectly, without the prior written consent of Digital Energy:")

        if self.t == "Distributor":
            self.p(f"(a) contact, solicit, deal with, or enter into any business relationship with any Associated Counterparty introduced by Digital Energy in connection with the Purpose or the Transaction;")
            self.p(f"(b) circumvent, avoid, or bypass Digital Energy in order to deal directly or indirectly with any Associated Counterparty; or")
            self.p(f"(c) attempt to divert or appropriate any business opportunity disclosed by Digital Energy in connection with the Purpose or the Transaction.")
        elif self.t == "StrategicSupplier":
            self.p(f"(a) contact, solicit, deal with, or enter into any business relationship with any Associated Counterparty introduced by Digital Energy in connection with this LOI or any Provider project; or")
            self.p(f"(b) circumvent, avoid, or bypass Digital Energy in order to deal directly or indirectly with any Associated Counterparty in connection with the development, ownership, or operation of colocation or energy infrastructure on or adjacent to a site identified by Digital Energy.")
        else:  # Wholesale — lighter scope
            self.p(f"(a) contact, solicit, deal with, or enter into any business relationship with any Associated Counterparty introduced by Digital Energy in connection with this LOI or the Services; or")
            self.p(f"(b) circumvent, avoid, or bypass Digital Energy in order to deal directly or indirectly with any Associated Counterparty in connection with the development, ownership, or operation of colocation or energy infrastructure on or adjacent to a site identified by Digital Energy.")

        self.bp("7.2 Duration. ", f"The obligations in this Clause 7 shall continue for {nc_dur} after the earlier of: (a) the expiry or termination of this LOI; or (b) the expiry or termination of the {downstream}, if one is executed.")
        self.bp("7.3 Deemed Introduction. ", f"Digital Energy's sharing of any Site Identifier with the {self.party} shall constitute a deemed introduction of all Associated Counterparties for that site. Digital Energy is not required to separately name each Associated Counterparty.")

        if self.t == "Distributor":
            self.bp("7.4 Scope \u2014 Private and Public Bodies. ", "The non-circumvention obligations in Clause 7.1 apply to:")
            self.p("(a) Private Associated Counterparties (landowners, greenhouse operators, heat offtakers, energy counterparties, EPCs): in all cases where Digital Energy has introduced or disclosed their identity or involvement; and")
            self.p("(b) Public Bodies (government agencies, municipalities, regulatory bodies, distribution system operators): only where Digital Energy has made a specific, named introduction to an individual or department within that body. General knowledge that Digital Energy engages with a public body does not trigger non-circumvention protection.")
        elif self.t == "StrategicSupplier":
            self.bp("7.4 Scope Limitation. ", f"The non-circumvention obligations in this Clause 7 are limited to Digital Energy's supply-side relationships (site partners, energy counterparties, and infrastructure providers). Nothing in this Clause 7 restricts the {self.party} from conducting its own business with its existing or future customers, including other AI colocation operators, cloud service providers, or enterprise customers.")
        else:
            self.bp("7.4 Scope Limitation. ", f"The non-circumvention obligations in this Clause 7 are limited to Digital Energy's supply-side relationships (site partners, energy counterparties, and infrastructure providers). Nothing in this Clause 7 restricts the {self.party} from conducting its own business with its existing or future end-user customers, cloud service customers, or compute buyers.")

        self.bp("7.5 Independent Knowledge Exception. ", f"The obligations in Clause 7.1 do not apply to any Associated Counterparty with whom the {self.party} can demonstrate, by contemporaneous written evidence, that it had an existing business relationship or substantive commercial contact before the date of this LOI or before Digital Energy's disclosure.")
        self.bp(f"7.6 {downstream} Supersession. ", f"If the Parties execute a {downstream}, the non-circumvention provisions of the {downstream} shall replace this Clause 7 upon execution of the {downstream}, except that the survival period in Clause 7.2 shall apply to any introduction made before the {downstream} effective date that is not separately covered by the {downstream}.")

    def clause_general(self):
        cl = "8" if self.t != "EndUser" else "7"
        self.h(f"{cl}. General Provisions (BINDING)")
        val_date = self.g("dates", "validity_date")
        if not val_date:
            val_date = "[12 months from LOI date]"

        # v3.3: SS refers to "Framework Agreement" as downstream binding document.
        downstream = "Framework Agreement" if self.t == "StrategicSupplier" else "MSA"

        self.bp(f"{cl}.1 Non-Binding Status. ", "")
        if self.t != "EndUser":
            # v3.6.0 bug 2: SS Cl. 5 is "Supply Chain and Delivery Commitment",
            # not "Project Finance and Assignment". EP has no Cl. 5 finance
            # either (IP and Deliverables). Branch on self.t.
            if self.t == "StrategicSupplier":
                cl5_label = "Supply Chain and Delivery Commitment"
            elif self.t == "EcosystemPartnership":
                cl5_label = "IP and Deliverables"
            else:
                cl5_label = "Project Finance and Assignment"
            self.p(f"(a) Non-binding provisions. Clauses 2 through 4 and Schedule 1 of this LOI are non-binding expressions of the Parties' current intentions. They do not create legally enforceable obligations and are subject to the negotiation and execution of the {downstream}.")
            self.p(f"(b) Binding provisions. Clauses 5 ({cl5_label}), 6 (Confidentiality), 7 (Non-Circumvention), and {cl} (General Provisions) are legally binding and enforceable obligations.")
        else:
            self.p("(a) Non-binding provisions. Clauses 2 through 4 of this LOI are non-binding expressions of the Parties' current intentions. They do not create legally enforceable obligations and are subject to the negotiation and execution of the MSA.")
            self.p(f"(b) Binding provisions. Clauses 5 (Project Finance and Assignment), 6 (Confidentiality), and {cl} (General Provisions) are legally binding and enforceable obligations.")

        scoping_phrase = "commercial scoping and negotiation" if self.t == "Distributor" else "technical scoping and commercial negotiation"
        # v3.6.0 bug 1: non-EndUser previously concatenated "..., t" + "t"
        # producing "tthe good faith". Collapse to single 't' on non-EU.
        self.bp(f"{cl}.2 Good Faith. ", f"The Parties agree to engage in the {scoping_phrase} process in good faith (te goeder trouw) and in accordance with the principles of reasonableness and fairness (redelijkheid en billijkheid) as contemplated by Article 6:248 of the Dutch Civil Code (Burgerlijk Wetboek). " + ("" if self.t == "EndUser" else "For the avoidance of doubt, ") + ("T" if self.t == "EndUser" else "t") + f"he good faith obligation does not oblige either Party to enter into the {downstream}. Either Party may discontinue negotiations at any time, provided it does so in good faith. Any liability arising from a breach of this good faith obligation shall be limited to verifiable reliance damages (negatief contractsbelang)" + ("." if self.t == "EndUser" else " and shall not extend to loss of profit or expectation damages (positief contractsbelang)."))

        survive = "Clauses 5.2, 5.3, 6, and 7" if self.t != "EndUser" else "Clauses 5.2, 5.3, and 6"
        self.bp(f"{cl}.3 Validity. ", f"This LOI shall remain valid until {val_date}, after which it shall lapse automatically unless extended by mutual written agreement. Upon lapse, {survive} shall survive for their respective stated periods.")
        self.bp(f"{cl}.4 Costs. ", "Each Party shall bear its own costs in connection with this LOI" + (f" and the negotiation of the {downstream}." if self.t != "EndUser" else "."))
        self.bp(f"{cl}.5 Counterparts. ", "This LOI may be executed in counterparts, including by electronic signature within the meaning of the eIDAS Regulation (EU) No 910/2014." + (" Each counterpart constitutes an original." if self.t != "EndUser" else ""))
        self.bp(f"{cl}.6 Notices. ", "All notices under this LOI shall be in writing (including email) and addressed to the contact details in the preamble. A notice is effective upon receipt." + (" Each Party shall promptly notify the other of any change to its contact details." if self.t != "EndUser" else ""))
        self.bp(f"{cl}.7 Governing Law. ", "This LOI shall be governed by and construed in accordance with the laws of the Netherlands. The United Nations Convention on Contracts for the International Sale of Goods (CISG) is expressly excluded.")
        self.bp(f"{cl}.8 Jurisdiction. ", "The courts of Amsterdam (Rechtbank Amsterdam) shall have exclusive jurisdiction over any dispute arising out of or in connection with the binding provisions of this LOI.")

        if self.t != "EndUser":
            if self.t == "Distributor":
                transaction_ref = "the Transaction"
            elif self.t == "StrategicSupplier":
                transaction_ref = "the proposed supply relationship"
            else:
                transaction_ref = "the proposed transaction"
            # v3.6.0 bug 3: "(ALT-A)" is a drafting-variant marker, never
            # counterparty-facing. Gate phrase on existing_nda; when there
            # is no prior NDA, drop the reference entirely.
            if self.choice("existing_nda"):
                ea_body = f"This LOI, together with the NDA referenced in Clause 6, constitutes the entire agreement between the Parties in relation to its subject matter and supersedes all prior negotiations, representations, and agreements relating to {transaction_ref}. Nothing in this Clause limits liability for fraud."
            else:
                ea_body = f"This LOI constitutes the entire agreement between the Parties in relation to its subject matter and supersedes all prior negotiations, representations, and agreements relating to {transaction_ref}. Nothing in this Clause limits liability for fraud."
            self.bp(f"{cl}.9 Entire Agreement. ", ea_body)
            self.bp(f"{cl}.10 Partnership Disclaimer. ", "Nothing in this LOI shall be construed as creating a partnership, joint venture, agency, or employment relationship between the Parties. Neither Party has authority to bind the other or to incur any obligation on the other's behalf.")
        else:
            self.bp(f"{cl}.9 Entire Agreement. ", "This LOI constitutes the entire agreement between the Parties in relation to its subject matter and supersedes all prior negotiations, representations, and agreements. Nothing in this Clause limits liability for fraud.")

    def signature(self):
        self.line()
        # v3.3: honours OPEN-1 (commit 2097f52) — bespoke_closing removed.
        # Operational near-signature content belongs in the letter body, not
        # the closing line. Hardcoded single-sentence closing matches main.
        # If choices.bespoke_closing is present in YAML, it is silently ignored
        # for backward compatibility.
        #
        # v3.5 scope A' (signature block cleanup): removed KvK / registration
        # lines (duplicate Parties Preamble — scope A''') and removed the
        # "ACKNOWLEDGED AND AGREED" header (unnecessary on bilateral
        # instruments; signing = agreement). Added Place: field per Dutch/EU
        # execution convention (jurisdictional + eIDAS relevance).
        self.p("We look forward to working with you.")
        self.p("")
        self.p("Yours faithfully,")
        self.p("")

        prov = self.d.get("provider", {})
        self.p(f"For and on behalf of {prov.get('legal_name', '')}", bold=True)
        self.p("")
        self.p("Signature: ____________________________")
        self.p(f"Name: {self._render_placeholder(prov.get('signatory_name'), 'sig_block_name')}")
        self.p(f"Title: {self._render_placeholder(prov.get('signatory_title'), 'sig_block_title')}")
        self.p("Date: ____________________________")
        self.p("Place: ____________________________")

        self.p("")
        self.line()
        self.p("")

        cp = self.d.get("counterparty", {})
        self.p(f"For and on behalf of {cp.get('name', '')}", bold=True)
        self.p("")
        self.p("Signature: ____________________________")
        self.p(f"Name: {self._render_placeholder(cp.get('signatory_name'), 'sig_block_name')}")
        self.p(f"Title: {self._render_placeholder(cp.get('signatory_title'), 'sig_block_title')}")
        self.p("Date: ____________________________")
        self.p("Place: ____________________________")

    def schedule(self):
        # Schedule starts on its own page.
        # v3.2: schedule titles no longer carry "(NON-BINDING)" suffix. Italic
        # prefatory note carries the non-binding signal. Authoritative
        # binding/non-binding assignment lives in Cl. 5.1 / Cl. 8.1.
        self.doc.add_page_break()
        prefatory = ("This schedule expresses the Parties' current intentions "
                      "and is non-binding, in accordance with Clause 5.1.")
        if self.t == "Distributor":
            self.line()
            self.h("Schedule 1 \u2014 Partnership Details")
            self.p(prefatory, italic=True, color=GREY, size=FONT_SMALL)
            comm = self.d.get("commercial", {})
            self.table(
                ["Item", "Detail"],
                [
                    ["Partnership Mode", self.d.get("partnership_mode", "combined").title()],
                    ["Target Customer Segments", comm.get("target_segments", "")],
                    ["Geographic Scope", comm.get("territory", "")],
                    ["Estimated Annual Capacity", f"{comm.get('indicative_mw', '')} MW IT"],
                    ["Commercial scoping complete", f"{self.g('dates', 'loi_date')} + 30 days"],
                    ["Draft MSA issued", f"{self.g('dates', 'loi_date')} + 60 days"],
                    ["MSA executed", f"{self.g('dates', 'loi_date')} + 90 days"],
                ]
            )
        elif self.t == "StrategicSupplier":
            self.line()
            self.h("Schedule 1 \u2014 Scope and Capability Matrix")
            self.p(prefatory, italic=True, color=GREY, size=FONT_SMALL)
            supplier = self.d.get("supplier", {})
            purposes = supplier.get("strategic_purposes", [])
            rows = [
                ["Capability Category", supplier.get("capability_category", "")],
                ["Core Capability", supplier.get("core_capability", "")],
                ["Strategic Purposes", ", ".join(purposes)],
                ["Geographic Coverage", supplier.get("geographic_coverage", "")],
            ]
            if "capacity_lock_in" in purposes:
                rows.append(["Lead-Time Target", supplier.get("lead_time_target", "")])
            if "pricing_volume" in purposes:
                rows.append(["Indicative Volume", supplier.get("volume_indicative", "")])
            if "engineering_integration" in purposes:
                rows.append(["Joint IP Allocation", self.g("choices", "joint_ip", default="background")])
            rows += [
                ["Scoping complete", f"{self.g('dates', 'loi_date')} + 30 days"],
                ["Framework Agreement issued", f"{self.g('dates', 'loi_date')} + 60 days"],
                ["Framework Agreement executed", f"{self.g('dates', 'loi_date')} + 90 days"],
            ]
            self.table(["Item", "Detail"], rows)
        elif self.t == "Wholesale":
            self.line()
            self.h("Schedule 1 \u2014 Capacity and Technical Requirements")
            self.p(prefatory, italic=True, color=GREY, size=FONT_SMALL)
            comm = self.d.get("commercial", {})
            tech = self.d.get("schedule_1", {}).get("technical", {}) or self.d.get("technical", {})
            # v3.5 scope N (Jonathan memo): Schedule 1 now reads technical
            # fields from intake YAML. Previously hardcoded "[To be confirmed]"
            # placeholders regardless of intake content — caused GPU platform
            # commitments established in email/user intake to be lost in the
            # Polarise LOI. Default values kept for backward-compat when
            # intake omits `schedule_1.technical.*`.
            gpu_platform = self._render_placeholder(
                tech.get("gpu_platform"), "body_clause"
            ) or "[To be confirmed during technical scoping]"
            rack_density = tech.get("rack_density_kw") or comm.get("rack_density_kw")
            rack_density_cell = (
                f"{rack_density} kW/rack"
                if rack_density and not self._is_tbc(rack_density)
                else "[To be confirmed]"
            )
            cooling = tech.get("cooling") or comm.get(
                "cooling_topology", "Direct-to-chip liquid cooling"
            )
            designated_sites = tech.get("designated_sites") or self._render_placeholder(
                tech.get("designated_sites"), "body_clause"
            ) or "To be confirmed during technical scoping"
            self.table(
                ["Item", "Detail"],
                [
                    ["Indicative Capacity", f"{comm.get('indicative_mw', '')} MW IT"],
                    ["Designated Sites", designated_sites],
                    ["GPU / Accelerator Type", gpu_platform],
                    ["Target Rack Density", rack_density_cell],
                    ["Cooling Requirement", cooling],
                    ["Technical scoping complete", f"{self.g('dates', 'loi_date')} + 30 days"],
                    ["Credit assessment complete", f"{self.g('dates', 'loi_date')} + 30 days"],
                    ["Sales Order Form issued", f"{self.g('dates', 'loi_date')} + 60 days"],
                    ["MSA executed", f"{self.g('dates', 'loi_date')} + 90 days"],
                    ["Expansion Target", f"{comm.get('expansion_mw', '')} MW IT"],
                ]
            )
        else:  # EndUser
            self.line()
            self.h("Schedule 1 \u2014 Service Requirements")
            self.p(prefatory, italic=True, color=GREY, size=FONT_SMALL)
            comm = self.d.get("commercial", {})
            tech = self.d.get("schedule_1", {}).get("technical", {}) or self.d.get("technical", {})
            types = comm.get("service_type", [])
            type_str = ", ".join(t.replace("_", " ").title() for t in types)
            gpu_platform = self._render_placeholder(
                tech.get("gpu_platform"), "body_clause"
            ) or "[To be confirmed during technical scoping]"
            data_sovereignty = self._render_placeholder(
                tech.get("data_sovereignty"), "body_clause"
            ) or "[To be confirmed]"
            self.table(
                ["Item", "Detail"],
                [
                    ["Service Model", type_str],
                    ["Indicative Capacity", comm.get("indicative_capacity", "")],
                    ["GPU / Accelerator Type", gpu_platform],
                    ["Data Sovereignty", data_sovereignty],
                    ["Technical scoping complete", f"{self.g('dates', 'loi_date')} + 30 days"],
                    ["Commercial proposal issued", f"{self.g('dates', 'loi_date')} + 60 days"],
                    ["MSA executed", f"{self.g('dates', 'loi_date')} + 90 days"],
                ]
            )

    def footer(self):
        # Real Word footers are set up in _setup(). Add version reference at end of body.
        self.p("")
        # v3.3: version string per-type. New types (SS, EP) are v1.0.
        version_by_type = {
            "EndUser": "v3.2",
            "Distributor": "v3.2",
            "Wholesale": "v3.2",
            "StrategicSupplier": "v1.0",
            "EcosystemPartnership": "v1.0",
        }
        vsn = version_by_type.get(self.t, "v3.2")
        self.p(
            f"DE-LOI-{self.t}-{vsn}",
            italic=True, size=FONT_SMALL, color=GREY
        )

    # ----- EP (Ecosystem Partnership) clause builders -----
    # v3.3: EP uses a dedicated, lighter structure. No Cl. 5 Project Finance,
    # no Cl. 7 NC. Cl. 5 is IP & Deliverables. Cl. 6 is Tier A mutual light
    # (7 sub-clauses). Cl. 7 is General (11 sub-clauses).

    def definitions_ep(self):
        self.h("1. Definitions")
        self.p("1.1 In this LOI, unless the context requires otherwise:")
        self.bp(
            '"Confidential Information" ',
            "means all information (whether technical, commercial, financial, or "
            "otherwise) disclosed by a Party to the other in connection with the "
            "Collaboration, whether in writing, orally, electronically, or by "
            "inspection, including any information that by its nature or the "
            "circumstances of disclosure would reasonably be understood to be "
            "confidential. The existence and contents of this LOI are Confidential "
            "Information."
        )
        self.bp(
            '"Collaboration" ',
            "means the ecosystem activities and joint work contemplated by this "
            "LOI as described in Clause 3."
        )
        self.bp(
            '"Representatives" ',
            "means, in relation to a Party, its Affiliates and its and their "
            "respective directors, officers, employees, agents, and professional "
            "advisers."
        )

    def clause2_ep(self):
        self.h("2. Purpose and Scope")
        eco = self.d.get("ecosystem", {})
        themes = ", ".join(eco.get("collaboration_themes", [])) or "[COLLABORATION_THEMES]"
        self.p(
            f"2.1 This LOI records the Parties' mutual interest in a "
            f"non-commercial ecosystem collaboration around {themes}."
        )
        self.p(
            "2.2 The collaboration is strictly non-commercial. No Party shall be "
            "entitled to any fee, payment, capacity, discount, or commercial "
            "benefit from the other Party by reason of this LOI or any activity "
            "under it."
        )
        self.p(
            "2.3 The commercial implications of any future activity \u2014 including "
            "any joint offering, research licensing, commercialisation of joint "
            "output, or provision of services by either Party to the other \u2014 shall "
            "be the subject of a separate agreement and are expressly outside the "
            "scope of this LOI."
        )
        self.p(
            "2.4 The confidentiality and general provisions in Clauses 6 and 7 "
            "are legally binding and enforceable from the date of execution. "
            "Clauses 3 through 5 are non-binding expressions of collaborative intent."
        )

    def clause3_ep(self):
        self.h("3. Collaboration Scope")
        eco = self.d.get("ecosystem", {})
        themes = eco.get("collaboration_themes", [])
        categories = eco.get("joint_activity_categories", [])

        self.bp("3.1 Collaboration Themes. ", "The Parties intend to collaborate on the following themes:")
        if themes:
            for theme in themes:
                self.p(f"\u2014 {theme}")
        else:
            self.p("\u2014 [THEMES TO BE CONFIRMED]")

        self.bp(
            "3.2 Joint Activity Categories. ",
            "Subject to resources and mutual agreement, the Parties intend to "
            "engage in one or more of the following categories of activity:"
        )
        if categories:
            cat_labels = {
                "publications": "publications (joint white papers, research articles, reports)",
                "events": "events (conferences, roundtables, workshops)",
                "pilots": "pilots (joint demonstrations, proofs of concept)",
                "advocacy": "advocacy (policy input, consultation responses)",
                "working_groups": "working groups (standards and technical coordination)",
            }
            for cat in categories:
                self.p(f"(\u2014) {cat_labels.get(cat, cat)}")
        else:
            self.p("\u2014 [ACTIVITY CATEGORIES TO BE CONFIRMED]")

        self.bp(
            "3.3 Governance and Cadence. ",
            "The Parties intend to establish a lightweight governance mechanism "
            "\u2014 typically an initial set-up meeting and quarterly check-ins \u2014 to "
            "review activity, align on priorities, and resolve any coordination "
            "issues. No binding governance structure is created by this LOI."
        )
        self.bp(
            "3.4 Working-Group Participation. ",
            "Where either Party hosts or convenes a working group relevant to the "
            "collaboration themes, the other Party shall have a standing "
            "invitation to participate on mutually agreed terms."
        )
        self.bp(
            "3.5 Non-Exclusivity. ",
            "The collaboration contemplated by this LOI is non-exclusive. Both "
            "Parties retain the right to enter into similar or competing "
            "collaborations with third parties."
        )

    def clause4_ep(self):
        self.h("4. Announcements and Branding")
        announcement = self.g("choices", "announcement_protocol", default="mutual_approval")
        logo_use = self.g("choices", "logo_use", default="reciprocal")

        if announcement == "notify_only":
            self.bp(
                "4.1 Announcement Protocol. ",
                "Each Party shall give the other Party reasonable prior notice "
                "(target: 5 Business Days) before making any public announcement "
                "referring to the collaboration. Consent is not required."
            )
        else:  # mutual_approval (default)
            self.bp(
                "4.1 Announcement Protocol. ",
                "Any public announcement, press release, published article, or "
                "social media post referring to the collaboration, the other "
                "Party, or joint activity under this LOI shall require the prior "
                "written consent of the other Party. Consent shall not be "
                "unreasonably withheld."
            )

        if logo_use == "none":
            self.bp(
                "4.2 Logo Use. ",
                "Neither Party shall use the other Party's name, logo, or trade "
                "mark in any public context without the prior written consent of "
                "the other Party."
            )
        elif logo_use == "one_way":
            self.bp(
                "4.2 Logo Use. ",
                "One-way logo use is contemplated between the Parties; the "
                "Party, scope, and terms will be specified in a separate "
                "schedule or letter."
            )
        else:  # reciprocal (default)
            self.bp(
                "4.2 Logo Use. ",
                "Each Party grants the other a limited, non-exclusive, revocable "
                "licence to use its name and logo solely in the context of the "
                "collaboration and subject to the branding guidelines each Party "
                "provides. No other use is permitted."
            )

        self.bp(
            "4.3 Attribution. ",
            "Where either Party refers to the collaboration in any public "
            "context, it shall: (a) describe the collaboration accurately and "
            "without overstatement; (b) refrain from implying any commercial "
            "relationship, endorsement, or joint venture; and (c) include the "
            "other Party's preferred attribution language when supplied."
        )

    def clause5_ep(self):
        self.h("5. Intellectual Property and Deliverables")
        self.bp(
            "5.1 Background IP. ",
            "Each Party retains all right, title, and interest in its background "
            "intellectual property. Nothing in this LOI transfers or licenses any "
            "intellectual property between the Parties except as expressly set "
            "out in Clause 4.2."
        )
        self.bp(
            "5.2 Joint Deliverables. ",
            "Any intellectual property jointly developed by the Parties in the "
            "course of the collaboration (including joint publications, "
            "working-group outputs, and jointly authored materials) shall be "
            "governed by a separate agreement to be executed before the joint "
            "activity produces any material output."
        )
        self.bp(
            "5.3 No Assignment by Participation. ",
            "Nothing in this LOI or in any activity conducted under it shall be "
            "construed as an assignment, licence, or transfer of intellectual "
            "property."
        )
        self.bp(
            "5.4 Publication. ",
            "Each Party retains the right to publish, present, and disseminate "
            "its own work, subject to the confidentiality provisions in Clause 6."
        )

    def clause6_ep(self):
        self.h("6. Confidentiality (BINDING)")
        surv = self.g("protection", "confidentiality_survival", default="2 years")
        self.p(
            "6.1 Each Party shall keep confidential all Confidential Information "
            "received from the other Party and shall use such information solely "
            "for the Collaboration."
        )
        self.p(
            "6.2 Each Party may disclose Confidential Information only to its "
            "Representatives who have a genuine need to know for the "
            "Collaboration and who are bound by confidentiality obligations no "
            "less restrictive than this Clause 6."
        )
        self.p(
            "6.3 The obligations in Clauses 6.1 and 6.2 do not apply to "
            "information that the receiving Party can demonstrate: (a) is or "
            "becomes publicly available through no fault of the receiving Party; "
            "(b) was already in the receiving Party's possession without "
            "restriction; (c) was independently developed without use of the "
            "Confidential Information; or (d) was received from a third party "
            "without breach of any confidentiality obligation."
        )
        self.p(
            "6.4 A Party may disclose Confidential Information to the extent "
            "required by applicable law, regulation, or court order, provided "
            "that (where legally permitted) the disclosing Party gives the other "
            "Party prior written notice and discloses only the minimum required."
        )
        self.p(
            "6.5 Upon written request or expiry of this LOI, the receiving Party "
            "shall promptly return or destroy all Confidential Information and "
            "certify such return or destruction in writing. Copies required by "
            "applicable law or internal compliance policies may be retained, "
            "subject to continued confidentiality."
        )
        self.p(
            "6.6 Public statements about this LOI are governed by Clause 4."
        )
        self.p(
            f"6.7 The obligations in this Clause 6 shall survive for {surv} from "
            f"the date of termination or expiry of this LOI."
        )

    def clause7_ep_general(self):
        self.h("7. General Provisions (BINDING)")
        val_date = self.g("dates", "validity_date") or "[12 months from LOI date]"

        self.bp("7.1 Non-Binding Status. ", "")
        self.p(
            "(a) Non-binding provisions. Clauses 2, 3, 4, and 5 of this LOI are "
            "non-binding expressions of collaborative intent."
        )
        self.p(
            "(b) Binding provisions. Clauses 6 (Confidentiality) and 7 (General "
            "Provisions) are legally binding and enforceable obligations."
        )
        self.bp(
            "7.2 No Commercial Commitment. ",
            "Nothing in this LOI creates any obligation on either Party to pay, "
            "purchase, supply, or grant anything of commercial value to the "
            "other. Any such arrangement requires a separate agreement."
        )
        self.bp(
            "7.3 Non-Exclusivity. ",
            "As per Clause 3.5."
        )
        self.bp(
            "7.4 Term and Validity. ",
            f"This LOI shall remain valid until {val_date}, after which it shall "
            f"lapse automatically unless extended by mutual written agreement. "
            f"Upon lapse, Clause 6 (Confidentiality) shall survive for its "
            f"stated period."
        )
        self.bp(
            "7.5 Costs. ",
            "Each Party shall bear its own costs in connection with this LOI "
            "and the Collaboration."
        )
        self.bp(
            "7.6 Counterparts. ",
            "This LOI may be executed in counterparts, including by electronic "
            "signature within the meaning of the eIDAS Regulation (EU) No 910/2014."
        )
        self.bp(
            "7.7 Notices. ",
            "All notices under this LOI shall be in writing (including email) "
            "and addressed to the contact details in the preamble. A notice is "
            "effective upon receipt."
        )
        self.bp(
            "7.8 Governing Law. ",
            "This LOI shall be governed by the laws of the Netherlands. The "
            "United Nations Convention on Contracts for the International Sale "
            "of Goods (CISG) is expressly excluded."
        )
        self.bp(
            "7.9 Jurisdiction. ",
            "The courts of Amsterdam (Rechtbank Amsterdam) shall have exclusive "
            "jurisdiction over any dispute arising out of or in connection with "
            "the binding provisions of this LOI."
        )
        self.bp(
            "7.10 Entire Agreement. ",
            "This LOI constitutes the entire agreement between the Parties in "
            "relation to its subject matter and supersedes all prior "
            "negotiations, representations, and agreements. Nothing in this "
            "Clause limits liability for fraud."
        )
        self.bp(
            "7.11 Partnership Disclaimer. ",
            "Nothing in this LOI shall be construed as creating a partnership, "
            "joint venture, agency, or employment relationship between the "
            "Parties. Neither Party has authority to bind the other or to incur "
            "any obligation on the other's behalf."
        )

    def schedule_ep(self):
        """v3.3: Optional Joint Activity Plan schedule for EP.
        Only rendered if intake provides ecosystem.joint_activity_plan (list of
        dicts with keys: activity, category, lead, target_date).
        """
        plan = self.d.get("ecosystem", {}).get("joint_activity_plan", [])
        if not plan:
            return
        self.doc.add_page_break()
        self.line()
        self.h("Schedule 1 \u2014 Joint Activity Plan")
        self.p(
            "This schedule expresses the Parties' current intentions and is "
            "non-binding, in accordance with Clause 2.4.",
            italic=True, color=GREY, size=FONT_SMALL,
        )
        rows = []
        for item in plan:
            rows.append([
                item.get("activity", ""),
                item.get("category", ""),
                item.get("lead", ""),
                item.get("target_date", ""),
            ])
        self.table(["Activity", "Category", "Lead", "Target Date"], rows)

    def build(self) -> Document:
        # v3.3: EP uses its own clause builders (will be added next); until wired,
        # fall back to partial engine output.
        if self.t == "EcosystemPartnership":
            return self._build_ep()

        self.letterhead()
        # add_cover() already ends with page break
        # v3.5.2 scope A''': Parties Preamble renders legal identification
        # in the body before recitals — cover page is presentation only.
        self.parties()
        self.recitals()
        self.definitions()
        self.clause2()

        if self.t == "Distributor":
            self.clause3_ds()
        elif self.t == "Wholesale":
            self.clause3_ws()
        elif self.t == "StrategicSupplier":
            self.clause3_ss()
        else:
            self.clause3_eu()

        if self.t == "StrategicSupplier":
            self.clause4_ss()
        else:
            self.clause4()

        # v3.4: SS uses clause5_ss (Supply Chain and Delivery Commitment).
        # EU/DS/WS use the generic clause5 (Project Finance and Assignment).
        if self.t == "StrategicSupplier":
            self.clause5_ss()
        else:
            self.clause5()
        self.clause6()
        self.clause7_nc()
        self.clause_general()
        self.signature()
        self.schedule()
        self.footer()
        return self.doc

    def _build_ep(self) -> Document:
        """v3.3: Ecosystem Partnership build pipeline.
        Different structure: no Cl. 5 Project Finance, no Cl. 7 NC.
        Cl. 5 is IP & Deliverables, Cl. 6 is Tier A light mutual, Cl. 7 is General.
        """
        self.letterhead()
        # v3.5.2 scope A''': Parties Preamble for EP too.
        self.parties()
        self.recitals()
        self.definitions_ep()
        self.clause2_ep()
        self.clause3_ep()
        self.clause4_ep()
        self.clause5_ep()
        self.clause6_ep()
        self.clause7_ep_general()
        self.signature()
        # Schedule 1 is optional for EP — only rendered if the intake provides one.
        if self.d.get("ecosystem", {}).get("joint_activity_plan"):
            self.schedule_ep()
        self.footer()
        return self.doc


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

import re  # noqa: E402 — used below by qa_lint


# v3.2 QA linter. Rules in sync with _shared/loi-qa-gate.md.

_ARROW_CHARS = "\u2192\u21D2\u279C\u27F6\u21A6\u27F9\u21E8"

_FAIL_RULES = {
    "R-1": (r"minimum commitment term of 5 years", "body",
            "'minimum commitment term of 5 years' banned; use 'approximately 5 years, indicative only'"),
    "R-2": (r"\b\d+\s+identified\s+sites\b", "Recital A",
            "Fixed site-count language in Recital A"),
    "R-3": (r"12 months of commercial commitment", "Recital A",
            "'12 months of commercial commitment' banned in Recital A"),
    "R-5": (r"We are confident that", "closing",
            "'We are confident that' banned"),
    "R-7": (f"[{_ARROW_CHARS}]", "body",
            "Unicode arrow detected in body"),
    "R-8": (r"Schedule \d+ [^\n]*\(NON-BINDING\)", "schedule-title",
            "'(NON-BINDING)' suffix in schedule title — use italic prefatory note instead"),
    "R-10": (r"4\.2 Revenue Chain", "Cl. 4.2",
             "Cl. 4.2 heading must be 'Contractual Sequence', not 'Revenue Chain'"),
    "R-20": (r"programme spans \d+", "Recital A",
             "'programme spans N...' language regression"),
    # R-23 (v3.4): fabrication gate — implemented as a custom check in qa_lint,
    # not a simple regex. See _check_fabrication_gate() below.
    # v3.5.2 scope 0 additions:
    "R-24": (
        r"\[[A-Za-z][A-Za-z0-9._-]*\.(?:com|eu|de|co\.uk|org|nl|io|ai|ch)\]",
        "Recital B",
        "Inline source citation in Recital B (e.g. [polarise.eu]) — source attribution lives in counterparty.source_map YAML, not in rendered prose. Strip brackets.",
    ),
    "R-25": (
        # Vanity-financial patterns — catches unattributed capital-raise language
        # and valuation vanity. Does NOT catch named-endorser financings like
        # "backed by Macquarie" or "strategic investment from NVIDIA".
        #
        # Patterns (any match = FAIL):
        #   • "Series A/B/C/D/E/F" (not followed by "from" — that would be a
        #     specific named investor round and should be Signal-Test-judged)
        #   • "valuation of" (standalone valuation phrase)
        #   • "at a €500m valuation" or "at EUR 500m valuation" (valuation adj)
        #   • "raised $150M" / "raised EUR 117m" (unattributed raise)
        #   • "growth commitment" (pure financial headline)
        #   • "SAFE round" / "convertible note"
        r"\b("
        r"Series\s+[A-F](?!\s+from\s+\w)"
        r"|valuation of"
        r"|at\s+(?:a\s+)?(?:\$|€|£|USD|EUR|GBP)?\s*\d[\d.,]*\s*(?:bn|m|million|billion)\s+valuation"
        r"|raised\s+(?:\$|€|£|USD|EUR|GBP)?\s*\d[\d.,]*(?:\s*(?:M|B|m|bn|million|billion))?"
        r"|growth commitment"
        r"|SAFE round"
        r"|convertible note"
        r")\b",
        "Recital B",
        "Vanity-financial claim in Recital B — valuation numbers / generic VC labels / unattributed capital-raise language fail Signal Test gate 1. Named-endorser financings (e.g. 'backed by Macquarie') are signal and remain allowed; pure vanity metrics are not.",
    ),
    "R-27": (
        r"(?:Name|Title):\s*\[TBC\]",
        "sig-block",
        "'[TBC]' rendered literally in signature-block Name or Title line — must route through _render_placeholder so the line becomes a fillable blank on external-facing drafts.",
    ),
}

_WARN_RULES = {
    "R-11": (r"\bISO \d{4,5}\b", "Recital B",
             "ISO certification in Recital B (set choices.cert_relevant=true to suppress)"),
    # v3.4: R-14 scope broadened from Recital B to body. Added "purpose-built"
    # and "state-of-the-art" as R-21 (kept R-14 list focused on the original
    # salesy adjectives for clarity in QA reports).
    "R-14": (r"\b(leading|innovative|cutting-edge|world-class|best-in-class)\b", "body",
             "Salesy adjective"),
    "R-15": (r"positioning (its|itself) as", "body",
             "'positioning (its|itself) as' — formulaic"),
    "R-19": (r"^\s*\d+(?:\.\d+)?\s+[^\n]*\(NON-BINDING\)", "heading",
             "'(NON-BINDING)' in a clause heading — v3.2 style regression"),
    # v3.4 additions:
    "R-21": (r"\b(purpose-built|state-of-the-art)\b", "body",
             "Marketing adjective — 'purpose-built' / 'state-of-the-art' banned body-wide (v3.4)"),
    "R-22": (
        # v3.5.3-cont scope E: patterns narrowed to clause-context signatures
        # to reduce false positives. Prior regex flagged bare phrases like
        # "Provider's ability to" and "will require the exchange of" which
        # can legitimately appear in operative clauses (e.g. information-
        # exchange provisions in Recital E of NDAs). Narrowed forms below
        # require the meta-commentary *verb* context that only meta-
        # commentary uses (secure/obtain/access for abilities; Confidential
        # Information/material for exchanges).
        #
        # Allowlist — these patterns are NOT meta-commentary and MUST NOT
        # fire R-22 even though they contain similar tokens:
        #   • "ability to deliver" / "ability to scale" (operational capability)
        #   • "will require the exchange of information" (without
        #     "Confidential Information" qualifier — could be legitimate)
        #   • "depends on" (without "in part" modifier)
        r"("
        r"(?:Provider's|Digital Energy's) ability to (?:secure|obtain|access)|"
        r"depends in part on|"
        r"is intended to evidence|"
        r"while non-binding in its commercial terms|"
        r"to support (?:the Provider's|Digital Energy's) financing|"
        r"will require the exchange of (?:Confidential Information|material)|"
        r"The Parties acknowledge that (?:the Provider|Digital Energy) intends|"
        r"is intended to form the basis"
        r")",
        "body",
        "Meta-commentary pattern — explains the LOI rather than creating obligations (v3.4; narrowed v3.5.3-cont)"
    ),
    # v3.5.2 scope 0 note: R-28 is implemented as a density custom-check in
    # qa_lint(), not a simple regex (needs occurrence count). See _check_tbc_density().
}

# R-23 fabrication gate: regex targets material numeric-metric claims in
# Recital B. Every trigger must be matched against counterparty.source_map
# or marked [TBC] in the intake; otherwise R-23 FAILS the build.
_R23_CLAIM_PATTERN = re.compile(
    r"\b\d+[\d,]*\s*"
    r"(MW|GW|customers|clients|sites|deployments|GPUs|operations|"
    r"offices|countries|years|employees|%)\b",
    re.IGNORECASE,
)


# v3.5.6 scope D: diagnostic accumulator populated by _check_fabrication_gate
# on PASS paths so qa_lint can surface pillar-attribution info in the QA
# report even when no findings fire. Module-level list (cleared at start of
# each qa_lint run) keeps the return-type of _check_fabrication_gate stable.
_R23_PILLAR_DIAGNOSTIC: list = []


def _split_recital_b_sentences(text: str) -> list:
    """v3.5.6 scope D.2: split Recital B on sentence / paragraph boundaries.

    A `[TBC]` marker only covers claims in the same split segment (not the
    whole Recital B as v3.5.x did). Splits on `[.?!]` + whitespace, and on
    blank-line (`\\n\\n`) paragraph boundaries.
    """
    if not text:
        return []
    # Split on sentence-ending punctuation followed by whitespace, OR on blank lines
    parts = re.split(r"(?<=[.?!])\s+|\n\n+", text)
    return [p.strip() for p in parts if p.strip()]


def _claim_is_tbc_covered(claim_start: int, recital_b: str) -> bool:
    """v3.5.6 scope D.2: `[TBC]` covers claim only in the same sentence-
    boundary segment, OR when `[TBC]` sits at Recital-B-end and the claim
    is in the final segment (common trailing-`[TBC]` pattern).
    """
    if not recital_b:
        return False
    segments = _split_recital_b_sentences(recital_b)
    if not segments:
        return False
    # Rebuild an offset map to find which segment contains claim_start.
    # Walk segments through the original text, tracking running offset.
    pos = 0
    claim_segment_idx = None
    for i, seg in enumerate(segments):
        # Find this segment in the remaining text (from pos onwards)
        idx = recital_b.find(seg, pos)
        if idx < 0:
            continue
        seg_end = idx + len(seg)
        if idx <= claim_start <= seg_end:
            claim_segment_idx = i
            break
        pos = seg_end
    if claim_segment_idx is None:
        return False
    # Primary: [TBC] (or [TO BE CONFIRMED]) in the same segment
    seg_text = segments[claim_segment_idx]
    if "[TBC]" in seg_text or "[TO BE CONFIRMED]" in seg_text:
        return True
    # Special case (design decision D.2): [TBC] in final segment covers
    # final-segment claims (common trailing-[TBC] drafting pattern).
    if claim_segment_idx == len(segments) - 1 and (
        "[TBC]" in segments[-1] or "[TO BE CONFIRMED]" in segments[-1]
    ):
        return True
    return False


def _pillar_with_urls(source_map: dict):
    """Return the first pillar key (e.g. 'pillar_3') that has at least one
    URL entry, or None if no pillar has any. Used for the permissive any-
    pillar match and the diagnostic emission (v3.5.6 D.1).
    """
    if not isinstance(source_map, dict):
        return None
    # Preserve iteration order — typically pillar_1..pillar_5.
    for pillar_key, pillar_val in source_map.items():
        if isinstance(pillar_val, list) and any(
            isinstance(u, str) and u.startswith(("http://", "https://"))
            for u in pillar_val
        ):
            return pillar_key
        if isinstance(pillar_val, str) and pillar_val.startswith(
            ("http://", "https://")
        ):
            return pillar_key
    return None


# v3.5.6 scope D.3: thin-reason + structured-short-code patterns for
# override-reason validation. See _validate_override_reason().
_THIN_OVERRIDE_REASON_RE = re.compile(
    r"^\s*(?:ok|fine|yes|done|sure|good|n/?a|tbd)\s*$",
    re.IGNORECASE,
)
_STRUCTURED_OVERRIDE_REASON_RE = re.compile(
    r"^(?:OK|FINE|APPROVED|PREAPPROVED)-\d{4}-\d{2}-\d{2}\s+[A-Z]{2,4}$"
)


def _validate_override_reason(reason: str):
    """v3.5.6 scope D.3: hybrid override-reason validator.

    Accepts:
      - free-text rationale (≥15 chars after strip, NOT a thin-pattern match), OR
      - structured audit short-code: `<STATUS>-<YYYY-MM-DD> <INITIALS>`
        where STATUS in (OK, FINE, APPROVED, PREAPPROVED) and INITIALS is
        2–4 uppercase letters.

    Returns (is_valid: bool, error_message: str). error_message is empty
    when is_valid is True.
    """
    if reason is None:
        return False, "override reason required when --override is set"
    r = str(reason).strip()
    if not r:
        return False, "override reason cannot be empty"
    # Reject thin patterns regardless of length
    if _THIN_OVERRIDE_REASON_RE.match(r):
        return False, (
            "override reason is a thin-pattern phrase (e.g. 'ok', 'fine'). "
            "Provide either:\n"
            "  - free-text rationale (minimum 15 characters), OR\n"
            "  - structured audit short-code: <STATUS>-<YYYY-MM-DD> <INITIALS>\n"
            "    where STATUS in (OK, FINE, APPROVED, PREAPPROVED) and\n"
            "    INITIALS is 2-4 uppercase letters.\n"
            "    Example: OK-2026-04-17 JG"
        )
    # Accept structured short-code regardless of length
    if _STRUCTURED_OVERRIDE_REASON_RE.match(r):
        return True, ""
    # Accept free-text if length >= 15
    if len(r) >= 15:
        return True, ""
    return False, (
        f"override reason too short ({len(r)} chars). Provide either:\n"
        "  - free-text rationale (minimum 15 characters), OR\n"
        "  - structured audit short-code: <STATUS>-<YYYY-MM-DD> <INITIALS>\n"
        "    Example: OK-2026-04-17 JG"
    )


def _check_fabrication_gate(text_recital_b: str, source_map: dict,
                             overrides: set) -> list:
    """R-23: every material numeric claim in Recital B must be attributable.

    Returns a list of (rule_id, scope, message) findings. Empty list = pass.

    v3.5.6 scope D.1 + D.2 updates:
    - `[TBC]` proximity is sentence-boundary-scoped (not wildcard across
      whole Recital B). See _claim_is_tbc_covered().
    - On PASS, populates the module-level `_R23_PILLAR_DIAGNOSTIC` list
      with (claim_text, pillar_matched_or_'TBC-covered') tuples so qa_lint
      can surface the attribution info in the QA report.

    A claim is considered attributed if ANY of:
    - the claim is `[TBC]`-covered within the same sentence segment, OR
    - counterparty.source_map has at least one URL entry in ANY pillar
      (pillar-level granularity; reviewer's Phase 7.5 handles pillar-to-claim
      mapping), OR
    - the rule R-23 is in overrides (user ran with --override R-23 ...).
    """
    _R23_PILLAR_DIAGNOSTIC.clear()
    findings = []
    if "R-23" in overrides:
        return findings  # overridden
    claims = list(_R23_CLAIM_PATTERN.finditer(text_recital_b))
    if not claims:
        return findings  # no material claims

    pillar_key = _pillar_with_urls(source_map)

    # Walk each claim; if neither sentence-scoped [TBC] nor any pillar URL,
    # it's an unattributed claim. Collect diagnostic on the pass path.
    unattributed_claims = []
    for c in claims:
        claim_text = c.group(0)
        if _claim_is_tbc_covered(c.start(), text_recital_b):
            _R23_PILLAR_DIAGNOSTIC.append((claim_text, "TBC-covered"))
            continue
        if pillar_key:
            _R23_PILLAR_DIAGNOSTIC.append((claim_text, pillar_key))
            continue
        unattributed_claims.append(claim_text)

    if unattributed_claims:
        # Clear diagnostic on FAIL — reviewer focuses on the failures
        _R23_PILLAR_DIAGNOSTIC.clear()
        sample = unattributed_claims[:3]
        findings.append((
            "R-23", "Recital B",
            f"Fabrication gate: {len(unattributed_claims)} material claim(s) "
            f"in Recital B without pillar attribution or sentence-scoped "
            f"[TBC] marker (e.g. {', '.join(repr(s) for s in sample)}). "
            f"Add counterparty.source_map pillar_N: ['https://tier-1-url'] "
            f"to intake YAML, mark claims [TBC] within the same sentence, "
            f"or pass --override R-23 --override-reason \"<verbose>\"."
        ))
    return findings


def _extract_text(doc) -> str:
    """Extract all paragraph text from a python-docx Document, including tables."""
    lines = []
    for p in doc.paragraphs:
        lines.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    lines.append(p.text)
    return "\n".join(lines)


def qa_lint(doc, data: dict, builder_findings: list, overrides: set,
            override_reason: str):
    """Run QA rules over the built Document. Returns (status, report_lines).

    status: "PASS" | "PASS_WITH_WARN" | "FAIL"
    """
    text = _extract_text(doc)
    lines = [
        "LOI QA Report",
        f"Counterparty: {data.get('counterparty', {}).get('short', 'Unknown')}",
        f"Generated: {datetime.now().isoformat()}Z",
        f"Variant: programme.recital_a_variant={data.get('programme', {}).get('recital_a_variant', 'default')}",
        f"Overrides: {','.join(sorted(overrides)) if overrides else 'none'}",
    ]
    if override_reason:
        lines.append(f"Override reason: {override_reason}")
    lines.append("")
    lines.append("Findings:")

    fail_count = 0
    warn_count = 0

    for rule, scope, msg in builder_findings:
        if rule in overrides:
            lines.append(f"  [OVRD] {rule}  {scope}   {msg}")
        else:
            lines.append(f"  [FAIL] {rule}  {scope}   {msg}")
            fail_count += 1

    for rid, (pattern, scope, msg) in _FAIL_RULES.items():
        if re.search(pattern, text, re.MULTILINE):
            if rid in overrides:
                lines.append(f"  [OVRD] {rid}  {scope}   {msg}")
            else:
                lines.append(f"  [FAIL] {rid}  {scope}   {msg}")
                fail_count += 1

    for rid, (pattern, scope, msg) in _WARN_RULES.items():
        if re.search(pattern, text, re.IGNORECASE | re.MULTILINE):
            if rid == "R-11" and data.get("choices", {}).get("cert_relevant"):
                continue
            lines.append(f"  [WARN] {rid}  {scope}   {msg}")
            warn_count += 1

    # R-23 (v3.4) — fabrication gate on Recital B
    # v3.5.3-cont scope F: multi-paragraph Recital B extraction. Prior regex
    # stopped at the first blank line (`\n\n`), which partially scanned any
    # Recital B that legitimately spanned multiple paragraphs (rare but
    # possible for counterparties with consortium / holdco-subsidiary
    # disclosure needs). New regex extracts until the next recital marker
    # "(C)" or "(D)" or a section header (line starting with "## ") or EOF.
    recital_b_text = ""
    m = re.search(
        r"\(B\)\s*(.*?)(?=\(C\)\s|\(D\)\s|\n##\s|\Z)",
        text,
        re.DOTALL,
    )
    if m:
        recital_b_text = m.group(1)
    source_map = data.get("counterparty", {}).get("source_map", {})
    r23_findings = _check_fabrication_gate(recital_b_text, source_map, overrides)
    for rid, scope, msg in r23_findings:
        lines.append(f"  [FAIL] {rid}  {scope}   {msg}")
        fail_count += 1

    # v3.5.6 scope D.1: on PASS, emit pillar-attribution diagnostic so the
    # QA report documents WHICH pillar matched each claim (or that it was
    # sentence-scoped [TBC]-covered). Makes the attribution auditable
    # without tightening the gate itself.
    if not r23_findings and _R23_PILLAR_DIAGNOSTIC:
        n = len(_R23_PILLAR_DIAGNOSTIC)
        lines.append(
            f"  [INFO] R-23  Recital B   attribution diagnostic: "
            f"{n} claim(s) matched"
        )
        for claim_text, pillar_or_tbc in _R23_PILLAR_DIAGNOSTIC[:5]:
            lines.append(
                f"         {pillar_or_tbc}: {claim_text[:80]}"
            )
        if n > 5:
            lines.append(f"         (and {n - 5} more — see fixture for full list)")

    # v3.5.6 scope D.3: log active overrides + validated reason + timestamp
    # in QA report so the audit trail captures "who bypassed what when".
    if overrides:
        from datetime import datetime as _dt, timezone as _tz
        lines.append(
            f"  [INFO] R-override  meta   overrides active: "
            f"{', '.join(sorted(overrides))}; reason: "
            f"{(override_reason or '<none>')[:200]}; logged: "
            f"{_dt.now(_tz.utc).isoformat()}"
        )

    # R-28 (v3.5.2) — [TBC] density check in body. A few [TBC] markers are
    # expected on drafts (signatory title, counterparty reg number pre-
    # signing); > 5 suggests the intake was not fully prepared and the draft
    # should loop back to Phase 4/5 for completion before external delivery.
    if "R-28" not in overrides:
        tbc_count = text.count("[TBC]")
        if tbc_count > 5:
            lines.append(
                f"  [WARN] R-28  body   [TBC] count ({tbc_count}) exceeds 5 "
                f"\u2014 intake likely incomplete; consider Phase 4/5 revision before external delivery"
            )
            warn_count += 1

    if not data.get("programme", {}).get("recital_a_variant"):
        lines.append("  [INFO] R-16  YAML   Recital A variant not set, used 'default'")
    if not data.get("choices", {}).get("bespoke_closing"):
        lines.append("  [INFO] R-17  YAML   No bespoke_closing, used default single-sentence")

    lines.append("")
    if fail_count > 0:
        status = "FAIL"
    elif warn_count > 0:
        status = "PASS_WITH_WARN"
    else:
        status = "PASS"
    lines.append(f"Status: {status} (warnings: {warn_count}, failures: {fail_count})")
    return status, lines


def _parse_arg(flag: str) -> str:
    if flag in sys.argv:
        idx = sys.argv.index(flag)
        if idx + 1 < len(sys.argv):
            return sys.argv[idx + 1]
    return ""


def _migrate_check(path: str) -> int:
    """v3.5.3 scope K: legacy intake YAML migration pre-flight.

    Inspects intake YAML for missing `counterparty.source_map` (required by
    v3.4 R-23 fabrication gate). If absent or empty, emits a ready-to-paste
    snippet with all 5 pillars marked `[TBC]` and exits 0 (non-blocking —
    the user copies the snippet, pastes into their YAML, and re-runs the
    generator). If `source_map` is already present, reports that and exits 0.

    Uses raw yaml.safe_load (bypasses full validate()) so legacy YAMLs that
    would fail other v3.4/v3.5 validators can still be migration-checked.
    """
    try:
        with open(path, "r") as f:
            raw = yaml.safe_load(f) or {}
    except Exception as e:
        print(f"[--migrate-check] Could not read {path}: {e}")
        return 0
    cp = raw.get("counterparty", {}) or {}
    source_map = cp.get("source_map")
    if source_map and isinstance(source_map, dict) and any(source_map.values()):
        print(f"[--migrate-check] {path}")
        print("  OK: counterparty.source_map present with entries.")
        print("  R-23 fabrication gate will evaluate URL attribution per pillar.")
        return 0
    # Missing or empty — emit snippet
    print(f"[--migrate-check] {path}")
    print("  counterparty.source_map NOT SET — legacy v3.3 intake.")
    print("  Paste the following into your YAML under `counterparty:`:")
    print()
    print("  source_map:")
    print("    pillar_1: \"[TBC]\"   # Identity & scale — own website, registry")
    print("    pillar_2: \"[TBC]\"   # Core business — own website")
    print("    pillar_3: \"[TBC]\"   # Track record / proof points — named customer press")
    print("    pillar_4: \"inferred from Phase 1 context\"")
    print("    pillar_5: \"[TBC]\"   # Forward plans — only if anchored to named third party")
    print()
    print("  Replace [TBC] with tier-1 URLs where available. See")
    print("  _shared/counterpart-description-framework.md for the Signal")
    print("  Test and tier hierarchy policy.")
    return 0


# v3.5.6 scope G: Phase 7.5 fail-closed enforcement — sentinel file with
# SHA-256 of the .docx being blessed. Consumed on --phase-7-5-pass after
# verifying hash matches current .docx (prevents replay + post-approval
# tampering). Opt-in via --enforce-phase-7-5 flag OR env var
# DE_LOI_ENFORCE_PHASE_7_5=1 (OR'd). Default: fail-open (v3.5.x preserved).

_PHASE_7_5_SENTINEL_SUFFIX = ".phase_7_5_required"


def _phase_7_5_enforce_enabled() -> bool:
    """v3.5.6 G.2: enforcement activates via CLI flag OR env variable."""
    if "--enforce-phase-7-5" in sys.argv:
        return True
    val = os.environ.get("DE_LOI_ENFORCE_PHASE_7_5", "").strip().lower()
    return val in ("1", "true", "yes", "on")


def _docx_sha256(docx_path: str) -> str:
    """Return hex SHA-256 of .docx file contents."""
    import hashlib
    h = hashlib.sha256()
    with open(docx_path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def _phase_7_5_sentinel_path(docx_path: str) -> str:
    return docx_path + _PHASE_7_5_SENTINEL_SUFFIX


def _write_phase_7_5_sentinel(docx_path: str) -> str:
    """v3.5.6 G.1: write sentinel blessing a specific .docx content hash.

    Written alongside the .docx as `<docx>.phase_7_5_required`. Contents
    include SHA-256 (for tamper-detection on --phase-7-5-pass), ISO
    timestamp, and a HOW TO RESOLVE block so the error is self-documenting.
    """
    sentinel_path = _phase_7_5_sentinel_path(docx_path)
    docx_sha = _docx_sha256(docx_path)
    ts = datetime.now(timezone.utc).isoformat()
    with open(sentinel_path, "w", encoding="utf-8") as f:
        f.write(
            f"# Phase 7.5 review required\n"
            f"# Auto-generated by generate_loi.py ({ts})\n"
            f"#\n"
            f"docx: {os.path.basename(docx_path)}\n"
            f"docx_sha256: {docx_sha}\n"
            f"created: {ts}\n"
            f"\n"
            f"# HOW TO RESOLVE:\n"
            f"#   1. Load the Phase 7.5 callee workflow in your Claude session:\n"
            f"#      legal-counsel/specializations/contract-review/loi-review-workflow.md\n"
            f"#   2. Run the 4-point structured review against this .docx\n"
            f"#      (clause-type appropriateness, meta-commentary scan,\n"
            f"#       cross-clause consistency, source-verification sample)\n"
            f"#   3. If the review returns PASS, re-run the generator with\n"
            f"#      --phase-7-5-pass. The sentinel will be consumed\n"
            f"#      (verifying the .docx hash hasn't changed) and delivery\n"
            f"#      will be permitted.\n"
            f"#   4. If review returns FLAG-FOR-REVISION or REJECT, edit the\n"
            f"#      intake YAML and re-run the generator (a new sentinel\n"
            f"#      with a fresh hash will be written).\n"
            f"#\n"
            f"# DO NOT edit this file or the .docx between Phase 7.5 review\n"
            f"# and --phase-7-5-pass — the hash check will detect any change.\n"
        )
    return sentinel_path


def _consume_phase_7_5_sentinel(docx_path: str):
    """v3.5.6 G.1: verify sentinel's hash matches current .docx, then
    consume (delete) the sentinel. Returns (ok, err) — on ok, sentinel
    is removed and caller may proceed with delivery.
    """
    sentinel_path = _phase_7_5_sentinel_path(docx_path)
    if not os.path.exists(sentinel_path):
        return False, (
            f"Phase 7.5 sentinel not found at {sentinel_path}. "
            f"Either Phase 7.5 was not run for this .docx, or the sentinel "
            f"was consumed by a prior --phase-7-5-pass invocation. Regenerate "
            f"the .docx first (which writes a fresh sentinel under "
            f"enforcement mode)."
        )
    import re as _re
    with open(sentinel_path, "r", encoding="utf-8") as f:
        sentinel_content = f.read()
    m = _re.search(r"docx_sha256:\s*([0-9a-fA-F]{64})", sentinel_content)
    if not m:
        return False, (
            f"Phase 7.5 sentinel at {sentinel_path} is malformed "
            f"(no valid docx_sha256 line). Regenerate."
        )
    sentinel_sha = m.group(1).lower()
    current_sha = _docx_sha256(docx_path)
    if sentinel_sha != current_sha:
        return False, (
            f"Phase 7.5 sentinel hash mismatch: the .docx was modified "
            f"after Phase 7.5 approval.\n"
            f"  sentinel blessed: {sentinel_sha[:16]}...\n"
            f"  current .docx:    {current_sha[:16]}...\n"
            f"Re-run Phase 7.5 review against the current .docx."
        )
    # Consume
    os.remove(sentinel_path)
    return True, ""


def main():
    if len(sys.argv) < 2:
        print("Usage: python generate_loi.py <intake.yaml>")
        print("  [--output path.docx]")
        print("  [--override R-11,R-14]")
        print("  [--override-reason \"...\"]")
        print("  [--migrate-check]          (v3.5.3: inspect YAML for v3.4 source_map; non-blocking)")
        print("  [--enforce-phase-7-5]      (v3.5.6: opt-in Phase 7.5 fail-closed gate)")
        print("  [--phase-7-5-pass]         (v3.5.6: consume Phase 7.5 sentinel after review)")
        print("")
        print("  Env: DE_LOI_ENFORCE_PHASE_7_5=1  also activates Phase 7.5 enforcement")
        sys.exit(1)

    # v3.5.3 scope K: --migrate-check runs before load_intake so legacy
    # YAMLs that would fail full validation can still be inspected.
    if "--migrate-check" in sys.argv:
        sys.exit(_migrate_check(sys.argv[1]))

    data = load_intake(sys.argv[1])

    override_str = _parse_arg("--override")
    if override_str:
        data["_overrides"] = [r.strip() for r in override_str.split(",") if r.strip()]
    override_reason = _parse_arg("--override-reason")
    # v3.5.6 scope D.3: validate override-reason shape when --override is used.
    # Hybrid rule: accept free-text ≥15 chars OR structured short code
    # `<STATUS>-<YYYY-MM-DD> <INITIALS>`. Reject thin patterns unconditionally.
    if override_str:
        ok, err = _validate_override_reason(override_reason)
        if not ok:
            print(f"ERROR (--override-reason): {err}", file=sys.stderr)
            sys.exit(1)
    if override_reason:
        data["_override_reason"] = override_reason

    output = _parse_arg("--output")

    if not output:
        loi_date = data.get("dates", {}).get("loi_date", "")
        try:
            dt = datetime.strptime(loi_date, "%d %B %Y")
            date_str = dt.strftime("%Y%m%d")
        except ValueError:
            date_str = datetime.now().strftime("%Y%m%d")
        cp = data.get("counterparty", {}).get("short", "Unknown")
        output = f"{date_str}_DEG_LOI-{data['type']}_{cp}_(DRAFT).docx"

    # v3.6.0 item b: auto-version output filenames. If target exists,
    # append _v{N} and increment until unique. Addresses file-confusion
    # observed in Cerebro/Armada/InfraPartners retrospectives where
    # operators regenerated iteratively and lost track of canonical draft.
    if os.path.exists(output):
        base, ext = os.path.splitext(output)
        n = 2
        while os.path.exists(f"{base}_v{n}{ext}"):
            n += 1
        output = f"{base}_v{n}{ext}"
        print(f"[auto-version] Target exists; writing to {output}", file=sys.stderr)

    loi = LOI(data)
    doc = loi.build()

    qa_report_path = output.replace(".docx", "_qa.txt")
    qa_status, qa_lines = qa_lint(doc, data, loi.qa_findings, loi.overrides,
                                   loi.override_reason)
    with open(qa_report_path, "w", encoding="utf-8") as f:
        f.write("\n".join(qa_lines) + "\n")

    if qa_status == "FAIL":
        print("QA FAIL — build blocked. Findings:")
        for line in qa_lines:
            if "[FAIL]" in line:
                print(f"  {line}")
        print(f"\nFull QA report: {qa_report_path}")
        print("To override: --override R-xx,R-yy --override-reason \"...\"")
        sys.exit(2)

    doc.save(output)

    # v3.5.6 scope G: Phase 7.5 fail-closed enforcement. Only active when
    # opt-in (CLI flag OR env var). Default path (fail-open) preserves
    # v3.5.x behaviour for existing workflows.
    if _phase_7_5_enforce_enabled():
        if "--phase-7-5-pass" in sys.argv:
            ok, err = _consume_phase_7_5_sentinel(output)
            if not ok:
                print(
                    f"Phase 7.5 enforcement: DELIVERY BLOCKED\n\n{err}",
                    file=sys.stderr,
                )
                sys.exit(3)
            # Sentinel consumed cleanly — fall through to success print
        else:
            sentinel_path = _write_phase_7_5_sentinel(output)
            print(
                f"Generated: {output}\n"
                f"Type: DE-LOI-{data['type']}-v3.2\n"
                f"Counterparty: {data.get('counterparty', {}).get('name', 'Unknown')}\n"
                f"QA: {qa_status} ({qa_report_path})\n"
                f"\n"
                f"⚠ PHASE 7.5 REVIEW REQUIRED — delivery blocked.\n"
                f"\n"
                f"Phase 7.5 enforcement is enabled. The .docx has been written\n"
                f"but delivery is gated on legal-counsel review.\n"
                f"\n"
                f"Next step:\n"
                f"  1. Load the callee workflow in your Claude session:\n"
                f"     legal-counsel/specializations/contract-review/loi-review-workflow.md\n"
                f"  2. Run the 4-point structured review against {output}\n"
                f"  3. If review returns PASS, re-run with --phase-7-5-pass\n"
                f"\n"
                f"Sentinel written: {sentinel_path}"
            )
            sys.exit(3)

    print(f"Generated: {output}")
    print(f"Type: DE-LOI-{data['type']}-v3.2")
    print(f"Counterparty: {data.get('counterparty', {}).get('name', 'Unknown')}")
    print(f"QA: {qa_status} ({qa_report_path})")
    print(f"Clauses: ALL (full document)")


if __name__ == "__main__":
    main()
