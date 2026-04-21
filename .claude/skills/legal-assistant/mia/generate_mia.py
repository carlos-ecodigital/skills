#!/usr/bin/env python3
"""
MIA v1.0 — Document Generation Script

Usage:
    python generate_mia.py intake.yaml [--output path/to/output.docx]

Takes a YAML intake file and produces a complete, branded .docx ready to sign.
Generates Master + activated Annex(es) from intake data.
"""

import sys
import os
import yaml
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import shared cover page and branding from document-factory
SCRIPT_DIR = Path(__file__).parent
sys.path.insert(0, str(SCRIPT_DIR.parent.parent / "document-factory"))
from common import (  # noqa: E402
    add_cover, Party,
    COBALT, SLATE, SLATE_800, SLATE_900, WHITE,
    setup_first_page_header, setup_cont_header,
    setup_first_footer, setup_cont_footer,
)


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

FONT_NAME = "Inter"
FONT_BODY = Pt(10)
FONT_HEADING1 = Pt(13)
FONT_HEADING2 = Pt(11)
FONT_SMALL = Pt(8)
LINE_SPACING = 1.15

NAVY = SLATE_900
BLACK = SLATE_800
GREY = SLATE


# ---------------------------------------------------------------------------
# YAML Loading + Validation
# ---------------------------------------------------------------------------

def load_intake(path: str) -> dict:
    with open(path, "r") as f:
        data = yaml.safe_load(f)
    validate(data)
    return data


def validate(d: dict):
    errors = []
    t = d.get("type", "")
    if t != "MIA":
        errors.append(f"type must be MIA (got: {t})")
    scope = d.get("scope", "")
    if scope not in ("master_only", "commercial", "capital", "both"):
        errors.append(f"scope must be master_only, commercial, capital, or both (got: {scope})")
    for s in ("provider", "introducer", "dates", "activation"):
        if s not in d:
            errors.append(f"Missing section: {s}")
    intro = d.get("introducer", {})
    for f in ("name", "short", "description"):
        if not intro.get(f):
            errors.append(f"introducer.{f} required")
    if not d.get("dates", {}).get("mia_date"):
        errors.append("dates.mia_date required")
    act = d.get("activation", {})
    if act.get("annex_a") and "annex_a" not in d:
        errors.append("activation.annex_a is true but annex_a section missing")
    if act.get("annex_b") and "annex_b" not in d:
        errors.append("activation.annex_b is true but annex_b section missing")
    if errors:
        print("VALIDATION ERRORS:")
        for e in errors:
            print(f"  - {e}")
        sys.exit(1)


# ---------------------------------------------------------------------------
# Document Builder
# ---------------------------------------------------------------------------

class MIA:
    def __init__(self, data: dict):
        self.d = data
        self.scope = data["scope"]
        self.annex_a = bool(data.get("activation", {}).get("annex_a"))
        self.annex_b = bool(data.get("activation", {}).get("annex_b"))
        self.doc = Document()
        self._setup()

    def _setup(self):
        style = self.doc.styles["Normal"]
        style.font.name = FONT_NAME
        style.font.size = FONT_BODY
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = LINE_SPACING
        for s in self.doc.sections:
            s.page_width = Mm(210)
            s.page_height = Mm(297)
            s.top_margin = Mm(20)
            s.bottom_margin = Mm(35)
            s.left_margin = Mm(25)
            s.right_margin = Mm(20)
            s.different_first_page_header_footer = True
            setup_first_page_header(s)
            setup_cont_header(s, title="Master Introduction Agreement")
            setup_first_footer(s, classification="Confidential")
            setup_cont_footer(s, classification="Confidential")

    # --- Helpers ---

    def g(self, *keys, default=""):
        v = self.d
        for k in keys:
            v = v.get(k, default) if isinstance(v, dict) else default
        return v or default

    def choice(self, k):
        return bool(self.g("choices", k))

    def p(self, text, bold=False, italic=False, size=None, color=None,
           align=None, space_after=None):
        para = self.doc.add_paragraph()
        if align:
            para.alignment = align
        if space_after is not None:
            para.paragraph_format.space_after = Pt(space_after)
        run = para.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = size or FONT_BODY
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        return para

    def bp(self, label, text):
        """Bold label + normal text in one paragraph."""
        para = self.doc.add_paragraph()
        r1 = para.add_run(label)
        r1.bold = True
        r1.font.name = FONT_NAME
        r1.font.size = FONT_BODY
        r2 = para.add_run(text)
        r2.font.name = FONT_NAME
        r2.font.size = FONT_BODY
        return para

    def h(self, text, level=2):
        heading = self.doc.add_heading(text, level=level)
        for r in heading.runs:
            r.font.name = FONT_NAME
            r.font.color.rgb = NAVY

    def table(self, headers, rows):
        t = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.name = FONT_NAME
                    r.font.size = FONT_BODY
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = t.rows[ri + 1].cells[ci]
                cell.text = str(val)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = FONT_NAME
                        r.font.size = FONT_BODY
        return t

    def line(self):
        self.p("—" * 50, color=GREY, size=Pt(6), space_after=0)

    def page_break(self):
        self.doc.add_page_break()

    # --- Cover page ---

    def cover(self):
        prov = self.d.get("provider", {})
        provider_party = Party(
            legal_name=prov.get("legal_name", ""),
            address=prov.get("address", ""),
            registration_type=prov.get("reg_type"),
            registration_number=prov.get("reg_number"),
        )
        intro = self.d.get("introducer", {})
        introducer_party = Party(
            legal_name=intro.get("name", ""),
            address=intro.get("address", ""),
            registration_type=intro.get("reg_type"),
            registration_number=intro.get("reg_number"),
        )
        add_cover(
            self.doc,
            agreement_type="Master Introduction Agreement",
            subject="Framework for Commercial and Capital Introductions",
            date_str=self.g("dates", "mia_date"),
            parties=[provider_party, introducer_party],
            formality="binding",
            classification="Confidential",
        )

    # --- Master sections ---

    def recitals(self):
        self.h("Recitals")
        prov = self.g("provider", "short_name")
        intro_short = self.g("introducer", "short")
        intro_desc = self.g("introducer", "description")

        self.p(f'(A) {prov} (the "Provider") develops and operates Digital Energy '
               f"Centers (\"DECs\"): purpose-built, liquid-cooled colocation facilities "
               f"for high-density accelerated compute. The Provider engages with "
               f"prospective customers (for colocation services) and prospective "
               f"investors (for capital), and from time to time receives introductions "
               f"to such parties from third parties.")

        self.p(f'(B) {intro_short} (the "Introducer") {intro_desc}')

        self.p(f"(C) The Parties wish to record a master framework governing "
               f"introductions made by the Introducer to the Provider. This framework "
               f"comprises (i) this master agreement (the \"Master\"), and (ii) one or "
               f"both of two severable annexes — Annex A (Commercial Introductions) "
               f"and Annex B (Capital Introductions) — activated as agreed by the "
               f"Parties at execution.")

        self.p(f"(D) Where this MIA is executed alongside, or following, a Letter of "
               f"Intent between the Parties, this MIA is the \"Referral Agreement\" "
               f"contemplated by Clause 3.4 of the Provider's Distributor Letter of "
               f"Intent template (Mode B — Referral Arrangement).")

        self.p(f"(E) The Parties recognise that the activities contemplated by this "
               f"MIA — and in particular by Annex B (Capital Introductions) — touch on "
               f"regulated subject matter, and have structured this MIA, including the "
               f"severability of Annex B, accordingly.")

    def definitions(self):
        self.h("1. Definitions")
        self.p("In this MIA, unless the context requires otherwise:")
        prov_name = self.g("provider", "legal_name")
        intro_name = self.g("introducer", "name")
        intro_juris = self.g("introducer", "jurisdiction")
        prov_reg = self.g("provider", "reg_type")
        prov_reg_num = self.g("provider", "reg_number")

        defs = [
            ('"Annex"', "means Annex A or Annex B (or both), as activated under Clause 2.2."),
            ('"Annex A"', "means the Commercial Introductions annex to this Master."),
            ('"Annex B"', "means the Capital Introductions annex to this Master."),
            ('"Business Day"', "means a day other than a Saturday, Sunday, or public holiday in the Netherlands."),
            ('"Confidential Information"', "means all information of a confidential or proprietary nature disclosed by one Party to the other in connection with this MIA."),
            ('"Effective Date"', f"means {self.g('dates', 'effective_date')}."),
            ('"Fee"', "means a fee payable by the Provider to the Introducer pursuant to an activated Annex."),
            ('"Introduced Party"', "means a Customer (in respect of Annex A) or an Investor (in respect of Annex B), as applicable."),
            ('"Introducer"', f"means {intro_name}, a {intro_juris} entity."),
            ('"Introduction"', "means an introduction effected by the Introducer in accordance with the terms of an activated Annex."),
            ('"Master"', "means this master agreement."),
            ('"MIA"', "means this Master together with any Annex activated under Clause 2.2."),
            ('"Parties"', f"means the Provider and the Introducer; \"Party\" means either of them."),
            ('"Provider"', f"means {prov_name}, registered under {prov_reg}: {prov_reg_num}."),
            ('"Qualifying Introduction"', "has the meaning given in the relevant activated Annex."),
            ('"Representative"', "means, in respect of a Party, its directors, officers, employees, professional advisers, bona fide financing parties, and potential co-investors."),
        ]
        for term, defn in defs:
            self.bp(f"{term} ", defn)

    def purpose_and_scope(self):
        self.h("2. Purpose, Scope, and Activation")
        self.p("2.1 Purpose. This MIA establishes a master framework under which the "
               "Introducer may introduce Customers and/or Investors to the Provider, "
               "and under which the Provider may pay Fees to the Introducer for "
               "Qualifying Introductions made under an activated Annex.")

        self.p("2.2 Activation. The following Annex(es) are activated:")
        a_active = "Yes" if self.annex_a else "No"
        b_active = "Yes" if self.annex_b else "No"
        self.table(
            ["Annex", "Activated?", "Subject Matter"],
            [
                ["Annex A — Commercial Introductions", a_active, "Introductions of prospective Customers"],
                ["Annex B — Capital Introductions", b_active, "Introductions of prospective Investors"],
            ],
        )

        self.p("2.3 Independence and severability of Annexes. Each activated Annex "
               "operates independently of the other. The termination, suspension, "
               "expiry, or invalidity of one Annex does not affect the validity or "
               "operation of the other Annex or of this Master.")

        self.p("2.4 Master is binding. The provisions of this Master are binding on "
               "the Parties from the Effective Date.")

    def relationship(self):
        self.h("3. Relationship of the Parties")
        self.p("3.1 Independent contractors. The Parties are independent contractors. "
               "Nothing in this MIA creates an agency, partnership, joint venture, "
               "employment, or fiduciary relationship between them.")
        self.p("3.2 No authority. Neither Party has authority to bind the other, to "
               "make representations on the other's behalf, or to incur liabilities "
               "on the other's behalf.")
        self.p("3.3 Non-exclusive. Each Party is free to engage other parties for "
               "similar services or activities, subject to the binding obligations of "
               "this MIA.")

    def referral_identification(self):
        self.h("4. Pre-existing Relationships and Referral Agreement Identification")
        existing_nda = self.choice("existing_nda")
        if existing_nda:
            self.p("4.1 Where this MIA is executed alongside, or following, a Letter "
                   "of Intent between the Parties, this MIA is the Referral Agreement "
                   "referenced in Clause 3.4 of that Letter of Intent.")
        else:
            self.p("4.1 Where this MIA is executed alongside, or following, a Letter "
                   "of Intent between the Parties, this MIA is the Referral Agreement "
                   "referenced in Clause 3.4 of that Letter of Intent.")
            self.p("4.2 Where this MIA is executed without a companion Letter of "
                   "Intent, this MIA stands alone, and the confidentiality and "
                   "non-circumvention provisions in Clauses 5 and 6 apply directly.")

    def confidentiality(self):
        self.h("5. Confidentiality")
        existing_nda = self.choice("existing_nda")
        if existing_nda:
            # ALT-A
            loi_date = self.g("dates", "loi_date")
            self.p(f"5A.1 The Parties acknowledge that they are bound by the "
                   f"confidentiality obligations set out in the Letter of Intent "
                   f"dated {loi_date} (the \"Existing Confidentiality Agreement\").")
            self.p("5A.2 The Existing Confidentiality Agreement applies in full to "
                   "all Confidential Information exchanged in connection with this MIA.")
            self.p("5A.3 In the event of conflict between the Existing Confidentiality "
                   "Agreement and this Clause 5, the more protective provision prevails.")
        else:
            # ALT-B — embedded confidentiality (16 clauses)
            survival = self.g("protection", "confidentiality_survival", default="3 years")
            self._confidentiality_alt_b(survival)

    def _confidentiality_alt_b(self, survival):
        clauses = [
            ("5.1 Purpose Limitation. ", "Each Party shall use the other Party's Confidential Information solely in connection with this MIA and any activated Annex (the \"Purpose\") and for no other purpose."),
            ("5.2 Non-Disclosure. ", "Each Party shall keep confidential all Confidential Information received from the other Party and shall not disclose such information to any person except as permitted under this Clause 5."),
            ("5.3 Standard of Care. ", "Each Party shall apply no less than reasonable care to protect the other Party's Confidential Information, and no less than the care it applies to its own confidential information of a similar nature."),
            ("5.4 Permitted Disclosures. ", "A Party may disclose Confidential Information to: (a) its Representatives who have a genuine need to know for the Purpose; (b) its bona fide financing parties and potential co-investors; and (c) to the extent required by applicable law, regulation, court order, or relevant regulatory authority, subject to prior written notice where legally permitted."),
            ("5.5 Liability for Representatives. ", "Each Party is responsible for any breach of this Clause 5 by its Representatives as if it were a breach by the Party itself."),
            ("5.6 Exclusions. ", "This Clause 5 does not apply to information that: (a) is or becomes public other than through breach; (b) was known prior to disclosure; (c) is independently developed; or (d) is rightfully received from a third party."),
            ("5.7 No Implied Rights. ", "No licence or other right in or to any Confidential Information is granted under this MIA except as expressly stated."),
            ("5.8 Return and Destruction. ", "On written request, the receiving Party shall within fifteen (15) Business Days return or destroy all Confidential Information, save for one archival copy, and certify in writing."),
            ("5.9 Onward-Sharing Controls. ", "A Party shall not confirm, deny, or characterise Confidential Information in response to third-party inquiries, except as permitted."),
            ("5.10 Compliance Confirmation. ", "On not more than one occasion per calendar year, a Party may request written confirmation of compliance."),
            ("5.11 Breach Notification. ", "A Party shall notify the other in writing within seventy-two (72) hours of becoming aware of any actual or suspected breach."),
            ("5.12 Disclaimer. ", "Confidential Information is provided \"as is\" and without warranty as to accuracy or completeness."),
            ("5.13 Metadata Protection. ", "The protection under this Clause 5 extends to embedded metadata (EXIF data, geolocation, timestamps, filenames, document properties)."),
            ("5.14 Survival. ", f"The obligations of this Clause 5 survive termination for a period of {survival}. Obligations for trade secrets continue indefinitely."),
            ("5.15 Announcements. ", "Neither Party shall make any public announcement concerning this MIA without prior written consent, except as required by law."),
            ("5.16 Remedies. ", "Each Party acknowledges that damages alone may not be adequate and the other is entitled to seek injunctive relief."),
        ]
        for label, text in clauses:
            self.bp(label, text)

    def non_circumvention(self):
        self.h("6. Non-Circumvention")
        nc_dur = self.g("protection", "nc_duration", default="12 months")
        self.bp("6.1 Non-contact obligation. ",
                "During the term of this MIA and for the duration set out in Clause 6.2, "
                "the Introducer shall not, directly or indirectly, without the prior "
                "written consent of the Provider: (a) contact, solicit, deal with, or "
                "enter into any business relationship with any Introduced Party; "
                "(b) circumvent, avoid, or bypass the Provider to deal with any "
                "Introduced Party; or (c) attempt to divert any business opportunity "
                "disclosed by the Provider.")
        self.bp("6.2 Duration. ",
                f"The obligations in this Clause 6 continue for {nc_dur} after the "
                f"earlier of (a) termination of this MIA, or (b) completion of the "
                f"relevant Qualifying Introduction and payment of any Fee due.")
        self.bp("6.3 Independent Knowledge Exception. ",
                "This Clause 6 does not restrict the Introducer from dealing with any "
                "party with whom the Introducer can demonstrate a substantive prior "
                "relationship independent of any disclosure by the Provider.")
        self.bp("6.4 MSA / Investment Supersession. ",
                "Where an MSA is signed or capital is committed, the obligations of "
                "this Clause 6 in respect of that party are subsumed by the Qualifying "
                "Introduction record.")

    def pipeline_carveout(self):
        self.h("7. Pipeline Carve-Out")
        lookback = self.g("protection", "lookback_months", default="12 months")
        self.bp("7.1 No Fee where pre-existing relationship demonstrable. ",
                f"No Fee is due where the Provider can demonstrate: (a) a direct "
                f"relationship with the Introduced Party at the date of the Introduction; "
                f"or (b) substantive contact within the {lookback} preceding the Introduction.")
        self.bp("7.2 Evidence. ",
                "The Provider's records (CRM, email, calendar, correspondence) "
                "constitute presumptive evidence. Burden of proof is on the Provider.")
        self.bp("7.3 Notification. ",
                "The Provider shall notify the Introducer within fifteen (15) Business "
                "Days of receiving an Introduction if relying on this Clause 7. Failure "
                "to notify waives reliance.")
        self.bp("7.4 Application across Annexes. ",
                "This Clause 7 applies to Introductions under both Annex A and Annex B.")

    def general_provisions(self):
        self.h("8. General Provisions")
        eff_date = self.g("dates", "effective_date")
        clauses = [
            ("8.1 Notices. ", "Notices must be in writing. Deemed received on the second Business Day after dispatch by courier, date of delivery if hand-delivered, or the Business Day after email transmission."),
            ("8.2 Counterparts. ", "This MIA may be executed in counterparts, including by electronic signature within the meaning of the eIDAS Regulation (EU) No 910/2014."),
            ("8.3 Term. ", f"This MIA takes effect on {eff_date} and continues until terminated under Clause 8.4."),
            ("8.4 Termination. ", "Either Party may terminate on thirty (30) days' written notice. Termination of the Master terminates each activated Annex."),
            ("8.5 Termination for cause. ", "A Party may terminate immediately on written notice if the other commits a material unremedied breach (30-day cure period) or becomes insolvent."),
            ("8.6 Survival. ", "Clauses 5 (Confidentiality), 6 (Non-Circumvention), 7 (Pipeline Carve-Out), 8.7 (Governing Law), and 8.8 (Jurisdiction) survive termination."),
            ("8.7 Governing Law. ", "This MIA is governed by the laws of the Netherlands. CISG is expressly excluded."),
            ("8.8 Jurisdiction. ", "The courts of Amsterdam (Rechtbank Amsterdam) have exclusive jurisdiction."),
            ("8.9 Severability. ", "If any provision is invalid, the remainder continues. Annex B severability is governed by Clause B.13."),
            ("8.10 Entire Agreement. ", "This MIA constitutes the entire agreement in respect of its subject matter."),
            ("8.11 Assignment. ", "Neither Party may assign without consent, save that the Provider may assign by way of security to financing parties."),
            ("8.12 Variation. ", "No variation is effective unless in writing and signed by both Parties."),
            ("8.13 No partnership. ", "Nothing in this MIA constitutes a partnership, joint venture, agency, or employment."),
        ]
        for label, text in clauses:
            self.bp(label, text)

    def signatures(self):
        self.h("Signatures")
        prov = self.g("provider", "legal_name")
        prov_reg = self.g("provider", "reg_type")
        prov_reg_num = self.g("provider", "reg_number")
        intro = self.g("introducer", "name")
        intro_reg = self.g("introducer", "reg_type")
        intro_reg_num = self.g("introducer", "reg_number")

        self.bp(f"For and on behalf of {prov}", "")
        self.p(f"{prov_reg}: {prov_reg_num}")
        self.p("Signature: ____________________________")
        self.p(f"Name: {self.g('provider', 'signatory_name')}")
        self.p(f"Title: {self.g('provider', 'signatory_title')}")
        self.p("Date: ____________________________")

        self.line()

        self.bp(f"For and on behalf of {intro}", "")
        self.p(f"{intro_reg}: {intro_reg_num}")
        self.p("Signature: ____________________________")
        self.p(f"Name: {self.g('introducer', 'signatory_name')}")
        self.p(f"Title: {self.g('introducer', 'signatory_title')}")
        self.p("Date: ____________________________")

    def annex_execution_record(self):
        self.line()
        self.h("Annex Execution Record", level=3)
        a_active = "Yes" if self.annex_a else "No"
        b_active = "Yes" if self.annex_b else "No"
        self.table(
            ["Annex", "Activated?", "Annex Execution Date", "Schedule Ref"],
            [
                ["Annex A — Commercial", a_active, "", "Schedule A-1"],
                ["Annex B — Capital", b_active, "", "Schedule B-1"],
            ],
        )

    # --- Annex A sections ---

    def annex_a_full(self):
        self.page_break()
        self.h("Annex A — Commercial Introductions")
        prov = self.g("provider", "legal_name")
        intro = self.g("introducer", "name")
        mia_date = self.g("dates", "mia_date")
        tail = self.g("annex_a", "tail", default="12 months")

        self.p(f"This Annex A is part of the Master Introduction Agreement dated "
               f"{mia_date} between {prov} (the \"Provider\") and {intro} (the "
               f"\"Introducer\").", italic=True)

        self.h("A.1 Activation and Integration", level=3)
        eff = self.g("annex_a", "effective_date",
                      default=self.g("dates", "effective_date"))
        self.p(f"A.1.1 This Annex A is activated by the Parties' execution alongside "
               f"the Master.")
        self.p(f"A.1.2 Takes effect from {eff}.")

        self.h("A.2 Definitions", level=3)
        self.bp("\"Customer\" ", "means a third party introduced by the Introducer "
                "with a view to becoming a paying customer of the Provider's colocation "
                "services under an MSA.")
        self.bp("\"MRR\" ", "means the monthly recurring revenue payable by a Customer "
                "under an MSA (EUR per kW per month, excl. pass-through energy, "
                "one-off fees, and VAT).")
        self.bp("\"Qualifying Commercial Introduction\" ", "has the meaning in Cl. A.3.")

        self.h("A.3 Qualifying Commercial Introduction", level=3)
        self.p("A.3.1 A Qualifying Commercial Introduction arises when: "
               "(a) the Introducer makes a written introduction accepted by the Provider; "
               "(b) the Customer executes an MSA; and (c) the Customer pays the first "
               "invoice in cleared funds.")
        self.p("A.3.2 The Pipeline Carve-Out (Master Cl. 7) applies.")

        self.h("A.4 Fee Schedule", level=3)
        self.p("A.4.1 Election. The Introducer elects one fee basis per Introduction "
               "at the time of the original written introduction:")

        self.p("Option 1 — Tiered Residual % MRR:", bold=True)
        self.table(
            ["Capacity Bucket", "MW IT", "Residual Fee"],
            [
                ["Bucket 1", "< 5 MW IT", "8–10% of MRR"],
                ["Bucket 2", "5–20 MW IT", "6–8% of MRR"],
                ["Bucket 3", "> 20 MW IT", "4–6% of MRR"],
            ],
        )
        self.p("Payable monthly in arrears for the initial MSA term.")

        self.p("Option 2 — Capitalised One-Time (default):", bold=True)
        self.table(
            ["Capacity Bucket", "MW IT", "Capitalised Fee"],
            [
                ["Bucket 1", "< 5 MW IT", "3 months MRR equivalent"],
                ["Bucket 2", "5–20 MW IT", "2 months MRR equivalent"],
                ["Bucket 3", "> 20 MW IT", "1.5 months MRR equivalent"],
            ],
        )
        self.p("Payable in a single lump sum within 30 days of first Customer invoice.")
        self.p("A.4.5 If no election is made, Capitalised One-Time applies by default.")

        self.h("A.5 Payment", level=3)
        self.p("A.5.1 Currency, invoicing entity, payment terms, and VAT treatment "
               "per Schedule A-1.")

        self.h("A.6 Tail", level=3)
        self.p(f"A.6.1 Fee entitlement survives termination for {tail} for any "
               f"Qualifying Commercial Introduction where conditions A.3.1(a)–(c) "
               f"are satisfied within that period.")

        self.h("A.7 Termination", level=3)
        self.p("A.7.1 Terminates on: (a) termination of the Master, (b) mutual "
               "agreement, or (c) 30 days' written notice. Termination does not "
               "affect the Master or Annex B.")

        # Schedule A-1
        self.line()
        self.h("Schedule A-1 — Per-Execution Fee Details", level=3)
        sched = self.d.get("annex_a", {}).get("schedule_a1", {})
        self.table(
            ["Field", "Value"],
            [
                ["Customer name", ""],
                ["Date of written introduction", ""],
                ["Date of Provider's acceptance", ""],
                ["Indicative MW IT", ""],
                ["Fee election", self.g("annex_a", "fee_election", default="capitalised")],
                ["Invoicing entity", sched.get("invoicing_entity", "")],
                ["Payable-to entity", sched.get("payable_to_entity", "")],
                ["Currency", sched.get("currency", "EUR")],
                ["Payment terms", sched.get("payment_terms", "Net 30")],
                ["VAT treatment", sched.get("vat_treatment", "")],
            ],
        )

        # Signatures
        self._annex_signatures("Annex A")

    # --- Annex B sections ---

    def annex_b_full(self):
        self.page_break()
        self.h("Annex B — Capital Introductions")
        prov = self.g("provider", "legal_name")
        # For Annex B, use provider_annex_b if present (AG entity)
        prov_b = self.d.get("provider_annex_b", self.d.get("provider", {}))
        prov_b_name = prov_b.get("legal_name", prov)
        intro = self.g("introducer", "name")
        mia_date = self.g("dates", "mia_date")
        tail = self.g("annex_b", "tail", default="12 months")
        approval = self.g("annex_b", "approval_window", default="10 Business Days")
        suspend = self.g("annex_b", "suspension_restructure_window", default="60 days")
        lookback = self.g("protection", "lookback_months", default="12 months")

        self.p(f"This Annex B is part of the Master Introduction Agreement dated "
               f"{mia_date} between {prov_b_name} (the \"Provider\") and {intro} (the "
               f"\"Introducer\").", italic=True)

        self.h("B.1 Activation and Regulatory Intent", level=3)
        eff = self.g("annex_b", "effective_date",
                      default=self.g("dates", "effective_date"))
        self.p(f"B.1.1 This Annex B is activated by the Parties' execution.")
        self.p(f"B.1.2 Takes effect from {eff}.")
        self.p("B.1.4 Regulatory intent. The Parties enter into this Annex B on the "
               "express basis that the Introducer's activities are limited to "
               "introducing Investors and do not constitute investment services, "
               "brokerage, placement, or any other regulated activity under MiFID II, "
               "the Wft, the FCA Handbook, the US Securities Exchange Act, FINMA "
               "rules, or equivalent.")

        self.h("B.2 Definitions", level=3)
        self.bp("\"Approved Investor\" ", "means an Investor approved by the Provider under Cl. B.6.")
        self.bp("\"Capital Committed\" ", "means the aggregate capital committed by an Approved Investor in EUR (ECB reference rate for conversions).")
        self.bp("\"Investor\" ", "means a third party introduced with a view to committing capital to the Provider.")
        self.bp("\"Named Investor List\" ", "means the list submitted under Cl. B.6.")
        self.bp("\"Qualifying Capital Introduction\" ", "has the meaning in Cl. B.7.")

        self.h("B.3 Permitted Activities", level=3)
        self.p("B.3.1 The Introducer may only: (a) provide the Provider's contact "
               "information; (b) arrange an initial meeting; (c) share publicly "
               "available information; and (d) confirm factual, non-evaluative "
               "information on request.")

        self.h("B.4 Prohibited Activities", level=3)
        prohibitions = [
            "(a) provide investment advice or recommendations;",
            "(b) solicit, promote, market, or offer the investment opportunity;",
            "(c) negotiate investment terms on behalf of either Party;",
            "(d) distribute offering materials, financial models, or non-public information;",
            "(e) hold, receive, or transmit funds or securities;",
            "(f) make representations about expected returns or risk profiles;",
            "(g) act as agent, representative, broker, dealer, placement agent, or financial intermediary; or",
            "(h) hold itself out as having any capacity listed in (g).",
        ]
        self.p("B.4.1 The Introducer shall not:")
        for item in prohibitions:
            self.p(f"    {item}")

        self.h("B.5 Regulatory Representations", level=3)
        self.p("B.5.1 The Introducer represents it is not a regulated investment firm, "
               "broker-dealer, or placement agent, and holds no such licence.")
        self.p("B.5.2 The Provider represents it has not engaged the Introducer as "
               "a placement agent or financial advisor.")
        self.p("B.5.3 Both Parties acknowledge the structure is intended to fall "
               "outside the regulatory perimeter and agree to restructure or suspend "
               "if regulatory circumstances change.")

        self.h("B.6 Named Investor List", level=3)
        self.p(f"B.6.1 The Introducer provides a written Named Investor List.")
        self.p(f"B.6.2 The Provider approves/rejects within {approval}. "
               f"Non-response within that period = deemed rejected.")
        self.p("B.6.3 Only Approved Investors qualify for Fees.")
        self.p(f"B.6.4 Investors with Provider contact within prior {lookback} are "
               f"excluded (Pipeline Carve-Out, Master Cl. 7).")
        self.p("B.6.5 Updates on demand; no fixed cadence.")

        self.h("B.7 Qualifying Capital Introduction", level=3)
        self.p("B.7.1 Arises when: (a) Approved Investor was on an approved list; "
               "(b) first substantive meeting held; and (c) Investor commits capital "
               "in cleared funds (or first tranche received).")

        self.h("B.8 Fee Schedule (Option 2 — Capped %)", level=3)
        self.p("B.8.1 Fee: 2% of Capital Committed, capped at EUR 250,000 per "
               "Qualifying Capital Introduction.", bold=True)
        self.p("B.8.2 Where capital is called in tranches, Fee accrues pro rata; "
               "EUR 250,000 cap applies in aggregate.")

        self.h("B.9 Payment", level=3)
        self.p("B.9.1 Per Schedule B-1. Default: EUR, Net 30.")
        self.p("B.9.2 Fees paid prior to suspension/termination are non-refundable, "
               "except per Cl. B.11.2 (clawback).")

        self.h("B.10 Tail", level=3)
        self.p(f"B.10.1 Fee entitlement survives termination for {tail}.")

        self.h("B.11 Termination for Cause and Clawback", level=3)
        self.p("B.11.1 The Provider may terminate Annex B immediately if the "
               "Introducer breaches Cl. B.3, B.4, or B.5.")
        self.p("B.11.2 Clawback: Fees paid in the 12 months prior to the breach "
               "event. Payable within 30 days of written demand.")

        self.h("B.12 Automatic Suspension on Regulatory Change", level=3)
        self.p("B.12.1 Annex B suspends automatically on any Regulatory Trigger: "
               "(a) regulatory letter/notice; (b) published guidance; "
               "(c) enforcement action against comparable arrangement; or "
               "(d) legislative change.")
        self.p("B.12.2 Either Party may declare suspension by written notice. "
               "Suspension takes effect on the date of the Regulatory Trigger.")
        self.p("B.12.3 During suspension: no new introductions trigger Fees; "
               "accrued Fees remain payable; Master + Annex A continue.")
        self.p(f"B.12.4 Restructure or terminate within {suspend}.")

        self.h("B.13 Enhanced Severability", level=3)
        self.p("B.13.1 If any Annex B provision constitutes regulated activity, "
               "Annex B terminates automatically without affecting Master or Annex A.")
        self.p("B.13.2 Prior Fees non-refundable (subject to Cl. B.11.2).")
        self.p("B.13.3 Master Cl. 5, 6, 7, 8 survive.")

        self.h("B.14 Termination", level=3)
        self.p("B.14.1 Also terminates on: (a) Master termination, "
               "(b) mutual agreement, or (c) 30 days' written notice.")
        self.p("B.14.3 Surviving clauses: B.5, B.9, B.10, B.11.2, B.13.")

        # Schedule B-1
        self.line()
        self.h("Schedule B-1 — Per-Execution Capital Intro Details", level=3)
        sched = self.d.get("annex_b", {}).get("schedule_b1", {})
        self.table(
            ["Field", "Value"],
            [
                ["Approved Investor name", ""],
                ["Investor jurisdiction", ""],
                ["Date on Named Investor List", ""],
                ["Date of Provider's approval", ""],
                ["Anticipated Capital Committed (EUR)", ""],
                ["Invoicing entity", sched.get("invoicing_entity", "")],
                ["Payable-to entity", sched.get("payable_to_entity", "")],
                ["Currency", sched.get("currency", "EUR")],
                ["Payment terms", sched.get("payment_terms", "Net 30")],
                ["VAT treatment", sched.get("vat_treatment", "")],
                ["Counsel-review confirmation", sched.get("counsel_review_confirmation", "")],
            ],
        )

        # Signatures
        self._annex_signatures("Annex B", provider_override=prov_b)

    def _annex_signatures(self, annex_label, provider_override=None):
        self.line()
        self.h(f"Signatures — {annex_label}", level=3)
        prov = provider_override or self.d.get("provider", {})
        prov_name = prov.get("legal_name", "")
        prov_reg = prov.get("reg_type", "")
        prov_reg_num = prov.get("reg_number", "")
        prov_sig = prov.get("signatory_name", "")
        prov_title = prov.get("signatory_title", "")

        intro = self.d.get("introducer", {})
        intro_name = intro.get("name", "")
        intro_reg = intro.get("reg_type", "")
        intro_reg_num = intro.get("reg_number", "")

        self.bp(f"For and on behalf of {prov_name}", "")
        self.p(f"{prov_reg}: {prov_reg_num}")
        self.p("Signature: ____________________________")
        self.p(f"Name: {prov_sig}")
        self.p(f"Title: {prov_title}")
        self.p("Date: ____________________________")
        self.line()
        self.bp(f"For and on behalf of {intro_name}", "")
        self.p(f"{intro_reg}: {intro_reg_num}")
        self.p("Signature: ____________________________")
        self.p(f"Name: {intro.get('signatory_name', '')}")
        self.p(f"Title: {intro.get('signatory_title', '')}")
        self.p("Date: ____________________________")

    # --- Build ---

    def build(self):
        self.cover()
        self.recitals()
        self.definitions()
        self.purpose_and_scope()
        self.relationship()
        self.referral_identification()
        self.confidentiality()
        self.non_circumvention()
        self.pipeline_carveout()
        self.general_provisions()
        self.signatures()
        self.annex_execution_record()

        if self.annex_a:
            self.annex_a_full()
        if self.annex_b:
            self.annex_b_full()

        return self.doc


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    if len(sys.argv) < 2:
        print("Usage: python generate_mia.py <intake.yaml> [--output <path.docx>]")
        sys.exit(1)

    intake_path = sys.argv[1]
    output_path = "output_mia.docx"
    if "--output" in sys.argv:
        idx = sys.argv.index("--output")
        if idx + 1 < len(sys.argv):
            output_path = sys.argv[idx + 1]

    data = load_intake(intake_path)
    mia = MIA(data)
    doc = mia.build()
    doc.save(output_path)
    print(f"MIA generated: {output_path}")

    # Summary
    scope = data.get("scope", "")
    intro = data.get("introducer", {}).get("short", "Unknown")
    annexes = []
    if data.get("activation", {}).get("annex_a"):
        annexes.append("Annex A (Commercial)")
    if data.get("activation", {}).get("annex_b"):
        annexes.append("Annex B (Capital)")
    annex_str = " + ".join(annexes) if annexes else "Master only"
    print(f"  Scope: {scope} — {annex_str}")
    print(f"  Introducer: {intro}")
    if data.get("activation", {}).get("annex_b"):
        print("  ⚠  Annex B activated — route through legal-counsel before issuing.")


if __name__ == "__main__":
    main()
