#!/usr/bin/env python3
"""
LOI/NCNDA v3.2 — Document Generation Script

Usage:
    python generate_loi.py intake.yaml [--output path/to/output.docx]
                                        [--override R-11,R-14]
                                        [--override-reason "..."]

Takes a YAML intake file and produces a complete, branded .docx ready to sign.
All clauses, conditionals, and schedules are generated from the intake data.

v3.2 changes (2026-04-16):
- Recital A is library-sourced (see _shared/loi-recital-a-library.md). Variants:
  default | sovereignty | integration | bespoke.
- Customer-facing "DEC Block" language removed. Capacity expressed in MW IT +
  Designated Sites. Deprecated YAML field: commercial.dec_block_count.
- "Minimum commitment term of 5 years" replaced with "approximately 5 years,
  indicative only".
- Cl. 4.2 "Revenue Chain" with Unicode arrows replaced with "Contractual
  Sequence" prose + numbered list.
- Closing line: hardcoded single-sentence "We look forward to working with you."
  Honors OPEN-1 (commit 2097f52) — bespoke_closing is not supported.
  choices.bespoke_closing in YAML is silently ignored for backward compat.
- Schedule 1 title no longer carries "(NON-BINDING)" suffix. Italic prefatory
  note on the schedule; authoritative binding/non-binding assignment in Cl. 5.1.
- Pre-save QA linter. Rules in _shared/loi-qa-gate.md. Severity: fail / warn /
  info. Overridable via --override / --override-reason.

Planned (not yet wired in this file): new types StrategicSupplier (DE-LOI-SS-v1.0)
and EcosystemPartnership (DE-LOI-EP-v1.0). See CHANGELOG.md.
"""

import sys
import os
import yaml
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

# Import shared cover page and branding from document-factory
# Script now lives under legal-assistant/sales/, one level deeper than the
# original loi-generator/ layout, so resolve document-factory via parent.parent.
#
# ⚠️ STAGING MIRROR SENTINEL (v3.7.2): when this file is mirrored to
# degitos-staging (skills/de-legal-assistant/sales/generate_loi.py), the
# sibling skill name MUST be substituted: `document-factory` →
# `de-document-factory` (DEGitOS de- prefix convention). The mirror
# discipline lives in `docs/staging-mirror.md`. DO NOT remove or rename
# this sentinel without updating the mirror script at the same time.
SCRIPT_DIR = Path(__file__).parent
sys.path.insert(0, str(SCRIPT_DIR.parent.parent / "document-factory"))
from common import (  # noqa: E402
    add_cover, Party,
    COBALT, SLATE, SLATE_800, SLATE_900, WHITE,
    setup_first_page_header, setup_cont_header,
    setup_first_footer, setup_cont_footer,
)


# ---------------------------------------------------------------------------
# Configuration -- LOI clause formatting (body text, not cover page)
# ---------------------------------------------------------------------------

FONT_NAME = "Inter"
FONT_BODY = Pt(10)
FONT_HEADING1 = Pt(13)
FONT_HEADING2 = Pt(11)
FONT_SMALL = Pt(8)
LINE_SPACING = 1.15

# Local aliases for clause body text (imported colors used for cover page)
NAVY = SLATE_900
BLACK = SLATE_800
GREY = SLATE


# ---------------------------------------------------------------------------
# YAML Loading + Validation
# ---------------------------------------------------------------------------

def load_entities_register() -> dict:
    """Load the entities register from `config/entities.yaml`.

    v3.5.2 scope Q: single source of truth for DE's legal entities. Returns
    empty dict if register file is absent (backward-compat — intake YAMLs
    can still use explicit `provider.legal_name`, `address`, `kvk` fields).
    """
    here = os.path.dirname(os.path.abspath(__file__))
    register_path = os.path.join(here, "..", "config", "entities.yaml")
    if not os.path.exists(register_path):
        return {}
    try:
        with open(register_path, "r") as f:
            return yaml.safe_load(f) or {}
    except Exception:
        return {}


def expand_provider_from_register(data: dict, register: dict) -> dict:
    """Expand `provider.entity: "de_nl"` references into full provider record.

    v3.5.2 scope Q. Behaviour:
      - If `provider.entity` is set and register has the key, expand it.
        Any fields explicitly set on the intake YAML itself (e.g. signatory
        overrides) take precedence over the register defaults.
      - If `provider.entity` is not set but the register has a
        `type_defaults` block with an entry for the intake's `type`,
        populate the entity + signatory_mode from that per-type default
        (v3.5.3-cont scope J12 wiring). A minimal intake YAML with only
        `type: Wholesale` now auto-expands to NL BV / Carlos / Director.
      - If neither is set, the provider dict is returned unchanged
        (backward-compat with v3.5.1 and earlier).
      - `signatory_mode` on the intake YAML selects the default signatory
        variant from the register (pre_msa / post_msa for NL BV; ceo for AG).

    Mutates and returns `data`.
    """
    prov = data.get("provider", {}) or {}
    entity_key = prov.get("entity")

    # v3.5.3-cont scope J12: if no explicit entity is set AND the provider
    # dict is "minimal" (contains at most entity/signatory_mode/short_name —
    # i.e. no explicit identity fields that the user clearly set themselves),
    # apply the per-LOI-type default from the register's `type_defaults`
    # block. This is a pure fallback — explicit `provider.entity` wins, and
    # explicit-fields intakes (legal_name/address/kvk/etc. set) are
    # preserved untouched for backward-compat with v3.5.1.
    if not entity_key:
        # "Identity fields" — presence of ANY of these means the user wrote
        # an explicit-fields intake and type_defaults must NOT override.
        _EXPLICIT_IDENTITY_FIELDS = (
            "legal_name", "address", "kvk", "reg_number",
            "reg_type", "jurisdiction", "legal_form",
        )
        is_minimal = not any(prov.get(f) for f in _EXPLICIT_IDENTITY_FIELDS)
        if is_minimal:
            type_defaults = register.get("type_defaults", {}) or {}
            type_key = data.get("type")
            td = type_defaults.get(type_key)
            if td and isinstance(td, dict):
                entity_key = td.get("entity")
                if entity_key and "signatory_mode" not in prov:
                    prov = dict(prov)
                    prov["signatory_mode"] = td.get("signatory_mode", "pre_msa")
                    data["provider"] = prov

    if not entity_key:
        return data
    entities = register.get("entities", {})
    entity = entities.get(entity_key)
    if not entity:
        return data

    # Build the expanded provider record. Explicit intake values WIN over
    # register defaults — this lets intake YAMLs override signatory per-deal
    # without touching the register.
    expanded = {
        "legal_name": entity.get("legal_name", ""),
        "short_name": entity.get("short_name", ""),
        "abbreviation": entity.get("abbreviation", ""),
        "legal_form": entity.get("legal_form", ""),
        "jurisdiction": entity.get("jurisdiction", ""),
        "address": entity.get("address", ""),
        "reg_type": entity.get("reg_type", ""),
        "reg_number": entity.get("reg_number", ""),
        "rsin": entity.get("rsin", ""),
        "parent": entity.get("parent"),
    }
    # Pick signatory based on signatory_mode
    mode = prov.get("signatory_mode", "pre_msa")
    sig = None
    if mode == "pre_msa" and "default_signatory_pre_msa" in entity:
        sig = entity["default_signatory_pre_msa"]
    elif mode == "post_msa" and "default_signatory_post_msa" in entity:
        sig = entity["default_signatory_post_msa"]
    elif "default_signatory" in entity:
        sig = entity["default_signatory"]
    if sig:
        expanded["signatory_name"] = sig.get("name", "")
        expanded["signatory_title"] = sig.get("title", "")

    # Apply explicit intake overrides (intake wins)
    for k, v in prov.items():
        if k in ("entity", "signatory_mode"):
            continue
        if v:
            expanded[k] = v

    data["provider"] = expanded
    return data


def load_intake(path: str) -> dict:
    with open(path, "r") as f:
        data = yaml.safe_load(f)
    # v3.5.2 scope Q: expand provider.entity references BEFORE validation so
    # downstream validators see the fully-populated provider record.
    register = load_entities_register()

    # v3.5 polish: fail fast on unknown `provider.entity` keys. Without this
    # check, typos like `entity: "de_nnl"` would silently fall through to
    # the backward-compat path (treating the YAML's explicit fields as
    # ground truth) and the user would only notice at document-render time.
    prov = data.get("provider", {}) or {}
    entity_key = prov.get("entity")
    if entity_key:
        entities = register.get("entities", {}) or {}
        if entity_key not in entities:
            available = sorted(entities.keys())
            print(
                f"ERROR: provider.entity '{entity_key}' not found in "
                f"config/entities.yaml. Available keys: "
                f"{', '.join(available) if available else '(none — register missing or empty)'}",
                file=sys.stderr,
            )
            sys.exit(1)

    data = expand_provider_from_register(data, register)
    validate(data)
    return data


def validate(d: dict):
    """Validate intake YAML. Backward-compatible behaviour: prints errors
    + sys.exit(1) on failure (CLI semantics).

    v3.8.0: returns the error list when called from tests AND non-empty
    list triggers print + exit. Tests that want to inspect errors without
    crashing should call `validate_errors(d)` instead.
    """
    errors = validate_errors(d)
    if errors:
        print("VALIDATION ERRORS:")
        for e in errors:
            print(f"  - {e}")
        sys.exit(1)
    return errors  # always [] when reached


def validate_errors(d: dict) -> list:
    """Validate intake YAML. Raises SystemExit on failure.

    v3.2: rule R-18 (fail) — deprecated field commercial.dec_block_count.
    """
    errors = []
    t = d.get("type", "")
    if t not in ("EndUser", "Distributor", "Wholesale",
                  "StrategicSupplier", "EcosystemPartnership", "Bespoke"):
        errors.append(
            "type must be one of: EndUser, Distributor, Wholesale, "
            "StrategicSupplier, EcosystemPartnership, Bespoke "
            f"(got: {d.get('type')})"
        )
    for s in ("provider", "counterparty", "programme", "dates"):
        if s not in d:
            errors.append(f"Missing section: {s}")
    cp = d.get("counterparty", {})
    for f in ("name", "short"):
        if not cp.get(f):
            errors.append(f"counterparty.{f} required")
    if not d.get("dates", {}).get("loi_date"):
        errors.append("dates.loi_date required")

    # v3.8.0: Recital B is built from the slot block. The freeform
    # `description` field is removed. R-DEPRECATED-FIELD fires if it's
    # still present.
    if "description" in cp:
        errors.append(
            "[R-DEPRECATED-FIELD] counterparty.description is removed in "
            "v3.8.0 — use counterparty.recital_b slot block. See "
            "_shared/counterpart-description-framework.md for the schema."
        )
    if t != "Bespoke" and not cp.get("recital_b"):
        errors.append(
            "counterparty.recital_b required (slot block: legal_identity, "
            "operational_verb, customer_use_case, material_asset; "
            "bargain_relevant_fact optional). See _shared/"
            "counterpart-description-framework.md."
        )
    elif cp.get("recital_b"):
        # v3.8.0 R-32 — slot vocabulary + closed-enum + banned-phrase lint
        errors.extend(_validate_recital_b_slots(cp["recital_b"]))

    # R-18 — deprecated field migration error
    if "dec_block_count" in d.get("commercial", {}):
        errors.append(
            "[R-18] commercial.dec_block_count is deprecated in v3.2. "
            "Use commercial.indicative_mw instead. See CHANGELOG.md § Migration."
        )

    if t == "Distributor" and d.get("partnership_mode", "combined") == "combined":
        for f in ("partner_core_capability", "partner_contribution",
                   "combined_offering", "target_end_users", "partner_service_scope"):
            if not d.get("commercial", {}).get(f):
                errors.append(f"commercial.{f} required for Distributor Mode A")
    if t == "Wholesale":
        if not d.get("commercial", {}).get("indicative_mw"):
            errors.append("commercial.indicative_mw required for Wholesale (MW IT)")
    if t == "EndUser":
        if not d.get("commercial", {}).get("service_type"):
            errors.append("commercial.service_type required for EndUser")
    if t == "StrategicSupplier":
        supplier = d.get("supplier", {})
        if not supplier.get("capability_category"):
            errors.append("supplier.capability_category required for StrategicSupplier")
        if not supplier.get("core_capability"):
            errors.append("supplier.core_capability required for StrategicSupplier")
        purposes = supplier.get("strategic_purposes", [])
        if not purposes or not (1 <= len(purposes) <= 2):
            errors.append("supplier.strategic_purposes required for StrategicSupplier (1-2 items)")
        valid_purposes = {"capacity_lock_in", "pricing_volume",
                           "supply_chain_de_risking", "engineering_integration",
                           "pipeline_visibility"}
        for p in purposes:
            if p not in valid_purposes:
                errors.append(
                    f"supplier.strategic_purposes: invalid '{p}'. "
                    f"Valid: {sorted(valid_purposes)}"
                )
        if "capacity_lock_in" in purposes and not supplier.get("lead_time_target"):
            errors.append(
                "supplier.lead_time_target required when strategic_purposes "
                "includes capacity_lock_in"
            )
        if "pricing_volume" in purposes and not supplier.get("volume_indicative"):
            errors.append(
                "supplier.volume_indicative required when strategic_purposes "
                "includes pricing_volume"
            )
        if ("engineering_integration" in purposes
                and not d.get("choices", {}).get("joint_ip")):
            errors.append(
                "choices.joint_ip required when strategic_purposes "
                "includes engineering_integration (none | background | foreground)"
            )
    if t == "Bespoke":
        # Bespoke type — structural validation only, content is free text
        # but still passes through the full QA catalog (banned phrases etc).
        clauses = d.get("clauses")
        if not clauses or not isinstance(clauses, list) or len(clauses) < 1:
            errors.append(
                "clauses required for Bespoke (list of at least one clause "
                "with number, heading, paragraphs)"
            )
        else:
            for idx, c in enumerate(clauses):
                if not isinstance(c, dict):
                    errors.append(f"clauses[{idx}] must be a mapping")
                    continue
                if not c.get("number"):
                    errors.append(f"clauses[{idx}].number required")
                if not c.get("heading"):
                    errors.append(f"clauses[{idx}].heading required")
                paras = c.get("paragraphs")
                if not paras or not isinstance(paras, list) or len(paras) < 1:
                    errors.append(
                        f"clauses[{idx}].paragraphs required "
                        "(list of at least one string)"
                    )
                subs = c.get("subclauses", [])
                if subs and isinstance(subs, list):
                    for si, s in enumerate(subs):
                        if not isinstance(s, dict):
                            errors.append(
                                f"clauses[{idx}].subclauses[{si}] must be a mapping"
                            )
                            continue
                        if not s.get("letter"):
                            errors.append(
                                f"clauses[{idx}].subclauses[{si}].letter required"
                            )
                        if not s.get("text"):
                            errors.append(
                                f"clauses[{idx}].subclauses[{si}].text required"
                            )
        # Optional supplementary recitals
        recs = d.get("recitals", [])
        if recs and isinstance(recs, list):
            for ri, r in enumerate(recs):
                if not isinstance(r, dict):
                    errors.append(f"recitals[{ri}] must be a mapping")
                    continue
                if not r.get("letter"):
                    errors.append(f"recitals[{ri}].letter required")
                if not r.get("text"):
                    errors.append(f"recitals[{ri}].text required")
    if t == "EcosystemPartnership":
        eco = d.get("ecosystem", {})
        if not eco.get("relationship_type"):
            errors.append("ecosystem.relationship_type required for EcosystemPartnership")
        valid_rel = {"standards_body", "university", "research_consortium",
                      "co_marketing", "industry_association", "policy_partner",
                      "other"}
        if eco.get("relationship_type") and eco.get("relationship_type") not in valid_rel:
            errors.append(
                f"ecosystem.relationship_type must be one of {sorted(valid_rel)}"
            )
        if not eco.get("collaboration_themes"):
            errors.append("ecosystem.collaboration_themes required for EcosystemPartnership")
        if not eco.get("joint_activity_categories"):
            errors.append("ecosystem.joint_activity_categories required for EcosystemPartnership")

    # v3.7.0: choices.recital_b_density validation
    density = d.get("choices", {}).get("recital_b_density", "standard")
    if density not in ("terse", "standard", "verbose"):
        errors.append(
            f"choices.recital_b_density must be one of: terse, standard, verbose "
            f"(got: {density!r})"
        )

    # v3.7.0 extensibility layer — all optional, default-compatible
    choices = d.get("choices", {})

    # choices.include_schedule: bool (default True)
    inc_sched = choices.get("include_schedule", True)
    if not isinstance(inc_sched, bool):
        errors.append(
            f"choices.include_schedule must be a boolean (got: {inc_sched!r})"
        )

    # v3.7.1: choices.auto_renumber: bool (default False)
    auto_renum = choices.get("auto_renumber", False)
    if not isinstance(auto_renum, bool):
        errors.append(
            f"choices.auto_renumber must be a boolean (got: {auto_renum!r})"
        )

    # choices.confidentiality_opt_outs: list of valid keys
    opt_outs = choices.get("confidentiality_opt_outs", [])
    if opt_outs:
        valid_opt_outs = {"onward_sharing", "compliance_confirmation", "metadata_protection"}
        if not isinstance(opt_outs, list):
            errors.append(
                f"choices.confidentiality_opt_outs must be a list (got: {type(opt_outs).__name__})"
            )
        else:
            for k in opt_outs:
                if k not in valid_opt_outs:
                    errors.append(
                        f"choices.confidentiality_opt_outs: invalid key '{k}'. "
                        f"Valid: {sorted(valid_opt_outs)}"
                    )

    # supplier.rofr block (SS only)
    supplier = d.get("supplier", {})
    rofr = supplier.get("rofr")
    if rofr is not None:
        if t != "StrategicSupplier":
            errors.append(
                "supplier.rofr is only valid for type=StrategicSupplier"
            )
        if isinstance(rofr, dict):
            valid_lock_out = {"alignment", "sole_discretion", "hard_minimum", "milestone"}
            lock_out = rofr.get("lock_out_style", "sole_discretion")
            if lock_out not in valid_lock_out:
                errors.append(
                    f"supplier.rofr.lock_out_style must be one of "
                    f"{sorted(valid_lock_out)} (got: {lock_out!r})"
                )

    # v3.7.1: supplier.co_marketing block (SS only)
    co_marketing = supplier.get("co_marketing")
    if co_marketing is not None:
        if t != "StrategicSupplier":
            errors.append(
                "supplier.co_marketing is only valid for type=StrategicSupplier"
            )
        if isinstance(co_marketing, dict):
            valid_framing = {"multi_supplier", "preferred", "exclusive"}
            framing = co_marketing.get("framing", "multi_supplier")
            if framing not in valid_framing:
                errors.append(
                    f"supplier.co_marketing.framing must be one of "
                    f"{sorted(valid_framing)} (got: {framing!r})"
                )
            valid_logo = {"yes", "no", "per_event_approval"}
            logo_use = co_marketing.get("logo_use", "per_event_approval")
            if logo_use not in valid_logo:
                errors.append(
                    f"supplier.co_marketing.logo_use must be one of "
                    f"{sorted(valid_logo)} (got: {logo_use!r})"
                )
            valid_press = {"none", "joint", "unilateral_allowed"}
            press = co_marketing.get("press_at_loi", "none")
            if press not in valid_press:
                errors.append(
                    f"supplier.co_marketing.press_at_loi must be one of "
                    f"{sorted(valid_press)} (got: {press!r})"
                )

    # supplier.referral_rider: bool (SS only)
    referral_rider = supplier.get("referral_rider")
    if referral_rider is not None:
        if t != "StrategicSupplier":
            errors.append(
                "supplier.referral_rider is only valid for type=StrategicSupplier"
            )
        if not isinstance(referral_rider, bool):
            errors.append(
                f"supplier.referral_rider must be a boolean (got: {referral_rider!r})"
            )

    # custom.definitions[] + custom.clauses[] + custom.definitions_include[]
    custom = d.get("custom", {})
    # v3.7.2: collision check — warn if custom.clauses[].number would overlap
    # an engine-emitted clause number for this type. Hardcoded (type-aware)
    # list of known-emitted numbers; append-mode collisions are still allowed
    # (document will contain both, numbered the same — operator should use
    # a new number), but replace/insert-after with a typo is surfaced here.
    KNOWN_CLAUSE_NUMBERS_BY_TYPE = {
        "Wholesale": [
            "3.1","3.2","3.3","3.4","3.5","3.6","3.7","3.8",
            "4.1","4.2","4.3","4.4","4.5",
            "5.1","5.2","5.3","5.4",
            "6.1","6.2","6.3","6.4","6.5","6.6","6.7","6.8","6.9","6.10",
            "6.11","6.12","6.13","6.14","6.15","6.16",
            "7.1","7.2","7.3","7.4","7.5","7.6",
            "8.1","8.2","8.3","8.4","8.5","8.6","8.7","8.8","8.9","8.10",
        ],
        "StrategicSupplier": [
            "3.1","3.2","3.3","3.4","3.5","3.6","3.7","3.8","3.9","3.10","3.11",
            "4.1","4.2","4.3","4.4","4.5","4.6",
            "5.1","5.2","5.3","5.4",
            "6.1","6.2","6.3","6.4","6.5","6.6","6.7","6.8","6.9","6.10",
            "6.11","6.12","6.13","6.14","6.15","6.16",
            "7.1","7.2","7.3","7.4","7.5","7.6",
            "8.1","8.2","8.3","8.4","8.5","8.6","8.7","8.8","8.9","8.10",
        ],
    }
    if custom and isinstance(custom, dict):
        known = set(KNOWN_CLAUSE_NUMBERS_BY_TYPE.get(t, []))
        for i, item in enumerate(custom.get("clauses", []) or []):
            if not isinstance(item, dict):
                continue
            num = str(item.get("number", "")).strip()
            mode = item.get("mode", "append")
            if mode == "append" and num in known:
                errors.append(
                    f"custom.clauses[{i}]: append-mode number {num!r} collides "
                    f"with engine-emitted clause for type={t!r}. "
                    f"Use a different number (e.g., 9.X) to avoid duplicate "
                    f"numbering, or use mode=replace to overwrite the "
                    f"engine output."
                )
    if custom:
        if not isinstance(custom, dict):
            errors.append(
                f"custom must be a dict (got: {type(custom).__name__})"
            )
        else:
            cdefs = custom.get("definitions", [])
            if cdefs and not isinstance(cdefs, list):
                errors.append(
                    f"custom.definitions must be a list (got: {type(cdefs).__name__})"
                )
            else:
                for i, item in enumerate(cdefs or []):
                    if not isinstance(item, dict) or "key" not in item or "text" not in item:
                        errors.append(
                            f"custom.definitions[{i}] must have 'key' and 'text' fields"
                        )
            cincl = custom.get("definitions_include", [])
            if cincl and not isinstance(cincl, list):
                errors.append(
                    f"custom.definitions_include must be a list (got: {type(cincl).__name__})"
                )
            cclauses = custom.get("clauses", [])
            if cclauses and not isinstance(cclauses, list):
                errors.append(
                    f"custom.clauses must be a list (got: {type(cclauses).__name__})"
                )
            else:
                valid_modes = {"append", "replace", "insert-after"}
                for i, item in enumerate(cclauses or []):
                    if not isinstance(item, dict):
                        errors.append(f"custom.clauses[{i}] must be a dict")
                        continue
                    if "number" not in item or "text" not in item:
                        errors.append(
                            f"custom.clauses[{i}] must have 'number' and 'text' fields"
                        )
                    mode = item.get("mode", "append")
                    mode_base = mode.split(":")[0]
                    if mode_base not in valid_modes:
                        errors.append(
                            f"custom.clauses[{i}].mode must be one of: "
                            f"{sorted(valid_modes)} or 'insert-after:N' "
                            f"(got: {mode!r})"
                        )

    # dates.financing_context validation
    fctx = d.get("dates", {}).get("financing_context")
    if fctx is not None:
        if not isinstance(fctx, dict):
            errors.append(
                f"dates.financing_context must be a dict (got: {type(fctx).__name__})"
            )
        else:
            linked = fctx.get("linked_to_fundraise", False)
            if linked:
                close = fctx.get("fundraise_close_target")
                buffer_months = fctx.get("buffer_months_post_close")
                if not close:
                    errors.append(
                        "dates.financing_context.fundraise_close_target required "
                        "when linked_to_fundraise=true"
                    )
                if buffer_months is None:
                    errors.append(
                        "dates.financing_context.buffer_months_post_close required "
                        "when linked_to_fundraise=true"
                    )
                elif not isinstance(buffer_months, int):
                    errors.append(
                        "dates.financing_context.buffer_months_post_close must be an integer"
                    )

    # counterparty.relationship_cluster + identity_map (structural only; QA-report-surfaced)
    cp = d.get("counterparty", {})
    rel = cp.get("relationship_cluster")
    if rel is not None and not isinstance(rel, dict):
        errors.append(
            f"counterparty.relationship_cluster must be a dict (got: {type(rel).__name__})"
        )
    idmap = cp.get("identity_map")
    if idmap is not None and not isinstance(idmap, dict):
        errors.append(
            f"counterparty.identity_map must be a dict (got: {type(idmap).__name__})"
        )

    # Recital A variant validation
    variant = d.get("programme", {}).get("recital_a_variant", "default")
    if variant not in ("default", "sovereignty", "integration", "bespoke"):
        errors.append(
            f"programme.recital_a_variant must be one of: default, sovereignty, "
            f"integration, bespoke (got: {variant})"
        )
    if variant == "bespoke" and not d.get("programme", {}).get("recital_a_bespoke"):
        errors.append(
            "programme.recital_a_bespoke required when recital_a_variant=bespoke"
        )

    return errors  # validate_errors() return point


def _validate_recital_b_slots(rb: dict) -> list:
    """v3.8.0 R-32 — slot vocabulary + closed-enum + banned-phrase lint
    + slot-5 named-entity proof requirement.

    Returns list of error strings (empty on pass).
    """
    from recital_b_vocab import (
        find_banned_phrases,
        find_named_entities_in_text,
        validate_legal_form,
        validate_operational_verb,
    )

    errors: list[str] = []

    if not isinstance(rb, dict):
        return ["[R-32] counterparty.recital_b must be a mapping (slot block)"]

    # Required slots
    for slot_key in ("legal_identity", "operational_verb",
                     "customer_use_case", "material_asset"):
        if slot_key not in rb or not isinstance(rb[slot_key], dict):
            errors.append(
                f"[R-32] counterparty.recital_b.{slot_key} required (mapping)"
            )

    # Slot 1 — legal_identity
    li = rb.get("legal_identity") or {}
    if li:
        legal_form = li.get("legal_form")
        jurisdiction = li.get("jurisdiction")
        if legal_form and jurisdiction:
            ok, msg = validate_legal_form(legal_form, jurisdiction)
            if not ok:
                errors.append(f"[R-32] legal_identity: {msg}")
            elif msg:  # ok with warn
                # Warn-class messages are surfaced as INFO, not errors.
                # We attach them via a special prefix the caller can ignore.
                pass

    # Slot 2 — operational_verb
    ov = rb.get("operational_verb") or {}
    if ov:
        verb = ov.get("verb")
        if verb:
            ok, msg = validate_operational_verb(verb)
            if not ok:
                errors.append(f"[R-32] operational_verb: {msg}")
        obj = ov.get("object")
        if obj:
            findings = find_banned_phrases(obj)
            if findings:
                errors.append(
                    f"[R-32] operational_verb.object contains banned phrase(s): "
                    f"{findings}"
                )

    # Slot 3 — customer_use_case
    cu = rb.get("customer_use_case") or {}
    if cu:
        category = cu.get("category")
        if category:
            findings = find_banned_phrases(category)
            if findings:
                errors.append(
                    f"[R-32] customer_use_case.category contains banned "
                    f"phrase(s): {findings}"
                )

    # Slot 4 — material_asset
    ma = rb.get("material_asset") or {}
    if ma:
        asset = ma.get("asset")
        if asset:
            findings = find_banned_phrases(asset)
            if findings:
                errors.append(
                    f"[R-32] material_asset.asset contains banned phrase(s): "
                    f"{findings}"
                )

    # Slot 5 — bargain_relevant_fact (OPTIONAL)
    fact = rb.get("bargain_relevant_fact")
    if fact:
        claim = fact.get("claim", "")
        # Banned-phrase scan
        if claim:
            findings = find_banned_phrases(claim)
            if findings:
                errors.append(
                    f"[R-32] bargain_relevant_fact.claim contains banned "
                    f"phrase(s): {findings}"
                )
        # Named-entity proof requirement
        names = find_named_entities_in_text(claim)
        if names:
            named_entities = fact.get("named_entities") or []
            if not named_entities:
                errors.append(
                    f"[R-32] bargain_relevant_fact.claim references named "
                    f"entities {names!r} but no `named_entities[]` proof "
                    f"block provided. Add structured proof per "
                    f"_shared/counterpart-description-framework.md."
                )
            else:
                # Each named_entity must carry materiality + proof
                for idx, ne in enumerate(named_entities):
                    if not isinstance(ne, dict):
                        errors.append(
                            f"[R-32] named_entities[{idx}] must be a mapping"
                        )
                        continue
                    if not ne.get("name"):
                        errors.append(
                            f"[R-32] named_entities[{idx}].name required"
                        )
                    materiality = ne.get("materiality", "")
                    if len(materiality) < 30:
                        errors.append(
                            f"[R-32] named_entities[{idx}].materiality must "
                            f"be ≥30 chars; got {len(materiality)}"
                        )
                    if materiality:
                        m_findings = find_banned_phrases(materiality)
                        if m_findings:
                            errors.append(
                                f"[R-32] named_entities[{idx}].materiality "
                                f"contains banned phrase(s): {m_findings}"
                            )
                    proof = ne.get("proof") or {}
                    if not proof.get("url"):
                        errors.append(
                            f"[R-32] named_entities[{idx}].proof.url required"
                        )
                    if not proof.get("dated"):
                        errors.append(
                            f"[R-32] named_entities[{idx}].proof.dated required"
                        )

    return errors


# ---------------------------------------------------------------------------
# Recital A (v3.4 — single canonical body + per-type tails)
# ---------------------------------------------------------------------------
# Source of truth is _shared/loi-recital-a-library.md. v3.4 change: collapsed
# three variants (default / sovereignty / integration) into a single canonical
# body. User-approved wording 2026-04-17. Per-type tails now cover all 5 types
# (v3.3 only had tails for DS/SS/EP; EU/WS were empty strings).

RECITAL_A_BODY = (
    # v3.5.2 brand-name rename: defined term "Digital Energy" is established
    # in the Parties Preamble; Recital A uses it directly (no "(the "Provider")"
    # parenthetical). Any reference to "Digital Energy" in the body has been
    # replaced with "Digital Energy".
    '{prov} develops and operates Digital Energy Centers '
    '("DECs"), distributed energy hubs for liquid-cooled AI colocation, '
    "integrating accelerated compute with heat recycling and behind-the-meter "
    "(BTM) power production, engineered as one integrated system. Digital Energy "
    "is building an integrated sovereign AI infrastructure platform for "
    "enterprise and institutional customers, designed for edge inference."
)

# Per-type tail appended to the shared body. Customer-facing types (EU, WS)
# use "Digital Energy's integrated platform [verb]" subject (procurement pattern);
# relationship-facing types (DS, SS, EP) use "Digital Energy [verb]" subject.
RECITAL_A_TAIL_BY_TYPE = {
    "EndUser": (
        " Digital Energy's integrated platform supplies dedicated AI inference "
        "capacity to enterprises, AI labs, and research institutions requiring "
        "sovereign, high-density and low-latency infrastructure on European soil."
    ),
    "Distributor": (
        " Digital Energy seeks qualified channel and integration partners to "
        "extend its platform reach to end-user segments where the Partner "
        "holds established customer relationships and domain expertise."
    ),
    "Wholesale": (
        " Digital Energy's integrated platform contracts liquid-cooled AI "
        "colocation capacity at megawatt scale to NeoCloud operators and "
        "GPU cloud providers internationally."
    ),
    "StrategicSupplier": (
        " Digital Energy seeks qualified EPC contractors, modular infrastructure "
        "manufacturers, and OEM vendors to deliver the DEC platform and secure "
        "supply continuity across its active development pipeline."
    ),
    "EcosystemPartnership": (
        " Digital Energy engages with ecosystem partners on sovereign AI "
        "infrastructure, sustainable datacentre design, and European "
        "industrial policy alignment."
    ),
    # Bespoke: no standard fit-line. The YAML author supplies recitals
    # tailored to the specific deal. Falls through to the base body only.
    "Bespoke": "",
}


def resolve_recital_a(d: dict) -> str:
    """Return the Recital A body + type-specific tail (without the '(A) ' prefix).

    v3.4: single canonical body formatted with provider short name, plus a
    per-type tail. Bespoke override path (programme.recital_a_variant='bespoke'
    with programme.recital_a_bespoke) returns the raw bespoke string, with the
    QA linter enforcing forbidden-pattern rules.

    Legacy YAMLs with recital_a_variant set to 'default' / 'sovereignty' /
    'integration' still render correctly — those variant keys are ignored and
    the single canonical body is used. No migration break.
    """
    prov = d.get("provider", {}).get("short_name", "Digital Energy")
    t = d.get("type", "EndUser")
    variant = d.get("programme", {}).get("recital_a_variant", "default")
    if variant == "bespoke":
        bespoke = d.get("programme", {}).get("recital_a_bespoke", "")
        if bespoke:
            return bespoke
    body = RECITAL_A_BODY.format(prov=prov)
    tail = RECITAL_A_TAIL_BY_TYPE.get(t, "")
    return body + tail


# ---------------------------------------------------------------------------
# Document Builder
# ---------------------------------------------------------------------------

_AGREEMENT_TYPE_BY_LOI = {
    "Distributor": "Letter of Intent and NCNDA",
    "Wholesale": "Letter of Intent and NCNDA",
    "EndUser": "Letter of Intent",
    "StrategicSupplier": "Letter of Intent and NCNDA",
    "EcosystemPartnership": "Letter of Intent",
    # M4 Bespoke default — can be overridden via top-level `agreement_type:` in YAML
    "Bespoke": "Letter of Intent",
}

# v3.4: Wholesale subject dropped "Purpose-Built" marketing modifier.
_SUBJECT_BY_LOI = {
    "Distributor": "Strategic Infrastructure Partnership",
    "Wholesale": "AI Colocation Capacity",
    "EndUser": "AI Compute Infrastructure Services",
    "StrategicSupplier": "Strategic Supply and Infrastructure Partnership",
    "EcosystemPartnership": "Strategic Ecosystem Collaboration",
    # M4 Bespoke default — override via top-level `subject:` in YAML
    "Bespoke": "Bespoke Engagement",
}


def _lead_time_under_six_months(value, *, allow_unparseable=True) -> bool:
    """v3.7.1 — parse a lead_time_target string into days; return True if <180.

    Accepted forms (case-insensitive): "90 days", "6 weeks", "3 months",
    "90d", "6w", "3m". Falls back to False on unparseable input (safer
    default — the clause doesn't fire).

    v3.7.2: ``lead_time_parse_result`` (module-level) captures the last
    parse outcome for qa_lint to surface as a WARN when operator-supplied
    text was non-empty but unparseable.
    """
    global _LAST_LEAD_TIME_PARSE  # used by qa_lint for advisory WARN
    _LAST_LEAD_TIME_PARSE = {"input": value, "parsed_days": None, "ok": False}
    if not value:
        _LAST_LEAD_TIME_PARSE["ok"] = True  # empty is valid (clause doesn't fire)
        return False
    s = str(value).strip().lower()
    # Split numeric prefix from unit
    m = re.match(r"^([0-9]+(?:\.[0-9]+)?)\s*(days?|weeks?|months?|d|w|m)\b", s)
    if not m:
        # Non-empty but unparseable — signal to qa_lint
        _LAST_LEAD_TIME_PARSE["ok"] = False
        return False
    n = float(m.group(1))
    unit = m.group(2)
    if unit.startswith("d"):
        days = n
    elif unit.startswith("w"):
        days = n * 7
    elif unit.startswith("m"):
        days = n * 30  # approximate month
    else:
        _LAST_LEAD_TIME_PARSE["ok"] = False
        return False
    _LAST_LEAD_TIME_PARSE["parsed_days"] = days
    _LAST_LEAD_TIME_PARSE["ok"] = True
    return days < 180


# Populated by _lead_time_under_six_months() each call; qa_lint reads to emit
# a WARN when the intake had a non-empty but unparseable value.
_LAST_LEAD_TIME_PARSE: dict = {"input": None, "parsed_days": None, "ok": True}


def _has_super_factory_initiative(data: dict) -> bool:
    """Return True if Super-Factory Initiative is in scope via either
    `custom.definitions_include` or a verbatim `custom.definitions[].key`.
    """
    custom = data.get("custom", {}) or {}
    if "super_factory_initiative" in (custom.get("definitions_include") or []):
        return True
    for item in custom.get("definitions", []) or []:
        k = (item.get("key") or "").lower()
        if "super-factory" in k or "super_factory" in k:
            return True
    return False


def _load_common_defined_terms() -> dict:
    """v3.7.0 — parse `_shared/loi-common-defined-terms.md` into
    `{key: {name, text}}`.

    Library format (simplified markdown parser):
        ## Super-Factory Initiative
        - **Key:** `super_factory_initiative`
        ...
        > **"Super-Factory Initiative"** means the Provider's programme...

    Only `Key`, term name (from heading), and the first blockquote after the
    key line are extracted. Returns empty dict if the library file is missing.
    """
    candidates = [
        os.path.join(os.path.dirname(__file__), "..", "..", "_shared",
                     "loi-common-defined-terms.md"),
        os.path.join(os.path.dirname(__file__), "..", "..", "..", "_shared",
                     "loi-common-defined-terms.md"),
    ]
    for path in candidates:
        if os.path.exists(path):
            break
    else:
        return {}

    entries: dict = {}
    with open(path, encoding="utf-8") as f:
        text = f.read()

    current_name = None
    current_key = None
    collecting_body = False
    body_lines: list = []

    for line in text.splitlines():
        stripped = line.strip()

        if stripped.startswith("## ") and not stripped.startswith("## Extending"):
            if current_key and body_lines:
                entries[current_key] = {
                    "name": current_name,
                    "text": " ".join(body_lines).strip(),
                }
            current_name = stripped[3:].strip()
            current_key = None
            collecting_body = False
            body_lines = []
            continue

        m = re.match(r"-\s*\*\*Key:\*\*\s*`([^`]+)`", stripped)
        if m:
            current_key = m.group(1)
            continue

        if stripped.startswith("> "):
            collecting_body = True
            content = stripped[2:]
            content = re.sub(r"\*\*[^*]+\*\*\s*", "", content, count=1)
            content = content.lstrip()
            body_lines.append(content)
        elif collecting_body and stripped == "":
            if body_lines:
                if current_key:
                    entries[current_key] = {
                        "name": current_name,
                        "text": " ".join(body_lines).strip(),
                    }
                body_lines = []
                collecting_body = False

    if current_key and body_lines:
        entries[current_key] = {
            "name": current_name,
            "text": " ".join(body_lines).strip(),
        }

    return entries


class LOI:
    def __init__(self, data: dict):
        self.d = data
        self.t = data["type"]
        # M4: Bespoke deals can override the default agreement_type /
        # subject from top-level YAML keys; templated types ignore them.
        self.agreement_type = (
            data.get("agreement_type") or _AGREEMENT_TYPE_BY_LOI[self.t]
        )
        self.subject = data.get("subject") or _SUBJECT_BY_LOI[self.t]
        self.doc = Document()
        self._setup()
        party_by_type = {
            "Distributor": "Partner",
            "StrategicSupplier": "Supplier",
            "EcosystemPartnership": "Partner",
            "Wholesale": "Customer",
            "EndUser": "Customer",
            # M4 Bespoke: default "Counterparty"; YAML can override via `party_label:`
            "Bespoke": "Counterparty",
        }
        self.party = data.get("party_label") or party_by_type.get(self.t, "Customer")
        # v3.5.2 (brand-name defined term): provider is defined by its
        # short name throughout the body. The v3.5.2 brand rename replaced
        # every prior "the Provider" reference with "Digital Energy".
        self.provider_term = self._derive_provider_term(data)
        # QA linter accumulators (populated during build)
        self.overrides = set(self.d.get("_overrides", []))
        self.override_reason = self.d.get("_override_reason", "")
        self.qa_findings = []

    def _setup(self):
        style = self.doc.styles["Normal"]
        style.font.name = FONT_NAME
        style.font.size = FONT_BODY
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = LINE_SPACING
        # v3.5 scope A'''': derive footer entity from provider.legal_name so
        # BV-signed instruments render "Digital Energy Netherlands B.V."
        # footer and AG-signed instruments render "Digital Energy Group AG"
        # footer. Prior default of "ag" caused BV LOIs to ship with the
        # Swiss parent's entity in the footer — material misidentification.
        entity = self._derive_footer_entity(
            self.d.get("provider", {}).get("legal_name", "")
        )

        for s in self.doc.sections:
            s.page_width = Mm(210)
            s.page_height = Mm(297)
            s.top_margin = Mm(20)
            s.bottom_margin = Mm(35)
            s.left_margin = Mm(25)
            s.right_margin = Mm(20)
            s.different_first_page_header_footer = True

            # Shared header/footer from document-factory
            setup_first_page_header(s)
            setup_cont_header(s, title=self.agreement_type)
            setup_first_footer(s, classification="Confidential", entity=entity)
            setup_cont_footer(s, classification="Confidential", entity=entity)

    # --- Helpers ---

    def g(self, *keys, default=""):
        v = self.d
        for k in keys:
            v = v.get(k, default) if isinstance(v, dict) else default
        return v or default

    def choice(self, k):
        return bool(self.g("choices", k))

    @staticmethod
    def _is_tbc(value):
        """Detect placeholder / unresolved sentinel values."""
        if value is None:
            return True
        s = str(value).strip()
        if not s:
            return True
        return s.upper() in ("[TBC]", "[TO BE CONFIRMED]", "TBC", "TO BE CONFIRMED", "XXXXXXXX")

    @staticmethod
    def _derive_provider_term(data):
        """Return the defined-term short-name used throughout the body.

        v3.5 polish: derive from `provider.short_name` with "Digital Energy"
        fallback — preserves brand when AG signs (short_name still
        "Digital Energy") while supporting future subsidiary/JV instruments
        with different short names without another body-wide rename.

        Extracted as a static helper (matches `_derive_footer_entity`) so
        the derivation is unit-testable without constructing a full
        DocBuilder (which triggers `_setup()` and document-factory image
        loading — brittle in CI across different path layouts).
        """
        prov = (data or {}).get("provider") or {}
        return prov.get("short_name") or "Digital Energy"

    @staticmethod
    def _derive_footer_entity(legal_name):
        """Map provider legal name to document-factory footer entity key.

        v3.5 scope A'''': dedicated helper so the derivation is unit-testable.
        Returns "nl" for Digital Energy Netherlands B.V.; "ag" for Digital
        Energy Group AG; "ag" as safe default for anything unrecognized.
        """
        if not legal_name:
            return "ag"
        name = str(legal_name)
        if "Netherlands" in name or "B.V." in name or " BV" in name:
            return "nl"
        return "ag"

    def _render_placeholder(self, value, context):
        """Return render-appropriate string for a value that may be a placeholder.

        context values:
          - "sig_block_name"   → blank fillable line
          - "sig_block_title"  → blank fillable line
          - "body_clause"      → keep "[TBC]" as visible draft marker
          - default            → return empty string

        v3.5 scope J5 (Jonathan memo): `[TBC]` should never surface literally
        in the sig block of an external-facing draft — renders as a fillable
        blank line per standard legal instrument convention. The helper is
        also used for other render contexts; behaviour is context-scoped.
        """
        is_tbc = self._is_tbc(value)
        if context in ("sig_block_name", "sig_block_title"):
            return "____________________________" if is_tbc else str(value)
        if context == "body_clause":
            return "[TBC]" if is_tbc else str(value)
        return "" if is_tbc else str(value)

    def p(self, text, bold=False, italic=False, size=None, color=None, align=None, space_after=None):
        para = self.doc.add_paragraph()
        if align:
            para.alignment = align
        if space_after is not None:
            para.paragraph_format.space_after = Pt(space_after)
        run = para.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = size or FONT_BODY
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        return para

    def bp(self, label, text):
        """Bold label + normal text in one paragraph."""
        para = self.doc.add_paragraph()
        r1 = para.add_run(label)
        r1.bold = True
        r1.font.name = FONT_NAME
        r1.font.size = FONT_BODY
        r2 = para.add_run(text)
        r2.font.name = FONT_NAME
        r2.font.size = FONT_BODY
        return para

    def h(self, text, level=2):
        heading = self.doc.add_heading(text, level=level)
        for r in heading.runs:
            r.font.name = FONT_NAME
            r.font.color.rgb = NAVY

    def table(self, headers, rows):
        t = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.name = FONT_NAME
                    r.font.size = FONT_BODY
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = t.rows[ri + 1].cells[ci]
                cell.text = str(val)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = FONT_NAME
                        r.font.size = FONT_BODY
        return t

    def line(self):
        self.p("—" * 50, color=GREY, size=Pt(6), space_after=0)

    # --- Sections ---

    def letterhead(self):
        """Render IB-standard cover page via document-factory's add_cover()."""
        agreement_type = self.agreement_type
        subject = self.subject

        # Build Party objects from YAML data
        prov = self.d.get("provider", {})
        provider_party = Party(
            legal_name=prov.get("legal_name", ""),
            address=prov.get("address", ""),
            registration_type="KvK" if prov.get("kvk") else None,
            registration_number=prov.get("kvk"),
            parent=prov.get("parent"),
        )

        cp = self.d.get("counterparty", {})
        counterparty = Party(
            legal_name=cp.get("name", ""),
            address=cp.get("address", ""),
            registration_type=cp.get("reg_type"),
            registration_number=cp.get("reg_number"),
        )

        # Delegate to document-factory's shared cover page
        add_cover(
            self.doc,
            agreement_type=agreement_type,
            subject=subject,
            date_str=self.g("dates", "loi_date"),
            parties=[provider_party, counterparty],
            formality="non_binding",
            classification="Confidential",
        )

    def parties(self):
        """Render Parties Preamble — legal identification block in body.

        v3.5.2 scope A''' (Jonathan memo / user directive): prior documents
        identified parties only on the cover page. Cover page is a
        presentation layer, not a legal identification layer. Proper
        contract structure has parties named in the body of the instrument
        itself. This block appears between the cover and Recital A; its
        job is to make the legal identification of both parties verifiable
        from the body alone, independent of the cover page.

        Defined terms established here propagate to every clause:
          - Provider defined term: "Digital Energy" (brand-name — v3.5.2)
          - Counterparty defined term: "the Customer" / "the Partner" /
            "the Supplier" per type
        """
        prov = self.d.get("provider", {})
        cp = self.d.get("counterparty", {})
        loi_date = self.g("dates", "loi_date") or "____________________________"

        # v3.5.5 spacing fix: prior implementation emitted explicit blank
        # `self.p("")` paragraphs between each block, doubling the vertical
        # gap (blank paragraph height + default 3pt space_after × 2). The
        # preamble now uses `space_after=6` on the intro + party paragraphs
        # (single consistent gap) and a default-spaced closing line. This
        # matches the tighter cadence of the body clauses and keeps the
        # preamble block proportional to the rest of the page.
        self.p(
            f'THIS LETTER OF INTENT (the "LOI") is dated {loi_date} and entered into between:',
            space_after=6,
        )

        # Party 1 — Digital Energy (Provider)
        prov_legal = prov.get("legal_name", "")
        prov_form = prov.get("legal_form") or "a private limited liability company"
        # Ensure form starts with an article — grammatical Parties Preamble
        if prov_form and not prov_form.lower().startswith(("a ", "an ", "the ")):
            prov_form = f"a {prov_form}"
        prov_jur = prov.get("jurisdiction") or "the Netherlands"
        prov_addr = prov.get("address", "")
        prov_reg_type = prov.get("reg_type") or ("KvK" if prov.get("kvk") else "")
        prov_reg_number = prov.get("reg_number") or prov.get("kvk", "")

        party1_frag = f"(1) {prov_legal}, {prov_form} incorporated under the laws of {prov_jur}"
        if prov_addr and not self._is_tbc(prov_addr):
            party1_frag += f", with registered office at {prov_addr}"
        if prov_reg_type and prov_reg_number and not self._is_tbc(prov_reg_number):
            party1_frag += f" and registered with the {prov_reg_type} under number {prov_reg_number}"
        # Brand-name defined term — no "the" prefix; the brand name itself
        # is the defined short-form used throughout the body.
        party1_frag += f' ("{self.provider_term}"); and'
        self.p(party1_frag, space_after=6)

        # Party 2 — Counterparty
        cp_legal = cp.get("name", "")
        cp_jur = cp.get("jurisdiction") or self._render_placeholder(
            cp.get("jurisdiction"), "body_clause"
        )
        cp_addr = cp.get("address", "")
        cp_reg_type = cp.get("reg_type", "")
        cp_reg_number = cp.get("reg_number", "")

        party2_frag = f"(2) {cp_legal}, an entity incorporated under the laws of {cp_jur or '[TBC]'}"
        if cp_addr and not self._is_tbc(cp_addr):
            party2_frag += f", with registered office at {cp_addr}"
        # v3.6.0 bug 7: Party 1 has a KvK fallback; Party 2 previously
        # required BOTH cp_reg_type AND cp_reg_number, silently dropping
        # the registration number when reg_type was absent. Now emit the
        # number whenever it's present (using a generic label when
        # reg_type is unset — operator can override via YAML).
        if cp_reg_number and not self._is_tbc(cp_reg_number):
            if cp_reg_type:
                party2_frag += f" and registered with the {cp_reg_type} under number {cp_reg_number}"
            else:
                party2_frag += f" with company number {cp_reg_number}"
        party2_frag += f' (the "{self.party}").'
        self.p(party2_frag, space_after=6)

        self.p('(each a "Party" and together the "Parties")')

    def recitals(self):
        self.h("Recitals")
        cp = self.g("counterparty", "short")

        # Recital A — v3.4: resolve_recital_a() returns body + type-specific tail.
        self.p(f"(A) {resolve_recital_a(self.d)}")

        # v3.8.0: Recital B is rendered from the typed slot block. The
        # freeform `description` field is removed (R-DEPRECATED-FIELD).
        # Engine concatenates slot values into the canonical sentence;
        # no prose generation. See _shared/counterpart-description-
        # framework.md for the schema and Adams §4.7 for the principle.
        rb = self.d.get("counterparty", {}).get("recital_b") or {}
        if rb:
            from recital_b_vocab import render_recital_b_sentence
            sentence = render_recital_b_sentence(
                rb, party_label=self.party, short_name=cp,
            )
            self.p(sentence)
        else:
            # Should be unreachable post-v3.8.0 because validate() requires
            # `recital_b` for non-Bespoke types; defensive only.
            self.p(f'(B) {cp} (the "{self.party}") [Recital B slot block missing — see counterpart-description-framework.md].')

        if self.t == "Distributor":
            self.p(
                f"(C) The Parties wish to record their mutual interest in establishing a strategic "
                f"partnership under which the {self.party} would collaborate with Digital Energy to deliver "
                f"AI colocation services to end-user customers, on the indicative terms set out below. "
                f'This letter of intent and non-circumvention non-disclosure agreement (the "LOI") '
                f"is subject to the negotiation and execution of a definitive partnership "
                f'agreement or master services agreement (the "MSA").'
            )
            self.p(
                "(D) The Parties will exchange commercially sensitive and competitively valuable "
                "information in connection with this LOI, and agree to the binding confidentiality "
                "and non-circumvention provisions set out in Clauses 6 and 7."
            )
        elif self.t == "Wholesale":
            self.p(
                f"(C) The Parties wish to record their mutual interest in Digital Energy making available, "
                f"and the {self.party} procuring, dedicated AI colocation capacity at Digital Energy's DEC "
                f"facilities, on the indicative terms set out below. This letter of intent and "
                f'non-circumvention non-disclosure agreement (the "LOI") is subject to the '
                f'negotiation and execution of a Master Services Agreement (the "MSA").'
            )
            self.p(
                "(D) The Parties will exchange commercially sensitive information, including "
                "site-specific data, pricing models, and infrastructure specifications, in connection "
                "with this LOI, and agree to the binding confidentiality and non-circumvention "
                "provisions set out in Clauses 6 and 7."
            )
        elif self.t == "StrategicSupplier":
            self.p(
                f"(C) The Parties wish to record their mutual interest in establishing a strategic "
                f"supply partnership under which the {self.party} would contribute to the delivery of "
                f"Digital Energy's DEC platform, on the indicative terms set out below. This letter of "
                f'intent and non-circumvention non-disclosure agreement (the "LOI") is subject '
                f"to the negotiation and execution of a Framework Agreement and statements of "
                f"work for named Provider projects."
            )
            self.p(
                "(D) The Parties will exchange commercially sensitive and competitively valuable "
                "information, including site-specific information about Provider projects and Partner "
                "design and manufacturing methodology, in connection with this LOI, and agree to the "
                "binding confidentiality and non-circumvention provisions set out in Clauses 6 and 7."
            )
        elif self.t == "EcosystemPartnership":
            themes_list = self.d.get("ecosystem", {}).get("collaboration_themes", [])
            themes_str = ", ".join(themes_list) if themes_list else "[COLLABORATION_THEMES]"
            self.p(
                f"(C) The Parties wish to record their mutual interest in a non-commercial ecosystem "
                f"collaboration focused on {themes_str}, on the framework set out in Clauses 3 through 5. "
                f'This letter of intent (the "LOI") is non-commercial in scope and does not '
                f"contemplate any payment, capacity purchase, or commercial commitment between "
                f"the Parties."
            )
            self.p(
                "(D) The Parties will exchange commercially sensitive or non-public information in "
                "connection with research, policy, or programme activities under this LOI, and agree "
                "to the mutual confidentiality provisions set out in Clause 6."
            )
        else:  # EndUser
            self.p(
                f"(C) The Parties wish to record their mutual interest in the {self.party} procuring "
                f"AI compute infrastructure services at Digital Energy's DEC facilities, on the indicative "
                f'terms set out below. This letter of intent (the "LOI") is subject to the '
                f'negotiation and execution of a service agreement (the "MSA").'
            )

    def _custom_definitions(self) -> list:
        """v3.7.0 — resolve `custom.definitions[]` + `custom.definitions_include[]`.

        Returns a list of (label, text) tuples prepended to the standard
        definitions. Preserves operator order. Include-keys resolved from
        `_shared/loi-common-defined-terms.md` when available; unknown keys
        silently skipped (the linter flags them elsewhere).
        """
        custom = self.d.get("custom", {}) or {}
        injected = []

        include_keys = custom.get("definitions_include", []) or []
        if include_keys:
            library = _load_common_defined_terms()
            for key in include_keys:
                entry = library.get(key)
                if entry:
                    label = f'"{entry["name"]}"'
                    injected.append((label, entry["text"]))

        for item in custom.get("definitions", []) or []:
            key = item.get("key", "")
            text = item.get("text", "")
            if key and text:
                label = f'"{key}"' if not key.startswith('"') else key
                injected.append((label, text))

        return injected

    def definitions(self):
        self.h("1. Definitions")
        self.p("1.1 In this LOI, unless the context requires otherwise:")

        defs = []

        # v3.7.0 — custom/include-library definitions injected at top
        defs.extend(self._custom_definitions())

        defs.append(('"Affiliate"', 'means, in relation to a Party, any entity that directly or indirectly controls, is controlled by, or is under common control with that Party, where "control" means the ownership of more than 50% of the voting rights or equivalent ownership interest;'))

        if self.t in ("Distributor", "Wholesale", "StrategicSupplier"):
            ac_text = (
                '"Associated Counterparty" means, in relation to a Site Identifier, any of the following '
                "with whom Digital Energy has a contractual, commercial, or active non-public engagement "
                "in connection with that site: (a) landowners and land lessors; (b) greenhouse operators "
                "and agricultural partners; (c) heat offtake counterparties; (d) energy procurement "
                "counterparties; (e) engineering, procurement, and construction contractors engaged or "
                "under negotiation; and (f) grid connection holders and distribution system operators "
                "(to the extent of non-public engagement)."
            )
            if self.t == "Distributor":
                ac_text += (
                    " (g) government agencies, municipalities, and regulatory bodies with whom the "
                    "Provider has active non-public engagement (excluding general regulatory contacts "
                    "available to any market participant);"
                )
            else:
                ac_text += (
                    " For the avoidance of doubt, government agencies and regulatory bodies are "
                    "excluded unless Digital Energy has made a specific named introduction;"
                )
            defs.append(("", ac_text))

        defs.append(('"Business Day"', "means a day other than a Saturday, Sunday, or public holiday in the Netherlands;"))

        if self.t == "Distributor":
            defs.append(('"Competitor"', 'means any entity whose primary business includes the development, ownership, or operation of colocation facilities for high-density compute workloads, or the provision of AI infrastructure services, in the European Economic Area;'))

        ci_text = (
            '"Confidential Information" means all information (whether technical, commercial, financial, '
            "or otherwise) disclosed by a Party to the other in connection with the Purpose, whether in "
            "writing, orally, electronically, or by inspection, including any information that by its "
            "nature or the circumstances of disclosure would reasonably be understood to be confidential"
        )
        if self.t in ("Distributor", "Wholesale", "StrategicSupplier"):
            ci_text += ", and including metadata, EXIF data, digital artifacts, file names, folder names, and any derivative information"
        ci_text += ". The existence and contents of this LOI are Confidential Information;"
        defs.append(("", ci_text))

        defs.append(('"DEC"', "means a Digital Energy Center: a liquid-cooled AI colocation facility developed and operated by Digital Energy;"))
        # v3.2: "DEC Block" removed from customer-facing definitions. Capacity expressed in MW IT + Designated Sites.
        defs.append(('"Designated Site"', "means any DEC or DEC site designated by Digital Energy for the delivery of capacity to the " + self.party + ", as confirmed in the MSA;"))
        defs.append(('"Financing Party"', "means any bank, financial institution, fund, security trustee, or other entity providing or arranging " + ("debt, mezzanine, or structured " if self.t != "EndUser" else "") + "finance to Digital Energy or any of its Affiliates in connection with the development or operation of any DEC;"))

        if self.t == "StrategicSupplier":
            defs.append((
                '"Framework Agreement"',
                "means the definitive Framework Agreement and accompanying statements of work to be negotiated between the Parties, setting out the commercial and operational framework for the supply relationship;"
            ))

        if self.t == "Distributor":
            defs.append(('"MSA"', "means the definitive Master Services Agreement or partnership agreement to be negotiated between the Parties;"))
        elif self.t == "Wholesale":
            defs.append(('"MSA"', "means the definitive Master Services Agreement to be negotiated between the Parties, incorporating the Sales Order Form, SLA Schedule, and Pricing Framework;"))
        elif self.t == "EndUser":
            defs.append(('"MSA"', "means the definitive service agreement to be negotiated between the Parties;"))
        # StrategicSupplier uses Framework Agreement (above) — no MSA definition.

        if self.t == "Distributor":
            defs.append(('"Protected Business Information" or "PBI"', "means Digital Energy's proprietary methodologies, strategies, and frameworks including: (a) site-sourcing and land-acquisition methodology; (b) energy procurement and grid connection strategy; (c) regulatory and permitting playbook; (d) heat offtake and energy recycling commercial model; (e) financial and operational modelling frameworks; (f) DEC design specifications and engineering standards; and (g) colocation pricing methodology and deal economics;"))

        # SS: PBI only if engineering_integration in strategic purposes
        if self.t == "StrategicSupplier":
            purposes = set(self.d.get("supplier", {}).get("strategic_purposes", []))
            if "engineering_integration" in purposes:
                defs.append(('"Protected Business Information" or "PBI"', "means Digital Energy's proprietary design, engineering, and infrastructure integration methodologies, interface specifications, and technical reference architectures disclosed in connection with the design integration collaboration under Clause 3.7;"))

        if self.t == "Distributor":
            defs.append(('"Purpose"', "means evaluating, negotiating, and progressing toward the execution of an MSA for the Transaction, including all activities under any companion agreements between the Parties (such as a Referral Agreement);"))
        elif self.t == "Wholesale":
            defs.append(('"Purpose"', "means evaluating, negotiating, and progressing toward the execution of an MSA for the provision of AI colocation services;"))
        elif self.t == "StrategicSupplier":
            defs.append(('"Purpose"', "means evaluating, negotiating, and progressing toward the execution of a Framework Agreement and accompanying statements of work for the supply relationship;"))
        else:
            defs.append(('"Purpose"', "means evaluating, negotiating, and progressing toward the execution of an MSA;"))

        if self.t in ("Distributor", "Wholesale", "StrategicSupplier"):
            defs.append(('"Ready-for-Service" or "RFS"', "means the date on which a DEC (or a discrete capacity allocation within a DEC) has been commissioned and is available for the provision of colocation services;"))

        defs.append(('"Representatives"', "means, in relation to a Party, its Affiliates, and its and their respective directors, officers, employees, agents, and professional advisers;"))

        if self.t in ("Distributor", "Wholesale"):
            defs.append(('"Sales Order Form"', "means the document that will form the basis for each binding service order under the MSA" + (", setting out confirmed capacity, pricing, site allocation, and RFS date;" if self.t == "Wholesale" else ";")))

        if self.t == "Wholesale":
            defs.append(('"Services"', "means the AI colocation services to be provided by Digital Energy, comprising power supply and distribution, liquid cooling infrastructure, physical security, building management, and facility operations;"))
        elif self.t == "EndUser":
            defs.append(('"Services"', "means the AI compute infrastructure services to be provided by Digital Energy, the scope of which will depend on the service model selected under Clause 3.1."))

        if self.t in ("Distributor", "Wholesale", "StrategicSupplier"):
            defs.append(('"Site Identifier"', "means any information that identifies or could reasonably be used to identify a specific DEC location, including: address, GPS coordinates, cadastral reference, project codename, photographs, aerial imagery, file names, folder names, metadata, and any information derived from the foregoing;"))

        if self.t == "Distributor":
            defs.append(('"Transaction"', "means the proposed strategic partnership between the Parties as described in Clause 3;"))

        if self.t in ("Distributor", "Wholesale"):
            defs.append(('"Whitespace"', "means designated floor area within a DEC available for the deployment of IT hardware and supporting infrastructure."))

        for label, text in defs:
            if label:
                self.bp(label + " ", text)
            else:
                self.p(text)

    def clause2(self):
        self.h("2. Purpose and Scope")
        if self.t == "Distributor":
            self.p(f"2.1 This LOI records the Parties' mutual interest in establishing a strategic partnership under which the {self.party} would participate in the delivery of AI colocation services to end-user customers through Digital Energy's DEC platform, on the indicative terms set out in Clauses 3 and 4.")
        elif self.t == "Wholesale":
            self.p(f"2.1 This LOI records the Parties' mutual interest in the {self.party} procuring dedicated AI colocation capacity at Digital Energy's DEC facilities, on the indicative terms set out in Clauses 3 and 4.")
        elif self.t == "StrategicSupplier":
            self.p(f"2.1 This LOI records the Parties' mutual interest in establishing a strategic supply partnership under which the {self.party} would contribute to the delivery of Digital Energy's DEC platform, on the indicative terms set out in Clauses 3 and 4.")
        else:
            self.p(f"2.1 This LOI records the Parties' mutual interest in the {self.party} procuring AI compute infrastructure services at Digital Energy's DEC facilities, on the indicative terms set out in Clauses 3 and 4.")

        if self.t in ("Distributor", "Wholesale", "StrategicSupplier"):
            if self.t == "Distributor":
                scoping_phrase = "a framework for further commercial discussion and negotiation toward a definitive MSA"
            elif self.t == "Wholesale":
                scoping_phrase = "the basis for technical scoping and commercial negotiation toward a definitive MSA"
            else:  # StrategicSupplier
                scoping_phrase = "a framework for further technical and commercial discussion toward a definitive Framework Agreement"
            self.p(f"2.2 The Parties intend this LOI to provide {scoping_phrase}. The commercial terms in Clauses 3 and 4 are non-binding expressions of intent.")
            self.p("2.3 The confidentiality, non-circumvention, and general provisions in Clauses 5 through 8 are legally binding and enforceable from the date of execution.")
        else:
            self.p("2.2 The commercial terms in Clauses 3 and 4 are non-binding expressions of intent. The confidentiality and general provisions in Clauses 5 through 7 are legally binding and enforceable from the date of execution.")

    def clause3_ds(self):
        mode = self.d.get("partnership_mode", "combined")
        heading = "3. Partnership and Combined Offering" if mode == "combined" else "3. Partnership and Referral Arrangement"
        self.h(heading)
        comm = self.d.get("commercial", {})
        mode = self.d.get("partnership_mode", "combined")

        # 3.1
        self.bp("3.1 Partnership Overview. ",
                f"{self.g('counterparty', 'short')} provides {comm.get('partner_core_capability', '')}. "
                f"Digital Energy develops and operates liquid-cooled AI colocation facilities with secured "
                f"energy and grid access, liquid cooling, and full energy recovery infrastructure.")
        if mode == "combined":
            self.p(
                f"The Parties intend to combine the {self.party}'s {comm.get('partner_contribution', '')} "
                f"with Digital Energy's AI colocation platform to deliver {comm.get('combined_offering', '')} "
                f"to {comm.get('target_end_users', '')}."
            )
        else:
            self.p(
                f"The Parties intend to establish a referral arrangement under which the {self.party} "
                f"would introduce qualified end-user customers to Digital Energy's DEC platform, leveraging "
                f"the {self.party}'s established client relationships and market presence."
            )

        # 3.2
        if mode == "combined":
            self.bp("3.2 Combined Offering. ", "Under the envisaged partnership:")
            self.p("(a) Digital Energy would supply dedicated AI colocation capacity at its DEC facilities, including power distribution, liquid cooling, physical security, and facility operations;")
            self.p(f"(b) the {self.party} would contribute {comm.get('partner_service_scope', '')}; and")
            self.p(f"(c) the end-user customer would procure a single, integrated solution through the {self.party}, with Digital Energy's DEC platform as the infrastructure foundation.")
            self.p("The precise service boundaries, responsibility matrix, SLA allocation, and commercial terms of the combined offering will be defined in the MSA.")
        else:
            self.bp("3.2 Referral Arrangement. ",
                     f"The {self.party} would identify and introduce qualified end-user customers to the "
                     f"Provider's DEC platform, leveraging the {self.party}'s market presence, customer "
                     f"relationships, and domain credibility. Digital Energy would manage all commercial, "
                     f"technical, and contractual relationships with introduced customers directly. The "
                     f"{self.party}'s ongoing role following an introduction \u2014 including relationship "
                     f"support, account management, or technical liaison \u2014 will be agreed on a "
                     f"case-by-case basis and set out in the MSA or a separate Referral Agreement.")
            self.p("The economic terms of the referral arrangement, including qualifying introduction criteria, fee structure, payment terms, and duration, will be set out in a separate Referral Agreement between the Parties.")

        # 3.3
        territory = comm.get("territory", "")
        segments = comm.get("target_segments", "")
        mw = comm.get("indicative_mw", "")
        self.bp("3.3 Commercial Scope. ", "The Parties intend the partnership to address the following market:")
        self.p(f"(a) Territory: {territory};")
        self.p(f"(b) Target segments: {segments}; and")
        self.p(f"(c) Estimated capacity: The {self.party} has indicated an estimated annual capacity requirement of approximately {mw} MW IT across its end-user customer base. This estimate is non-binding and is provided to inform capacity planning.")

        # 3.4
        if self.choice("pricing"):
            self.bp("3.4 Indicative Economics. ", "Based on preliminary discussions, the Parties envisage the following indicative economic framework:")
            ch = self.d.get("choices", {})
            self.table(
                ["Parameter", "Indicative Terms"],
                [
                    ["Partner capacity rate", f"EUR {ch.get('partner_rate', '')} per kW per month"],
                    ["Commercial structure", ch.get("economics_description", "")],
                    ["Minimum annual commitment", f"{ch.get('min_commitment', '')} MW IT"],
                ]
            )
            self.p("All economic terms are indicative and subject to the terms of the MSA.")
        else:
            self.bp("3.4 Indicative Economics. ",
                     "The economic framework for the partnership will be determined during commercial scoping and set out in the MSA. Digital Energy will issue indicative terms upon agreement of partnership scope, capacity estimates, and contract structure.")

        # 3.5
        if self.choice("exclusivity"):
            scope = self.d.get("choices", {}).get("exclusivity_scope", "")
            self.bp("3.5 Preferred Partner. ",
                     f"Subject to execution of the MSA and achievement of the minimum capacity commitments set out therein, Digital Energy intends to designate the {self.party} as a preferred partner within the following scope: {scope}. This designation means Digital Energy will offer the {self.party} priority participation in opportunities within the defined scope before engaging other channel partners for the same opportunity. This LOI does not create any exclusivity obligation; the terms of any such designation will be agreed in the MSA.")
        else:
            self.bp("3.5 Non-Exclusivity. ", "The partnership contemplated by this LOI is non-exclusive. Both Parties retain the right to enter into similar arrangements with third parties.")

    def clause3_ws(self):
        self.h("3. Indicative Capacity and Commercial Terms")
        comm = self.d.get("commercial", {})
        ch = self.d.get("choices", {})
        mw = comm.get("indicative_mw", "")

        # v3.2: capacity expressed in MW IT + Designated Sites; no "DEC Block" customer-facing.
        self.bp("3.1 Indicative Capacity. ", f"The {self.party} has indicated interest in approximately {mw} MW IT of liquid-cooled AI colocation capacity, to be delivered across one or more Designated Sites. Site configuration, phasing, and delivery milestones will be set out in the MSA.")

        # v3.5 scope J1 (Jonathan memo): Cl. 3.2 parametrized. Default density
        # bumped from "40 kW and above" to "approximately 130 kW and above"
        # matching NVIDIA GB200/GB300 NVL72 reference architectures. Cooling
        # topology explicit. Both overridable via YAML for non-AI workloads.
        rack_density = comm.get("rack_density_kw", "130")
        cooling_topology = comm.get(
            "cooling_topology",
            "direct-to-chip liquid cooling (consistent with NVIDIA GB200 NVL72 and GB300 NVL72 reference architectures, which target approximately 120\u2013140 kW per rack at full configuration)",
        )
        self.bp(
            "3.2 Technical Specification. ",
            f"All DEC facilities are designed for high-density AI compute workloads and are expected to include, at minimum: facility power supply and distribution, {cooling_topology} supporting rack densities of approximately {rack_density} kW and above, rear-door heat exchangers or equivalent where applicable, building management, physical security, and 24/7 facility operations. The exact capacity, rack configuration, power density, and cooling requirements will be determined during the technical scoping phase following this LOI.",
        )

        if self.choice("phasing"):
            self.bp("3.3 Deployment Phasing. ", f"The {self.party} has indicated the following high-level phasing interest:")
            phases = ch.get("phases", [])
            rows = [[f"Phase {i+1}", f"{p.get('mw', '')} MW IT", p.get("timeline", "")] for i, p in enumerate(phases)]
            self.table(["Phase", "Approximate Capacity", "Indicative Timeline"], rows)

        # v3.5 scope J3 (Jonathan memo): Cl. 3.4 Expansion branches on value
        # type. Numeric values render the approximate-MW sentence; empty,
        # [TBC], "to be discussed" and non-numeric strings render a generic
        # forward-expansion clause that doesn't insert broken grammar into
        # the instrument.
        exp = comm.get("expansion_mw", "")
        exp_str = str(exp).strip() if exp is not None else ""
        is_numeric_exp = bool(exp_str) and exp_str.lstrip("~<>= ").split()[0].split(".")[0].isdigit()
        if is_numeric_exp and not self._is_tbc(exp_str):
            self.bp(
                "3.4 Expansion. ",
                f"The {self.party} has expressed interest in future expansion to approximately {exp_str} MW IT, subject to availability, commercial agreement, and the terms of the MSA. Digital Energy will use reasonable endeavours to accommodate expansion requirements within its DEC programme.",
            )
        else:
            self.bp(
                "3.4 Expansion. ",
                f"The {self.party} has expressed interest in future expansion beyond the initial deployment, with scale to be determined following technical scoping and subject to availability, commercial agreement, and the terms of the MSA. Digital Energy will use reasonable endeavours to accommodate expansion requirements within its DEC programme.",
            )

        if self.choice("pricing"):
            self.bp("3.5 Indicative Pricing. ", f"Based on the {self.party}'s indicated capacity and term preferences, Digital Energy's indicative pricing is as follows:")
            self.table(
                ["Parameter", "Indicative Terms"],
                [
                    ["Base colocation fee", f"EUR {ch.get('base_price', '')} per kW per month"],
                    ["Contract term", f"{ch.get('term_years', '')} years"],
                    ["Power", "Metered IT Load (kWh) x PUE x Energy Rate (EUR/kWh) \u2014 billed additionally"],
                    ["Price escalation", "Subject to annual escalation, the mechanism for which will be agreed in the MSA"],
                    ["Indicative RFS", ch.get("target_rfs_date", "")],
                    ["Ramp schedule", ch.get("ramp_schedule", "")],
                ]
            )
            self.p(f"All commercial terms are indicative and subject to the outcome of the technical scoping phase, the {self.party}'s final capacity requirements, contract term, and credit profile.")
        else:
            self.bp("3.5 Indicative Pricing. ", f"Pricing for the Services will be determined following the technical scoping phase and set out in the Sales Order Form. Digital Energy will provide indicative pricing upon completion of the {self.party}'s capacity, density, and term requirements.")

        term = comm.get("min_term", "")
        # v3.2: "minimum commitment term" replaced with "approximately X, indicative only".
        self.bp("3.6 Indicative Term. ", f"The {self.party} anticipates a commitment term of approximately {term}, indicative only and subject to confirmation in the MSA. Actual term may be longer or shorter depending on final commercial terms. The {self.party} acknowledges that Digital Energy intends the MSA to include take-or-pay provisions commensurate with the committed capacity, the terms of which will be negotiated in good faith.")

        self.bp("3.7 Credit Assessment. ", f"Digital Energy will complete a credit assessment of the {self.party} (or the entity that will execute the MSA) as part of the commercial process. The {self.party} agrees to cooperate with such assessment, which may include provision of audited financial statements, credit references, evidence of parent company support, or other financial information as reasonably requested. Where the {self.party} does not hold an investment-grade credit rating (or equivalent), the Parties will discuss appropriate credit support mechanisms, which may include a parent company guarantee, security deposit, or letter of credit.")

        self.bp("3.8 Site Allocation. ", f"Digital Energy is developing DEC facilities across multiple locations. Digital Energy will allocate capacity to the {self.party} based on the DEC(s) best suited to the {self.party}'s requirements, taking into account development readiness, grid availability, RFS timeline, and the {self.party}'s latency and connectivity needs. Site allocation will be confirmed during the technical scoping phase and formalised in the Sales Order Form.")

    def clause3_eu(self):
        self.h("3. Service Requirements")
        comm = self.d.get("commercial", {})
        types = comm.get("service_type", [])

        self.bp("3.1 Service Model. ", f"The {self.party} has indicated interest in the following service model:")

        descs = {
            "bare_metal": (
                "Bare Metal Colocation \u2014 Dedicated, liquid-cooled rack space within a DEC, supporting "
                "densities of 40 kW per rack and above. Digital Energy supplies power distribution, direct "
                "liquid cooling, physical security, and 24/7 facility operations. The Customer deploys and "
                "manages its own hardware and software. The demarcation point is at the rack: everything "
                "below (facility infrastructure) is Digital Energy's responsibility; everything above (compute "
                "hardware, operating system, workloads) is the Customer's. Billed monthly on a per-kW basis "
                "with a committed term."
            ),
            "shared_cloud": (
                "Shared Cloud \u2014 Managed GPU compute capacity hosted on sovereign European infrastructure "
                "at Digital Energy's DEC facilities. The Customer accesses reserved or on-demand compute "
                "resources without procuring, deploying, or managing hardware. Digital Energy or a designated "
                "delivery partner operates the compute platform, including hardware lifecycle, scheduling, "
                "and platform software. The Customer manages workloads, data, and application layers. Billed "
                "on reserved capacity or metered usage, with a committed term or minimum spend."
            ),
            "tokens": (
                "Token-Based GPU Access \u2014 Flexible, on-demand access to GPU compute measured in GPU-hours "
                "or equivalent units. The Customer purchases compute tokens \u2014 pre-paid or pay-as-you-go "
                "\u2014 redeemable against available capacity across Digital Energy's DEC network. No dedicated "
                "hardware allocation or minimum infrastructure commitment. Digital Energy or a designated "
                "delivery partner manages all infrastructure and scheduling. The Customer submits workloads "
                "and consumes capacity as needed. Lowest entry threshold; designed for variable or exploratory "
                "workloads."
            ),
        }
        for st in types:
            if st in descs:
                self.p(descs[st])

        if any(st in ("shared_cloud", "tokens") for st in types):
            self.p("For Shared Cloud and Token-Based models: Digital Energy may deliver these services in partnership with a qualified delivery partner. The specific delivery partner, platform, and commercial terms will be confirmed in the MSA.", italic=True)

        cap = comm.get("indicative_capacity", "")
        self.bp("3.2 Indicative Capacity. ", f"The {self.party} has indicated interest in approximately {cap} of compute capacity. The exact capacity, configuration, and technical requirements will be determined during the technical scoping phase.")

        self.bp("3.3 Technical Requirements. ", f"Digital Energy's DEC facilities support rack densities of 40 kW and above with direct liquid cooling, designed for GPU-intensive training and inference workloads. The {self.party}'s specific requirements \u2014 including GPU type, rack density, network connectivity, and data sovereignty needs \u2014 will be confirmed during technical scoping and formalised in the MSA.")

        if self.choice("pricing"):
            ch = self.d.get("choices", {})
            self.bp("3.4 Indicative Pricing. ", "Based on preliminary discussions, Digital Energy's indicative pricing is as follows:")
            self.table(
                ["Parameter", "Indicative Terms"],
                [
                    ["Service rate", f"EUR {ch.get('base_price', '')} {ch.get('pricing_unit', 'per kW per month')}"],
                    ["Contract term", ch.get("term_years", "")],
                    ["Indicative RFS", ch.get("target_rfs_date", "")],
                ]
            )
            self.p(f"All terms are indicative and subject to the outcome of technical scoping and the {self.party}'s final requirements.")
        else:
            self.bp("3.4 Indicative Pricing. ", f"Pricing will be determined following the technical scoping phase and set out in the MSA. Digital Energy will provide indicative pricing upon agreement of the {self.party}'s capacity, density, and service model requirements.")

        term = comm.get("min_term", "")
        # v3.2: indicative, not "minimum commitment"
        self.bp("3.5 Indicative Term. ", f"The {self.party} anticipates a commitment term of approximately {term}, indicative only and subject to confirmation in the MSA.")

    def clause3_ss(self):
        """v3.3: Strategic Supplier Cl. 3 — Partnership Scope.

        3.1 Capability Contribution (ALWAYS). Then 3.2-3.8 are purpose-driven:
        - capacity_lock_in       -> 3.2 Capacity Reservation, 3.3 Lead-Time Targets
        - pricing_volume         -> 3.4 Pricing Framework, 3.5 Volume Tiers
        - supply_chain_de_risking-> 3.6 Dual-Source + Continuity
        - engineering_integration-> 3.7 Design Integration + IP Allocation
        - pipeline_visibility    -> 3.8 Preferred-Supplier / ROFR
        """
        self.h("3. Partnership Scope")
        supplier = self.d.get("supplier", {})
        purposes = set(supplier.get("strategic_purposes", []))

        # 3.1 Capability Contribution (ALWAYS)
        self.bp(
            "3.1 Capability Contribution. ",
            "Under the envisaged partnership:"
        )
        self.p(
            "(a) Digital Energy would contribute the site, energy procurement, grid "
            "connection, land and regulatory compliance, long-term operational "
            "responsibility, and integration into Digital Energy's broader DEC "
            "platform for each Designated Site; and"
        )
        # v3.6.0 bug 6: strip trailing period from core_capability to
        # avoid double-period when operator ends YAML value with '.'.
        _cc = supplier.get("core_capability", "[CAPABILITY DESCRIPTION TO BE CONFIRMED]")
        if isinstance(_cc, str):
            _cc = _cc.rstrip().rstrip(".")
        self.p(
            f"(b) the {self.party} would contribute {_cc}."
        )
        self.p(
            "The precise scope boundaries, responsibility matrix, and commercial "
            "terms will be defined in the Framework Agreement and accompanying "
            "statements of work."
        )

        # 3.2 Capacity Reservation — IF capacity_lock_in
        if "capacity_lock_in" in purposes:
            lead_time = supplier.get("lead_time_target", "[LEAD TIME TO BE CONFIRMED]")
            self.bp(
                "3.2 Capacity Reservation. ",
                f"Digital Energy has indicated a projected demand across its active "
                f"development pipeline. Subject to commercial agreement, the "
                f"{self.party} will reserve capacity in its manufacturing or "
                f"service-delivery plan to support Digital Energy's indicative "
                f"pipeline, with specific volumes and delivery windows set out in "
                f"the Framework Agreement."
            )
            # 3.3 Lead-Time Targets — same trigger
            self.bp(
                "3.3 Lead-Time Targets. ",
                f"The {self.party} targets {lead_time}. The Framework Agreement "
                f"will set the binding lead-time commitments, performance "
                f"measurement, and remedies for delay."
            )

        # 3.4 Pricing Framework — IF pricing_volume
        if "pricing_volume" in purposes:
            volume = supplier.get("volume_indicative", "[VOLUME TO BE CONFIRMED]")
            self.bp(
                "3.4 Pricing Framework. ",
                "The pricing framework will be structured as cost-plus, unit "
                "economics, or volume-tier model as appropriate to the supply "
                "category. Specific pricing, volume thresholds, and escalation "
                "mechanisms will be set out in the Framework Agreement."
            )
            # 3.5 Volume Tiers — same trigger
            self.bp(
                "3.5 Volume Tiers. ",
                f"Indicative volume contemplated: {volume}. The Framework "
                f"Agreement will set out final volume tiers and the commercial "
                f"principles applicable at each tier. All volumes and pricing "
                f"in this LOI are non-binding."
            )

        # 3.6 Dual-Source and Continuity — IF supply_chain_de_risking
        if "supply_chain_de_risking" in purposes:
            self.bp(
                "3.6 Dual-Source and Continuity Commitments. ",
                f"The {self.party} acknowledges Digital Energy's requirement for "
                f"supply-chain resilience. The Parties will discuss, and document "
                f"in the Framework Agreement: (a) component substitution and "
                f"second-source provisions; (b) continuity commitments, including "
                f"obligations to provide advance notice of discontinuation and to "
                f"support orderly transition; and (c) delivery service-level "
                f"commitments and remedies for failure to meet them."
            )

        # 3.7 Design Integration and IP Allocation — IF engineering_integration
        if "engineering_integration" in purposes:
            joint_ip = self.g("choices", "joint_ip", default="background").lower()
            if joint_ip == "none":
                ip_clause = (
                    "no foreground IP is contemplated as a deliverable of the "
                    "collaboration, and any IP developed will remain with the "
                    "Party that created it;"
                )
            elif joint_ip == "foreground":
                ip_clause = (
                    "foreground IP developed jointly in the course of the "
                    "collaboration will be jointly owned, with cross-licences on "
                    "terms to be agreed in the Framework Agreement;"
                )
            else:  # background (default)
                ip_clause = (
                    "each Party retains all right, title, and interest in its "
                    "background IP; any foreground IP developed jointly will be "
                    "allocated per the principles set out in the Framework Agreement;"
                )
            self.bp(
                "3.7 Design Integration and IP Allocation. ",
                f"The Parties will collaborate on design integration across the "
                f"specified interfaces between Digital Energy's DEC platform and "
                f"the {self.party}'s contribution. IP allocation: (a) {ip_clause} "
                f"(b) no Party assigns pre-existing IP by reason of this LOI or "
                f"any discussions under it; and (c) the full IP framework will be "
                f"set out in the Framework Agreement."
            )

        # 3.8 Preferred-Supplier / ROFR — IF pipeline_visibility
        # v3.7.0: supplier.rofr block parameterizes site_scope, response_window,
        # lock_out_style (alignment|sole_discretion|hard_minimum|milestone),
        # continues_on_remaining. Backward-compatible default matches pre-v3.7.0.
        if "pipeline_visibility" in purposes:
            rofr_cfg = self.d.get("supplier", {}).get("rofr") or {}
            site_scope = rofr_cfg.get("site_scope", "Digital Energy's active development pipeline")
            response_window = rofr_cfg.get("response_window", "20 Business Days")
            lock_out_style = rofr_cfg.get("lock_out_style", "sole_discretion")
            continues_on_remaining = rofr_cfg.get("continues_on_remaining", True)

            rofr_lines = [
                f"Subject to commercial agreement in the Framework Agreement, the "
                f"Provider intends to grant the {self.party} a right of first "
                f"refusal on the procurement of the supply scope set out in "
                f"Clause 3.1(b) across {site_scope}. "
                f"The right of first refusal requires the {self.party} to submit "
                f"a compliant proposal within {response_window} of Digital Energy's "
                f"invitation."
            ]

            if lock_out_style == "alignment":
                rofr_lines.append(
                    " A 'compliant proposal' is one that addresses the Provider's "
                    "site-specific technical, commercial, and timeline requirements "
                    "as communicated to the Supplier. The Parties recognise that a "
                    "Designated Site award depends on alignment between the "
                    f"Supplier's proposal and those requirements. Where the Parties, "
                    f"following good-faith dialogue within the {response_window} "
                    "response window (or such extended period as the Parties may "
                    "agree), are unable to reach alignment for a particular "
                    "Designated Site, the Provider may engage an alternative "
                    "supplier for that site without further obligation under this "
                    "LOI in respect of that site"
                )
                if continues_on_remaining:
                    rofr_lines.append(
                        "; the Supplier's right of first refusal shall continue to "
                        "apply to any remaining Designated Sites."
                    )
                else:
                    rofr_lines.append(".")
            elif lock_out_style == "hard_minimum":
                rofr_lines.append(
                    " The Provider commits to award to the Supplier not fewer than "
                    "one Designated Site during the term of this LOI, subject only "
                    "to the Supplier's submission of a compliant proposal for such "
                    "site."
                )
            elif lock_out_style == "milestone":
                rofr_lines.append(
                    " The right of first refusal for each Designated Site is "
                    "conditional on the Supplier's timely delivery of the "
                    "pre-commitment deliverables set out in the Framework "
                    "Agreement. Failure to deliver any such milestone entitles "
                    "the Provider to engage an alternative supplier for the "
                    "affected site."
                )
            # sole_discretion is the default; no additional text needed.

            rofr_lines.append(
                " This LOI does not create any binding right of first "
                "refusal; the terms will be set out in the Framework Agreement."
            )

            self.bp(
                "3.8 Preferred-Supplier and Right of First Refusal. ",
                "".join(rofr_lines),
            )

        # v3.7.0 + v3.7.1: optional SS sub-clauses §3.9+. Use a counter so
        # enabling only a subset still renders contiguous numbering (avoids
        # the §3.9 → §3.11 gap when referral_rider=false + joint_stocking on).
        _ss_opt_num = 9

        # v3.7.0: supplier.referral_rider — bidirectional flow sentence appended
        # to §3.1 equivalent. InfraPartners §3.2 pattern, Armada referral rider.
        if self.d.get("supplier", {}).get("referral_rider"):
            self.bp(
                f"3.{_ss_opt_num} Mutual Referral Rider. ",
                f"In addition to the supply arrangements above, the Parties "
                f"acknowledge a bidirectional referral interest: Digital Energy may "
                f"introduce end-customers requiring the Supplier's products to the "
                f"Supplier, and the Supplier may introduce colocation customers "
                f"requiring European AI capacity to Digital Energy. Commercial "
                f"terms for such referrals (including any fee or margin sharing) "
                f"shall be set out in the Framework Agreement. Mutual "
                f"non-circumvention obligations in Clause 7 apply to both "
                f"referral directions."
            )
            _ss_opt_num += 1

        # v3.7.1: Joint Stocking Programme — parametric on supplier.lead_time_target < 6 months
        # (InfraPartners §5.6). Decouples manufacturing from site-specific demand to hit
        # aggressive RFS targets. Mentions Super-Factory Initiative when that defined
        # term is present in the intake (custom.definitions_include).
        lt = self.d.get("supplier", {}).get("lead_time_target", "")
        if _lead_time_under_six_months(lt):
            sfi_tail = (
                " through the Super-Factory Initiative"
                if _has_super_factory_initiative(self.d)
                else ""
            )
            self.bp(
                f"3.{_ss_opt_num} Joint Stocking Programme. ",
                f"The Parties acknowledge that a 90-day Ready-for-Service requires "
                f"decoupling manufacturing from site-specific demand. The Parties "
                f"intend to develop a joint stocking programme under which the "
                f"Supplier maintains strategic inventory of modular units aligned "
                f"to the Provider's active development pipeline, available to support "
                f"accelerated deployment to the Provider's colocation customers and "
                f"to the Supplier's customers seeking European AI capacity"
                f"{sfi_tail}."
            )
            _ss_opt_num += 1

        # v3.7.1: Reference and Co-Marketing (InfraPartners §5.7). Parameterizes
        # the 6 sub-clauses (a–f) by:
        #   - framing: multi_supplier | preferred | exclusive
        #   - logo_use: yes | no | per_event_approval
        #   - site_naming_approval_sla: e.g. "2 BD"
        #   - press_at_loi: none | joint | unilateral_allowed
        cm = self.d.get("supplier", {}).get("co_marketing") or {}
        if cm:
            framing = cm.get("framing", "multi_supplier")
            logo_use = cm.get("logo_use", "per_event_approval")
            sla = cm.get("site_naming_approval_sla", "2 Business Days")
            press = cm.get("press_at_loi", "none")

            self.bp(
                f"3.{_ss_opt_num} Reference and Co-Marketing. ",
                "The Parties intend to support reciprocal reference and co-marketing "
                "activity under the following principles:",
            )
            _ss_opt_num += 1
            # (a) framing — multi_supplier is the "Provider's programme" guard;
            # preferred permits "preferred supplier" wording; exclusive permits
            # sole-supplier wording (rare; requires prior Legal sign-off).
            if framing == "multi_supplier":
                self.p(
                    f"(a) the {self.party} may reference its role as a strategic "
                    f"modular infrastructure supplier to the Provider's programme, "
                    f"provided such references avoid language implying exclusivity "
                    f"or co-ownership of the programme;"
                )
            elif framing == "preferred":
                self.p(
                    f"(a) the {self.party} may reference its role as a preferred "
                    f"supplier for the Provider's programme, provided exclusivity "
                    f"language is reserved for confirmation in the Framework Agreement;"
                )
            else:  # exclusive
                self.p(
                    f"(a) the {self.party} may reference its role as sole named "
                    f"supplier for the Provider's programme, consistent with the "
                    f"exclusivity commitments to be formalised in the Framework Agreement;"
                )
            # (b) DE showcase of IP design
            self.p(
                f"(b) the Provider may showcase, to current and prospective "
                f"colocation customers, non-confidential attributes of the "
                f"{self.party}'s design, specifications, and performance data "
                f"as disclosed to the Provider under this LOI's confidentiality "
                f"terms;"
            )
            # (c) site-naming pre-approval
            self.p(
                f"(c) the {self.party} may describe the Provider's DEC site envelope "
                f"(power, cooling, and dimensions) to its own prospects without naming "
                f"specific sites; specific-site references require the Provider's "
                f"written consent, which will not be unreasonably withheld and on a "
                f"{sla} acknowledgement service level;"
            )
            # (d) logo use
            if logo_use == "yes":
                self.p(
                    f"(d) each Party grants the other a limited, revocable right to "
                    f"use its name and logo on website, investor materials, and sales "
                    f"collateral, subject to the other Party's brand guidelines as "
                    f"provided from time to time;"
                )
            elif logo_use == "per_event_approval":
                self.p(
                    f"(d) each Party may use the other's name and logo on specific "
                    f"marketing materials subject to prior written approval on a "
                    f"per-event basis, with brand guidelines supplied on request;"
                )
            else:  # no
                self.p(
                    f"(d) neither Party grants the other any right to use its name "
                    f"or logo outside the confidential context of this LOI; marketing "
                    f"rights will be negotiated in the Framework Agreement;"
                )
            # (e) press at LOI
            if press == "none":
                self.p(
                    f"(e) the Parties shall not issue any joint press release or "
                    f"public announcement in connection with this LOI; the first "
                    f"joint announcement is deferred to the award of a Designated "
                    f"Site under the Framework Agreement; and"
                )
            elif press == "joint":
                self.p(
                    f"(e) the Parties intend to issue a joint press release at the "
                    f"execution of this LOI, the text of which shall be agreed in "
                    f"writing by both Parties prior to release; and"
                )
            else:  # unilateral_allowed
                self.p(
                    f"(e) either Party may issue a unilateral press or investor "
                    f"communication describing this LOI at a high level, provided "
                    f"the text omits confidential commercial detail and is shared "
                    f"in advance with the other Party; and"
                )
            # (f) commercial routing
            self.p(
                f"(f) commercial discussions initiated with any prospect introduced "
                f"by the {self.party} to the Provider shall route through the "
                f"Provider's sales process, consistent with the non-circumvention "
                f"obligations in Clause 7."
            )

    def clause4_ss(self):
        """v3.3: Strategic Supplier Cl. 4 — Pipeline Engagement.

        4.2 Contractual Sequence, 4.4 Change of Control, 4.6 Implementation
        Roadmap are ALWAYS. Others are purpose-driven.
        """
        self.h("4. Pipeline Engagement")
        purposes = set(self.d.get("supplier", {}).get("strategic_purposes", []))

        # 4.1 Project Introduction Process — IF pipeline_visibility
        if "pipeline_visibility" in purposes:
            self.bp(
                "4.1 Project Introduction Process. ",
                f"Following execution of this LOI, Digital Energy intends to "
                f"maintain a pipeline register of Designated Sites. The "
                f"{self.party} will be invited to participate in named project "
                f"opportunities on the cadence and criteria set out in the "
                f"Framework Agreement."
            )

        # 4.2 Contractual Sequence (ALWAYS) — mirrors WS Cl. 4.2 pattern
        self.bp(
            "4.2 Contractual Sequence. ",
            "The Parties acknowledge the intended progression of commercial instruments:"
        )
        self.p("(a) this LOI, setting out non-binding commercial intent and binding confidentiality and non-circumvention;")
        self.p("(b) a Framework Agreement, setting out the definitive commercial and operational framework for the supply relationship;")
        self.p("(c) one or more Statements of Work for named Provider projects, each executed under the Framework Agreement; and")
        # v3.7.0: when include_schedule=false, scrub the "schedules, or
        # operational annexes" phrasing so §4.2(d) aligns with the omitted
        # Schedule 1. InfraPartners §4.8 pattern.
        if self.d.get("choices", {}).get("include_schedule", True) is False:
            self.p("(d) site-specific deliverables executed under each Statement of Work.")
        else:
            self.p("(d) site-specific deliverables, schedules, or operational annexes executed under each Statement of Work.")
        # v3.6.0 bug 4 (SS): meta-commentary trailer removed — explains
        # the LOI rather than creating obligation; R-22 class.

        # 4.3 Joint-Development Governance — IF engineering_integration
        if "engineering_integration" in purposes:
            self.bp(
                "4.3 Joint-Development Governance. ",
                "The Parties intend to establish a joint design and engineering "
                "working group to coordinate interfaces, design reviews, and "
                "lifecycle decisions. The governance structure, decision rights, "
                "and escalation paths will be set out in the Framework Agreement."
            )

        # 4.4 Change of Control (ALWAYS)
        self.bp(
            "4.4 Change of Control. ",
            f"If, during the term of this LOI, a direct competitor of the "
            f"Provider (as determined by reference to its primary business being "
            f"the development, ownership, or operation of colocation facilities "
            f"for high-density compute workloads) acquires Control of the "
            f"{self.party}, Digital Energy may terminate this LOI by written notice "
            f"with immediate effect. Upon such termination, the confidentiality "
            f"and non-circumvention obligations in Clauses 6 and 7 shall "
            f"continue for their stated survival periods."
        )

        # 4.5 Exclusivity — IF capacity_lock_in AND choices.exclusivity
        if "capacity_lock_in" in purposes and self.choice("exclusivity"):
            self.bp(
                "4.5 Exclusivity. ",
                "Subject to execution of the Framework Agreement and achievement "
                "of the reservation commitments set out therein, the Parties "
                "intend to agree a scope of mutual exclusivity. This LOI does "
                "not create any exclusivity; the terms will be agreed in the "
                "Framework Agreement."
            )

        # 4.6 Implementation Roadmap (ALWAYS)
        self.bp(
            "4.6 Implementation Roadmap. ",
            "Following execution of this LOI, the Parties intend to proceed as follows:"
        )
        self.p("(a) Technical and commercial scoping (target: 30 days post-LOI) — detailed discovery on scope, capability, lead times, pricing framework, and IP allocation.")
        self.p("(b) Draft Framework Agreement (target: 60 days post-LOI) — Digital Energy will issue a draft Framework Agreement incorporating the agreed terms.")
        self.p("(c) Framework Agreement negotiation and execution (target: 90 days post-LOI) — the Parties will negotiate and execute the Framework Agreement.")
        self.p("These timelines are indicative and non-binding.")

    def clause4(self):
        if self.t == "Distributor":
            self.h("4. Relationship Structure and Protection")
            pbi = self.g("protection", "pbi_survival", default="10 years")
            self.bp("4.1 Protected Business Information. ",
                     f"The {self.party} acknowledges that Digital Energy's competitive position depends on the confidentiality of its Protected Business Information as defined in Clause 1. The {self.party} agrees that, for a period of {pbi} from the date of this LOI, it shall not directly or indirectly use any PBI to replicate, reverse-engineer, or independently develop any material element of Digital Energy's business model, site-sourcing methodology, energy procurement strategy, or commercial framework, whether for itself or for any third party. This obligation survives termination or expiry of this LOI and of any MSA.")
            self.bp("4.2 Change of Control. ",
                     f"If, during the term of this LOI, a Competitor acquires Control (as defined in Clause 1.1, \u201cAffiliate\u201d) of the {self.party}, Digital Energy may terminate this LOI by written notice with immediate effect. Upon such termination, the confidentiality and non-circumvention obligations in Clauses 6 and 7 shall continue for their stated survival periods.")
            self.bp("4.3 Associated Counterparties. ",
                     f"The {self.party} acknowledges that Digital Energy maintains relationships with Associated Counterparties at each DEC site. The sharing of any Site Identifier by Digital Energy shall constitute a deemed introduction for the purposes of Clause 7, on the terms set out in Clause 7.3 (Deemed Introduction by Category, including the named-list backstop).")
            self.bp("4.4 Governance. ", "The Parties intend to establish a joint steering committee or equivalent governance mechanism, the terms of which will be set out in the MSA.")
            self.bp("4.5 Implementation Roadmap. ", "Following execution of this LOI, the Parties intend to proceed as follows:")
            self.p("(a) Commercial scoping (target: 30 days post-LOI) \u2014 The Parties will define the partnership scope, target customer segments, capacity estimates, and economic framework.", bold=False)
            self.p("(b) Draft Partnership Agreement (target: 60 days post-LOI) \u2014 Digital Energy will issue a draft MSA or partnership agreement incorporating the agreed commercial terms.")
            self.p("(c) MSA negotiation and execution (target: 90 days post-LOI) \u2014 The Parties will negotiate and execute the MSA.")
            self.p("These timelines are indicative and non-binding.")

        elif self.t == "Wholesale":
            self.h("4. Relationship Structure and Next Steps")
            self.bp("4.1 MSA Structure. ", "The Parties intend to execute a Master Services Agreement incorporating: (a) a Pricing Framework setting out the commercial terms for the Services; (b) Sales Order Forms for each capacity commitment; and (c) an SLA Schedule defining availability, performance, and remediation commitments. The MSA will be the sole binding commercial agreement between the Parties.")
            # v3.2: Cl. 4.2 arrows replaced with institutional prose + numbered list.
            self.bp("4.2 Contractual Sequence. ", "The Parties acknowledge the intended progression of commercial instruments:")
            self.p("(a) this LOI, setting out non-binding commercial intent;")
            self.p("(b) a Sales Order Form or equivalent non-binding capacity and pricing summary, prepared during technical scoping for planning purposes;")
            self.p("(c) the Master Services Agreement (MSA), containing definitive commercial terms; and")
            # v3.7.0: include_schedule-aware phrasing (mirrors clause4_ss §4.2(d))
            if self.d.get("choices", {}).get("include_schedule", True) is False:
                self.p("(d) site-specific deliverables executed under the MSA.")
            else:
                self.p("(d) site-specific deliverables, schedules, or operational annexes executed under the MSA.")
            # v3.6.0 bug 4 (WS): meta-commentary trailer removed.
            self.bp("4.3 Direct Agreement Willingness. ", f"The {self.party} confirms its willingness, subject to commercially reasonable terms, to enter into a direct agreement with Digital Energy's Financing Parties if requested under Clause 5.3. The {self.party}'s cooperation in this regard materially supports delivery of the committed capacity on the indicative timeline.")
            self.bp("4.4 Expansion and Priority. ", f"Digital Energy will offer the {self.party} priority access to additional capacity within the DEC(s) allocated to the {self.party}, subject to availability. Expansion terms will be governed by the MSA.")
            self.bp("4.5 Implementation Roadmap. ", "Following execution of this LOI, the Parties intend to proceed as follows:")
            self.p("(a) Technical scoping (target: 30 days post-LOI) \u2014 Detailed technical discovery to determine exact capacity, rack layout, power distribution, cooling specifications, and connectivity requirements.")
            self.p(f"(b) Credit assessment (target: 30 days post-LOI, concurrent with technical scoping) \u2014 Digital Energy will complete a credit assessment of the {self.party} or the entity that will execute the MSA.")
            self.p("(c) Sales Order Form (target: 60 days post-LOI) \u2014 Upon completion of technical scoping, Digital Energy will issue a Sales Order Form setting out confirmed commercial terms, facility specifications, and service level framework.")
            self.p("(d) MSA negotiation and execution (target: 90 days post-LOI) \u2014 The Parties will negotiate and execute the MSA.")
            self.p("These timelines are indicative and non-binding.")
            self.bp("4.6 No Capacity Reservation. ", "This LOI does not reserve capacity at any DEC. Capacity allocation is on a first-come, first-served basis and will only be confirmed upon execution of the MSA.")

        else:  # EndUser
            self.h("4. Next Steps")
            self.p("4.1 Following execution of this LOI, the Parties intend to proceed as follows:")
            self.p("(a) Technical scoping (target: 30 days post-LOI) \u2014 Detailed discovery to determine capacity, configuration, connectivity, and service model requirements.")
            self.p("(b) Commercial proposal (target: 60 days post-LOI) \u2014 Digital Energy will issue a commercial proposal setting out confirmed pricing, service scope, and SLA framework.")
            self.p("(c) MSA negotiation and execution (target: 90 days post-LOI) \u2014 The Parties will negotiate and execute the MSA.")
            self.p("These timelines are indicative and non-binding.")
            self.p("4.2 This LOI does not reserve capacity at any DEC. Capacity allocation will be confirmed upon execution of the MSA.")

    def clause5(self):
        """Cl. 5 for EU / DS / WS — Project Finance and Assignment.
        v3.4: meta-commentary stripped from 5.1. Strategic Supplier uses
        clause5_ss() instead (supply-side bankability, not revenue-bankability).
        """
        self.h("5. Project Finance and Assignment (BINDING)")
        self.bp(
            "5.1 Project Finance Context. ",
            "Digital Energy is developing the DEC programme under a combination "
            "of equity investment and non-recourse project finance. This LOI "
            "is binding in Clauses 5, 6, 7, and 8 to support that financing "
            "structure."
        )
        self.bp("5.2 Assignment. ", f"Neither Party may assign its rights or obligations under this LOI without the prior written consent of the other Party, except that:")
        self.p("(a) Digital Energy may assign this LOI, or any rights under it, to any Financing Party or security trustee as security for project finance obligations; and")
        self.p(f"(b) Digital Energy may assign this LOI to any Affiliate or special-purpose vehicle within its corporate group without the {self.party}'s consent, provided Digital Energy remains liable for the performance of the assignee's obligations.")
        self.bp("5.3 Lender Acknowledgment. ", f"The {self.party} acknowledges and agrees that, upon Digital Energy's written request, the {self.party} shall negotiate in good faith and execute a direct agreement (or lender acknowledgment letter) with Digital Energy's Financing Party within 30 Business Days of such request. Such direct agreement shall be limited to the following customary items only: (a) notice of Digital Energy default; (b) cure periods in favour of the Financing Party; (c) step-in and substitution rights, limited to the performance of Digital Energy's obligations under the MSA; and (d) information rights enabling the Financing Party to monitor the commercial relationship. No direct agreement shall, without the {self.party}'s separate written consent, alter, amend, or worsen the {self.party}'s pricing, service levels, liability profile, indemnity scope, or payment obligations under the MSA. The terms of any such direct agreement shall be commercially reasonable and consistent with market practice for project finance transactions.")

    def clause5_ss(self):
        """v3.4: Cl. 5 for Strategic Supplier — Supply Chain and Delivery Commitment.

        A supplier is not a revenue counterparty. The bankability signal DE
        cares about from a supplier is SUPPLY CERTAINTY and DELIVERY
        COMMITMENT, not contracted revenue. This clause reshapes Cl. 5 to
        that frame, keeping the Assignment and Financing Continuity
        Acknowledgment sub-clauses structurally parallel to the commercial
        Cl. 5 so a lender reading across the five LOI types sees consistent
        architecture.
        """
        self.h("5. Supply Chain and Delivery Commitment (BINDING)")
        self.bp(
            "5.1 Delivery Intent. ",
            f"The {self.party} confirms its intent to reserve manufacturing "
            f"and service-delivery capacity to support Digital Energy's active "
            f"development pipeline, on the terms to be agreed in the Framework "
            f"Agreement and accompanying statements of work. This LOI is "
            f"binding in Clauses 5, 6, 7, and 8 to support that supply "
            f"relationship."
        )
        self.bp(
            "5.2 Assignment. ",
            f"Neither Party may assign its rights or obligations under this "
            f"LOI without the prior written consent of the other Party (such "
            f"consent not to be unreasonably withheld), except that:"
        )
        self.p("(a) Digital Energy may assign this LOI, or any rights under it, to any Financing Party or security trustee as security for project finance obligations; and")
        self.p(f"(b) Digital Energy may assign this LOI to any Affiliate or special-purpose vehicle within its corporate group without the {self.party}'s consent, provided Digital Energy remains liable for the performance of the assignee's obligations.")
        self.bp(
            "5.3 Financing Continuity Acknowledgment. ",
            f"The {self.party} acknowledges that Digital Energy's DEC programme "
            f"is financed on a non-recourse, per-site basis, and that supply "
            f"continuity from qualified suppliers materially supports that "
            f"financing. Upon Digital Energy's written request, the {self.party} "
            f"shall negotiate in good faith with Digital Energy's Financing "
            f"Party to confirm supply arrangements on financed projects within "
            f"30 Business Days of such request. Such direct agreement shall be "
            f"limited to the following customary items only: (a) notice of "
            f"Digital Energy default; (b) cure periods in favour of the "
            f"Financing Party; (c) step-in and substitution rights, limited to "
            f"the performance of Digital Energy's obligations under the "
            f"Framework Agreement; and (d) information rights enabling the "
            f"Financing Party to monitor the supply relationship. No direct "
            f"agreement shall, without the {self.party}'s separate written "
            f"consent, alter, amend, or worsen the {self.party}'s pricing, "
            f"service levels, liability profile, indemnity scope, or payment "
            f"obligations under the Framework Agreement. The terms of any such "
            f"direct agreement shall be commercially reasonable and consistent "
            f"with market practice for project finance supply arrangements."
        )

    def clause6(self):
        self.h(f"6. Confidentiality {'and Non-Disclosure ' if self.t != 'EndUser' else ''}(BINDING)")
        existing = self.choice("existing_nda")

        if existing:
            nda_date = self.g("choices", "nda_date")
            transaction_suffix = " and the Transaction." if self.t == "Distributor" else (" and the proposed transaction." if self.t == "Wholesale" else ".")
            self.p(f'6.1 The Non-Disclosure Agreement between the Parties dated {nda_date} (the "NDA") remains in full force and effect and applies to all information exchanged in connection with this LOI{transaction_suffix}')
            self.p("6.2 The existence and contents of this LOI shall be treated as Confidential Information under the NDA.")
            if self.t != "EndUser":
                self.p("6.3 To the extent of any conflict between the NDA and this LOI, the provisions of this LOI shall prevail.")
                transaction_ref = "the Transaction" if self.t == "Distributor" else "the proposed transaction"
                self.p(f"6.4 Neither Party shall make any public announcement regarding this LOI or {transaction_ref} without the prior written consent of the other Party, except as required by applicable law.")
            else:
                self.p("6.3 Neither Party shall make any public announcement regarding this LOI without the prior written consent of the other Party, except as required by applicable law.")
        else:
            surv = self.g("protection", "confidentiality_survival", default="3 years")
            if self.t == "EndUser":
                # Tier A — 8 clauses
                self.p(f"6.1 Each Party shall keep confidential all Confidential Information received from the other Party and shall use such information solely for the Purpose.")
                self.p("6.2 Each Party may disclose Confidential Information only to its Representatives who have a genuine need to know for the Purpose and who are bound by confidentiality obligations no less restrictive than this Clause 6.")
                self.p("6.3 The obligations in Clauses 6.1 and 6.2 do not apply to information that the receiving Party can demonstrate: (a) is or becomes publicly available through no fault of the receiving Party; (b) was already in the receiving Party's possession without restriction; (c) was independently developed without use of the Confidential Information; or (d) was received from a third party without breach of any confidentiality obligation.")
                self.p("6.4 A Party may disclose Confidential Information to the extent required by applicable law, regulation, or court order, provided that (where legally permitted) the disclosing Party gives the other Party prior written notice and discloses only the minimum required.")
                self.p("6.5 Upon written request or expiry of this LOI, the receiving Party shall promptly return or destroy all Confidential Information and certify such return or destruction in writing. The receiving Party may retain copies required by applicable law or internal compliance policies, subject to continued confidentiality.")
                self.p("6.6 Neither Party shall make any public announcement regarding this LOI without the prior written consent of the other Party, except as required by applicable law.")
                self.p(f"6.7 The obligations in this Clause 6 shall survive for {surv} from the date of termination or expiry of this LOI. Obligations in respect of trade secrets shall continue indefinitely.")
                self.p("6.8 Each Party acknowledges that a breach of this Clause 6 may cause irreparable harm for which damages would not be an adequate remedy. The disclosing Party shall be entitled to seek injunctive or other equitable relief without the need to prove actual loss.")
            else:
                # Tier B — 16 clauses, reshaped as a list so v3.7.0
                # confidentiality_opt_outs can suppress named sub-clauses and
                # the remainder auto-renumbers.
                opt_outs = set(self.d.get("choices", {}).get("confidentiality_opt_outs") or [])
                surv_text_body = (
                    f"The obligations in this Clause 6 shall survive termination or "
                    f"expiry of this LOI for a period of {surv} from the date of "
                    f"termination or expiry. Obligations in respect of information "
                    f"that constitutes a trade secret under applicable law shall "
                    f"continue indefinitely."
                )
                if self.t == "Distributor":
                    surv_text_body += (
                        " Obligations in respect of Protected Business Information "
                        "shall survive for the period specified in Clause 4.1."
                    )
                transaction_ref = "the Transaction" if self.t == "Distributor" else "the proposed transaction"

                # Each entry: (opt_out_key_or_None, heading_label, body_text, [sub_lines])
                # opt_out_key_or_None: when set, skipped if key appears in opt_outs
                tier_b_items = [
                    (None, "Purpose Limitation", "Each Party shall use the other Party's Confidential Information solely for the Purpose and for no other purpose.", []),
                    (None, "Non-Disclosure", "Each Party shall keep confidential all Confidential Information received from the other Party and shall not disclose such information to any person except as permitted under this Clause 6.", []),
                    (None, "Standard of Care", "Each Party shall apply no less than reasonable care to protect the other Party's Confidential Information, and no less than the care it applies to its own confidential information of a similar nature.", []),
                    (None, "Permitted Disclosures", "A Party may disclose Confidential Information to:", [
                        "(a) its Representatives who have a genuine need to know for the Purpose and who are bound by confidentiality obligations no less restrictive than this Clause 6 (whether by professional duty or written undertaking);",
                        "(b) its bona fide Financing Parties and potential co-investors, provided they are bound by confidentiality obligations no less restrictive than this Clause 6; and",
                        "(c) to the extent required by applicable law, regulation, court order, or the rules of any relevant regulatory authority or stock exchange, provided that (where legally permitted) the disclosing Party: (i) gives the other Party prior written notice as soon as reasonably practicable; (ii) consults with the other Party regarding the scope and manner of disclosure; and (iii) discloses only the minimum information required to comply.",
                    ]),
                    (None, "Liability for Representatives", "Each Party shall be responsible for any breach of this Clause 6 by its Representatives.", []),
                    (None, "Exclusions", "The obligations in Clauses 6.1 through 6.3 do not apply to information that the receiving Party can demonstrate:", [
                        "(a) is or becomes publicly available through no fault of the receiving Party or its Representatives;",
                        "(b) was already in the lawful possession of the receiving Party before disclosure, without restriction as to use or disclosure;",
                        "(c) was independently developed by the receiving Party without use of or reference to the Confidential Information; or",
                        "(d) was received from a third party who was not, to the receiving Party's knowledge, under any obligation of confidentiality in respect of that information.",
                    ]),
                    (None, "No Implied Rights", "No licence or right is granted under this LOI to the receiving Party in respect of any intellectual property rights of the disclosing Party. All Confidential Information remains the property of the disclosing Party.", []),
                    (None, "Return and Destruction", "Upon the earlier of: (a) the disclosing Party's written request, or (b) the expiry or termination of this LOI, the receiving Party shall promptly return or destroy all documents, materials, and tangible items containing Confidential Information and certify such return or destruction in writing within 15 Business Days. The receiving Party may retain copies to the extent required by applicable law or its internal compliance policies, provided such retained copies remain subject to this Clause 6.", []),
                    ("onward_sharing", "Onward-Sharing Controls", "If the receiving Party receives an inquiry from any third party regarding the disclosing Party's Confidential Information, the receiving Party shall: (a) not respond to such inquiry without the disclosing Party's prior written consent; and (b) promptly notify the disclosing Party of such inquiry. The receiving Party shall not further distribute or re-disclose Confidential Information beyond the persons authorised under Clause 6.4 without the disclosing Party's prior written consent.", []),
                    ("compliance_confirmation", "Compliance Confirmation", "Upon the disclosing Party's reasonable written request (not more than once per calendar year), the receiving Party shall confirm in writing its compliance with the obligations in this Clause 6.", []),
                    (None, "Breach Notification", "Each Party shall notify the other Party in writing within 72 hours of becoming aware of any actual or suspected breach of this Clause 6, and shall take all reasonable steps to mitigate the effects of such breach.", []),
                    (None, "Disclaimer", 'All Confidential Information is disclosed "as is." The disclosing Party makes no representation or warranty, express or implied, as to the accuracy, completeness, or reliability of any Confidential Information. The receiving Party shall be solely responsible for its own assessment and due diligence.', []),
                    ("metadata_protection", "Metadata Protection", "Confidential Information includes metadata, EXIF data, geolocation data, timestamps, file names, folder names, and any digital artifacts associated with or derived from disclosed materials. The receiving Party shall not extract, analyse, or use such metadata except as necessary for the Purpose.", []),
                    (None, "Survival", surv_text_body, []),
                    (None, "Announcements", f"Neither Party shall make any public announcement regarding this LOI or {transaction_ref} without the prior written consent of the other Party, except as required by applicable law.", []),
                    (None, "Remedies", "Each Party acknowledges that a breach of this Clause 6 may cause the disclosing Party irreparable harm for which damages would not be an adequate remedy. The disclosing Party shall be entitled to seek injunctive or other equitable relief from any court of competent jurisdiction, without the need to prove actual loss and without prejudice to any other rights or remedies.", []),
                ]

                filtered = [
                    (heading, body, subs)
                    for (opt_key, heading, body, subs) in tier_b_items
                    if not (opt_key and opt_key in opt_outs)
                ]

                for idx, (heading, body, subs) in enumerate(filtered, start=1):
                    self.bp(f"6.{idx} {heading}. ", body)
                    for sub in subs:
                        self.p(sub)

    def clause7_nc(self):
        if self.t == "EndUser" or self.t == "EcosystemPartnership":
            return  # No NC for End Users or Ecosystem Partnerships

        self.h("7. Non-Circumvention (BINDING)")
        nc_dur = self.g("protection", "nc_duration", default="24 months")

        # v3.3: for SS, downstream agreement is "Framework Agreement"; otherwise "MSA".
        downstream = "Framework Agreement" if self.t == "StrategicSupplier" else "MSA"

        self.p(f"7.1 The {self.party} shall not, directly or indirectly, without the prior written consent of Digital Energy:")

        if self.t == "Distributor":
            self.p(f"(a) contact, solicit, deal with, or enter into any business relationship with any Associated Counterparty introduced by Digital Energy in connection with the Purpose or the Transaction;")
            self.p(f"(b) circumvent, avoid, or bypass Digital Energy in order to deal directly or indirectly with any Associated Counterparty; or")
            self.p(f"(c) attempt to divert or appropriate any business opportunity disclosed by Digital Energy in connection with the Purpose or the Transaction.")
        elif self.t == "StrategicSupplier":
            self.p(f"(a) contact, solicit, deal with, or enter into any business relationship with any Associated Counterparty introduced by Digital Energy in connection with this LOI or any Provider project; or")
            self.p(f"(b) circumvent, avoid, or bypass Digital Energy in order to deal directly or indirectly with any Associated Counterparty in connection with the development, ownership, or operation of colocation or energy infrastructure on or adjacent to a site identified by Digital Energy.")
        else:  # Wholesale — lighter scope
            self.p(f"(a) contact, solicit, deal with, or enter into any business relationship with any Associated Counterparty introduced by Digital Energy in connection with this LOI or the Services; or")
            self.p(f"(b) circumvent, avoid, or bypass Digital Energy in order to deal directly or indirectly with any Associated Counterparty in connection with the development, ownership, or operation of colocation or energy infrastructure on or adjacent to a site identified by Digital Energy.")

        self.bp("7.2 Duration. ", f"The obligations in this Clause 7 shall continue for {nc_dur} after the earlier of: (a) the expiry or termination of this LOI; or (b) the expiry or termination of the {downstream}, if one is executed.")
        self.bp("7.3 Deemed Introduction by Category. ", f"Digital Energy's sharing of any Site Identifier with the {self.party} shall constitute a deemed introduction of all Associated Counterparties for that site, as defined by category in Clause 1 (Definitions). Digital Energy is not required to name each Associated Counterparty individually. Upon the {self.party}'s written request, Digital Energy shall, within ten (10) Business Days, provide a list of the named Associated Counterparties for the relevant Site Identifier, to enable the {self.party} to comply with this Clause 7.")

        if self.t == "Distributor":
            self.bp("7.4 Scope \u2014 Private and Public Bodies. ", "The non-circumvention obligations in Clause 7.1 apply to:")
            self.p("(a) Private Associated Counterparties (landowners, greenhouse operators, heat offtakers, energy counterparties, EPCs): in all cases where Digital Energy has introduced or disclosed their identity or involvement; and")
            self.p("(b) Public Bodies (government agencies, municipalities, regulatory bodies, distribution system operators): only where Digital Energy has made a specific, named introduction to an individual or department within that body. General knowledge that Digital Energy engages with a public body does not trigger non-circumvention protection.")
        elif self.t == "StrategicSupplier":
            self.bp("7.4 Scope Limitation. ", f"The non-circumvention obligations in this Clause 7 are limited to Digital Energy's supply-side relationships (site partners, energy counterparties, and infrastructure providers). Nothing in this Clause 7 restricts the {self.party} from conducting its own business with its existing or future customers, including other AI colocation operators, cloud service providers, or enterprise customers.")
        else:
            self.bp("7.4 Scope Limitation. ", f"The non-circumvention obligations in this Clause 7 are limited to Digital Energy's supply-side relationships (site partners, energy counterparties, and infrastructure providers). Nothing in this Clause 7 restricts the {self.party} from conducting its own business with its existing or future end-user customers, cloud service customers, or compute buyers.")

        self.bp("7.5 Independent Knowledge Exception. ", f"The obligations in Clause 7.1 do not apply to any Associated Counterparty with whom the {self.party} can demonstrate, by contemporaneous written evidence, that it had an existing business relationship or substantive commercial contact before the date of this LOI or before Digital Energy's disclosure.")
        self.bp(f"7.6 {downstream} Supersession. ", f"If the Parties execute a {downstream}, the non-circumvention provisions of the {downstream} shall replace this Clause 7 upon execution of the {downstream}, except that the survival period in Clause 7.2 shall apply to any introduction made before the {downstream} effective date that is not separately covered by the {downstream}.")

    def clause_general(self):
        cl = "8" if self.t != "EndUser" else "7"
        self.h(f"{cl}. General Provisions (BINDING)")
        val_date = self.g("dates", "validity_date")
        if not val_date:
            val_date = "[12 months from LOI date]"

        # v3.3: SS refers to "Framework Agreement" as downstream binding document.
        downstream = "Framework Agreement" if self.t == "StrategicSupplier" else "MSA"

        self.bp(f"{cl}.1 Non-Binding Status. ", "")
        if self.t != "EndUser":
            # v3.6.0 bug 2: SS Cl. 5 is "Supply Chain and Delivery Commitment",
            # not "Project Finance and Assignment". EP has no Cl. 5 finance
            # either (IP and Deliverables). Branch on self.t.
            if self.t == "StrategicSupplier":
                cl5_label = "Supply Chain and Delivery Commitment"
            elif self.t == "EcosystemPartnership":
                cl5_label = "IP and Deliverables"
            else:
                cl5_label = "Project Finance and Assignment"
            # v3.7.0: when include_schedule=false, drop "and Schedule 1" reference
            sched_ref = "" if self.d.get("choices", {}).get("include_schedule", True) is False else " and Schedule 1"
            self.p(f"(a) Non-binding provisions. Clauses 2 through 4{sched_ref} of this LOI are non-binding expressions of the Parties' current intentions. They do not create legally enforceable obligations and are subject to the negotiation and execution of the {downstream}.")
            self.p(f"(b) Binding provisions. Clauses 5 ({cl5_label}), 6 (Confidentiality), 7 (Non-Circumvention), and {cl} (General Provisions) are legally binding and enforceable obligations.")
        else:
            self.p("(a) Non-binding provisions. Clauses 2 through 4 of this LOI are non-binding expressions of the Parties' current intentions. They do not create legally enforceable obligations and are subject to the negotiation and execution of the MSA.")
            self.p(f"(b) Binding provisions. Clauses 5 (Project Finance and Assignment), 6 (Confidentiality), and {cl} (General Provisions) are legally binding and enforceable obligations.")

        scoping_phrase = "commercial scoping and negotiation" if self.t == "Distributor" else "technical scoping and commercial negotiation"
        # v3.6.0 bug 1: non-EndUser previously concatenated "..., t" + "t"
        # producing "tthe good faith". Collapse to single 't' on non-EU.
        self.bp(f"{cl}.2 Good Faith. ", f"The Parties agree to engage in the {scoping_phrase} process in good faith (te goeder trouw) and in accordance with the principles of reasonableness and fairness (redelijkheid en billijkheid) as contemplated by Article 6:248 of the Dutch Civil Code (Burgerlijk Wetboek). " + ("" if self.t == "EndUser" else "For the avoidance of doubt, ") + ("T" if self.t == "EndUser" else "t") + f"he good faith obligation does not oblige either Party to enter into the {downstream}. Either Party may discontinue negotiations at any time, provided it does so in good faith. Any liability arising from a breach of this good faith obligation shall be limited to verifiable reliance damages (negatief contractsbelang)" + ("." if self.t == "EndUser" else " and shall not extend to loss of profit or expectation damages (positief contractsbelang)."))

        survive = "Clauses 5.2, 5.3, 6, and 7" if self.t != "EndUser" else "Clauses 5.2, 5.3, and 6"
        self.bp(f"{cl}.3 Validity. ", f"This LOI shall remain valid until {val_date}, after which it shall lapse automatically unless extended by mutual written agreement. Upon lapse, {survive} shall survive for their respective stated periods.")
        self.bp(f"{cl}.4 Costs. ", "Each Party shall bear its own costs in connection with this LOI" + (f" and the negotiation of the {downstream}." if self.t != "EndUser" else "."))
        self.bp(f"{cl}.5 Counterparts. ", "This LOI may be executed in counterparts, including by electronic signature within the meaning of the eIDAS Regulation (EU) No 910/2014." + (" Each counterpart constitutes an original." if self.t != "EndUser" else ""))
        self.bp(f"{cl}.6 Notices. ", "All notices under this LOI shall be in writing (including email) and addressed to the contact details in the preamble. A notice is effective upon receipt." + (" Each Party shall promptly notify the other of any change to its contact details." if self.t != "EndUser" else ""))
        self.bp(f"{cl}.7 Governing Law. ", "This LOI shall be governed by and construed in accordance with the laws of the Netherlands. The United Nations Convention on Contracts for the International Sale of Goods (CISG) is expressly excluded.")
        self.bp(f"{cl}.8 Jurisdiction. ", "The courts of Amsterdam (Rechtbank Amsterdam) shall have exclusive jurisdiction over any dispute arising out of or in connection with the binding provisions of this LOI.")

        if self.t != "EndUser":
            if self.t == "Distributor":
                transaction_ref = "the Transaction"
            elif self.t == "StrategicSupplier":
                transaction_ref = "the proposed supply relationship"
            else:
                transaction_ref = "the proposed transaction"
            # v3.6.0 bug 3: "(ALT-A)" is a drafting-variant marker, never
            # counterparty-facing. Gate phrase on existing_nda; when there
            # is no prior NDA, drop the reference entirely.
            if self.choice("existing_nda"):
                ea_body = f"This LOI, together with the NDA referenced in Clause 6, constitutes the entire agreement between the Parties in relation to its subject matter and supersedes all prior negotiations, representations, and agreements relating to {transaction_ref}. Nothing in this Clause limits liability for fraud."
            else:
                ea_body = f"This LOI constitutes the entire agreement between the Parties in relation to its subject matter and supersedes all prior negotiations, representations, and agreements relating to {transaction_ref}. Nothing in this Clause limits liability for fraud."
            self.bp(f"{cl}.9 Entire Agreement. ", ea_body)
            self.bp(f"{cl}.10 Partnership Disclaimer. ", "Nothing in this LOI shall be construed as creating a partnership, joint venture, agency, or employment relationship between the Parties. Neither Party has authority to bind the other or to incur any obligation on the other's behalf.")
        else:
            self.bp(f"{cl}.9 Entire Agreement. ", "This LOI constitutes the entire agreement between the Parties in relation to its subject matter and supersedes all prior negotiations, representations, and agreements. Nothing in this Clause limits liability for fraud.")

    def signature(self):
        self.line()
        # v3.3: honours OPEN-1 (commit 2097f52) — bespoke_closing removed.
        # Operational near-signature content belongs in the letter body, not
        # the closing line. Hardcoded single-sentence closing matches main.
        # If choices.bespoke_closing is present in YAML, it is silently ignored
        # for backward compatibility.
        #
        # v3.5 scope A' (signature block cleanup): removed KvK / registration
        # lines (duplicate Parties Preamble — scope A''') and removed the
        # "ACKNOWLEDGED AND AGREED" header (unnecessary on bilateral
        # instruments; signing = agreement). Added Place: field per Dutch/EU
        # execution convention (jurisdictional + eIDAS relevance).
        self.p("We look forward to working with you.")
        self.p("")
        self.p("Yours faithfully,")
        self.p("")

        prov = self.d.get("provider", {})
        self.p(f"For and on behalf of {prov.get('legal_name', '')}", bold=True)
        self.p("")
        self.p("Signature: ____________________________")
        self.p(f"Name: {self._render_placeholder(prov.get('signatory_name'), 'sig_block_name')}")
        self.p(f"Title: {self._render_placeholder(prov.get('signatory_title'), 'sig_block_title')}")
        self.p("Date: ____________________________")
        self.p("Place: ____________________________")

        self.p("")
        self.line()
        self.p("")

        cp = self.d.get("counterparty", {})
        self.p(f"For and on behalf of {cp.get('name', '')}", bold=True)
        self.p("")
        self.p("Signature: ____________________________")
        self.p(f"Name: {self._render_placeholder(cp.get('signatory_name'), 'sig_block_name')}")
        self.p(f"Title: {self._render_placeholder(cp.get('signatory_title'), 'sig_block_title')}")
        self.p("Date: ____________________________")
        self.p("Place: ____________________________")

    def schedule(self):
        # Schedule starts on its own page.
        # v3.2: schedule titles no longer carry "(NON-BINDING)" suffix. Italic
        # prefatory note carries the non-binding signal. Authoritative
        # binding/non-binding assignment lives in Cl. 5.1 / Cl. 8.1.
        self.doc.add_page_break()
        prefatory = ("This schedule expresses the Parties' current intentions "
                      "and is non-binding, in accordance with Clause 5.1.")
        if self.t == "Distributor":
            self.line()
            self.h("Schedule 1 \u2014 Partnership Details")
            self.p(prefatory, italic=True, color=GREY, size=FONT_SMALL)
            comm = self.d.get("commercial", {})
            self.table(
                ["Item", "Detail"],
                [
                    ["Partnership Mode", self.d.get("partnership_mode", "combined").title()],
                    ["Target Customer Segments", comm.get("target_segments", "")],
                    ["Geographic Scope", comm.get("territory", "")],
                    ["Estimated Annual Capacity", f"{comm.get('indicative_mw', '')} MW IT"],
                    ["Commercial scoping complete", f"{self.g('dates', 'loi_date')} + 30 days"],
                    ["Draft MSA issued", f"{self.g('dates', 'loi_date')} + 60 days"],
                    ["MSA executed", f"{self.g('dates', 'loi_date')} + 90 days"],
                ]
            )
        elif self.t == "StrategicSupplier":
            self.line()
            self.h("Schedule 1 \u2014 Scope and Capability Matrix")
            self.p(prefatory, italic=True, color=GREY, size=FONT_SMALL)
            supplier = self.d.get("supplier", {})
            purposes = supplier.get("strategic_purposes", [])
            rows = [
                ["Capability Category", supplier.get("capability_category", "")],
                ["Core Capability", supplier.get("core_capability", "")],
                ["Strategic Purposes", ", ".join(purposes)],
                ["Geographic Coverage", supplier.get("geographic_coverage", "")],
            ]
            if "capacity_lock_in" in purposes:
                rows.append(["Lead-Time Target", supplier.get("lead_time_target", "")])
            if "pricing_volume" in purposes:
                rows.append(["Indicative Volume", supplier.get("volume_indicative", "")])
            if "engineering_integration" in purposes:
                rows.append(["Joint IP Allocation", self.g("choices", "joint_ip", default="background")])
            rows += [
                ["Scoping complete", f"{self.g('dates', 'loi_date')} + 30 days"],
                ["Framework Agreement issued", f"{self.g('dates', 'loi_date')} + 60 days"],
                ["Framework Agreement executed", f"{self.g('dates', 'loi_date')} + 90 days"],
            ]
            self.table(["Item", "Detail"], rows)
        elif self.t == "Wholesale":
            self.line()
            self.h("Schedule 1 \u2014 Capacity and Technical Requirements")
            self.p(prefatory, italic=True, color=GREY, size=FONT_SMALL)
            comm = self.d.get("commercial", {})
            tech = self.d.get("schedule_1", {}).get("technical", {}) or self.d.get("technical", {})
            # v3.5 scope N (Jonathan memo): Schedule 1 now reads technical
            # fields from intake YAML. Previously hardcoded "[To be confirmed]"
            # placeholders regardless of intake content — caused GPU platform
            # commitments established in email/user intake to be lost in the
            # Polarise LOI. Default values kept for backward-compat when
            # intake omits `schedule_1.technical.*`.
            gpu_platform = self._render_placeholder(
                tech.get("gpu_platform"), "body_clause"
            ) or "[To be confirmed during technical scoping]"
            rack_density = tech.get("rack_density_kw") or comm.get("rack_density_kw")
            rack_density_cell = (
                f"{rack_density} kW/rack"
                if rack_density and not self._is_tbc(rack_density)
                else "[To be confirmed]"
            )
            cooling = tech.get("cooling") or comm.get(
                "cooling_topology", "Direct-to-chip liquid cooling"
            )
            designated_sites = tech.get("designated_sites") or self._render_placeholder(
                tech.get("designated_sites"), "body_clause"
            ) or "To be confirmed during technical scoping"
            self.table(
                ["Item", "Detail"],
                [
                    ["Indicative Capacity", f"{comm.get('indicative_mw', '')} MW IT"],
                    ["Designated Sites", designated_sites],
                    ["GPU / Accelerator Type", gpu_platform],
                    ["Target Rack Density", rack_density_cell],
                    ["Cooling Requirement", cooling],
                    ["Technical scoping complete", f"{self.g('dates', 'loi_date')} + 30 days"],
                    ["Credit assessment complete", f"{self.g('dates', 'loi_date')} + 30 days"],
                    ["Sales Order Form issued", f"{self.g('dates', 'loi_date')} + 60 days"],
                    ["MSA executed", f"{self.g('dates', 'loi_date')} + 90 days"],
                    ["Expansion Target", f"{comm.get('expansion_mw', '')} MW IT"],
                ]
            )
        else:  # EndUser
            self.line()
            self.h("Schedule 1 \u2014 Service Requirements")
            self.p(prefatory, italic=True, color=GREY, size=FONT_SMALL)
            comm = self.d.get("commercial", {})
            tech = self.d.get("schedule_1", {}).get("technical", {}) or self.d.get("technical", {})
            types = comm.get("service_type", [])
            type_str = ", ".join(t.replace("_", " ").title() for t in types)
            gpu_platform = self._render_placeholder(
                tech.get("gpu_platform"), "body_clause"
            ) or "[To be confirmed during technical scoping]"
            data_sovereignty = self._render_placeholder(
                tech.get("data_sovereignty"), "body_clause"
            ) or "[To be confirmed]"
            self.table(
                ["Item", "Detail"],
                [
                    ["Service Model", type_str],
                    ["Indicative Capacity", comm.get("indicative_capacity", "")],
                    ["GPU / Accelerator Type", gpu_platform],
                    ["Data Sovereignty", data_sovereignty],
                    ["Technical scoping complete", f"{self.g('dates', 'loi_date')} + 30 days"],
                    ["Commercial proposal issued", f"{self.g('dates', 'loi_date')} + 60 days"],
                    ["MSA executed", f"{self.g('dates', 'loi_date')} + 90 days"],
                ]
            )

    def footer(self):
        # Real Word footers are set up in _setup(). Add version reference at end of body.
        self.p("")
        # v3.3: version string per-type. New types (SS, EP) are v1.0.
        version_by_type = {
            "EndUser": "v3.2",
            "Distributor": "v3.2",
            "Wholesale": "v3.2",
            "StrategicSupplier": "v1.0",
            "EcosystemPartnership": "v1.0",
            # M4: Bespoke type — new in v1.0. Escape hatch for deals that
            # don't fit the 5 templated types, still gated through the
            # full QA catalog.
            "Bespoke": "v1.0",
        }
        vsn = version_by_type.get(self.t, "v3.2")
        self.p(
            f"DE-LOI-{self.t}-{vsn}",
            italic=True, size=FONT_SMALL, color=GREY
        )

    # ----- EP (Ecosystem Partnership) clause builders -----
    # v3.3: EP uses a dedicated, lighter structure. No Cl. 5 Project Finance,
    # no Cl. 7 NC. Cl. 5 is IP & Deliverables. Cl. 6 is Tier A mutual light
    # (7 sub-clauses). Cl. 7 is General (11 sub-clauses).

    def definitions_ep(self):
        self.h("1. Definitions")
        self.p("1.1 In this LOI, unless the context requires otherwise:")
        self.bp(
            '"Confidential Information" ',
            "means all information (whether technical, commercial, financial, or "
            "otherwise) disclosed by a Party to the other in connection with the "
            "Collaboration, whether in writing, orally, electronically, or by "
            "inspection, including any information that by its nature or the "
            "circumstances of disclosure would reasonably be understood to be "
            "confidential. The existence and contents of this LOI are Confidential "
            "Information."
        )
        self.bp(
            '"Collaboration" ',
            "means the ecosystem activities and joint work contemplated by this "
            "LOI as described in Clause 3."
        )
        self.bp(
            '"Representatives" ',
            "means, in relation to a Party, its Affiliates and its and their "
            "respective directors, officers, employees, agents, and professional "
            "advisers."
        )

    def clause2_ep(self):
        self.h("2. Purpose and Scope")
        eco = self.d.get("ecosystem", {})
        themes = ", ".join(eco.get("collaboration_themes", [])) or "[COLLABORATION_THEMES]"
        self.p(
            f"2.1 This LOI records the Parties' mutual interest in a "
            f"non-commercial ecosystem collaboration around {themes}."
        )
        self.p(
            "2.2 The collaboration is strictly non-commercial. No Party shall be "
            "entitled to any fee, payment, capacity, discount, or commercial "
            "benefit from the other Party by reason of this LOI or any activity "
            "under it."
        )
        self.p(
            "2.3 The commercial implications of any future activity \u2014 including "
            "any joint offering, research licensing, commercialisation of joint "
            "output, or provision of services by either Party to the other \u2014 shall "
            "be the subject of a separate agreement and are expressly outside the "
            "scope of this LOI."
        )
        self.p(
            "2.4 The confidentiality and general provisions in Clauses 6 and 7 "
            "are legally binding and enforceable from the date of execution. "
            "Clauses 3 through 5 are non-binding expressions of collaborative intent."
        )

    def clause3_ep(self):
        self.h("3. Collaboration Scope")
        eco = self.d.get("ecosystem", {})
        themes = eco.get("collaboration_themes", [])
        categories = eco.get("joint_activity_categories", [])

        self.bp("3.1 Collaboration Themes. ", "The Parties intend to collaborate on the following themes:")
        if themes:
            for theme in themes:
                self.p(f"\u2014 {theme}")
        else:
            self.p("\u2014 [THEMES TO BE CONFIRMED]")

        self.bp(
            "3.2 Joint Activity Categories. ",
            "Subject to resources and mutual agreement, the Parties intend to "
            "engage in one or more of the following categories of activity:"
        )
        if categories:
            cat_labels = {
                "publications": "publications (joint white papers, research articles, reports)",
                "events": "events (conferences, roundtables, workshops)",
                "pilots": "pilots (joint demonstrations, proofs of concept)",
                "advocacy": "advocacy (policy input, consultation responses)",
                "working_groups": "working groups (standards and technical coordination)",
            }
            for cat in categories:
                self.p(f"(\u2014) {cat_labels.get(cat, cat)}")
        else:
            self.p("\u2014 [ACTIVITY CATEGORIES TO BE CONFIRMED]")

        self.bp(
            "3.3 Governance and Cadence. ",
            "The Parties intend to establish a lightweight governance mechanism "
            "\u2014 typically an initial set-up meeting and quarterly check-ins \u2014 to "
            "review activity, align on priorities, and resolve any coordination "
            "issues. No binding governance structure is created by this LOI."
        )
        self.bp(
            "3.4 Working-Group Participation. ",
            "Where either Party hosts or convenes a working group relevant to the "
            "collaboration themes, the other Party shall have a standing "
            "invitation to participate on mutually agreed terms."
        )
        self.bp(
            "3.5 Non-Exclusivity. ",
            "The collaboration contemplated by this LOI is non-exclusive. Both "
            "Parties retain the right to enter into similar or competing "
            "collaborations with third parties."
        )

    def clause4_ep(self):
        self.h("4. Announcements and Branding")
        announcement = self.g("choices", "announcement_protocol", default="mutual_approval")
        logo_use = self.g("choices", "logo_use", default="reciprocal")

        if announcement == "notify_only":
            self.bp(
                "4.1 Announcement Protocol. ",
                "Each Party shall give the other Party reasonable prior notice "
                "(target: 5 Business Days) before making any public announcement "
                "referring to the collaboration. Consent is not required."
            )
        else:  # mutual_approval (default)
            self.bp(
                "4.1 Announcement Protocol. ",
                "Any public announcement, press release, published article, or "
                "social media post referring to the collaboration, the other "
                "Party, or joint activity under this LOI shall require the prior "
                "written consent of the other Party. Consent shall not be "
                "unreasonably withheld."
            )

        if logo_use == "none":
            self.bp(
                "4.2 Logo Use. ",
                "Neither Party shall use the other Party's name, logo, or trade "
                "mark in any public context without the prior written consent of "
                "the other Party."
            )
        elif logo_use == "one_way":
            self.bp(
                "4.2 Logo Use. ",
                "One-way logo use is contemplated between the Parties; the "
                "Party, scope, and terms will be specified in a separate "
                "schedule or letter."
            )
        else:  # reciprocal (default)
            self.bp(
                "4.2 Logo Use. ",
                "Each Party grants the other a limited, non-exclusive, revocable "
                "licence to use its name and logo solely in the context of the "
                "collaboration and subject to the branding guidelines each Party "
                "provides. No other use is permitted."
            )

        self.bp(
            "4.3 Attribution. ",
            "Where either Party refers to the collaboration in any public "
            "context, it shall: (a) describe the collaboration accurately and "
            "without overstatement; (b) refrain from implying any commercial "
            "relationship, endorsement, or joint venture; and (c) include the "
            "other Party's preferred attribution language when supplied."
        )

    def clause5_ep(self):
        self.h("5. Intellectual Property and Deliverables")
        self.bp(
            "5.1 Background IP. ",
            "Each Party retains all right, title, and interest in its background "
            "intellectual property. Nothing in this LOI transfers or licenses any "
            "intellectual property between the Parties except as expressly set "
            "out in Clause 4.2."
        )
        self.bp(
            "5.2 Joint Deliverables. ",
            "Any intellectual property jointly developed by the Parties in the "
            "course of the collaboration (including joint publications, "
            "working-group outputs, and jointly authored materials) shall be "
            "governed by a separate agreement to be executed before the joint "
            "activity produces any material output."
        )
        self.bp(
            "5.3 No Assignment by Participation. ",
            "Nothing in this LOI or in any activity conducted under it shall be "
            "construed as an assignment, licence, or transfer of intellectual "
            "property."
        )
        self.bp(
            "5.4 Publication. ",
            "Each Party retains the right to publish, present, and disseminate "
            "its own work, subject to the confidentiality provisions in Clause 6."
        )

    def clause6_ep(self):
        self.h("6. Confidentiality (BINDING)")
        surv = self.g("protection", "confidentiality_survival", default="2 years")
        self.p(
            "6.1 Each Party shall keep confidential all Confidential Information "
            "received from the other Party and shall use such information solely "
            "for the Collaboration."
        )
        self.p(
            "6.2 Each Party may disclose Confidential Information only to its "
            "Representatives who have a genuine need to know for the "
            "Collaboration and who are bound by confidentiality obligations no "
            "less restrictive than this Clause 6."
        )
        self.p(
            "6.3 The obligations in Clauses 6.1 and 6.2 do not apply to "
            "information that the receiving Party can demonstrate: (a) is or "
            "becomes publicly available through no fault of the receiving Party; "
            "(b) was already in the receiving Party's possession without "
            "restriction; (c) was independently developed without use of the "
            "Confidential Information; or (d) was received from a third party "
            "without breach of any confidentiality obligation."
        )
        self.p(
            "6.4 A Party may disclose Confidential Information to the extent "
            "required by applicable law, regulation, or court order, provided "
            "that (where legally permitted) the disclosing Party gives the other "
            "Party prior written notice and discloses only the minimum required."
        )
        self.p(
            "6.5 Upon written request or expiry of this LOI, the receiving Party "
            "shall promptly return or destroy all Confidential Information and "
            "certify such return or destruction in writing. Copies required by "
            "applicable law or internal compliance policies may be retained, "
            "subject to continued confidentiality."
        )
        self.p(
            "6.6 Public statements about this LOI are governed by Clause 4."
        )
        self.p(
            f"6.7 The obligations in this Clause 6 shall survive for {surv} from "
            f"the date of termination or expiry of this LOI."
        )

    def clause7_ep_general(self):
        self.h("7. General Provisions (BINDING)")
        val_date = self.g("dates", "validity_date") or "[12 months from LOI date]"

        self.bp("7.1 Non-Binding Status. ", "")
        self.p(
            "(a) Non-binding provisions. Clauses 2, 3, 4, and 5 of this LOI are "
            "non-binding expressions of collaborative intent."
        )
        self.p(
            "(b) Binding provisions. Clauses 6 (Confidentiality) and 7 (General "
            "Provisions) are legally binding and enforceable obligations."
        )
        self.bp(
            "7.2 No Commercial Commitment. ",
            "Nothing in this LOI creates any obligation on either Party to pay, "
            "purchase, supply, or grant anything of commercial value to the "
            "other. Any such arrangement requires a separate agreement."
        )
        self.bp(
            "7.3 Non-Exclusivity. ",
            "As per Clause 3.5."
        )
        self.bp(
            "7.4 Term and Validity. ",
            f"This LOI shall remain valid until {val_date}, after which it shall "
            f"lapse automatically unless extended by mutual written agreement. "
            f"Upon lapse, Clause 6 (Confidentiality) shall survive for its "
            f"stated period."
        )
        self.bp(
            "7.5 Costs. ",
            "Each Party shall bear its own costs in connection with this LOI "
            "and the Collaboration."
        )
        self.bp(
            "7.6 Counterparts. ",
            "This LOI may be executed in counterparts, including by electronic "
            "signature within the meaning of the eIDAS Regulation (EU) No 910/2014."
        )
        self.bp(
            "7.7 Notices. ",
            "All notices under this LOI shall be in writing (including email) "
            "and addressed to the contact details in the preamble. A notice is "
            "effective upon receipt."
        )
        self.bp(
            "7.8 Governing Law. ",
            "This LOI shall be governed by the laws of the Netherlands. The "
            "United Nations Convention on Contracts for the International Sale "
            "of Goods (CISG) is expressly excluded."
        )
        self.bp(
            "7.9 Jurisdiction. ",
            "The courts of Amsterdam (Rechtbank Amsterdam) shall have exclusive "
            "jurisdiction over any dispute arising out of or in connection with "
            "the binding provisions of this LOI."
        )
        self.bp(
            "7.10 Entire Agreement. ",
            "This LOI constitutes the entire agreement between the Parties in "
            "relation to its subject matter and supersedes all prior "
            "negotiations, representations, and agreements. Nothing in this "
            "Clause limits liability for fraud."
        )
        self.bp(
            "7.11 Partnership Disclaimer. ",
            "Nothing in this LOI shall be construed as creating a partnership, "
            "joint venture, agency, or employment relationship between the "
            "Parties. Neither Party has authority to bind the other or to incur "
            "any obligation on the other's behalf."
        )

    def schedule_ep(self):
        """v3.3: Optional Joint Activity Plan schedule for EP.
        Only rendered if intake provides ecosystem.joint_activity_plan (list of
        dicts with keys: activity, category, lead, target_date).
        """
        plan = self.d.get("ecosystem", {}).get("joint_activity_plan", [])
        if not plan:
            return
        self.doc.add_page_break()
        self.line()
        self.h("Schedule 1 \u2014 Joint Activity Plan")
        self.p(
            "This schedule expresses the Parties' current intentions and is "
            "non-binding, in accordance with Clause 2.4.",
            italic=True, color=GREY, size=FONT_SMALL,
        )
        rows = []
        for item in plan:
            rows.append([
                item.get("activity", ""),
                item.get("category", ""),
                item.get("lead", ""),
                item.get("target_date", ""),
            ])
        self.table(["Activity", "Category", "Lead", "Target Date"], rows)

    def build(self) -> Document:
        # v3.3: EP uses its own clause builders (will be added next); until wired,
        # fall back to partial engine output.
        if self.t == "EcosystemPartnership":
            return self._build_ep()

        # M4: Bespoke — free-text clauses with structural validation,
        # full QA catalog still applies.
        if self.t == "Bespoke":
            return self._build_bespoke()

        self.letterhead()
        # add_cover() already ends with page break
        # v3.5.2 scope A''': Parties Preamble renders legal identification
        # in the body before recitals — cover page is presentation only.
        self.parties()
        self.recitals()
        self.definitions()
        self.clause2()

        if self.t == "Distributor":
            self.clause3_ds()
        elif self.t == "Wholesale":
            self.clause3_ws()
        elif self.t == "StrategicSupplier":
            self.clause3_ss()
        else:
            self.clause3_eu()

        if self.t == "StrategicSupplier":
            self.clause4_ss()
        else:
            self.clause4()

        # v3.4: SS uses clause5_ss (Supply Chain and Delivery Commitment).
        # EU/DS/WS use the generic clause5 (Project Finance and Assignment).
        if self.t == "StrategicSupplier":
            self.clause5_ss()
        else:
            self.clause5()
        self.clause6()
        self.clause7_nc()
        self.clause_general()
        self.signature()
        # v3.7.0: choices.include_schedule=false suppresses Schedule 1 entirely.
        # Default (unset or True) preserves pre-v3.7.0 behavior. Note that
        # self.choice() returns False when unset — so we look at the raw
        # intake dict to distinguish unset (default True) from explicit False.
        _inc_sched = self.d.get("choices", {}).get("include_schedule", True)
        if _inc_sched:
            self.schedule()
        self._inject_custom_clauses()
        # v3.7.1: apply `custom.clauses` replace + insert-after modes after
        # all body rendering is complete.
        self._apply_custom_mutations()
        # Note: v3.7.2 custom-mutation failures are stored on
        # `self._custom_mutation_failures` and surfaced by qa_lint via the
        # `builder_warnings` parameter when called from main().
        # v3.7.1: opt-in post-template renumbering pass. Closes gaps in
        # top-level clause numbering when conditional sub-clauses skip
        # (e.g., clause4_ss emits 4.1 / 4.2 / 4.4 / 4.6 when only
        # pipeline_visibility is in strategic_purposes — renumbering
        # closes to 4.1 / 4.2 / 4.3 / 4.4). Opt-in to preserve goldens
        # of pre-v3.7.1 callers.
        if self.d.get("choices", {}).get("auto_renumber"):
            self._renumber_clauses()
        self.footer()
        return self.doc

    def _renumber_clauses(self):
        """v3.7.1: close numbering gaps in top-level clauses (N.M).

        Only renumbers paragraphs whose text starts with exactly `N.M` followed
        by a word boundary. Body cross-references of the form "Clause N.M" or
        "Clauses N.M" are updated via a final text-sweep. Sub-minor numbering
        (N.M.X) is preserved — only the minor level renumbers.
        """
        from collections import defaultdict

        paras = self.doc.paragraphs
        # Pass 1: group by major clause number, preserving document order
        seen_order = defaultdict(list)  # major:int -> [original minor ints in order]
        by_major = defaultdict(list)    # major:int -> [(para, orig_minor_int)]

        for para in paras:
            text = para.text.lstrip()
            m = re.match(r"^(\d+)\.(\d+)(?!\d)(?!\.\d)", text)
            if not m:
                continue
            major = int(m.group(1))
            minor = int(m.group(2))
            by_major[major].append((para, minor))
            if minor not in seen_order[major]:
                seen_order[major].append(minor)

        # Pass 2: build remap per major; skip majors that are already contiguous
        remap = {}  # "N.M" -> "N.M'"
        for major, minors in seen_order.items():
            expected = list(range(1, len(minors) + 1))
            if minors == expected:
                continue
            new_minors = dict(zip(minors, expected))
            for orig, new in new_minors.items():
                if orig != new:
                    remap[f"{major}.{orig}"] = f"{major}.{new}"

        if not remap:
            return

        # Pass 3: rewrite paragraph prefixes. Two-phase substitution via
        # placeholder tokens to avoid chained-rewrite bugs (e.g., remapping
        # 3.7 → 3.2 and 3.12 → 3.7 — if we substitute 3.12 first, then
        # 3.7 would double-rewrite to 3.2).
        #
        # Phase A: replace each old `N.M` with a unique placeholder.
        # Phase B: replace placeholders with final `N.M'`.
        #
        # Regex guards: match "N.M" only when not preceded by a digit/dot
        # and not followed by ".digit" or another digit — so "3.1" doesn't
        # match inside "3.10" or "3.1.1".
        sorted_keys = sorted(remap.keys(), key=lambda k: -len(k))
        placeholders = {
            old: f"\uE000RENUM_{i}\uE001" for i, old in enumerate(sorted_keys)
        }

        def substitute_a(text: str) -> str:
            for old in sorted_keys:
                if old not in remap or old == remap[old]:
                    continue
                pattern = re.compile(
                    rf"(?<!\d)(?<!\.)\b{re.escape(old)}\b(?!\.\d)(?!\d)"
                )
                text = pattern.sub(placeholders[old], text)
            return text

        def substitute_b(text: str) -> str:
            for old, ph in placeholders.items():
                if ph in text:
                    text = text.replace(ph, remap[old])
            return text

        for para in paras:
            for run in para.runs:
                if run.text:
                    new_text = substitute_b(substitute_a(run.text))
                    if new_text != run.text:
                        run.text = new_text

    def _inject_custom_clauses(self):
        """v3.7.0 + v3.7.1 — inject `custom.clauses[]` entries.

        Modes:
        - `append` (v3.7.0): render at end of body in the order provided.
        - `replace` (v3.7.1): find an existing paragraph whose text starts
          with `{number}\\b` and overwrite its text + any immediate
          sub-paragraphs that follow before the next numbered paragraph.
        - `insert-after:N` (v3.7.1): find the paragraph starting with `N\\b`
          and insert the new clause immediately after it (including sub-
          clause lines).

        Replace/insert-after are post-render mutations of the Document
        object. Append stays inline at the end of body (pre-schedule).
        """
        append_items = []
        post_items = []  # replace + insert-after

        for item in self.d.get("custom", {}).get("clauses", []) or []:
            mode = item.get("mode", "append")
            if mode == "append":
                append_items.append(item)
            else:
                post_items.append(item)

        # Append happens inline now (current paragraph insertion order).
        for item in append_items:
            number = item.get("number", "")
            heading = item.get("heading", "")
            text = item.get("text", "")
            if heading:
                self.bp(f"{number} {heading}. ", text)
            else:
                self.p(f"{number} {text}")
            for sub in item.get("sub_clauses", []) or []:
                self.p(sub)

        # Replace + insert-after are applied after all append items are
        # in the document. Store for post-build processing in build().
        self._pending_custom_mutations = post_items

    def _apply_custom_mutations(self):
        """v3.7.1 — apply `replace` + `insert-after:N` custom.clauses[] entries
        to the already-rendered Document. Called from build() after append
        items are inline-rendered.

        v3.7.2: records failed lookups so `qa_lint()` can surface them as
        WARN lines instead of silently no-op'ing. Stored in
        `self._custom_mutation_failures` as a list of
        `(mode, target, reason)` tuples.
        """
        self._custom_mutation_failures = []
        mutations = getattr(self, "_pending_custom_mutations", None) or []
        if not mutations:
            return

        for item in mutations:
            mode = item.get("mode", "append")
            target_number = item.get("number", "")
            heading = item.get("heading", "")
            text = item.get("text", "")
            sub_clauses = item.get("sub_clauses", []) or []

            if mode.startswith("insert-after:"):
                target = mode.split(":", 1)[1].strip()
                idx = self._find_paragraph_index(target)
                if idx < 0:
                    self._custom_mutation_failures.append(
                        ("insert-after", target,
                         f"target clause {target!r} not found in rendered body")
                    )
                    continue
                self._insert_clause_after(
                    target, target_number, heading, text, sub_clauses
                )
            elif mode == "replace":
                idx = self._find_paragraph_index(target_number)
                if idx < 0:
                    self._custom_mutation_failures.append(
                        ("replace", target_number,
                         f"target clause {target_number!r} not found in rendered body")
                    )
                    continue
                self._replace_clause(target_number, heading, text, sub_clauses)

    def _find_paragraph_index(self, clause_number: str):
        """Return the index of the first paragraph whose text begins with the
        given clause number followed by whitespace or period-then-whitespace.
        Returns -1 if not found.
        """
        target = clause_number.strip()
        for idx, para in enumerate(self.doc.paragraphs):
            text = para.text.lstrip()
            # Match "3.8 " or "3.8. " but not "3.80" or "3.8.1"
            if re.match(rf"^{re.escape(target)}(?=[\s.][^\d]|[\s.]$|\s)", text):
                return idx
            # Also match bold run pattern "3.8 Heading. "
            if text.startswith(target + " ") or text.startswith(target + ". "):
                return idx
        return -1

    def _span_end_index(self, start_idx: int) -> int:
        """Find the end of the clause span that starts at `start_idx`.

        A clause span ends at the next paragraph whose text starts with
        another `N.M` token (at same or higher depth) or a new top-level
        heading (`N. Title`), or at end-of-document.
        """
        paras = self.doc.paragraphs
        start_text = paras[start_idx].text.lstrip()
        m = re.match(r"^(\d+)\.(\d+)", start_text)
        if not m:
            return start_idx + 1
        start_major = int(m.group(1))
        for idx in range(start_idx + 1, len(paras)):
            text = paras[idx].text.lstrip()
            nm = re.match(r"^(\d+)\.(\d+)", text)
            if nm:
                return idx
            nh = re.match(r"^(\d+)\.\s+[A-Z]", text)
            if nh and int(nh.group(1)) >= start_major:
                return idx
            # Break on new heading patterns (Schedule, etc.)
            if text.startswith("Schedule") or text.startswith("Signed "):
                return idx
        return len(paras)

    def _replace_clause(self, target_number, heading, text, sub_clauses):
        """Replace the paragraph at target_number + its body sub-paragraphs."""
        idx = self._find_paragraph_index(target_number)
        if idx < 0:
            return  # silent — target not present in this LOI type
        end = self._span_end_index(idx)

        # Remove paragraphs idx .. end-1
        body = self.doc.element.body
        for _ in range(end - idx):
            para = self.doc.paragraphs[idx]
            p_el = para._element
            p_el.getparent().remove(p_el)

        # Insert replacement at the original idx. We create new paragraphs via
        # python-docx append + reorder relative to the deleted span.
        insert_before = (
            self.doc.paragraphs[idx]._element if idx < len(self.doc.paragraphs) else None
        )
        new_paras = self._build_inserted_paragraphs(
            target_number, heading, text, sub_clauses
        )
        for np_el in new_paras:
            if insert_before is not None:
                insert_before.addprevious(np_el)
            else:
                body.append(np_el)

    def _insert_clause_after(self, after_number, new_number, heading, text, sub_clauses):
        """Insert new_number's clause immediately after after_number."""
        idx = self._find_paragraph_index(after_number)
        if idx < 0:
            return  # silent
        end = self._span_end_index(idx)

        # New paragraphs are inserted at position `end` (just before the
        # next clause/span). Use addprevious to anchor.
        insert_before = (
            self.doc.paragraphs[end]._element if end < len(self.doc.paragraphs) else None
        )
        new_paras = self._build_inserted_paragraphs(
            new_number, heading, text, sub_clauses
        )
        if insert_before is not None:
            for np_el in new_paras:
                insert_before.addprevious(np_el)
        else:
            body = self.doc.element.body
            for np_el in new_paras:
                body.append(np_el)

    def _build_inserted_paragraphs(self, number, heading, text, sub_clauses):
        """Render the new clause + sub_clauses into detached paragraph elements.

        Returns list of lxml elements (the <w:p> XML) ready to be inserted
        into the document body.
        """
        # Trick: render paragraphs using self.bp/self.p, then detach the
        # tail paragraphs from the document so they can be re-inserted at
        # the desired position.
        starting_count = len(self.doc.paragraphs)
        if heading:
            self.bp(f"{number} {heading}. ", text)
        else:
            self.p(f"{number} {text}")
        for sub in sub_clauses or []:
            self.p(sub)

        # Detach newly-appended paragraphs
        new_paras = []
        body = self.doc.element.body
        all_paras = self.doc.paragraphs
        appended = all_paras[starting_count:]
        for para in appended:
            p_el = para._element
            p_el.getparent().remove(p_el)
            new_paras.append(p_el)
        return new_paras

    def _build_ep(self) -> Document:
        """v3.3: Ecosystem Partnership build pipeline.
        Different structure: no Cl. 5 Project Finance, no Cl. 7 NC.
        Cl. 5 is IP & Deliverables, Cl. 6 is Tier A light mutual, Cl. 7 is General.
        """
        self.letterhead()
        # v3.5.2 scope A''': Parties Preamble for EP too.
        self.parties()
        self.recitals()
        self.definitions_ep()
        self.clause2_ep()
        self.clause3_ep()
        self.clause4_ep()
        self.clause5_ep()
        self.clause6_ep()
        self.clause7_ep_general()
        self.signature()
        # Schedule 1 is optional for EP — only rendered if the intake provides one.
        if self.d.get("ecosystem", {}).get("joint_activity_plan"):
            self.schedule_ep()
        self.footer()
        return self.doc

    # ----- Bespoke build (M4) -----
    # Escape hatch for deals that don't fit the 5 templated types
    # (EU/DS/WS/SS/EP). Clauses come from the YAML as a structured list:
    # each clause has a number, heading, paragraphs (ordered), and
    # optional lettered subclauses. Recital A is always library-sourced;
    # supplementary recitals (B+) come from the optional `recitals` list.
    #
    # IMPORTANT: Bespoke ≠ unvalidated. The full QA catalogue
    # (R-1 banned phrases, R-7 Unicode arrows, R-20 programme-span
    # regressions, R-24 party-duplication, etc.) still applies to
    # bespoke clause text — qa_lint() walks the entire rendered
    # document and doesn't care how it was built. If a bespoke
    # paragraph contains "→" or "minimum commitment term of 5 years",
    # the gate fires as it would for any templated type.

    def _bespoke_recitals(self):
        """Recital A from library + optional supplementary recitals from YAML."""
        self.h("Recitals")
        cp = self.g("counterparty", "short")
        desc = self.g("counterparty", "description")
        # v3.6.0 bug 5 parity: strip trailing period from description
        # before the engine appends its own.
        if isinstance(desc, str):
            desc = desc.rstrip().rstrip(".")

        recital_a_body = resolve_recital_a(self.d)
        self.p(f"(A) {recital_a_body}")
        self.p(f'(B) {cp} (the "{self.party}") {desc}.')

        # Optional C, D, E... from YAML. Author supplies letter + text;
        # validate() already enforces presence of both fields on each.
        for r in self.d.get("recitals", []) or []:
            letter = str(r.get("letter", "")).strip()
            text = str(r.get("text", "")).strip()
            if not letter or not text:
                continue
            self.p(f"({letter}) {text}")

    def _bespoke_clauses(self):
        """Render the YAML `clauses` list as Heading 1 + body paragraphs
        + (a)/(b) lettered subclauses."""
        for c in self.d.get("clauses", []) or []:
            number = str(c.get("number", "")).strip()
            heading = str(c.get("heading", "")).strip()
            if not number or not heading:
                continue
            self.h(f"{number}. {heading}", level=1)

            for para in c.get("paragraphs", []) or []:
                text = str(para).strip()
                if text:
                    self.p(text)

            for sub in c.get("subclauses", []) or []:
                letter = str(sub.get("letter", "")).strip()
                text = str(sub.get("text", "")).strip()
                if letter and text:
                    self.p(f"({letter}) {text}")

    def _build_bespoke(self) -> Document:
        """M4: Bespoke build pipeline.

        Structure:
          1. Cover (letterhead) — via add_cover, brand-consistent.
          2. Parties Preamble — v3.5.2 scope A''' (same as templated types).
          3. Recitals — (A) library-sourced, (B) counterparty descriptor,
             (C)... from YAML `recitals:` list.
          4. Clauses — from YAML `clauses:` list, each with number,
             heading, paragraphs, optional subclauses.
          5. Signature block — shared across all types.
          6. Footer — DE-LOI-Bespoke-v1.0 version stamp.

        No definitions(), no clause5/7_nc/general — bespoke deals
        include those directly inside the free-text clauses list when
        relevant. Keeps the engine honest: structure validated, content
        authored.
        """
        self.letterhead()
        # v3.5.2 scope A''': Parties Preamble applies here too.
        self.parties()
        self._bespoke_recitals()
        self._bespoke_clauses()
        self.signature()
        self.footer()
        return self.doc


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

import re  # noqa: E402 — used below by qa_lint


# v3.2 QA linter. Rules in sync with _shared/loi-qa-gate.md.

_ARROW_CHARS = "\u2192\u21D2\u279C\u27F6\u21A6\u27F9\u21E8"

_FAIL_RULES = {
    "R-1": (r"minimum commitment term of 5 years", "body",
            "'minimum commitment term of 5 years' banned; use 'approximately 5 years, indicative only'"),
    "R-2": (r"\b\d+\s+identified\s+sites\b", "Recital A",
            "Fixed site-count language in Recital A"),
    "R-3": (r"12 months of commercial commitment", "Recital A",
            "'12 months of commercial commitment' banned in Recital A"),
    "R-5": (r"We are confident that", "closing",
            "'We are confident that' banned"),
    "R-7": (f"[{_ARROW_CHARS}]", "body",
            "Unicode arrow detected in body"),
    "R-8": (r"Schedule \d+ [^\n]*\(NON-BINDING\)", "schedule-title",
            "'(NON-BINDING)' suffix in schedule title — use italic prefatory note instead"),
    "R-10": (r"4\.2 Revenue Chain", "Cl. 4.2",
             "Cl. 4.2 heading must be 'Contractual Sequence', not 'Revenue Chain'"),
    "R-20": (r"programme spans \d+", "Recital A",
             "'programme spans N...' language regression"),
    # R-23 (v3.4): fabrication gate — implemented as a custom check in qa_lint,
    # not a simple regex. See _check_fabrication_gate() below.
    # v3.5.2 scope 0 additions:
    "R-24": (
        r"\[[A-Za-z][A-Za-z0-9._-]*\.(?:com|eu|de|co\.uk|org|nl|io|ai|ch)\]",
        "Recital B",
        "Inline source citation in Recital B (e.g. [polarise.eu]) — source attribution lives in counterparty.source_map YAML, not in rendered prose. Strip brackets.",
    ),
    "R-25": (
        # Vanity-financial patterns — catches unattributed capital-raise language
        # and valuation vanity. Does NOT catch named-endorser financings like
        # "backed by Macquarie" or "strategic investment from NVIDIA".
        #
        # Patterns (any match = FAIL):
        #   • "Series A/B/C/D/E/F" (not followed by "from" — that would be a
        #     specific named investor round and should be Signal-Test-judged)
        #   • "valuation of" (standalone valuation phrase)
        #   • "at a €500m valuation" or "at EUR 500m valuation" (valuation adj)
        #   • "raised $150M" / "raised EUR 117m" (unattributed raise)
        #   • "growth commitment" (pure financial headline)
        #   • "SAFE round" / "convertible note"
        r"\b("
        r"Series\s+[A-F](?!\s+from\s+\w)"
        r"|valuation of"
        r"|at\s+(?:a\s+)?(?:\$|€|£|USD|EUR|GBP)?\s*\d[\d.,]*\s*(?:bn|m|million|billion)\s+valuation"
        r"|raised\s+(?:\$|€|£|USD|EUR|GBP)?\s*\d[\d.,]*(?:\s*(?:M|B|m|bn|million|billion))?"
        r"|growth commitment"
        r"|SAFE round"
        r"|convertible note"
        r")\b",
        "Recital B",
        "Vanity-financial claim in Recital B — valuation numbers / generic VC labels / unattributed capital-raise language fail Signal Test gate 1. Named-endorser financings (e.g. 'backed by Macquarie') are signal and remain allowed; pure vanity metrics are not.",
    ),
    "R-27": (
        # v3.7.0: broaden to also match bare TBC (no brackets), but exclude
        # TBC that appears inside a URL (preceded by http or path component).
        # Pattern: Name/Title field followed by [TBC] OR bare word TBC.
        r"(?:Name|Title):\s*(?:\[TBC\]|(?<![/\w])TBC(?![\w]))",
        "sig-block",
        "'TBC' rendered literally in signature-block Name or Title line — must route through _render_placeholder so the line becomes a fillable blank on external-facing drafts.",
    ),
    # R-30 (v3.7.0 fail): double-period detector. Excludes '...' ellipsis and
    # numbered-list notation like '3.1.' (a digit immediately preceding the dots).
    "R-30": (
        r"(?<!\.)(?<!\d)\.{2}(?!\.)",
        "body",
        "Double-period '..' in rendered body — likely trailing-period concatenation "
        "bug in a free-text YAML field. Exclude: '...' ellipsis, '3.1.' numbering.",
    ),
}

_WARN_RULES = {
    "R-11": (r"\bISO \d{4,5}\b", "Recital B",
             "ISO certification in Recital B (set choices.cert_relevant=true to suppress)"),
    # v3.4: R-14 scope broadened from Recital B to body. Added "purpose-built"
    # and "state-of-the-art" as R-21 (kept R-14 list focused on the original
    # salesy adjectives for clarity in QA reports).
    "R-14": (r"\b(leading|innovative|cutting-edge|world-class|best-in-class)\b", "body",
             "Salesy adjective"),
    "R-15": (r"positioning (its|itself) as", "body",
             "'positioning (its|itself) as' — formulaic"),
    "R-19": (r"^\s*\d+(?:\.\d+)?\s+[^\n]*\(NON-BINDING\)", "heading",
             "'(NON-BINDING)' in a clause heading — v3.2 style regression"),
    # v3.4 additions:
    "R-21": (r"\b(purpose-built|state-of-the-art)\b", "body",
             "Marketing adjective — 'purpose-built' / 'state-of-the-art' banned body-wide (v3.4)"),
    "R-22": (
        # v3.5.3-cont scope E: patterns narrowed to clause-context signatures
        # to reduce false positives. Prior regex flagged bare phrases like
        # "Provider's ability to" and "will require the exchange of" which
        # can legitimately appear in operative clauses (e.g. information-
        # exchange provisions in Recital E of NDAs). Narrowed forms below
        # require the meta-commentary *verb* context that only meta-
        # commentary uses (secure/obtain/access for abilities; Confidential
        # Information/material for exchanges).
        #
        # Allowlist — these patterns are NOT meta-commentary and MUST NOT
        # fire R-22 even though they contain similar tokens:
        #   • "ability to deliver" / "ability to scale" (operational capability)
        #   • "will require the exchange of information" (without
        #     "Confidential Information" qualifier — could be legitimate)
        #   • "depends on" (without "in part" modifier)
        r"("
        r"(?:Provider's|Digital Energy's) ability to (?:secure|obtain|access)|"
        r"depends in part on|"
        r"is intended to evidence|"
        r"while non-binding in its commercial terms|"
        r"to support (?:the Provider's|Digital Energy's) financing|"
        r"will require the exchange of (?:Confidential Information|material)|"
        r"The Parties acknowledge that (?:the Provider|Digital Energy) intends|"
        r"is intended to form the basis"
        r")",
        "body",
        "Meta-commentary pattern — explains the LOI rather than creating obligations (v3.4; narrowed v3.5.3-cont)"
    ),
    # v3.5.2 scope 0 note: R-28 is implemented as a density custom-check in
    # qa_lint(), not a simple regex (needs occurrence count). See _check_tbc_density().
    # R-31 (v3.7.0 warn): contact_name == signatory_name (case-insensitive, trimmed).
    # Not a regex — evaluated as a custom check in qa_lint() like R-28.
    # Sentinel entry here so the rule ID appears in the registry.
}

# R-23 fabrication gate: regex targets material numeric-metric claims in
# Recital B. Every trigger must be matched against counterparty.source_map
# or marked [TBC] in the intake; otherwise R-23 FAILS the build.
_R23_CLAIM_PATTERN = re.compile(
    r"\b\d+[\d,]*\s*"
    r"(MW|GW|customers|clients|sites|deployments|GPUs|operations|"
    r"offices|countries|years|employees|%)\b",
    re.IGNORECASE,
)


# v3.5.6 scope D: diagnostic accumulator populated by _check_fabrication_gate
# on PASS paths so qa_lint can surface pillar-attribution info in the QA
# report even when no findings fire. Module-level list (cleared at start of
# each qa_lint run) keeps the return-type of _check_fabrication_gate stable.
_R23_PILLAR_DIAGNOSTIC: list = []


def _split_recital_b_sentences(text: str) -> list:
    """v3.5.6 scope D.2: split Recital B on sentence / paragraph boundaries.

    A `[TBC]` marker only covers claims in the same split segment (not the
    whole Recital B as v3.5.x did). Splits on `[.?!]` + whitespace, and on
    blank-line (`\\n\\n`) paragraph boundaries.
    """
    if not text:
        return []
    # Split on sentence-ending punctuation followed by whitespace, OR on blank lines
    parts = re.split(r"(?<=[.?!])\s+|\n\n+", text)
    return [p.strip() for p in parts if p.strip()]


def _claim_is_tbc_covered(claim_start: int, recital_b: str) -> bool:
    """v3.5.6 scope D.2: `[TBC]` covers claim only in the same sentence-
    boundary segment, OR when `[TBC]` sits at Recital-B-end and the claim
    is in the final segment (common trailing-`[TBC]` pattern).
    """
    if not recital_b:
        return False
    segments = _split_recital_b_sentences(recital_b)
    if not segments:
        return False
    # Rebuild an offset map to find which segment contains claim_start.
    # Walk segments through the original text, tracking running offset.
    pos = 0
    claim_segment_idx = None
    for i, seg in enumerate(segments):
        # Find this segment in the remaining text (from pos onwards)
        idx = recital_b.find(seg, pos)
        if idx < 0:
            continue
        seg_end = idx + len(seg)
        if idx <= claim_start <= seg_end:
            claim_segment_idx = i
            break
        pos = seg_end
    if claim_segment_idx is None:
        return False
    # Primary: [TBC] (or [TO BE CONFIRMED]) in the same segment
    seg_text = segments[claim_segment_idx]
    if "[TBC]" in seg_text or "[TO BE CONFIRMED]" in seg_text:
        return True
    # Special case (design decision D.2): [TBC] in final segment covers
    # final-segment claims (common trailing-[TBC] drafting pattern).
    if claim_segment_idx == len(segments) - 1 and (
        "[TBC]" in segments[-1] or "[TO BE CONFIRMED]" in segments[-1]
    ):
        return True
    return False


def _pillar_with_urls(source_map: dict):
    """Return the first pillar key (e.g. 'pillar_3') that has at least one
    URL entry, or None if no pillar has any. Used for the permissive any-
    pillar match and the diagnostic emission (v3.5.6 D.1).
    """
    if not isinstance(source_map, dict):
        return None
    # Preserve iteration order — typically pillar_1..pillar_5.
    for pillar_key, pillar_val in source_map.items():
        if isinstance(pillar_val, list) and any(
            isinstance(u, str) and (
                u.startswith(("http://", "https://"))
                or re.match(r"^internal:brochure_\d{8}_\w+$", u)
            )
            for u in pillar_val
        ):
            return pillar_key
        if isinstance(pillar_val, str) and (
            pillar_val.startswith(("http://", "https://"))
            or re.match(r"^internal:brochure_\d{8}_\w+$", pillar_val)
        ):
            return pillar_key
    return None


def certifications_in_source(intake: dict) -> list:
    """v3.7.0 R-11 helper: return list of certification strings detected in
    counterparty source material.

    v3.8.0: scans the assembled slot text (legal_identity + operational_verb
    object + customer_use_case category + material_asset asset + slot 5
    claim) instead of the removed `description` field.

    Detects ISO N{4,5} and common named certs (SOC 2, PCI-DSS, etc.).
    Phase 5 consumes this list when deciding include/omit for Recital B.
    Returns empty list when none found.
    """
    cp = intake.get("counterparty", {}) or {}
    rb = cp.get("recital_b") or {}
    parts: list[str] = []
    # Each slot may contribute text
    for slot_key in ("legal_identity", "operational_verb",
                     "customer_use_case", "material_asset"):
        s = rb.get(slot_key) or {}
        for v_key in ("legal_form", "verb", "object", "category", "asset"):
            v = s.get(v_key)
            if isinstance(v, str):
                parts.append(v)
    fact = rb.get("bargain_relevant_fact") or {}
    if isinstance(fact.get("claim"), str):
        parts.append(fact["claim"])
    # Slot-source quotes also count as source material
    for slot_key in ("operational_verb", "customer_use_case", "material_asset"):
        src = (rb.get(slot_key) or {}).get("source") or {}
        if isinstance(src.get("source_quote"), str):
            parts.append(src["source_quote"])
    # Fallback to legacy description if present (won't reach this in v3.8.0
    # because validate() rejects it; defensive).
    if cp.get("description"):
        parts.append(cp["description"])
    text = "\n".join(parts)
    found = []
    # ISO NNN patterns
    for m in re.finditer(r"\bISO\s*\d{4,5}\b", text, re.IGNORECASE):
        cert = m.group(0).replace("  ", " ").strip()
        if cert not in found:
            found.append(cert)
    # Named certs
    _NAMED_CERTS = re.compile(
        r"\b(SOC\s*[12]|PCI[- ]DSS|ISO\s*\d{4,5}|GDPR|HIPAA|FedRAMP|CSA\s*STAR)\b",
        re.IGNORECASE,
    )
    for m in _NAMED_CERTS.finditer(text):
        cert = m.group(0).strip()
        if cert not in found:
            found.append(cert)
    return found


# v3.7.0: brochure source_map token pattern. Tokens matching this pattern
# are accepted as tier-2 sources by R-23 (pass) and trigger R-24 (warn).
_BROCHURE_TOKEN_RE = re.compile(r"^internal:brochure_\d{8}_\w+$")


def _source_map_has_brochure(source_map: dict) -> bool:
    """Return True if any pillar value is an internal:brochure_* token."""
    if not isinstance(source_map, dict):
        return False
    for val in source_map.values():
        if isinstance(val, str) and _BROCHURE_TOKEN_RE.match(val):
            return True
        if isinstance(val, list):
            if any(isinstance(v, str) and _BROCHURE_TOKEN_RE.match(v) for v in val):
                return True
    return False


def _check_url_content(url: str, keyword: str, *, fetcher=None) -> bool:
    """v3.7.0 R-29 helper: fetch URL and check keyword appears in >=500 chars.

    Returns True (ok) if content is >=500 chars and keyword found.
    Returns False if content is short, keyword missing, or fetch fails.

    The `fetcher` argument is an object with a .fetch(url) -> str method.
    When None (production), uses urllib; in tests, pass a FakeFetcher.
    """
    if fetcher is not None:
        content = fetcher.fetch(url)
    else:
        try:
            import urllib.request as _req
            with _req.urlopen(url, timeout=10) as resp:
                content = resp.read(8192).decode("utf-8", errors="replace")
        except Exception:
            content = ""
    if len(content) < 500:
        return False
    return keyword.lower() in content.lower()


def _parse_flag(flag: str) -> bool:
    """v3.7.0: boolean CLI flag parser — returns True if flag in sys.argv."""
    return flag in sys.argv


# v3.5.6 scope D.3: thin-reason + structured-short-code patterns for
# override-reason validation. See _validate_override_reason().
_THIN_OVERRIDE_REASON_RE = re.compile(
    r"^\s*(?:ok|fine|yes|done|sure|good|n/?a|tbd)\s*$",
    re.IGNORECASE,
)
_STRUCTURED_OVERRIDE_REASON_RE = re.compile(
    r"^(?:OK|FINE|APPROVED|PREAPPROVED)-\d{4}-\d{2}-\d{2}\s+[A-Z]{2,4}$"
)


def _validate_override_reason(reason: str):
    """v3.5.6 scope D.3: hybrid override-reason validator.

    Accepts:
      - free-text rationale (≥15 chars after strip, NOT a thin-pattern match), OR
      - structured audit short-code: `<STATUS>-<YYYY-MM-DD> <INITIALS>`
        where STATUS in (OK, FINE, APPROVED, PREAPPROVED) and INITIALS is
        2–4 uppercase letters.

    Returns (is_valid: bool, error_message: str). error_message is empty
    when is_valid is True.
    """
    if reason is None:
        return False, "override reason required when --override is set"
    r = str(reason).strip()
    if not r:
        return False, "override reason cannot be empty"
    # Reject thin patterns regardless of length
    if _THIN_OVERRIDE_REASON_RE.match(r):
        return False, (
            "override reason is a thin-pattern phrase (e.g. 'ok', 'fine'). "
            "Provide either:\n"
            "  - free-text rationale (minimum 15 characters), OR\n"
            "  - structured audit short-code: <STATUS>-<YYYY-MM-DD> <INITIALS>\n"
            "    where STATUS in (OK, FINE, APPROVED, PREAPPROVED) and\n"
            "    INITIALS is 2-4 uppercase letters.\n"
            "    Example: OK-2026-04-17 JG"
        )
    # Accept structured short-code regardless of length
    if _STRUCTURED_OVERRIDE_REASON_RE.match(r):
        return True, ""
    # Accept free-text if length >= 15
    if len(r) >= 15:
        return True, ""
    return False, (
        f"override reason too short ({len(r)} chars). Provide either:\n"
        "  - free-text rationale (minimum 15 characters), OR\n"
        "  - structured audit short-code: <STATUS>-<YYYY-MM-DD> <INITIALS>\n"
        "    Example: OK-2026-04-17 JG"
    )


def _check_fabrication_gate(text_recital_b: str, source_map: dict,
                             overrides: set) -> list:
    """R-23: every material numeric claim in Recital B must be attributable.

    Returns a list of (rule_id, scope, message) findings. Empty list = pass.

    v3.5.6 scope D.1 + D.2 updates:
    - `[TBC]` proximity is sentence-boundary-scoped (not wildcard across
      whole Recital B). See _claim_is_tbc_covered().
    - On PASS, populates the module-level `_R23_PILLAR_DIAGNOSTIC` list
      with (claim_text, pillar_matched_or_'TBC-covered') tuples so qa_lint
      can surface the attribution info in the QA report.

    A claim is considered attributed if ANY of:
    - the claim is `[TBC]`-covered within the same sentence segment, OR
    - counterparty.source_map has at least one URL entry in ANY pillar
      (pillar-level granularity; reviewer's Phase 7.5 handles pillar-to-claim
      mapping), OR
    - the rule R-23 is in overrides (user ran with --override R-23 ...).
    """
    _R23_PILLAR_DIAGNOSTIC.clear()
    findings = []
    if "R-23" in overrides:
        return findings  # overridden
    claims = list(_R23_CLAIM_PATTERN.finditer(text_recital_b))
    if not claims:
        return findings  # no material claims

    pillar_key = _pillar_with_urls(source_map)

    # Walk each claim; if neither sentence-scoped [TBC] nor any pillar URL,
    # it's an unattributed claim. Collect diagnostic on the pass path.
    unattributed_claims = []
    for c in claims:
        claim_text = c.group(0)
        if _claim_is_tbc_covered(c.start(), text_recital_b):
            _R23_PILLAR_DIAGNOSTIC.append((claim_text, "TBC-covered"))
            continue
        if pillar_key:
            _R23_PILLAR_DIAGNOSTIC.append((claim_text, pillar_key))
            continue
        unattributed_claims.append(claim_text)

    if unattributed_claims:
        # Clear diagnostic on FAIL — reviewer focuses on the failures
        _R23_PILLAR_DIAGNOSTIC.clear()
        sample = unattributed_claims[:3]
        findings.append((
            "R-23", "Recital B",
            f"Fabrication gate: {len(unattributed_claims)} material claim(s) "
            f"in Recital B without pillar attribution or sentence-scoped "
            f"[TBC] marker (e.g. {', '.join(repr(s) for s in sample)}). "
            f"Add counterparty.source_map pillar_N: ['https://tier-1-url'] "
            f"to intake YAML, mark claims [TBC] within the same sentence, "
            f"or pass --override R-23 --override-reason \"<verbose>\"."
        ))
    return findings


def _extract_text(doc) -> str:
    """Extract all paragraph text from a python-docx Document, including tables."""
    lines = []
    for p in doc.paragraphs:
        lines.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    lines.append(p.text)
    return "\n".join(lines)


def qa_lint(doc, data: dict, builder_findings: list, overrides: set,
            override_reason: str, *, verify_urls=None, url_fetcher=None,
            builder_warnings=None):
    """Run QA rules over the built Document. Returns (status, report_lines).

    status: "PASS" | "PASS_WITH_WARN" | "FAIL"

    v3.7.2: R-29 URL content verification is now default-on. Behavior:
    - `verify_urls=None` (default) auto-detects from env (`LOI_NO_NETWORK=1`
      disables) and CLI (`--no-network` disables, `--verify-source-urls`
      explicitly enables).
    - `verify_urls=True|False` overrides the auto-detection (tests pass
      explicit values with a FakeFetcher).
    - `url_fetcher` injects a fake fetcher for tests; None = real urllib.
    """
    text = _extract_text(doc)
    lines = [
        "LOI QA Report",
        f"Counterparty: {data.get('counterparty', {}).get('short', 'Unknown')}",
        f"Generated: {datetime.now().isoformat()}Z",
        f"Variant: programme.recital_a_variant={data.get('programme', {}).get('recital_a_variant', 'default')}",
        f"Overrides: {','.join(sorted(overrides)) if overrides else 'none'}",
    ]
    if override_reason:
        lines.append(f"Override reason: {override_reason}")
    lines.append("")
    lines.append("Findings:")

    fail_count = 0
    warn_count = 0

    for rule, scope, msg in builder_findings:
        if rule in overrides:
            lines.append(f"  [OVRD] {rule}  {scope}   {msg}")
        else:
            lines.append(f"  [FAIL] {rule}  {scope}   {msg}")
            fail_count += 1

    # v3.7.2: emit engine-side WARN findings (e.g., silent-no-op mutations)
    for rule, scope, msg in (builder_warnings or []):
        lines.append(f"  [WARN] {rule}  {scope}   {msg}")
        warn_count += 1

    for rid, (pattern, scope, msg) in _FAIL_RULES.items():
        if re.search(pattern, text, re.MULTILINE):
            if rid in overrides:
                lines.append(f"  [OVRD] {rid}  {scope}   {msg}")
            else:
                lines.append(f"  [FAIL] {rid}  {scope}   {msg}")
                fail_count += 1

    for rid, (pattern, scope, msg) in _WARN_RULES.items():
        if re.search(pattern, text, re.IGNORECASE | re.MULTILINE):
            if rid == "R-11" and data.get("choices", {}).get("cert_relevant"):
                continue
            lines.append(f"  [WARN] {rid}  {scope}   {msg}")
            warn_count += 1

    # R-23 (v3.4) — fabrication gate on Recital B
    # v3.5.3-cont scope F: multi-paragraph Recital B extraction. Prior regex
    # stopped at the first blank line (`\n\n`), which partially scanned any
    # Recital B that legitimately spanned multiple paragraphs (rare but
    # possible for counterparties with consortium / holdco-subsidiary
    # disclosure needs). New regex extracts until the next recital marker
    # "(C)" or "(D)" or a section header (line starting with "## ") or EOF.
    recital_b_text = ""
    m = re.search(
        r"\(B\)\s*(.*?)(?=\(C\)\s|\(D\)\s|\n##\s|\Z)",
        text,
        re.DOTALL,
    )
    if m:
        recital_b_text = m.group(1)
    source_map = data.get("counterparty", {}).get("source_map", {})

    # v3.7.2: R-29 URL content verification — default-on, env-escape-hatched.
    # Fetch one URL per pillar, check Recital B keyword presence, downgrade
    # insufficient pillars before R-23 fabrication gate runs. Sits upstream
    # of R-23 so downgrades propagate into R-23's coverage check.
    if verify_urls is None:
        _no_net_env = os.environ.get("LOI_NO_NETWORK") == "1"
        _no_net_flag = "--no-network" in sys.argv
        _explicit_on = "--verify-source-urls" in sys.argv
        verify_urls = (not (_no_net_env or _no_net_flag)) or _explicit_on

    effective_source_map = source_map
    if verify_urls and source_map and "R-29" not in overrides:
        effective_source_map = dict(source_map)
        downgrades = []
        short_name = data.get("counterparty", {}).get("short", "") or \
                     data.get("counterparty", {}).get("name", "")

        for pillar_key, pillar_val in source_map.items():
            urls = pillar_val if isinstance(pillar_val, list) else [pillar_val]
            first_url = next(
                (u for u in urls
                 if isinstance(u, str) and u.startswith(("http://", "https://"))),
                None,
            )
            if not first_url:
                continue

            # Pick a representative keyword: counterparty short name is the
            # most load-bearing claim on every pillar; if Recital B is too
            # short to yield a meaningful keyword, fall back to the short name.
            _kw = short_name or "Digital Energy"

            try:
                ok = _check_url_content(first_url, _kw, fetcher=url_fetcher)
            except Exception:
                # Network error — log but don't downgrade (avoid flaky runs)
                lines.append(
                    f"  [INFO] R-29  source_map   {pillar_key}: "
                    f"URL fetch raised; skipping verification (re-check "
                    f"before signing)."
                )
                continue

            if not ok:
                downgrades.append((pillar_key, first_url, _kw))
                effective_source_map[pillar_key] = ["[TBC]"]

        for pillar_key, url, kw in downgrades:
            lines.append(
                f"  [WARN] R-29  source_map   {pillar_key}: URL {url!r} "
                f"returned <500 chars or missing keyword {kw!r} \u2014 pillar "
                f"downgraded to [TBC]. Re-verify + replace before signing."
            )
            warn_count += 1

    r23_findings = _check_fabrication_gate(recital_b_text, effective_source_map, overrides)
    for rid, scope, msg in r23_findings:
        lines.append(f"  [FAIL] {rid}  {scope}   {msg}")
        fail_count += 1

    # v3.5.6 scope D.1: on PASS, emit pillar-attribution diagnostic so the
    # QA report documents WHICH pillar matched each claim (or that it was
    # sentence-scoped [TBC]-covered). Makes the attribution auditable
    # without tightening the gate itself.
    if not r23_findings and _R23_PILLAR_DIAGNOSTIC:
        n = len(_R23_PILLAR_DIAGNOSTIC)
        lines.append(
            f"  [INFO] R-23  Recital B   attribution diagnostic: "
            f"{n} claim(s) matched"
        )
        for claim_text, pillar_or_tbc in _R23_PILLAR_DIAGNOSTIC[:5]:
            lines.append(
                f"         {pillar_or_tbc}: {claim_text[:80]}"
            )
        if n > 5:
            lines.append(f"         (and {n - 5} more — see fixture for full list)")

    # v3.5.6 scope D.3: log active overrides + validated reason + timestamp
    # in QA report so the audit trail captures "who bypassed what when".
    if overrides:
        from datetime import datetime as _dt, timezone as _tz
        lines.append(
            f"  [INFO] R-override  meta   overrides active: "
            f"{', '.join(sorted(overrides))}; reason: "
            f"{(override_reason or '<none>')[:200]}; logged: "
            f"{_dt.now(_tz.utc).isoformat()}"
        )

    # R-28 (v3.5.2) — [TBC] density check in body. A few [TBC] markers are
    # expected on drafts (signatory title, counterparty reg number pre-
    # signing); > 5 suggests the intake was not fully prepared and the draft
    # should loop back to Phase 4/5 for completion before external delivery.
    # v3.7.0: count BOTH [TBC] bracketed AND bare word TBC (excluding TBC
    # inside URLs). This normalises the threshold across all intake styles.
    if "R-28" not in overrides:
        tbc_count = text.count("[TBC]")
        # Count bare TBC not inside a URL (preceded by space/newline or start)
        tbc_count += len(re.findall(r"(?<![/\w])TBC(?![\w\]])", text))
        if tbc_count > 5:
            lines.append(
                f"  [WARN] R-28  body   [TBC] count ({tbc_count}) exceeds 5 "
                f"\u2014 intake likely incomplete; consider Phase 4/5 revision before external delivery"
            )
            warn_count += 1

    # R-31 (v3.7.0 warn): contact_name == signatory_name suggests a single
    # point of contact — unusual and worth confirming intentional.
    cp_data = data.get("counterparty", {}) or {}
    _contact = (cp_data.get("contact_name") or "").strip().lower()
    _signatory = (cp_data.get("signatory_name") or "").strip().lower()
    if _contact and _signatory and _contact == _signatory and "R-31" not in overrides:
        lines.append(
            "  [WARN] R-31  counterparty   contact_name equals signatory_name "
            "— confirm intentional single-point-of-contact arrangement."
        )
        warn_count += 1

    # R-24 (v3.7.0 warn): brochure-sourced pillars require tier-1 corroboration.
    # This is the NEW R-24 for brochure tokens. The old R-24 (inline citation
    # in Recital B prose) lives in _FAIL_RULES as a regex rule under "R-24".
    # We use key "R-24B" to avoid ID collision while staying in the same family.
    _source_map = data.get("counterparty", {}).get("source_map", {}) or {}
    if _source_map_has_brochure(_source_map) and "R-24" not in overrides:
        lines.append(
            "  [WARN] R-24  source_map   Brochure-sourced material claims require "
            "tier-1 public corroboration before signing. "
            "Replace internal:brochure_* tokens with https:// URLs before external delivery."
        )
        warn_count += 1

    # R-21 scope narrowing (v3.7.0): 'purpose-built' is allowed inside Clause 3
    # product-capability paragraphs. Strip Cl. 3 text before R-21 regex scan
    # so the rule only fires outside Cl. 3.
    # The _WARN_RULES R-21 regex was already applied above. If it fired but the
    # match is ONLY inside Cl. 3, remove the false-positive warning.
    _cl3_text = ""
    _cl3_match = re.search(
        r"(?:^|\n)3\.\s+[A-Z].*?(?=\n[4-9]\.\s+|\Z)", text, re.DOTALL
    )
    if _cl3_match:
        _cl3_text = _cl3_match.group(0)
    _purpose_built_outside_cl3 = False
    for _pm in re.finditer(r"\b(purpose-built|state-of-the-art)\b", text, re.IGNORECASE):
        _match_pos = _pm.start()
        if _cl3_text:
            _cl3_start = text.find(_cl3_text)
            _cl3_end = _cl3_start + len(_cl3_text) if _cl3_start >= 0 else -1
            if _cl3_start >= 0 and _cl3_start <= _match_pos < _cl3_end:
                continue  # inside Cl. 3 — skip
        _purpose_built_outside_cl3 = True
        break
    # If R-21 fired but all matches are in Cl. 3, retroactively remove the warn line
    if not _purpose_built_outside_cl3:
        _new_lines = [l for l in lines if "R-21" not in l or "[WARN]" not in l]
        _removed = len(lines) - len(_new_lines)
        warn_count -= _removed
        lines = _new_lines

    if not data.get("programme", {}).get("recital_a_variant"):
        lines.append("  [INFO] R-16  YAML   Recital A variant not set, used 'default'")
    if not data.get("choices", {}).get("bespoke_closing"):
        lines.append("  [INFO] R-17  YAML   No bespoke_closing, used default single-sentence")

    # v3.7.2: custom.clauses silent no-op surfacing — when replace or
    # insert-after targets weren't found in the rendered body, report it.
    # (Common cause: operator typo'd the clause number or the target clause
    # doesn't exist in this LOI type.)
    _lb_cm_failures = []
    # `builder_findings` carries engine-emitted findings keyed by rule id;
    # we attach mutation failures via the loi object when qa_lint is called
    # from main(). Fall back gracefully if not available.
    try:
        # Best-effort: pull from the build-time LOI object if the caller
        # threaded it through. Current signature doesn't, but we can find
        # the failures list on any attribute the user might have plumbed in.
        # For now we don't have access — skip. This hook is populated
        # properly via the SESSION_LOG + qa_lint call in main().
        pass
    except Exception:
        pass

    # v3.7.2: lead_time unparseable advisory
    _lt_state = _LAST_LEAD_TIME_PARSE
    _lt_input = _lt_state.get("input")
    if _lt_input and not _lt_state.get("ok"):
        lines.append(
            f"  [WARN] R-lead-time  supplier   supplier.lead_time_target "
            f"{_lt_input!r} could not be parsed as days/weeks/months "
            f"(accepted: '90 days', '6 weeks', '3 months', '90d', '6w', '3m'). "
            f"Joint Stocking Programme clause did NOT fire — consider "
            f"reformatting the value."
        )
        warn_count += 1

    # v3.7.2: Recital B density advisory — measure actual word count against
    # the chosen density band and emit INFO when out of band. Engine does not
    # auto-edit description; this tells operator their density choice is
    # inconsistent with their supplied text.
    _density = data.get("choices", {}).get("recital_b_density", "standard")
    _recital_b_wc = len(re.findall(r"\b\w+\b", recital_b_text or ""))
    if _recital_b_wc > 0:
        _bands = {
            "terse": (40, 100, "~80w logo-drop"),
            "standard": (60, 150, "~120w default"),
            "verbose": (110, 200, "~150w contextual"),
        }
        low, high, label = _bands.get(_density, (0, 10_000, ""))
        if _recital_b_wc < low or _recital_b_wc > high:
            lines.append(
                f"  [INFO] R-density  Recital B   word count {_recital_b_wc} "
                f"outside {_density!r} band ({low}-{high}, {label}). "
                f"Either adjust description length or change "
                f"choices.recital_b_density."
            )

    # v3.7.0: surface structured metadata from the intake — relationship_cluster,
    # identity_map, and financing_context do not appear in the LOI body but
    # are valuable for the QA audit trail and the cover-email cross-check.
    _rel = data.get("counterparty", {}).get("relationship_cluster")
    if _rel:
        _gid = _rel.get("group_id", "<unnamed>")
        _aff = _rel.get("affiliated_entities", []) or []
        lines.append(
            f"  [INFO] V3-7-meta  counterparty   relationship_cluster: "
            f"group_id={_gid}, affiliated_entities={len(_aff)}"
        )
        for entity in _aff[:3]:
            lines.append(
                f"         - {entity.get('name', '<unnamed>')}: "
                f"{entity.get('role', '<role unspecified>')}"
            )

    _idmap = data.get("counterparty", {}).get("identity_map")
    if _idmap:
        lines.append(
            f"  [INFO] V3-7-meta  counterparty   identity_map: "
            f"{len(_idmap)} person(s) tracked across multiple domains"
        )

    _fctx = data.get("dates", {}).get("financing_context")
    if _fctx:
        _linked = _fctx.get("linked_to_fundraise", False)
        _close = _fctx.get("fundraise_close_target", "<unset>")
        _buffer = _fctx.get("buffer_months_post_close", "<unset>")
        lines.append(
            f"  [INFO] V3-7-meta  dates   financing_context: "
            f"linked={_linked}, close_target={_close}, buffer_months={_buffer}"
        )

    # v3.7.0: certifications_in_source (R-11 helper) — surface the detected
    # list so Phase 5 can make the include/omit decision explicitly.
    _certs = certifications_in_source(data)
    if _certs:
        lines.append(
            f"  [INFO] R-11  source   certifications_detected: {', '.join(_certs)}"
        )
        lines.append(
            "         Phase 5 decision point: include in Recital B (differentiator) "
            "or omit (conservative). Default: omit."
        )

    lines.append("")
    if fail_count > 0:
        status = "FAIL"
    elif warn_count > 0:
        status = "PASS_WITH_WARN"
    else:
        status = "PASS"
    lines.append(f"Status: {status} (warnings: {warn_count}, failures: {fail_count})")
    return status, lines


def _parse_arg(flag: str) -> str:
    if flag in sys.argv:
        idx = sys.argv.index(flag)
        if idx + 1 < len(sys.argv):
            return sys.argv[idx + 1]
    return ""


def _migrate_check(path: str) -> int:
    """v3.5.3 scope K: legacy intake YAML migration pre-flight.

    Inspects intake YAML for missing `counterparty.source_map` (required by
    v3.4 R-23 fabrication gate). If absent or empty, emits a ready-to-paste
    snippet with all 5 pillars marked `[TBC]` and exits 0 (non-blocking —
    the user copies the snippet, pastes into their YAML, and re-runs the
    generator). If `source_map` is already present, reports that and exits 0.

    Uses raw yaml.safe_load (bypasses full validate()) so legacy YAMLs that
    would fail other v3.4/v3.5 validators can still be migration-checked.
    """
    try:
        with open(path, "r") as f:
            raw = yaml.safe_load(f) or {}
    except Exception as e:
        print(f"[--migrate-check] Could not read {path}: {e}")
        return 0
    cp = raw.get("counterparty", {}) or {}
    source_map = cp.get("source_map")
    if source_map and isinstance(source_map, dict) and any(source_map.values()):
        print(f"[--migrate-check] {path}")
        print("  OK: counterparty.source_map present with entries.")
        print("  R-23 fabrication gate will evaluate URL attribution per pillar.")
        return 0
    # Missing or empty — emit snippet
    print(f"[--migrate-check] {path}")
    print("  counterparty.source_map NOT SET — legacy v3.3 intake.")
    print("  Paste the following into your YAML under `counterparty:`:")
    print()
    print("  source_map:")
    print("    pillar_1: \"[TBC]\"   # Identity & scale — own website, registry")
    print("    pillar_2: \"[TBC]\"   # Core business — own website")
    print("    pillar_3: \"[TBC]\"   # Track record / proof points — named customer press")
    print("    pillar_4: \"inferred from Phase 1 context\"")
    print("    pillar_5: \"[TBC]\"   # Forward plans — only if anchored to named third party")
    print()
    print("  Replace [TBC] with tier-1 URLs where available. See")
    print("  _shared/counterpart-description-framework.md for the Signal")
    print("  Test and tier hierarchy policy.")
    return 0


# v3.5.6 scope G: Phase 7.5 fail-closed enforcement — sentinel file with
# SHA-256 of the .docx being blessed. Consumed on --phase-7-5-pass after
# verifying hash matches current .docx (prevents replay + post-approval
# tampering). Opt-in via --enforce-phase-7-5 flag OR env var
# DE_LOI_ENFORCE_PHASE_7_5=1 (OR'd). Default: fail-open (v3.5.x preserved).

_PHASE_7_5_SENTINEL_SUFFIX = ".phase_7_5_required"


def _phase_7_5_enforce_enabled() -> bool:
    """v3.5.6 G.2: enforcement activates via CLI flag OR env variable."""
    if "--enforce-phase-7-5" in sys.argv:
        return True
    val = os.environ.get("DE_LOI_ENFORCE_PHASE_7_5", "").strip().lower()
    return val in ("1", "true", "yes", "on")


def _docx_sha256(docx_path: str) -> str:
    """Return hex SHA-256 of .docx file contents."""
    import hashlib
    h = hashlib.sha256()
    with open(docx_path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def _phase_7_5_sentinel_path(docx_path: str) -> str:
    return docx_path + _PHASE_7_5_SENTINEL_SUFFIX


def _write_phase_7_5_sentinel(docx_path: str) -> str:
    """v3.5.6 G.1: write sentinel blessing a specific .docx content hash.

    Written alongside the .docx as `<docx>.phase_7_5_required`. Contents
    include SHA-256 (for tamper-detection on --phase-7-5-pass), ISO
    timestamp, and a HOW TO RESOLVE block so the error is self-documenting.
    """
    sentinel_path = _phase_7_5_sentinel_path(docx_path)
    docx_sha = _docx_sha256(docx_path)
    ts = datetime.now(timezone.utc).isoformat()
    with open(sentinel_path, "w", encoding="utf-8") as f:
        f.write(
            f"# Phase 7.5 review required\n"
            f"# Auto-generated by generate_loi.py ({ts})\n"
            f"#\n"
            f"docx: {os.path.basename(docx_path)}\n"
            f"docx_sha256: {docx_sha}\n"
            f"created: {ts}\n"
            f"\n"
            f"# HOW TO RESOLVE:\n"
            f"#   1. Load the Phase 7.5 callee workflow in your Claude session:\n"
            f"#      legal-counsel/specializations/contract-review/loi-review-workflow.md\n"
            f"#   2. Run the 4-point structured review against this .docx\n"
            f"#      (clause-type appropriateness, meta-commentary scan,\n"
            f"#       cross-clause consistency, source-verification sample)\n"
            f"#   3. If the review returns PASS, re-run the generator with\n"
            f"#      --phase-7-5-pass. The sentinel will be consumed\n"
            f"#      (verifying the .docx hash hasn't changed) and delivery\n"
            f"#      will be permitted.\n"
            f"#   4. If review returns FLAG-FOR-REVISION or REJECT, edit the\n"
            f"#      intake YAML and re-run the generator (a new sentinel\n"
            f"#      with a fresh hash will be written).\n"
            f"#\n"
            f"# DO NOT edit this file or the .docx between Phase 7.5 review\n"
            f"# and --phase-7-5-pass — the hash check will detect any change.\n"
        )
    return sentinel_path


def _consume_phase_7_5_sentinel(docx_path: str):
    """v3.5.6 G.1: verify sentinel's hash matches current .docx, then
    consume (delete) the sentinel. Returns (ok, err) — on ok, sentinel
    is removed and caller may proceed with delivery.
    """
    sentinel_path = _phase_7_5_sentinel_path(docx_path)
    if not os.path.exists(sentinel_path):
        return False, (
            f"Phase 7.5 sentinel not found at {sentinel_path}. "
            f"Either Phase 7.5 was not run for this .docx, or the sentinel "
            f"was consumed by a prior --phase-7-5-pass invocation. Regenerate "
            f"the .docx first (which writes a fresh sentinel under "
            f"enforcement mode)."
        )
    import re as _re
    with open(sentinel_path, "r", encoding="utf-8") as f:
        sentinel_content = f.read()
    m = _re.search(r"docx_sha256:\s*([0-9a-fA-F]{64})", sentinel_content)
    if not m:
        return False, (
            f"Phase 7.5 sentinel at {sentinel_path} is malformed "
            f"(no valid docx_sha256 line). Regenerate."
        )
    sentinel_sha = m.group(1).lower()
    current_sha = _docx_sha256(docx_path)
    if sentinel_sha != current_sha:
        return False, (
            f"Phase 7.5 sentinel hash mismatch: the .docx was modified "
            f"after Phase 7.5 approval.\n"
            f"  sentinel blessed: {sentinel_sha[:16]}...\n"
            f"  current .docx:    {current_sha[:16]}...\n"
            f"Re-run Phase 7.5 review against the current .docx."
        )
    # Consume
    os.remove(sentinel_path)
    return True, ""


def main():
    if len(sys.argv) < 2:
        print("Usage: python generate_loi.py <intake.yaml>")
        print("  [--output path.docx]")
        print("  [--override R-11,R-14]")
        print("  [--override-reason \"...\"]")
        print("  [--migrate-check]          (v3.5.3: inspect YAML for v3.4 source_map; non-blocking)")
        print("  [--enforce-phase-7-5]      (v3.5.6: opt-in Phase 7.5 fail-closed gate)")
        print("  [--phase-7-5-pass]         (v3.5.6: consume Phase 7.5 sentinel after review)")
        print("")
        print("  Env: DE_LOI_ENFORCE_PHASE_7_5=1  also activates Phase 7.5 enforcement")
        sys.exit(1)

    # v3.5.3 scope K: --migrate-check runs before load_intake so legacy
    # YAMLs that would fail full validation can still be inspected.
    if "--migrate-check" in sys.argv:
        sys.exit(_migrate_check(sys.argv[1]))

    # v3.7.0: --audit-only — read prior_loi_path (or --prior) and run linter
    if "--audit-only" in sys.argv:
        prior_path = _parse_arg("--prior") or sys.argv[1]
        _audit_only_mode(prior_path)
        sys.exit(0)

    # v3.7.0: --phase-8-auto-execute — stored/accepted; Specialist C wires the action
    _phase8_auto = "--phase-8-auto-execute" in sys.argv

    # v3.7.0: --verify-source-urls — activates R-29 URL content verification
    _verify_urls = "--verify-source-urls" in sys.argv

    data = load_intake(sys.argv[1])

    override_str = _parse_arg("--override")
    if override_str:
        data["_overrides"] = [r.strip() for r in override_str.split(",") if r.strip()]
    override_reason = _parse_arg("--override-reason")
    # v3.5.6 scope D.3: validate override-reason shape when --override is used.
    # Hybrid rule: accept free-text ≥15 chars OR structured short code
    # `<STATUS>-<YYYY-MM-DD> <INITIALS>`. Reject thin patterns unconditionally.
    if override_str:
        ok, err = _validate_override_reason(override_reason)
        if not ok:
            print(f"ERROR (--override-reason): {err}", file=sys.stderr)
            sys.exit(1)
    if override_reason:
        data["_override_reason"] = override_reason

    output = _parse_arg("--output")

    if not output:
        loi_date = data.get("dates", {}).get("loi_date", "")
        try:
            dt = datetime.strptime(loi_date, "%d %B %Y")
            date_str = dt.strftime("%Y%m%d")
        except ValueError:
            date_str = datetime.now().strftime("%Y%m%d")
        cp = data.get("counterparty", {}).get("short", "Unknown")
        output = f"{date_str}_DEG_LOI-{data['type']}_{cp}_(DRAFT).docx"

    # v3.6.0 item b: auto-version output filenames. If target exists,
    # append _v{N} and increment until unique. Addresses file-confusion
    # observed in Cerebro/Armada/InfraPartners retrospectives where
    # operators regenerated iteratively and lost track of canonical draft.
    if os.path.exists(output):
        base, ext = os.path.splitext(output)
        n = 2
        while os.path.exists(f"{base}_v{n}{ext}"):
            n += 1
        output = f"{base}_v{n}{ext}"
        print(f"[auto-version] Target exists; writing to {output}", file=sys.stderr)

    loi = LOI(data)
    doc = loi.build()

    qa_report_path = output.replace(".docx", "_qa.txt")
    # v3.7.2: build the builder_warnings list from engine-side advisories
    # (currently: custom.clauses silent-no-op failures)
    _builder_warnings = [
        ("R-custom-mut", "custom.clauses",
         f"{mode} mode targeting {target!r}: {reason}. "
         f"The mutation was not applied; check intake for typos "
         f"or wrong-type targets.")
        for mode, target, reason in getattr(loi, "_custom_mutation_failures", [])
    ]
    qa_status, qa_lines = qa_lint(doc, data, loi.qa_findings, loi.overrides,
                                   loi.override_reason,
                                   builder_warnings=_builder_warnings)
    with open(qa_report_path, "w", encoding="utf-8") as f:
        f.write("\n".join(qa_lines) + "\n")

    if qa_status == "FAIL":
        print("QA FAIL — build blocked. Findings:")
        for line in qa_lines:
            if "[FAIL]" in line:
                print(f"  {line}")
        print(f"\nFull QA report: {qa_report_path}")
        print("To override: --override R-xx,R-yy --override-reason \"...\"")
        sys.exit(2)

    # v3.7.0: --recital-b-only — replace Recital B paragraph in existing .docx
    recital_b_only_path = _parse_arg("--recital-b-only")
    if recital_b_only_path:
        _recital_b_only_replace(recital_b_only_path, loi, data)
        print(f"[--recital-b-only] Recital B replaced in {recital_b_only_path}")

    doc.save(output)

    # v3.5.6 scope G: Phase 7.5 fail-closed enforcement. Only active when
    # opt-in (CLI flag OR env var). Default path (fail-open) preserves
    # v3.5.x behaviour for existing workflows.
    if _phase_7_5_enforce_enabled():
        if "--phase-7-5-pass" in sys.argv:
            ok, err = _consume_phase_7_5_sentinel(output)
            if not ok:
                print(
                    f"Phase 7.5 enforcement: DELIVERY BLOCKED\n\n{err}",
                    file=sys.stderr,
                )
                sys.exit(3)
            # Sentinel consumed cleanly — fall through to success print
        else:
            sentinel_path = _write_phase_7_5_sentinel(output)
            print(
                f"Generated: {output}\n"
                f"Type: DE-LOI-{data['type']}-v3.2\n"
                f"Counterparty: {data.get('counterparty', {}).get('name', 'Unknown')}\n"
                f"QA: {qa_status} ({qa_report_path})\n"
                f"\n"
                f"⚠ PHASE 7.5 REVIEW REQUIRED — delivery blocked.\n"
                f"\n"
                f"Phase 7.5 enforcement is enabled. The .docx has been written\n"
                f"but delivery is gated on legal-counsel review.\n"
                f"\n"
                f"Next step:\n"
                f"  1. Load the callee workflow in your Claude session:\n"
                f"     legal-counsel/specializations/contract-review/loi-review-workflow.md\n"
                f"  2. Run the 4-point structured review against {output}\n"
                f"  3. If review returns PASS, re-run with --phase-7-5-pass\n"
                f"\n"
                f"Sentinel written: {sentinel_path}"
            )
            sys.exit(3)

    print(f"Generated: {output}")
    print(f"Type: DE-LOI-{data['type']}-v3.2")
    print(f"Counterparty: {data.get('counterparty', {}).get('name', 'Unknown')}")
    print(f"QA: {qa_status} ({qa_report_path})")
    print(f"Clauses: ALL (full document)")

    # v3.7.0: emit SESSION_LOG.md alongside the .docx
    _emit_session_log(output, data, qa_status, qa_lines)
    # v3.7.1: emit 84-item audit checklist alongside the .docx
    _emit_audit_checklist(output, doc, data)
    # v3.7.2: Phase 8 auto-execute — runs selected actions. Local-only
    # actions (artifact_storage_push, domain_card_create) execute inline;
    # MCP actions (HubSpot, ClickUp) emit a JSON dispatch file for the
    # orchestrator session to consume. Action selection via
    # `--phase-8-actions=key1,key2,...`; dispatch shape controlled by
    # `--phase-8-auto-execute` (default dry_run=True).
    _run_phase_8_wiring(output, data)


def _audit_only_mode(prior_loi_path: str):
    """v3.7.0 --audit-only: extract text from a prior .docx and run the full
    linter, emitting a {basename}_audit.txt compliance-delta report.
    """
    import importlib
    try:
        import docx as _docx_module
    except ImportError:
        print("[--audit-only] python-docx required. Install: pip install python-docx",
              file=sys.stderr)
        return
    if not os.path.exists(prior_loi_path):
        print(f"[--audit-only] Path not found: {prior_loi_path}", file=sys.stderr)
        return
    from docx import Document as _Doc
    _doc = _Doc(prior_loi_path)
    _text_lines = [p.text for p in _doc.paragraphs if p.text]
    for tbl in _doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                _text_lines += [p.text for p in cell.paragraphs if p.text]
    text = "\n".join(_text_lines)

    # Run rules over raw text with empty data (no YAML context)
    findings = []
    for rid, (pattern, scope, msg) in _FAIL_RULES.items():
        if re.search(pattern, text, re.MULTILINE):
            findings.append(f"  [FAIL] {rid}  {scope}   {msg}")
    for rid, (pattern, scope, msg) in _WARN_RULES.items():
        if re.search(pattern, text, re.IGNORECASE | re.MULTILINE):
            findings.append(f"  [WARN] {rid}  {scope}   {msg}")

    base = os.path.splitext(os.path.basename(prior_loi_path))[0]
    out_dir = os.path.dirname(os.path.abspath(prior_loi_path))
    audit_path = os.path.join(out_dir, f"{base}_audit.txt")
    report = [
        f"Audit Report — {os.path.basename(prior_loi_path)}",
        f"Generated: {datetime.now().isoformat()}Z",
        f"Rules: {len(_FAIL_RULES)} fail-rules, {len(_WARN_RULES)} warn-rules",
        "",
        "Findings:",
        *(findings if findings else ["  (none — prior LOI passes current linter)"]),
        "",
        f"Total findings: {len(findings)}",
    ]
    with open(audit_path, "w", encoding="utf-8") as f:
        f.write("\n".join(report) + "\n")
    print(f"[--audit-only] Audit written to {audit_path}")
    for line in findings:
        print(line)


def _recital_b_only_replace(prior_path: str, loi, data: dict):
    """v3.7.0 --recital-b-only: replace Recital B paragraph in an existing
    .docx. Writes a new versioned file (_v{N}.docx) alongside the original.

    Errors clearly when prior_path is missing.
    """
    if not prior_path:
        print("[--recital-b-only] ERROR: path argument required after flag.",
              file=sys.stderr)
        return
    if not os.path.exists(prior_path):
        print(f"[--recital-b-only] ERROR: path not found: {prior_path}",
              file=sys.stderr)
        return
    from docx import Document as _Doc
    _doc = _Doc(prior_path)

    # Find the (B) recital paragraph
    new_b_text = ""
    for p in loi.doc.paragraphs:
        if p.text.startswith("(B) "):
            new_b_text = p.text
            break
    if not new_b_text:
        print("[--recital-b-only] Could not locate (B) Recital in generated LOI.",
              file=sys.stderr)
        return

    for p in _doc.paragraphs:
        if p.text.startswith("(B) "):
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = new_b_text
            break

    base, ext = os.path.splitext(prior_path)
    n = 2
    while os.path.exists(f"{base}_v{n}{ext}"):
        n += 1
    out_path = f"{base}_v{n}{ext}"
    _doc.save(out_path)
    print(f"[--recital-b-only] Written: {out_path}")


def _run_phase_8_wiring(docx_path: str, data: dict) -> None:
    """v3.7.2: Phase 8 auto-execute wiring — invoked from main() post-save.

    Reads CLI flags:
      --phase-8-actions=<comma-separated>  # defaults to all 5 local+MCP
      --phase-8-auto-execute               # flips dry_run to False

    Local-file actions (artifact_storage_push, domain_card_create) execute
    inline when dry_run=False. MCP actions (hubspot_upsert_company,
    clickup_create_task) emit a JSON dispatch file
    (`{stem}_PHASE8_DISPATCH.json`) which the Claude Code orchestrator
    session reads + invokes via the real MCP tool calls. This split keeps
    the generator CLI runnable in non-Claude environments (it can write
    files, but it can't invoke MCP tools — those are session-scoped).

    A `{stem}_PHASE8_DISPATCH.json` is ALWAYS emitted (even in dry-run)
    so operators can inspect the planned payloads before running for real.
    """
    import json as _json

    # Parse action selection
    actions_arg = _parse_arg("--phase-8-actions")
    if actions_arg:
        actions = [a.strip() for a in actions_arg.split(",") if a.strip()]
    else:
        # default: local-only (safer: never dispatches MCP writes without explicit opt-in)
        actions = ["artifact_storage_push", "domain_card_create"]
        if "--phase-8-auto-execute" in sys.argv:
            # Full opt-in also includes MCP dispatch
            actions = [
                "hubspot_upsert_company",
                "clickup_create_task",
                "artifact_storage_push",
                "domain_card_create",
                "cover_email_cross_check",
            ]

    auto_execute = "--phase-8-auto-execute" in sys.argv
    dry_run = not auto_execute

    # Import the dispatcher module
    try:
        sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        from scripts import phase8_actions as _p8
        from scripts import artifact_storage as _artstore
    except Exception as exc:
        print(f"[phase-8] scripts module unavailable: {exc!r} — skipping", file=sys.stderr)
        return

    result = _p8.run_phase_8_actions(data, docx_path, actions, dry_run=dry_run)

    # Write the dispatch JSON
    stem = os.path.splitext(os.path.basename(docx_path))[0]
    parent = os.path.dirname(os.path.abspath(docx_path))
    dispatch_path = os.path.join(parent, f"{stem}_PHASE8_DISPATCH.json")
    with open(dispatch_path, "w", encoding="utf-8") as f:
        _json.dump(result, f, indent=2, default=str)

    # Execute local-file actions inline when opted in
    executed_local = []
    if auto_execute:
        for action_result in result["actions"]:
            key = action_result.get("action", "")
            if key == "artifact_storage_push":
                try:
                    target = _artstore.upload_artifact(docx_path, data, dry_run=False)
                    executed_local.append(("artifact_storage_push", target))
                except Exception as exc:
                    executed_local.append(("artifact_storage_push", f"error: {exc!r}"))
            elif key == "domain_card_create":
                try:
                    card_path = _write_domain_card(docx_path, data)
                    executed_local.append(("domain_card_create", card_path))
                except Exception as exc:
                    executed_local.append(("domain_card_create", f"error: {exc!r}"))

    # Print a concise report to stdout so operators see what happened
    if auto_execute:
        print(f"Phase 8: dispatch written to {dispatch_path}")
        for key, target in executed_local:
            print(f"  executed: {key} \u2192 {target}")
        if not executed_local:
            print("  (no local actions executed; MCP actions await orchestrator)")
    else:
        print(f"Phase 8: dry-run dispatch written to {dispatch_path}")


def _write_domain_card(docx_path: str, data: dict) -> str:
    """v3.7.2: real implementation of domain card creation (E1).

    Writes `/domains/counterparties/{slug}/overview.md` (repo-relative)
    with a structured template. The base path defaults to a sibling
    `domains/` directory under the current working directory; operators
    in the DEGitOS repo work out of the project root where this path
    resolves correctly. In other contexts (CI, ad-hoc), writes to a
    fallback under the docx's directory.
    """
    # Reuse artifact_storage slug helper
    from scripts.artifact_storage import _slugify

    cp = data.get("counterparty", {}) or {}
    slug = _slugify(cp.get("name", ""))
    loi_type = data.get("type", "Unknown")

    # Prefer DEGitOS repo-root /domains/ if present; else fallback next to docx
    docx_parent = os.path.dirname(os.path.abspath(docx_path))
    repo_root_guess = os.getcwd()
    primary = os.path.join(repo_root_guess, "domains", "counterparties", slug)
    fallback = os.path.join(docx_parent, "domain_cards", slug)

    if os.path.isdir(os.path.join(repo_root_guess, "domains")):
        target_dir = primary
    else:
        target_dir = fallback

    os.makedirs(target_dir, exist_ok=True)
    card_path = os.path.join(target_dir, "overview.md")

    # Don't overwrite an existing card — append a version suffix instead
    if os.path.exists(card_path):
        from datetime import datetime as _dt
        stamp = _dt.now().strftime("%Y%m%d_%H%M%S")
        card_path = os.path.join(target_dir, f"overview_v3.7.2_{stamp}.md")

    # Assemble the card
    sup = data.get("supplier", {}) or {}
    rel = cp.get("relationship_cluster") or {}
    idm = cp.get("identity_map") or {}
    fctx = data.get("dates", {}).get("financing_context") or {}

    content_lines = [
        f"# {cp.get('name', slug)}",
        "",
        f"- **Type (this LOI):** {loi_type}",
        f"- **LOI path:** {docx_path}",
        f"- **LOI date:** {data.get('dates', {}).get('loi_date', '[TBC]')}",
        f"- **Validity:** {data.get('dates', {}).get('validity_date', '[TBC]')}",
        f"- **Signatory (counterparty):** {cp.get('signatory_name', '[TBC]')}",
        f"- **Contact (counterparty):** {cp.get('contact_name', '[TBC]')}",
        f"- **Short name:** {cp.get('short', slug)}",
        f"- **Jurisdiction:** {cp.get('jurisdiction', '[TBC]')}",
        f"- **Registration:** {cp.get('reg_type', '[TBC]')} {cp.get('reg_number', '[TBC]')}",
        "",
        "## HubSpot / ClickUp / Drive links",
        "",
        "- HubSpot company: [pending — populated by Phase 8 MCP dispatch]",
        "- HubSpot deal: [pending]",
        "- ClickUp task: [pending]",
        "- Drive audit folder: [pending]",
        "",
    ]

    if rel:
        content_lines += [
            "## Relationship cluster",
            "",
            f"- **Group ID:** {rel.get('group_id', '[unnamed]')}",
            f"- **Primary entity:** {rel.get('primary_entity', cp.get('name', ''))}",
        ]
        aff = rel.get("affiliated_entities", []) or []
        if aff:
            content_lines.append("- **Affiliated entities:**")
            for e in aff:
                content_lines.append(
                    f"  - {e.get('name', '[unnamed]')} "
                    f"\u2014 {e.get('role', '[role unspecified]')}"
                )
        content_lines.append("")

    if idm:
        content_lines += [
            "## Identity map",
            "",
        ]
        for key, person in idm.items():
            domains = ", ".join(
                f"{d.get('domain', '')} ({d.get('role', '')})"
                for d in (person.get("email_domains") or [])
            )
            content_lines.append(
                f"- **{person.get('display_name', key)}** "
                f"\u2014 preferred domain: "
                f"{person.get('preferred_email_for_this_loi', '[TBC]')}"
            )
            if domains:
                content_lines.append(f"  - Known domains: {domains}")
        content_lines.append("")

    if fctx:
        content_lines += [
            "## Financing context",
            "",
            f"- Linked to fundraise: {fctx.get('linked_to_fundraise', False)}",
            f"- Fundraise close target: {fctx.get('fundraise_close_target', '[TBC]')}",
            f"- Buffer months post-close: {fctx.get('buffer_months_post_close', '[TBC]')}",
            "",
        ]

    content_lines += [
        "## Status",
        "",
        "- LOI: DRAFT (as of emission)",
        "- Next step: legal-counsel Phase 7.5 review + cover-email drafting",
        "",
    ]

    with open(card_path, "w", encoding="utf-8") as f:
        f.write("\n".join(content_lines) + "\n")

    return card_path


def _emit_session_log(docx_path: str, data: dict, qa_status: str,
                      qa_lines: list) -> str:
    """v3.7.0: emit a SESSION_LOG.md file alongside the generated .docx.

    Captures all non-default intake decisions, active CLI flags, QA summary,
    and custom definition/clause counts for session auditability.
    Returns the path to the written file.
    """
    stem = os.path.splitext(os.path.basename(docx_path))[0]
    log_dir = os.path.dirname(os.path.abspath(docx_path))
    log_path = os.path.join(log_dir, f"{stem}_SESSION_LOG.md")

    from datetime import datetime as _dt, timezone as _tz
    ts = _dt.now(_tz.utc).isoformat()
    cp = data.get("counterparty", {}) or {}
    choices = data.get("choices", {}) or {}
    custom = data.get("custom", {}) or {}

    # Count R-rules triggered
    triggered = [l.split()[1] for l in qa_lines if l.strip().startswith("[FAIL]") or
                 l.strip().startswith("[WARN]")]
    fail_count = sum(1 for l in qa_lines if "[FAIL]" in l)
    warn_count = sum(1 for l in qa_lines if "[WARN]" in l)
    info_count = sum(1 for l in qa_lines if "[INFO]" in l)

    # Non-default choices
    DEFAULTS = {
        "recital_b_density": "standard",
        "joint_ip": None,
        "bespoke_closing": None,
        "cert_relevant": None,
        "existing_nda": None,
        "include_schedule": None,
    }
    intake_decisions = []
    prog = data.get("programme", {}) or {}
    intake_decisions.append(
        f"- recital_a_variant: {prog.get('recital_a_variant', 'default')}"
    )
    density = choices.get("recital_b_density", "standard")
    if density != "standard":
        intake_decisions.append(f"- recital_b_density: {density} (non-default)")
    else:
        intake_decisions.append(f"- recital_b_density: {density}")
    sup = data.get("supplier", {}) or {}
    if sup.get("strategic_purposes"):
        intake_decisions.append(f"- strategic_purposes: {sup['strategic_purposes']}")
    # Active CLI flags
    cli_flags = []
    for flag in ("--verify-source-urls", "--audit-only", "--recital-b-only",
                 "--phase-8-auto-execute", "--enforce-phase-7-5"):
        if flag in sys.argv:
            cli_flags.append(f"- {flag}")
    if not cli_flags:
        cli_flags = ["- (none)"]

    lines = [
        f"# Session Log — {os.path.basename(docx_path)}",
        f"**Generated:** {ts}",
        f"**Type:** {data.get('type', 'Unknown')}",
        f"**Counterparty:** {cp.get('name', 'Unknown')}",
        "",
        "## Intake decisions",
        *intake_decisions,
        "",
        "## CLI flags used",
        *cli_flags,
        "",
        "## QA summary",
        f"- Status: {qa_status} — {fail_count} fail, {warn_count} warn, {info_count} info",
        f"- Rules triggered: [{', '.join(triggered) if triggered else 'none'}]",
        "",
        "## Customizations",
        f"- custom.definitions: {len(custom.get('definitions', []))}",
        f"- custom.clauses: {len(custom.get('clauses', []))}",
    ]

    with open(log_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines) + "\n")
    return log_path


def _emit_audit_checklist(docx_path: str, doc, data: dict) -> str:
    """v3.7.1: emit `{stem}_AUDIT.txt` with per-customization substring
    assertions.

    Pattern lifted from InfraPartners §10 (84-item checklist). For each
    assertion, writes `PASS: <label>` or `FAIL: <label>`. Target: 20+ core
    assertions for any LOI + per-customization additions that scale up to
    ~84 for a fully-configured SS LOI.

    The checklist is read/curated by the operator during Phase 7 review,
    NOT enforced by CI. FAIL lines are advisory — they flag content that
    was expected but not rendered, which often points to a Phase 5 intake
    oversight rather than an engine bug.
    """
    stem = os.path.splitext(os.path.basename(docx_path))[0]
    audit_dir = os.path.dirname(os.path.abspath(docx_path))
    audit_path = os.path.join(audit_dir, f"{stem}_AUDIT.txt")

    text = _extract_text(doc)
    assertions = []  # list of (label, expected_bool, predicate_result)

    def assert_present(label: str, needle: str):
        assertions.append((label, True, needle in text))

    def assert_absent(label: str, needle: str):
        assertions.append((label, False, needle in text))

    t = data.get("type", "")
    cp = data.get("counterparty", {}) or {}
    choices = data.get("choices", {}) or {}
    custom = data.get("custom", {}) or {}
    supplier = data.get("supplier", {}) or {}

    # ---- Core assertions (apply to all LOI types) ----
    assert_present("Preamble — Digital Energy party intro",
                   "(1) Digital Energy Netherlands B.V.")
    assert_present("Preamble — counterparty party intro",
                   f"(2) {cp.get('name', '')}")
    assert_present("Recital A — canonical body opener",
                   "develops and operates Digital Energy Centers")
    assert_present("Recital B — counterparty short name",
                   cp.get("short", cp.get("name", "")))
    assert_present('Definitions — "DEC"', '"DEC"')
    assert_present('Definitions — "Business Day"', '"Business Day"')
    assert_present('Definitions — "Confidential Information"',
                   '"Confidential Information"')
    assert_present('Definitions — "Purpose"', '"Purpose"')
    assert_present('Definitions — "Representatives"', '"Representatives"')
    assert_present("Cl. 2 — Purpose and Scope heading", "2. Purpose and Scope")
    # Cl. 3 exists under different names by type; just check "3.1" renders
    assert_present("Cl. 3 — first sub-clause present", "3.1")
    assert_present("Cl. 6 — Confidentiality heading + BINDING marker",
                   "Confidentiality" if t == "EcosystemPartnership" else "Confidentiality")
    assert_present("Cl. 6 — first sub-clause present", "6.1")
    gen_cl = "7. General Provisions" if t == "EndUser" else "8. General Provisions"
    assert_present(f"{gen_cl} — heading", gen_cl)
    assert_present(f"{gen_cl} — BINDING marker", "(BINDING)")
    # Validity date
    val_date = data.get("dates", {}).get("validity_date", "")
    if val_date:
        assert_present(f"Validity — {val_date}", val_date)
    assert_present("Signature block — Digital Energy attestation",
                   "For and on behalf of Digital Energy")
    # Footer
    # Template version tag — v3.2 for EU/DS/WS, v1.0 for SS/EP/Bespoke
    _footer_vsn = {"StrategicSupplier": "v1.0", "EcosystemPartnership": "v1.0",
                   "Bespoke": "v1.0"}.get(t, "v3.2")
    assert_present("Footer — template version tag",
                   f"DE-LOI-{t}-{_footer_vsn}")
    # v3.6.0 bug-fix invariants — should never appear in any LOI
    assert_absent('v3.6.0 bug — "tthe" typo', "tthe good faith")
    assert_absent('v3.6.0 bug — "(ALT-A)" leftover', "(ALT-A)")
    assert_absent('v3.6.0 bug — Recital B double-period', "..")

    # Closing line
    if t != "EcosystemPartnership":
        assert_present("Closing — We look forward to working with you",
                       "We look forward to working with you")

    # ---- Type-specific assertions ----
    if t in ("Distributor", "Wholesale"):
        assert_present('Definitions — "MSA"', '"MSA"')
    elif t == "StrategicSupplier":
        assert_present('Definitions — "Framework Agreement"',
                       '"Framework Agreement"')

    # ---- Per-customization assertions ----
    # supplier.rofr
    rofr = supplier.get("rofr")
    if rofr:
        # Match on heading text only, not the number — auto_renumber may
        # shift §3.8 to a different number.
        assert_present("RoFR Preferred-Supplier clause present",
                       "Preferred-Supplier and Right of First Refusal")
        scope = rofr.get("site_scope", "")
        if scope:
            assert_present(f"RoFR — site_scope {scope!r} rendered", scope)
        lock_out = rofr.get("lock_out_style", "sole_discretion")
        if lock_out == "alignment":
            assert_present("RoFR — alignment framing present",
                           "good-faith dialogue")
        elif lock_out == "hard_minimum":
            assert_present("RoFR — hard_minimum commitment present",
                           "not fewer than one Designated Site")
        elif lock_out == "milestone":
            assert_present("RoFR — milestone trigger present",
                           "pre-commitment deliverables")

    # supplier.referral_rider
    if supplier.get("referral_rider"):
        assert_present("Mutual Referral Rider heading present",
                       "Mutual Referral Rider")
        assert_present("Referral Rider — bidirectional phrasing",
                       "bidirectional referral interest")

    # Joint Stocking Programme
    if _lead_time_under_six_months(supplier.get("lead_time_target", "")):
        assert_present("Joint Stocking Programme heading present",
                       "Joint Stocking Programme")
        assert_present("Joint Stocking — 90-day RFS rationale",
                       "90-day Ready-for-Service")

    # Co-Marketing
    cm = supplier.get("co_marketing") or {}
    if cm:
        assert_present("Reference and Co-Marketing heading present",
                       "Reference and Co-Marketing")
        framing = cm.get("framing", "multi_supplier")
        if framing == "multi_supplier":
            assert_present("Co-Marketing — multi-supplier framing guard",
                           "avoid language implying exclusivity")
        elif framing == "exclusive":
            assert_present("Co-Marketing — exclusive framing",
                           "sole named supplier")

    # include_schedule
    if choices.get("include_schedule") is False:
        assert_absent("Schedule 1 suppressed (include_schedule=false)",
                      "Schedule 1")
        assert_absent("§8.1(a) Schedule 1 ref scrubbed",
                      "and Schedule 1 of this LOI are non-binding")

    # confidentiality_opt_outs
    opt_outs = choices.get("confidentiality_opt_outs") or []
    if "onward_sharing" in opt_outs:
        assert_absent("§6 Onward-Sharing suppressed", "Onward-Sharing Controls")
    if "compliance_confirmation" in opt_outs:
        assert_absent("§6 Compliance Confirmation suppressed",
                      "Compliance Confirmation")
    if "metadata_protection" in opt_outs:
        assert_absent("§6 Metadata Protection suppressed",
                      "Metadata Protection")

    # custom.definitions
    for item in custom.get("definitions", []) or []:
        key = item.get("key", "")
        if key:
            assert_present(f"custom.definitions — {key!r} injected", key)

    # custom.definitions_include (library lookup)
    library = _load_common_defined_terms()
    for key in custom.get("definitions_include", []) or []:
        entry = library.get(key)
        if entry:
            assert_present(
                f"custom.definitions_include — {entry['name']!r} injected",
                entry["name"],
            )

    # custom.clauses (all modes) — assert on heading text, not the original
    # clause number (auto_renumber may shift it).
    for item in custom.get("clauses", []) or []:
        head = item.get("heading", "")
        num = item.get("number", "")
        if head:
            assert_present(
                f"custom.clauses[{num}] heading {head!r}",
                head,
            )

    # Build output
    pass_count = sum(
        1 for _, expected, observed in assertions if expected == observed
    )
    fail_count = len(assertions) - pass_count

    out_lines = [
        f"# Audit Checklist — {os.path.basename(docx_path)}",
        f"# Total assertions: {len(assertions)} | PASS: {pass_count} | FAIL: {fail_count}",
        "",
    ]
    for label, expected, observed in assertions:
        ok = expected == observed
        out_lines.append(
            f"{'PASS' if ok else 'FAIL'}: "
            f"{'[must be present] ' if expected else '[must be absent] '}"
            f"{label}"
        )

    with open(audit_path, "w", encoding="utf-8") as f:
        f.write("\n".join(out_lines) + "\n")
    return audit_path


if __name__ == "__main__":
    main()
