"""Visual layout invariant tests — v3.5.8 PRINCIPLES.md tripwire #6.

Asserts concrete `paragraph_format` / `alignment` / section-geometry
values on a reference rendered LOI. These invariants express design
decisions that are easy to break accidentally and that neither the QA
linter (pattern-based) nor golden-file tests (fingerprint-based) would
explain in a human-readable way.

Observed failure mode this guards against: the v3.5.5 Parties Preamble
spacing bug (blank `self.p("")` paragraphs doubled vertical gap to ~15pt
— "not proportional" per user). If this test existed at v3.5.2, it would
have caught the bug on first CI run, not three revisions later.

Reference fixture: the Polarise Wholesale regression intake (exercises
every v3.5.x render path).
"""
from __future__ import annotations

import subprocess
import sys
from pathlib import Path

import pytest

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:  # pragma: no cover
    Document = None
    WD_ALIGN_PARAGRAPH = None


_TESTS_DIR = Path(__file__).resolve().parent
_COLOCATION_DIR = _TESTS_DIR.parent
_REGRESSION_DIR = _COLOCATION_DIR / "regression" / "v3.5"
_POLARISE_INTAKE = _REGRESSION_DIR / "polarise_wholesale_intake.yaml"


def _regen_polarise(tmp_path: Path) -> Path:
    out_docx = tmp_path / "polarise_visual.docx"
    cmd = [
        sys.executable,
        str(_COLOCATION_DIR / "generate_loi.py"),
        str(_POLARISE_INTAKE),
        "--output", str(out_docx),
    ]
    result = subprocess.run(
        cmd, cwd=_COLOCATION_DIR, capture_output=True, text=True,
    )
    assert result.returncode == 0, (
        f"generate_loi.py failed:\n"
        f"stdout:\n{result.stdout}\nstderr:\n{result.stderr}"
    )
    return out_docx


@pytest.fixture(scope="module")
def polarise_doc(tmp_path_factory):
    """Generate one reference Polarise Wholesale LOI shared across tests."""
    if Document is None:
        pytest.skip("python-docx not available")
    tmp = tmp_path_factory.mktemp("visual_layout")
    path = _regen_polarise(tmp)
    return Document(str(path))


# -----------------------------------------------------------------------------
# Section geometry — page setup regression guard
# -----------------------------------------------------------------------------

class TestSectionGeometry:
    """v3.5.1 A'''' + v3.5.5 spacing fix live in section + paragraph format.
    Regressions are silent without explicit assertions."""

    def test_page_size_a4(self, polarise_doc):
        s = polarise_doc.sections[0]
        assert round(s.page_width.mm) == 210, "A4 width"
        assert round(s.page_height.mm) == 297, "A4 height"

    def test_margins_mm(self, polarise_doc):
        s = polarise_doc.sections[0]
        assert round(s.top_margin.mm) == 20, "top margin (v3.5.1 preserved)"
        assert round(s.bottom_margin.mm) == 35, "bottom margin (footer room)"
        assert round(s.left_margin.mm) == 25, "left margin"
        assert round(s.right_margin.mm) == 20, "right margin"

    def test_different_first_page_header_footer(self, polarise_doc):
        s = polarise_doc.sections[0]
        assert s.different_first_page_header_footer is True


# -----------------------------------------------------------------------------
# Footer — v3.5.1 A'''' regression guard
# -----------------------------------------------------------------------------

class TestFooterLayout:
    """Footer must be CENTER-aligned (v3.5.5 style fix) and render the NL
    entity for a BV-signed instrument (v3.5.1 A'''' entity-derivation fix)."""

    def test_footer_center_aligned(self, polarise_doc):
        s = polarise_doc.sections[0]
        for footer in (s.first_page_footer, s.footer):
            for p in footer.paragraphs:
                assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER, (
                    "Footer paragraphs must be CENTER-aligned (v3.5.5 fix)"
                )

    def test_footer_renders_nl_entity_for_bv_signer(self, polarise_doc):
        """Polarise fixture's provider.legal_name = 'Digital Energy Netherlands B.V.'
        → footer must show NL entity (Mijnsherenweg / KvK 98580086), NOT AG."""
        s = polarise_doc.sections[0]
        footer_text = "\n".join(
            p.text for p in s.first_page_footer.paragraphs
        )
        assert "Digital Energy Netherlands B.V." in footer_text
        assert "Mijnsherenweg" in footer_text
        assert "KvK 98580086" in footer_text
        assert "Group AG" not in footer_text  # v3.5.1 A'''' fix: no AG leak
        assert "Zug" not in footer_text       # v3.5.1 A'' fix: no AG address


# -----------------------------------------------------------------------------
# Parties Preamble — v3.5.2 A''' + v3.5.5 spacing fix
# -----------------------------------------------------------------------------

class TestPartiesPreambleVisualLayout:
    """The Parties Preamble block appears between the cover and Recital A.
    v3.5.5 replaced blank-paragraph spacers with `space_after=6` on the
    intro + party paragraphs. This is the exact bug the user flagged as
    'not proportional'."""

    def _preamble_paragraphs(self, polarise_doc):
        """Return the 4 Parties Preamble paragraphs (intro, party 1,
        party 2, closing)."""
        ps = []
        capture = False
        for p in polarise_doc.paragraphs:
            txt = p.text
            if txt.startswith("THIS LETTER OF INTENT"):
                capture = True
            if capture:
                ps.append(p)
            if capture and txt.startswith("(each a"):
                break
        return ps

    def test_preamble_present(self, polarise_doc):
        ps = self._preamble_paragraphs(polarise_doc)
        assert len(ps) == 4, "Parties Preamble should be exactly 4 lines"

    def test_preamble_opener(self, polarise_doc):
        ps = self._preamble_paragraphs(polarise_doc)
        assert ps[0].text.startswith("THIS LETTER OF INTENT")

    def test_preamble_party_enumeration(self, polarise_doc):
        ps = self._preamble_paragraphs(polarise_doc)
        assert ps[1].text.startswith("(1) ")
        assert ps[2].text.startswith("(2) ")

    def test_preamble_closing(self, polarise_doc):
        ps = self._preamble_paragraphs(polarise_doc)
        assert ps[3].text.startswith("(each a")

    def test_preamble_spacing_is_6pt_between_blocks(self, polarise_doc):
        """v3.5.5 spacing fix: intro + party-1 + party-2 paragraphs use
        space_after=6pt. Closing uses default (3pt from Normal style).
        v3.5.2 bug was blank `self.p('')` spacers → effective gap ~15pt."""
        ps = self._preamble_paragraphs(polarise_doc)
        # First 3 paragraphs (intro + 2 parties) = 6pt
        for i, p in enumerate(ps[:3]):
            sa = p.paragraph_format.space_after
            assert sa is not None, f"preamble[{i}] must set explicit space_after"
            assert sa.pt == 6.0, (
                f"preamble[{i}] space_after should be 6pt (v3.5.5 fix), "
                f"got {sa.pt}pt"
            )


# -----------------------------------------------------------------------------
# Signature block — v3.5.1 A' + J5 regression guards
# -----------------------------------------------------------------------------

class TestSignatureBlockVisualLayout:
    """Sig block was cleaned in v3.5.1 A': no KvK line, no 'ACKNOWLEDGED
    AND AGREED' header, no reg-number line; Place: field added per
    Dutch/EU execution convention."""

    def _all_paragraphs_text(self, polarise_doc) -> str:
        return "\n".join(p.text for p in polarise_doc.paragraphs)

    def test_no_kvk_line_in_body(self, polarise_doc):
        """v3.5.1 A': KvK line removed from sig block."""
        text = self._all_paragraphs_text(polarise_doc)
        # Footer contains "KvK 98580086" legitimately — we want to ensure
        # the body (paragraph stream, not footer) doesn't have a `KvK:` line.
        assert "KvK:" not in text, "v3.5.1 A' removed the KvK: label from sig block"

    def test_no_acknowledged_and_agreed(self, polarise_doc):
        """v3.5.1 A': 'ACKNOWLEDGED AND AGREED' header removed."""
        text = self._all_paragraphs_text(polarise_doc)
        assert "ACKNOWLEDGED AND AGREED" not in text.upper()

    def test_place_field_present_both_sides(self, polarise_doc):
        """v3.5.1 A': 'Place:' line added per Dutch/EU execution convention."""
        text = self._all_paragraphs_text(polarise_doc)
        # One for Provider, one for Counterparty
        assert text.count("Place: ____________________________") == 2


# -----------------------------------------------------------------------------
# Brand rename — v3.5.2 regression guard
# -----------------------------------------------------------------------------

class TestBrandRename:
    """v3.5.2 replaced `the Provider` / `The Provider` with `Digital Energy`
    across 99 occurrences in generator templates + Recital A constants."""

    def test_no_the_provider_in_body(self, polarise_doc):
        text = "\n".join(p.text for p in polarise_doc.paragraphs)
        assert "the Provider" not in text, "v3.5.2 brand rename regression"
        assert "The Provider" not in text, "v3.5.2 brand rename regression"

    def test_digital_energy_appears_multiple_times(self, polarise_doc):
        text = "\n".join(p.text for p in polarise_doc.paragraphs)
        # Polarise fixture should show "Digital Energy" many times (Recital A,
        # Parties Preamble, clause bodies). At least 20 occurrences is
        # conservative; actual is ~40+.
        assert text.count("Digital Energy") >= 20, (
            "v3.5.2 brand rename should make 'Digital Energy' prominent"
        )
