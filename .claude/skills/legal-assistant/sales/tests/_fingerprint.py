"""Deterministic .docx fingerprint extractor for golden-file testing.

v3.5.8 tripwire #5 (PRINCIPLES.md). Extracts a structural digest of a
rendered LOI .docx that:
  - captures paragraph count + per-paragraph text hash (structural + content)
  - captures table dimensions (# tables, # rows per table, # cells per row)
  - captures section margins (page geometry regression protection)
  - captures footer entity + alignment (v3.5.1 A'''' regression protection)
  - does NOT capture volatile data (timestamps, absolute file paths, UUIDs)

A fingerprint is stable across runs against the same intake YAML + generator
code; it changes when either the generator logic or intake data changes.
That stability is what lets us commit goldens to the repo and diff on PR.
"""
from __future__ import annotations

import hashlib
import json
from typing import Any


def _text_hash(s: str) -> str:
    """Short stable hash for a paragraph body — first 16 hex chars of SHA-256.

    Using a hash (not the raw text) keeps the golden file small and makes
    diffs unambiguous: a changed text produces a changed hash. Prefix of
    text is kept alongside for human-readable diff context.
    """
    return hashlib.sha256(s.encode("utf-8")).hexdigest()[:16]


def _mm(val) -> float | None:
    """Return length in millimeters (rounded) from a python-docx Length."""
    if val is None:
        return None
    try:
        # Length has .mm property
        return round(float(val.mm), 2)
    except Exception:
        return None


def _align_name(v):
    """Return the alignment name (e.g. 'CENTER') or None."""
    if v is None:
        return None
    try:
        return getattr(v, "name", None) or str(v)
    except Exception:
        return None


def fingerprint_docx(docx_path: str) -> dict[str, Any]:
    """Extract deterministic fingerprint from a rendered .docx.

    Returns a JSON-serialisable dict. Fields are stable across runs against
    the same intake+code; intentionally coarse (paragraph-level hashes, not
    full text) so small cosmetic diffs are visible without diffing an entire
    embedded document.
    """
    from docx import Document

    doc = Document(docx_path)

    paragraphs: list[dict[str, Any]] = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text
        pf = p.paragraph_format
        paragraphs.append({
            "i": i,
            "text_prefix": text[:80] if text else "",
            "text_len": len(text),
            "text_hash": _text_hash(text) if text else None,
            "space_before_pt": (pf.space_before.pt if pf.space_before is not None else None),
            "space_after_pt": (pf.space_after.pt if pf.space_after is not None else None),
            "alignment": _align_name(p.alignment),
        })

    tables: list[dict[str, Any]] = []
    for t_i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([
                # First 60 chars per cell — enough to detect structural swaps
                cell.text[:60] if cell.text else ""
                for cell in row.cells
            ])
        tables.append({
            "i": t_i,
            "n_rows": len(table.rows),
            "n_cols": len(table.rows[0].cells) if table.rows else 0,
            "rows": rows,
        })

    # Section geometry — regression protection for page setup
    sections = []
    for s in doc.sections:
        sections.append({
            "page_width_mm": _mm(s.page_width),
            "page_height_mm": _mm(s.page_height),
            "top_margin_mm": _mm(s.top_margin),
            "bottom_margin_mm": _mm(s.bottom_margin),
            "left_margin_mm": _mm(s.left_margin),
            "right_margin_mm": _mm(s.right_margin),
            "diff_first_page_header_footer": bool(
                getattr(s, "different_first_page_header_footer", False)
            ),
        })

    # Footer — v3.5.1 A'''' regression protection
    # First-page footer alignment + text (entity should match provider.legal_name)
    footer_data = []
    for s in doc.sections:
        for footer_name, footer in (
            ("first_page_footer", s.first_page_footer),
            ("footer", s.footer),
        ):
            if footer is None:
                continue
            paragraphs_info = []
            for p in footer.paragraphs:
                paragraphs_info.append({
                    "alignment": _align_name(p.alignment),
                    "text": p.text,
                })
            footer_data.append({
                "which": footer_name,
                "paragraphs": paragraphs_info,
            })

    return {
        "n_paragraphs": len(paragraphs),
        "paragraphs": paragraphs,
        "n_tables": len(tables),
        "tables": tables,
        "sections": sections,
        "footer": footer_data,
    }


def fingerprint_digest(fingerprint: dict[str, Any]) -> str:
    """Return a stable SHA-256 hex digest of a fingerprint — useful for
    quick equality checks without full diff."""
    canonical = json.dumps(fingerprint, sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest()


def diff_fingerprints(expected: dict[str, Any], actual: dict[str, Any]) -> list[str]:
    """Return a list of human-readable diff lines. Empty list = identical.

    Shows paragraph-level changes (added / removed / changed) by index so
    the diff is scoped and grep-able in CI logs.
    """
    diffs: list[str] = []

    # Paragraph count
    if expected["n_paragraphs"] != actual["n_paragraphs"]:
        diffs.append(
            f"paragraph count: expected {expected['n_paragraphs']} "
            f"got {actual['n_paragraphs']}"
        )

    # Per-paragraph content
    exp_ps = {p["i"]: p for p in expected["paragraphs"]}
    act_ps = {p["i"]: p for p in actual["paragraphs"]}
    all_idx = sorted(set(exp_ps) | set(act_ps))
    for i in all_idx:
        e = exp_ps.get(i)
        a = act_ps.get(i)
        if e is None:
            diffs.append(f"[+] paragraph {i}: new — {a['text_prefix']!r}")
            continue
        if a is None:
            diffs.append(f"[-] paragraph {i}: removed — was {e['text_prefix']!r}")
            continue
        if e["text_hash"] != a["text_hash"]:
            diffs.append(
                f"[~] paragraph {i} text changed:\n"
                f"    expected: {e['text_prefix']!r}\n"
                f"    actual:   {a['text_prefix']!r}"
            )
        if e["space_after_pt"] != a["space_after_pt"]:
            diffs.append(
                f"[~] paragraph {i} space_after_pt: "
                f"expected {e['space_after_pt']} got {a['space_after_pt']}"
            )
        if e["alignment"] != a["alignment"]:
            diffs.append(
                f"[~] paragraph {i} alignment: "
                f"expected {e['alignment']} got {a['alignment']}"
            )

    # Table shape
    if expected["n_tables"] != actual["n_tables"]:
        diffs.append(
            f"table count: expected {expected['n_tables']} got {actual['n_tables']}"
        )
    for i, (e_t, a_t) in enumerate(zip(expected["tables"], actual["tables"])):
        if e_t["n_rows"] != a_t["n_rows"] or e_t["n_cols"] != a_t["n_cols"]:
            diffs.append(
                f"[~] table {i} shape: expected {e_t['n_rows']}x{e_t['n_cols']} "
                f"got {a_t['n_rows']}x{a_t['n_cols']}"
            )
        elif e_t["rows"] != a_t["rows"]:
            diffs.append(f"[~] table {i} cell contents differ")

    # Section geometry
    for i, (e_s, a_s) in enumerate(zip(expected["sections"], actual["sections"])):
        for k in (
            "page_width_mm", "page_height_mm",
            "top_margin_mm", "bottom_margin_mm",
            "left_margin_mm", "right_margin_mm",
        ):
            if e_s[k] != a_s[k]:
                diffs.append(
                    f"[~] section {i} {k}: expected {e_s[k]} got {a_s[k]}"
                )

    # Footer
    if expected["footer"] != actual["footer"]:
        diffs.append("[~] footer content changed")
        # Surface the first textual difference to aid debugging
        for i, (e_f, a_f) in enumerate(zip(expected["footer"], actual["footer"])):
            for j, (e_p, a_p) in enumerate(zip(e_f["paragraphs"], a_f["paragraphs"])):
                if e_p != a_p:
                    diffs.append(
                        f"    footer[{i}].paragraph[{j}]: "
                        f"expected {e_p} got {a_p}"
                    )
                    break

    return diffs
