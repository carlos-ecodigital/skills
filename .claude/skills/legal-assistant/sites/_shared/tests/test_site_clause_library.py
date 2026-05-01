"""Tests for ``site_clause_library`` — YAML loader + bilingual section renderer.

Coverage targets:
  1. load_clauses returns a dict keyed by section_id
  2. load_clauses raises on duplicate section_id
  3. load_clauses raises on len(en) != len(nl)
  4. get_clause returns EN list when lang="en"
  5. get_clause returns NL list when lang="nl"
  6. get_clause returns ["[TBC]"] for missing section
  7. render_bilingual_section substitutes {{provider.legal_name}} correctly
  8. render_bilingual_section calls through to render_bilingual_clause
     (verified by counting paragraph rows in the resulting docx table)
  9. render_bilingual_section substitutes through normalise_placeholder for
     None values (None → [TBC])
 10. render_bilingual_section handles dotted paths with missing keys → [TBC]
"""

from __future__ import annotations

import sys
from pathlib import Path

import pytest
from docx import Document

# render_bilingual_section uses a deferred import of bilingual_body from
# document-factory; ensure that path is on sys.path before any test runs.
_SHARED = Path(__file__).resolve().parents[1]
if str(_SHARED) not in sys.path:
    sys.path.insert(0, str(_SHARED))
_FACTORY = _SHARED.resolve().parents[2] / "document-factory"
if str(_FACTORY) not in sys.path:
    sys.path.insert(0, str(_FACTORY))

from site_clause_library import (  # noqa: E402
    Clause,
    get_clause,
    load_clauses,
    render_bilingual_section,
)
from site_doc_base import TBC_TOKEN  # noqa: E402


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


def _write_yaml(tmp_path: Path, body: str) -> Path:
    """Write ``body`` as ``clauses.yaml`` in ``tmp_path`` and return its Path."""
    p = tmp_path / "clauses.yaml"
    p.write_text(body, encoding="utf-8")
    return p


_VALID_YAML = """\
schema_version: "1.0"
template: "TEST-v1.0"
clauses:
  - section_id: "1.1"
    heading_en: "Parties"
    heading_nl: "Partijen"
    en:
      - "1.1 {{provider.legal_name}} is a Dutch B.V."
    nl:
      - "1.1 {{provider.legal_name}} is een Nederlandse B.V."
    asset_gate: null
    render_order: 11

  - section_id: "1.2"
    heading_en: null
    heading_nl: null
    en:
      - "1.2 The other party."
      - "1.2a Continuation paragraph."
    nl:
      - "1.2 De andere partij."
      - "1.2a Vervolgalinea."
    asset_gate: null
    render_order: 12
"""


# ---------------------------------------------------------------------------
# 1. load_clauses returns a dict keyed by section_id
# ---------------------------------------------------------------------------


def test_load_clauses_returns_dict_keyed_by_section_id(tmp_path: Path) -> None:
    path = _write_yaml(tmp_path, _VALID_YAML)
    clauses = load_clauses(path)
    assert isinstance(clauses, dict)
    assert set(clauses.keys()) == {"1.1", "1.2"}
    assert isinstance(clauses["1.1"], Clause)
    assert clauses["1.1"].section_id == "1.1"
    assert clauses["1.1"].heading_en == "Parties"
    assert clauses["1.1"].render_order == 11
    assert clauses["1.2"].heading_en is None


# ---------------------------------------------------------------------------
# 2. load_clauses raises on duplicate section_id
# ---------------------------------------------------------------------------


def test_load_clauses_raises_on_duplicate_section_id(tmp_path: Path) -> None:
    body = """\
schema_version: "1.0"
template: "TEST"
clauses:
  - section_id: "1.1"
    en: ["a"]
    nl: ["b"]
  - section_id: "1.1"
    en: ["c"]
    nl: ["d"]
"""
    path = _write_yaml(tmp_path, body)
    with pytest.raises(ValueError, match="duplicate section_id"):
        load_clauses(path)


# ---------------------------------------------------------------------------
# 3. load_clauses raises on len(en) != len(nl)
# ---------------------------------------------------------------------------


def test_load_clauses_raises_on_paragraph_count_mismatch(tmp_path: Path) -> None:
    body = """\
schema_version: "1.0"
template: "TEST"
clauses:
  - section_id: "2.1"
    en:
      - "one"
      - "two"
    nl:
      - "een"
"""
    path = _write_yaml(tmp_path, body)
    with pytest.raises(ValueError, match="EN/NL paragraph-count mismatch"):
        load_clauses(path)


# ---------------------------------------------------------------------------
# 4. get_clause returns EN list when lang="en"
# ---------------------------------------------------------------------------


def test_get_clause_returns_en_when_lang_en(tmp_path: Path) -> None:
    clauses = load_clauses(_write_yaml(tmp_path, _VALID_YAML))
    out = get_clause(clauses, "1.2", "en")
    assert out == ["1.2 The other party.", "1.2a Continuation paragraph."]


# ---------------------------------------------------------------------------
# 5. get_clause returns NL list when lang="nl"
# ---------------------------------------------------------------------------


def test_get_clause_returns_nl_when_lang_nl(tmp_path: Path) -> None:
    clauses = load_clauses(_write_yaml(tmp_path, _VALID_YAML))
    out = get_clause(clauses, "1.2", "nl")
    assert out == ["1.2 De andere partij.", "1.2a Vervolgalinea."]


# ---------------------------------------------------------------------------
# 6. get_clause returns ["[TBC]"] for missing section
# ---------------------------------------------------------------------------


def test_get_clause_returns_tbc_for_missing_section(tmp_path: Path) -> None:
    clauses = load_clauses(_write_yaml(tmp_path, _VALID_YAML))
    out = get_clause(clauses, "9.9", "en")
    assert out == [TBC_TOKEN]


# ---------------------------------------------------------------------------
# 7. render_bilingual_section substitutes {{provider.legal_name}} correctly
# ---------------------------------------------------------------------------


def test_render_substitutes_dotted_placeholder(tmp_path: Path) -> None:
    clauses = load_clauses(_write_yaml(tmp_path, _VALID_YAML))
    doc = Document()
    render_bilingual_section(
        doc, clauses, ["1.1"],
        heading_en="1. Parties",
        heading_nl="1. Partijen",
        placeholder_subs={"provider": {"legal_name": "Acme B.V."}},
    )
    # The bilingual clause is rendered as a docx table; flatten its text.
    table = doc.tables[0]
    text = "\n".join(
        p.text for row in table.rows for cell in row.cells
        for p in cell.paragraphs
    )
    assert "Acme B.V." in text
    assert "{{provider.legal_name}}" not in text


# ---------------------------------------------------------------------------
# 8. render_bilingual_section calls through to render_bilingual_clause
#    (verified by paragraph-row count in the resulting table)
# ---------------------------------------------------------------------------


def test_render_calls_render_bilingual_clause(tmp_path: Path) -> None:
    clauses = load_clauses(_write_yaml(tmp_path, _VALID_YAML))
    doc = Document()
    render_bilingual_section(
        doc, clauses, ["1.1", "1.2"],
        heading_en="Test",
        heading_nl="Test NL",
        placeholder_subs={"provider": {"legal_name": "Acme B.V."}},
    )
    # 1 heading row + 1 paragraph (1.1) + 2 paragraphs (1.2) = 4 rows;
    # 2 columns (EN, NL).
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert len(table.rows) == 4
    assert len(table.columns) == 2


# ---------------------------------------------------------------------------
# 9. render_bilingual_section substitutes None values via normalise_placeholder
# ---------------------------------------------------------------------------


def test_render_substitutes_none_via_normalise_placeholder(tmp_path: Path) -> None:
    clauses = load_clauses(_write_yaml(tmp_path, _VALID_YAML))
    doc = Document()
    # provider.legal_name is explicitly None — normalise_placeholder
    # returns the [TBC] token.
    render_bilingual_section(
        doc, clauses, ["1.1"],
        heading_en=None,
        heading_nl=None,
        placeholder_subs={"provider": {"legal_name": None}},
    )
    table = doc.tables[0]
    text = "\n".join(
        p.text for row in table.rows for cell in row.cells
        for p in cell.paragraphs
    )
    assert TBC_TOKEN in text
    # Confirm we didn't render the literal Python "None" string.
    assert "None" not in text.replace(TBC_TOKEN, "")


# ---------------------------------------------------------------------------
# 10. render_bilingual_section handles dotted paths with missing keys → [TBC]
# ---------------------------------------------------------------------------


def test_render_missing_dotted_key_renders_tbc(tmp_path: Path) -> None:
    clauses = load_clauses(_write_yaml(tmp_path, _VALID_YAML))
    doc = Document()
    # provider.legal_name path can't be resolved (no provider key in
    # context dict at all) → [TBC] substitution.
    render_bilingual_section(
        doc, clauses, ["1.1"],
        heading_en=None,
        heading_nl=None,
        placeholder_subs={"unrelated": "value"},
    )
    table = doc.tables[0]
    text = "\n".join(
        p.text for row in table.rows for cell in row.cells
        for p in cell.paragraphs
    )
    assert TBC_TOKEN in text
    assert "{{provider.legal_name}}" not in text
