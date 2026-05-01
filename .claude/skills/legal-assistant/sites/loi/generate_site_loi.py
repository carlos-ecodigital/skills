"""
Site LOI engine — v0.1.

Reads a ``deal.yaml`` file, renders a bilingual EN/NL two-column ``.docx``
Letter of Intent derived from ``templates/DE-LOI-Site-v1.0_TEMPLATE.md`` and
the Van Gog LOI structure.

This v0.1 is a **working end-to-end pipeline** that exercises the
document-factory bilingual_body + signature_block modules shipped in
Phase B2. Integration points for cross-doc gate, HubSpot sync, and
document parsers are stubbed with ``TODO(stream_ref)`` markers so the
merge sequence remains auditable as later phases ship.

Usage::

    python3 generate_site_loi.py path/to/deal.yaml --out-dir /tmp/

CLI shape mirrors legal-assistant/sales/generate_loi.py (colocation).
"""

from __future__ import annotations

import argparse
import sys
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any, Dict, List, Optional

import yaml
from docx import Document

# The document-factory modules live as siblings of legal-assistant/sites/...
# We add the factory path at runtime so callers can invoke us from any cwd.
_FACTORY_PATH = (
    Path(__file__).resolve().parents[3] / "document-factory"
)
if str(_FACTORY_PATH) not in sys.path:
    sys.path.insert(0, str(_FACTORY_PATH))
# sites/_shared — for site_doc_base (role labels) + cross_doc_gate
_SHARED_PATH = Path(__file__).resolve().parents[1] / "_shared"
if str(_SHARED_PATH) not in sys.path:
    sys.path.insert(0, str(_SHARED_PATH))

from bilingual_body import render_bilingual_clause  # noqa: E402
from format_validators import run_all as run_format_validators  # noqa: E402
from generate import DE_ENTITIES, COBALT, SLATE_800, SLATE_900, FONT, Party, add_cover  # noqa: E402
from signature_block import SigParty, render_signature_page  # noqa: E402
import site_doc_base as sdb  # noqa: E402
from site_doc_base import derive_role_labels as sdb_derive_role_labels  # noqa: E402
from site_clause_library import load_clauses, render_bilingual_section  # noqa: E402
from docx.shared import Mm, Pt  # noqa: E402


# ---------------------------------------------------------------------------
# Placeholder sanitisation — chassis-authoritative (rc3.1)
# ---------------------------------------------------------------------------

# Values treated as missing / render-as-[TBC]. Engine never leaks raw
# ``None``, ``"TODO(...)"`` or similar author-placeholder strings into
# the generated docx. The authoritative implementation lives on the
# chassis so LOI + HoT agree; this alias preserves the existing call-site
# signature ``_sanitise(value, fallback=...)`` across the module.
from site_doc_base import TBC_TOKEN as _TBC_TOKEN  # noqa: E402
from site_doc_base import normalise_placeholder as _sanitise  # noqa: E402


# ---------------------------------------------------------------------------
# Page-layout helper — narrow margins so bilingual tables fit
# ---------------------------------------------------------------------------

def _set_narrow_margins(doc: Document) -> None:
    """Set 20 mm L/R and 25 mm T/B on every section so bilingual tables
    (USABLE_WIDTH_MM=165) render without overflow. Default python-docx
    margins are 31.75 mm which leaves only 146.5 mm usable — too narrow."""
    for section in doc.sections:
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)
        section.top_margin = Mm(25)
        section.bottom_margin = Mm(20)

# Phase B6 integration complete — imports from site_doc_base.
# Back-compat alias so existing tests that reference ``engine._derive_role_labels``
# keep working; new code should use ``sdb_derive_role_labels`` directly.
_derive_role_labels = sdb_derive_role_labels


# ---------------------------------------------------------------------------
# Clause library — rc3.3 Phase 3 step 1: §1 Parties externalised to YAML.
# §2–§7 follow in subsequent sub-PRs (see plans/rc3-migration-remaining-sections.md).
# Loaded once per process; YAML parsing is cheap but the cache avoids
# re-parsing on every build_document() call.
# ---------------------------------------------------------------------------

_CLAUSES = load_clauses(
    Path(__file__).parent / "templates" / "clauses" / "de-loi-site-v1.0.yaml"
)


# ---------------------------------------------------------------------------
# Clause content — Van Gog verbatim for LOI v1.0.
#
# rc3.3 (Phase 3) externalised §1 Parties to clauses/de-loi-site-v1.0.yaml.
# §2–§7 still live inline; each one carries a TODO(rc3-migration) marker
# below tracking the remaining migration. See:
#   - sites/_shared/site_clause_library.py  (loader + renderer)
#   - plans/rc3-migration-remaining-sections.md  (per-section checklist)
# ---------------------------------------------------------------------------


# Section 1 Parties is now sourced from clauses/de-loi-site-v1.0.yaml via
# the clause library; the inline ``_clause_1_parties()`` helper was deleted
# in rc3.3.


# TODO(rc3-migration): migrate to clauses/de-loi-site-v1.0.yaml.
# See sites/_shared/site_clause_library.py for the loader.
def _clause_2_background() -> tuple[List[str], List[str]]:
    en = [
        "2.1 Digital Energy has developed the Digital Energy Center "
        "(\u201cDEC\u201d), which captures and reuses waste heat from "
        "AI-computing data centers.",
        "2.2 The Parties wish to explore the development of a DEC project "
        "(the \u201cProject\u201d) at the location(s) specified in the LOI "
        "Schedule (Section L). The roles and asset contributions are as set "
        "out in Section R. The Project is intended to be implemented through "
        "a project-specific limited liability company (besloten vennootschap "
        "met beperkte aansprakelijkheid) to be incorporated by Digital Energy "
        "as the special-purpose vehicle for the Project "
        "(\u201cProjectBV\u201d) and the Site Partner(s).",
    ]
    nl = [
        "2.1 Digital Energy heeft het Digital Energy Center (\u201cDEC\u201d) "
        "ontwikkeld, dat restwarmte van AI-datacenters opvangt en hergebruikt.",
        "2.2 De Partijen wensen de ontwikkeling van een DEC-project (het "
        "\u201cProject\u201d) te verkennen op de locatie(s) vermeld in de "
        "Bijlage (Sectie L). De rollen en activa-bijdragen zijn uiteengezet "
        "in Sectie R. Het Project is bedoeld om te worden uitgevoerd via een "
        "projectspecifieke besloten vennootschap met beperkte "
        "aansprakelijkheid die door Digital Energy zal worden opgericht als "
        "special purpose vehicle voor het Project (\u201cProjectBV\u201d) en "
        "de Locatiepartner(s).",
    ]
    return en, nl


# TODO(rc3-migration): migrate to clauses/de-loi-site-v1.0.yaml.
# See sites/_shared/site_clause_library.py for the loader.
def _clause_3_1_dec_development() -> tuple[List[str], List[str]]:
    en = [
        "3.1 DEC Development. Digital Energy intends to design, build, "
        "finance, and operate a DEC at the Project location(s). Digital "
        "Energy expects to bear the cost and risk of the DEC, including, at "
        "the end of the term, removal and site restoration. The DEC is "
        "expected to be owned and operated by the ProjectBV.",
    ]
    nl = [
        "3.1 DEC-ontwikkeling. Digital Energy is voornemens een DEC te "
        "ontwerpen, bouwen, financieren en exploiteren op de "
        "projectlocatie(s). Digital Energy verwacht de kosten en het risico "
        "van het DEC te dragen, met inbegrip van verwijdering en "
        "locatieherstel aan het einde van de looptijd. Het DEC zal naar "
        "verwachting eigendom zijn van en worden geëxploiteerd door de "
        "ProjectBV.",
    ]
    return en, nl


# TODO(rc3-migration): migrate to clauses/de-loi-site-v1.0.yaml.
# See sites/_shared/site_clause_library.py for the loader.
def _clause_3_2_bess(details: dict) -> tuple[List[str], List[str]]:
    mw = details.get("mw", "[TBC]")
    mwh = details.get("mwh", "[TBC]")
    chem = details.get("chemistry", "[TBC]")
    split = details.get("co_invest_pct", 50)
    other = 100 - split if isinstance(split, int) else "[TBC]"
    en = [
        "3.2 BESS Co-Development. The Parties anticipate that the Project "
        "will include the co-development of a Battery Energy Storage System "
        "(\u201cBESS\u201d) as an initial development phase. The BESS is "
        "intended to utilise the Grid Contributor's electrical connections "
        "as described in the LOI Schedule (Section L) and to be located on "
        f"the Landowner's property. The anticipated configuration is "
        f"approximately {mw} MW / {mwh} MWh utilising {chem} technology, "
        "subject to detailed engineering and procurement.",
        "The BESS is intended as a joint investment between Digital Energy "
        f"and the Grid Contributor. The Parties anticipate a {split}/{other} "
        "equity joint venture structure, with the BESS to be owned and "
        "operated by a dedicated BESS SPV. The BESS is expected to generate "
        "revenue from energy arbitrage, frequency containment reserve (FCR), "
        "and other ancillary services on the Dutch electricity balancing "
        "market.",
        "The BESS is intended to be operational prior to the DEC, providing "
        "the Grid Contributor with standalone investment returns before the "
        "DEC development phase commences. At DC financial close, the Parties "
        "anticipate that the Grid Contributor's BESS equity interest will "
        "convert to an equity interest in the ProjectBV, with the conversion "
        "terms to be agreed in the HoT.",
        "All BESS terms, including final capacity, total investment, equity "
        "and debt structure, revenue arrangements, grid-sharing with the DEC, "
        "operational management, and equity conversion mechanics, are to be "
        "agreed in the HoT.",
    ]
    nl = [
        "3.2 BESS Co-Ontwikkeling. De Partijen voorzien dat het Project de "
        "gezamenlijke ontwikkeling omvat van een Battery Energy Storage "
        "System (\u201cBESS\u201d) als initiële ontwikkelingsfase. Het BESS "
        "is bedoeld om de elektrische aansluitingen van de Netbijdrager te "
        "benutten zoals beschreven in de Bijlage (Sectie L) en te worden "
        "gesitueerd op het eigendom van de Grondeigenaar. De beoogde "
        f"configuratie is circa {mw} MW / {mwh} MWh gebruikmakend van "
        f"{chem}-technologie, onder voorbehoud van gedetailleerde "
        "engineering en inkoop.",
        "Het BESS is bedoeld als gezamenlijke investering tussen Digital "
        "Energy en de Netbijdrager. De Partijen voorzien een joint-venture-"
        f"aandelenstructuur van {split}/{other}, waarbij het BESS eigendom "
        "zal zijn van en wordt geëxploiteerd door een speciaal daarvoor "
        "opgerichte BESS-SPV. Het BESS zal naar verwachting inkomsten "
        "genereren uit energie-arbitrage, frequentiereserve (FCR) en andere "
        "ondersteunende diensten op de Nederlandse elektriciteitsbalanceringsmarkt.",
        "Het BESS is bedoeld om operationeel te zijn vóór het DEC, waardoor "
        "de Netbijdrager zelfstandige investeringsrendementen ontvangt "
        "voordat de DEC-ontwikkelingsfase aanvangt. Bij financial close van "
        "het DC voorzien de Partijen dat het BESS-aandelenbelang van de "
        "Netbijdrager wordt omgezet in een aandelenbelang in de ProjectBV, "
        "met de conversievoorwaarden overeen te komen in de HoT.",
        "Alle BESS-voorwaarden, waaronder definitieve capaciteit, totale "
        "investering, eigen vermogen- en schuldstructuur, opbrengstregelingen, "
        "netdeling met het DEC, operationeel beheer en conversievoorwaarden, "
        "zijn overeen te komen in de HoT.",
    ]
    return en, nl


# TODO(rc3-migration): migrate to clauses/de-loi-site-v1.0.yaml.
# See sites/_shared/site_clause_library.py for the loader.
def _clause_3_3_heat_supply() -> tuple[List[str], List[str]]:
    en = [
        "3.3 Heat Supply. The Project's objective includes recovering and "
        "supplying waste heat from the DEC to the Heat Offtaker designated "
        "in Section R. The DEC converts electricity into compute capacity "
        "and, as a by-product, generates thermal energy. The facility is "
        "designed to capture this heat and deliver it to the adjacent Heat "
        "Offtaker's operations via a direct pipeline connection.",
        "The Parties anticipate that the heat supply will be formalised "
        "through a long-term heat offtake agreement between the ProjectBV "
        "and the Heat Offtaker, with pricing indexed to inflation (CPI). "
        "The specific heat supply terms, including price per MWh, volume "
        "commitments, delivery specifications, indexation mechanism, minimum "
        "offtake obligations, credit support, and contract duration, are to "
        "be agreed in the HoT.",
        "The Parties anticipate that the economic benefit of the heat "
        "revenue will be shared between Digital Energy and the Grid "
        "Contributor, in recognition of the Grid Contributor's contribution "
        "of the electrical grid connections that enable the DEC's "
        "operations. The terms of such revenue sharing are to be agreed in "
        "the HoT.",
    ]
    nl = [
        "3.3 Warmtelevering. Het Project heeft mede tot doel restwarmte van "
        "het DEC terug te winnen en te leveren aan de Warmteafnemer "
        "aangewezen in Sectie R. Het DEC zet elektriciteit om in "
        "rekencapaciteit en genereert als bijproduct thermische energie. De "
        "installatie is ontworpen om deze warmte op te vangen en via een "
        "directe pijpleidingverbinding te leveren aan de aangrenzende "
        "activiteiten van de Warmteafnemer.",
        "De Partijen voorzien dat de warmtelevering wordt geformaliseerd "
        "door middel van een langlopende warmteafnameovereenkomst tussen de "
        "ProjectBV en de Warmteafnemer, met prijsindexatie aan inflatie "
        "(CPI). De specifieke warmteleveringsvoorwaarden zijn overeen te "
        "komen in de HoT.",
        "De Partijen voorzien dat het economisch voordeel van de "
        "warmteopbrengst wordt gedeeld tussen Digital Energy en de "
        "Netbijdrager. De voorwaarden van een dergelijke opbrengstverdeling "
        "zijn overeen te komen in de HoT.",
    ]
    return en, nl


# TODO(rc3-migration): migrate to clauses/de-loi-site-v1.0.yaml.
# See sites/_shared/site_clause_library.py for the loader.
def _clause_3_4_grid() -> tuple[List[str], List[str]]:
    en = [
        "3.4 Grid Connection. The Parties anticipate that the Grid "
        "Contributor will make available (through sharing, transfer, or "
        "new-connection application as determined in Section R) the "
        "electrical grid connection(s) required for the DEC's operation. "
        "The specific connection arrangement terms are to be determined in "
        "the HoT.",
    ]
    nl = [
        "3.4 Netaansluiting. De Partijen voorzien dat de Netbijdrager de "
        "voor de werking van het DEC benodigde elektrische "
        "netaansluiting(en) beschikbaar zal stellen (via deling, overdracht "
        "of nieuwe aansluitaanvraag zoals bepaald in Sectie R). De "
        "specifieke aansluitvoorwaarden worden vastgelegd in de HoT.",
    ]
    return en, nl


# TODO(rc3-migration): migrate to clauses/de-loi-site-v1.0.yaml.
# See sites/_shared/site_clause_library.py for the loader.
def _clause_3_5_land() -> tuple[List[str], List[str]]:
    en = [
        "3.5 Land. Where a Site Partner provides land, the Parties "
        "contemplate that the Landowner would make land available for the "
        "DEC, to be secured by a right of superficies (recht van opstal) or "
        "similar arrangement under Dutch law. The anticipated land area for "
        "a DEC is approximately 300 m\u00B2 per MW of electrical capacity. "
        "The terms and conditions of any such land arrangement, including "
        "any fee or compensation, are to be determined in the HoT.",
    ]
    nl = [
        "3.5 Grond. Waar een Locatiepartner grond beschikbaar stelt, beogen "
        "de Partijen dat de Grondeigenaar grond ter beschikking stelt voor "
        "het DEC, te verzekeren door een recht van opstal of vergelijkbare "
        "regeling naar Nederlands recht. De verwachte grondoppervlakte voor "
        "een DEC bedraagt circa 300 m\u00B2 per MW elektrisch vermogen. De "
        "voorwaarden van een dergelijke grondregeling worden vastgelegd in "
        "de HoT.",
    ]
    return en, nl


# TODO(rc3-migration): migrate to clauses/de-loi-site-v1.0.yaml.
# See sites/_shared/site_clause_library.py for the loader.
def _clause_3_6_3_9() -> tuple[List[str], List[str]]:
    en = [
        "3.6 Multiple Locations. The Project may comprise one or more "
        "locations as listed in Section L of the LOI Schedule. Digital "
        "Energy will assess each location and the Parties will agree in the "
        "HoT which locations (if any) proceed to development.",
        "3.7 Separate Contributions. Where different Site Partners fill "
        "different roles as designated in Section R, Digital Energy shall "
        "coordinate the integration of the respective contributions.",
        "3.8 Term. The Parties contemplate an initial term of thirty (30) "
        "years with automatic renewal periods.",
        "3.9 Costs. Each Party bears its own costs in connection with this "
        "LOI and the pre-feasibility assessment. For the avoidance of doubt, "
        "neither Party is entitled to reimbursement of any costs from any "
        "other Party in connection with this LOI or the pre-feasibility "
        "assessment. Digital Energy expects to finance the DEC and the "
        "ProjectBV.",
    ]
    nl = [
        "3.6 Meerdere Locaties. Het Project kan een of meer locaties "
        "omvatten zoals vermeld in Sectie L van de Bijlage. Digital Energy "
        "zal elke locatie beoordelen en de Partijen zullen in de HoT "
        "overeenkomen welke locaties (indien van toepassing) tot "
        "ontwikkeling overgaan.",
        "3.7 Afzonderlijke Bijdragen. Waar verschillende Locatiepartners "
        "verschillende rollen vervullen zoals aangewezen in Sectie R, zal "
        "Digital Energy de integratie van de respectieve bijdragen "
        "coördineren.",
        "3.8 Looptijd. De Partijen beogen een initiële looptijd van dertig "
        "(30) jaar met automatische verlengingsperioden.",
        "3.9 Kosten. Elke Partij draagt haar eigen kosten in verband met "
        "deze LOI en het pre-haalbaarheidsonderzoek. Voor de duidelijkheid, "
        "geen der Partijen heeft recht op vergoeding van kosten van een "
        "andere Partij in verband met deze LOI of het "
        "pre-haalbaarheidsonderzoek. Digital Energy verwacht het DEC en de "
        "ProjectBV te financieren.",
    ]
    return en, nl


# TODO(rc3-migration): migrate to clauses/de-loi-site-v1.0.yaml.
# See sites/_shared/site_clause_library.py for the loader.
def _clause_4_prefeas_hot() -> tuple[List[str], List[str]]:
    en = [
        "4.1 Digital Energy intends to conduct a pre-feasibility assessment "
        "covering, among other matters, electrical engineering, thermal "
        "integration, site planning, heat-demand analysis, site "
        "investigation(s), and any location-specific assessments.",
        "4.2 If the pre-feasibility assessment is positive, Digital Energy "
        "expects to invite the Site Partner(s) to enter into binding HoT. "
        "The HoT is expected to include an exclusivity period, a Variable "
        "Schedule specifying site-specific and deal-specific parameters, "
        "and data carry-forward from this LOI where applicable.",
        "4.3 Target timeline: execution of the HoT within ninety (90) days "
        "of this LOI, subject to satisfactory pre-feasibility assessment, "
        "Digital Energy's timely completion thereof, and the Site "
        "Partner(s)' reasonable cooperation.",
        "4.4 Each Site Partner intends, subject to reasonable advance "
        "notice, to provide Digital Energy and its advisors access to the "
        "relevant site(s) and technical documentation relevant to the "
        "assets described in the LOI Schedule during normal business hours, "
        "to support the pre-feasibility assessment.",
    ]
    nl = [
        "4.1 Digital Energy is voornemens een pre-haalbaarheidsonderzoek uit "
        "te voeren dat onder meer betrekking heeft op elektrotechniek, "
        "thermische integratie, locatieplanning, warmtevraaganalyse, "
        "locatieonderzoek(en) en eventuele locatiespecifieke beoordelingen.",
        "4.2 Indien het pre-haalbaarheidsonderzoek positief is, verwacht "
        "Digital Energy de Locatiepartner(s) uit te nodigen om bindende HoT "
        "aan te gaan. De HoT zullen naar verwachting een "
        "exclusiviteitsperiode, een Variabelenoverzicht met locatiespecifieke "
        "en dealspecifieke parameters, en gegevensovername uit deze LOI "
        "bevatten waar van toepassing.",
        "4.3 Streeftijdlijn: ondertekening van de HoT binnen negentig (90) "
        "dagen na deze LOI, onder voorbehoud van een bevredigend "
        "pre-haalbaarheidsonderzoek, tijdige afronding daarvan door Digital "
        "Energy, en redelijke medewerking van de Locatiepartner(s).",
        "4.4 Elke Locatiepartner is voornemens, met inachtneming van "
        "redelijke voorafgaande kennisgeving, Digital Energy en haar "
        "adviseurs toegang te verlenen tot de relevante locatie(s) en "
        "technische documentatie tijdens normale kantooruren, ter "
        "ondersteuning van het pre-haalbaarheidsonderzoek.",
    ]
    return en, nl


# TODO(rc3-migration): migrate to clauses/de-loi-site-v1.0.yaml.
# See sites/_shared/site_clause_library.py for the loader.
def _clause_5_term() -> tuple[List[str], List[str]]:
    en = [
        "5.1 This LOI enters into force on the date of signature by the "
        "last Party and continues until the earlier of (a) the execution of "
        "the HoT; or (b) twelve (12) months from the last-signature date, "
        "unless extended by written agreement of the Parties.",
        "5.2 Upon expiry, all obligations cease except Section 6.1 "
        "(Confidentiality) and Section 6.2 (Governing Law and Jurisdiction), "
        "which survive in accordance with their terms.",
    ]
    nl = [
        "5.1 Deze LOI treedt in werking op de datum van ondertekening door "
        "de laatste Partij en blijft van kracht tot de vroegste van (a) de "
        "ondertekening van de HoT; of (b) twaalf (12) maanden na de "
        "laatste-ondertekeningsdatum, tenzij verlengd bij schriftelijke "
        "overeenkomst van de Partijen.",
        "5.2 Na het vervallen eindigen alle verplichtingen, met uitzondering "
        "van Artikel 6.1 (Vertrouwelijkheid) en Artikel 6.2 (Toepasselijk "
        "Recht en Bevoegde Rechter), die van kracht blijven overeenkomstig "
        "hun voorwaarden.",
    ]
    return en, nl


# TODO(rc3-migration): migrate to clauses/de-loi-site-v1.0.yaml.
# See sites/_shared/site_clause_library.py for the loader.
def _clause_6_binding() -> tuple[List[str], List[str]]:
    """Binding Provisions — §6.1 confidentiality (with §6.1.6 self-
    supersession), §6.2 governing law, §6.3 information right, §6.4 misc."""
    en = [
        "6.1.1 Each Party shall keep strictly confidential the existence, "
        "content, and terms of this LOI and all information exchanged in "
        "connection with the Project (\u201cConfidential Information\u201d).",
        "6.1.2 Disclosure of Confidential Information is permitted: (a) to "
        "professional advisors, financiers, affiliates, investors or "
        "potential investors, and any special-purpose vehicle incorporated "
        "by Digital Energy for the Project, in each case on a need-to-know "
        "basis and bound by confidentiality obligations no less protective "
        "than this Section 6.1; and (b) as required by law, court order, or "
        "regulatory authority.",
        "6.1.3 Confidential Information does not include information that: "
        "(a) is or becomes part of the public domain other than through "
        "breach by the receiving Party; (b) was known to the receiving "
        "Party before disclosure by the disclosing Party; (c) is "
        "independently developed by the receiving Party without use of the "
        "disclosing Party's Confidential Information; or (d) is received "
        "from a third party without any obligation of confidentiality.",
        "6.1.4 Confidential Information shall not be used by the receiving "
        "Party for any purpose other than the evaluation and execution of "
        "the Project.",
        "6.1.5 Upon expiry or termination of this LOI, each Party shall, at "
        "another Party's request, return or destroy all Confidential "
        "Information received from such Party, except to the extent "
        "retention is required by applicable law or for bona fide "
        "record-keeping purposes.",
        "6.1.6 This confidentiality obligation continues for two (2) years "
        "after the date of this LOI or, if the Parties enter into the HoT, "
        "is superseded by the confidentiality provisions therein.",
        "6.2.1 This LOI is governed by Dutch law.",
        "6.2.2 Any dispute arising from this LOI shall be submitted "
        "exclusively to the District Court of Amsterdam (Rechtbank "
        "Amsterdam).",
        "6.3 Information Right. During the period from execution of this "
        "LOI until the earlier of its expiry or execution of the HoT, each "
        "Site Partner shall inform Digital Energy before entering into "
        "binding discussions with any third party that would materially "
        "affect the assets for which that Site Partner is designated in "
        "Section R: (a) the Grid Contributor shall notify regarding "
        "allocation, sharing, or encumbrance of the grid connection(s); "
        "(b) the Landowner shall notify regarding use, lease, encumbrance, "
        "or development of the land; (c) the Heat Offtaker shall notify "
        "regarding any third-party heat supply or offtake arrangement.",
        "6.4.1 This LOI constitutes the entire understanding of the Parties "
        "regarding the matters set forth in Section 6 and supersedes any "
        "prior arrangements of like effect.",
        "6.4.2 If any provision of this LOI is held invalid or "
        "unenforceable, the remaining provisions continue in full force and "
        "effect.",
        "6.4.3 Amendments to the binding provisions of this LOI require "
        "written agreement signed by all Parties.",
        "6.4.4 Each Party shall process any personal data received in "
        "connection with this LOI in accordance with applicable data "
        "protection legislation, including Regulation (EU) 2016/679 (GDPR).",
        "6.4.5 Nothing in this LOI grants any Party any rights to any other "
        "Party's intellectual property.",
        "6.4.6 In the event of any discrepancy between the English and "
        "Dutch text of this LOI, the English text shall prevail.",
    ]
    nl = [
        "6.1.1 Elke Partij houdt het bestaan, de inhoud en de voorwaarden "
        "van deze LOI en alle in verband met het Project uitgewisselde "
        "informatie (\u201cVertrouwelijke Informatie\u201d) strikt "
        "vertrouwelijk.",
        "6.1.2 Openbaarmaking van Vertrouwelijke Informatie is toegestaan: "
        "(a) aan professionele adviseurs, financiers, gelieerde "
        "ondernemingen, investeerders of potentiële investeerders, en elke "
        "door Digital Energy voor het Project opgerichte special purpose "
        "vehicle, in elk geval op need-to-know basis en gebonden door "
        "geheimhoudingsverplichtingen die niet minder beschermend zijn dan "
        "dit Artikel 6.1; en (b) indien vereist door wet, rechterlijk "
        "bevel, of toezichthoudende instantie.",
        "6.1.3 Vertrouwelijke Informatie omvat geen informatie die: (a) "
        "publiek bekend is of wordt anders dan door schending door de "
        "ontvangende Partij; (b) bekend was bij de ontvangende Partij vóór "
        "openbaarmaking; (c) zelfstandig ontwikkeld is door de ontvangende "
        "Partij zonder gebruik van de Vertrouwelijke Informatie; of (d) "
        "ontvangen is van een derde zonder geheimhoudingsverplichting.",
        "6.1.4 Vertrouwelijke Informatie wordt door de ontvangende Partij "
        "niet gebruikt voor enig doel anders dan de evaluatie en uitvoering "
        "van het Project.",
        "6.1.5 Bij het vervallen of beëindiging van deze LOI zal elke "
        "Partij, op verzoek van een andere Partij, alle van die Partij "
        "ontvangen Vertrouwelijke Informatie retourneren of vernietigen, "
        "behalve voor zover bewaring wettelijk vereist is of voor legitieme "
        "archiveringsdoeleinden.",
        "6.1.6 Deze geheimhoudingsverplichting duurt voort gedurende twee "
        "(2) jaar na de datum van deze LOI of, indien de Partijen de HoT "
        "aangaan, wordt deze vervangen door de daarin opgenomen "
        "vertrouwelijkheidsbepalingen.",
        "6.2.1 Op deze LOI is Nederlands recht van toepassing.",
        "6.2.2 Geschillen voortvloeiend uit deze LOI worden exclusief "
        "voorgelegd aan de Rechtbank Amsterdam.",
        "6.3 Informatierecht. Gedurende de periode vanaf ondertekening van "
        "deze LOI tot de vroegste van het vervallen of ondertekening van de "
        "HoT, stelt elke Locatiepartner Digital Energy in kennis alvorens "
        "bindende besprekingen aan te gaan met derden die de activa "
        "wezenlijk zouden beïnvloeden: (a) de Netbijdrager over toewijzing, "
        "deling of bezwaring van de netaansluiting(en); (b) de "
        "Grondeigenaar over gebruik, verhuur, bezwaring of ontwikkeling "
        "van de grond; (c) de Warmteafnemer over enige warmtelevering of "
        "-afname aan/van derden.",
        "6.4.1 Deze LOI vormt de volledige overeenkomst van de Partijen "
        "inzake de onderwerpen in Artikel 6 en vervangt eerdere regelingen "
        "van gelijke strekking.",
        "6.4.2 Indien enige bepaling van deze LOI ongeldig of "
        "niet-afdwingbaar wordt geacht, blijven de overige bepalingen "
        "volledig van kracht.",
        "6.4.3 Wijzigingen van de bindende bepalingen van deze LOI vereisen "
        "schriftelijke overeenstemming ondertekend door alle Partijen.",
        "6.4.4 Elke Partij verwerkt persoonsgegevens die in verband met "
        "deze LOI worden ontvangen in overeenstemming met de toepasselijke "
        "wetgeving inzake gegevensbescherming, waaronder Verordening (EU) "
        "2016/679 (AVG).",
        "6.4.5 Niets in deze LOI verleent een Partij enig recht op het "
        "intellectuele eigendom van enige andere Partij.",
        "6.4.6 In geval van discrepantie tussen de Engelse en Nederlandse "
        "tekst van deze LOI, prevaleert de Engelse tekst.",
    ]
    return en, nl


# TODO(rc3-migration): migrate to clauses/de-loi-site-v1.0.yaml.
# See sites/_shared/site_clause_library.py for the loader.
def _clause_7_execution() -> tuple[List[str], List[str]]:
    en = [
        "Each signatory represents and warrants that they have full "
        "authority to execute this LOI on behalf of the Party they "
        "represent.",
        "LOI Schedule: integral part.",
    ]
    nl = [
        "Elke ondertekenaar verklaart en garandeert dat hij/zij volledig "
        "bevoegd is om deze LOI te ondertekenen namens de Partij die "
        "hij/zij vertegenwoordigt.",
        "Bijlage: integraal onderdeel.",
    ]
    return en, nl


# ---------------------------------------------------------------------------
# Asset-gate predicates
# ---------------------------------------------------------------------------

def _any_partner_contributes(site_partners: List[dict], asset: str) -> bool:
    for sp in site_partners:
        for contrib in sp.get("contributions") or []:
            if contrib.get("asset") == asset:
                return True
    return False


def _any_partner_returns(site_partners: List[dict], value: str) -> bool:
    for sp in site_partners:
        for ret in sp.get("returns") or []:
            if ret.get("value") == value:
                return True
    return False


def _find_bess_details(site_partners: List[dict]) -> dict:
    for sp in site_partners:
        for contrib in sp.get("contributions") or []:
            if contrib.get("asset") == "equipment_bess":
                return contrib.get("details") or {}
    return {}


# ---------------------------------------------------------------------------
# Core pipeline
# ---------------------------------------------------------------------------

def load_deal(deal_yaml_path: Path) -> dict:
    with open(deal_yaml_path, "r", encoding="utf-8") as f:
        deal = yaml.safe_load(f)
    schema_version = deal.get("deal_yaml_schema_version")
    if schema_version != "1.0":
        raise ValueError(
            f"deal.yaml schema version mismatch: got {schema_version!r}, "
            "expected '1.0'"
        )
    return deal


# ---------------------------------------------------------------------------
# Pipeline hooks — delegate to the chassis singleton (rc3.1)
# ---------------------------------------------------------------------------
#
# Historical note: rc1 + rc2 had hydrate_from_hubspot / parse_documents /
# run_cross_doc_gate implemented as module-level functions with their
# parser map + hubspot round-trip inlined. rc3.1 absorbed those into
# ``SiteDocBase`` so LOI + HoT agree.  Module-level wrappers stay — they
# keep the existing CLI contract (``engine.hydrate_from_hubspot(deal)``)
# and the existing test API working.


def _engine() -> "SiteLOI":
    """Return a lazily-constructed singleton ``SiteLOI`` for the module
    wrappers. Separate from the class so tests and callers that prefer
    explicit instantiation can still do ``SiteLOI()`` directly."""
    global _SINGLETON
    if _SINGLETON is None:
        _SINGLETON = SiteLOI()
    return _SINGLETON


_SINGLETON: Optional["SiteLOI"] = None


def hydrate_from_hubspot(deal: dict, client=None) -> dict:
    """Module-level wrapper → ``SiteLOI.hydrate_from_hubspot``."""
    return _engine().hydrate_from_hubspot(deal, client=client)


def parse_documents(deal: dict, documents_dir: Path) -> dict:
    """Module-level wrapper → ``SiteLOI.parse_documents``."""
    return _engine().parse_documents(deal, documents_dir)


def run_cross_doc_gate(deal: dict) -> list[dict]:
    """Module-level wrapper → ``SiteLOI.run_cross_doc_gate``.

    LOI stage fires only LOI-applicable rules; Gap-4/Gap-5 and Con-* are
    HoT-stage and are quiet here.
    """
    return _engine().run_cross_doc_gate(deal)


def build_document(deal: dict) -> Document:
    """Render the LOI as a python-docx Document."""
    site_partners = deal.get("site_partners") or []
    provider_party = DE_ENTITIES["nl"]
    # Provider dict view for placeholder substitution
    provider = {
        "legal_name": provider_party.legal_name,
        "address": provider_party.address,
        "registration_number": provider_party.registration_number,
    }

    doc = Document()
    _set_narrow_margins(doc)

    loi_date = deal.get("timeline", {}).get("loi_drafted_date") or date.today().isoformat()

    # Cover page — delegate to document-factory.add_cover for the IB-standard
    # title hierarchy (28pt bold title + 14pt subject + 11pt date + party
    # blocks), then append a small bilingual subtitle line so the Dutch
    # translation is visible on the cover. This replaces the inline
    # `doc.add_paragraph()` runs that had no branding, no hierarchy, and
    # bled into Section L on the first page.
    party_blocks = [
        Party(
            legal_name=provider_party.legal_name,
            address=provider_party.address,
            registration_type=provider_party.registration_type,
            registration_number=provider_party.registration_number,
        ),
    ]
    for sp in site_partners:
        party_blocks.append(Party(
            legal_name=_sanitise(sp.get("legal_name")),
            address=_sanitise(sp.get("registered_address"), fallback=""),
            registration_type="KvK" if _sanitise(sp.get("kvk"), fallback="") else None,
            registration_number=_sanitise(sp.get("kvk"), fallback=None)
                if _sanitise(sp.get("kvk"), fallback="") else None,
        ))
    add_cover(
        doc,
        agreement_type="Letter of Intent",
        subject="for a Digital Energy Center Project",
        date_str=loi_date,
        parties=party_blocks,
        formality="non_binding",
        classification="Confidential",
        cover_title="Letter of Intent",
    )

    # Bilingual subtitle: the Dutch translations of Title + Subject, set just
    # below the English block so Dutch readers see the header at a glance.
    nl_subtitle = doc.add_paragraph()
    nl_subtitle.paragraph_format.space_before = Pt(0)
    nl_subtitle.paragraph_format.space_after = Pt(2)
    nl_r = nl_subtitle.add_run("Intentieverklaring")
    nl_r.italic = True
    nl_r.font.name = FONT
    nl_r.font.size = Pt(16)
    nl_r.font.color.rgb = SLATE_900

    nl_sub2 = doc.add_paragraph()
    nl_sub2.paragraph_format.space_before = Pt(0)
    nl_sub2.paragraph_format.space_after = Pt(12)
    nl_sub2_r = nl_sub2.add_run("voor een Digital Energy Center-project")
    nl_sub2_r.italic = True
    nl_sub2_r.font.name = FONT
    nl_sub2_r.font.size = Pt(11)
    nl_sub2_r.font.color.rgb = SLATE_800

    # Page break so Section 1 starts on its own page
    from docx.enum.text import WD_BREAK
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # §1 Parties
    render_bilingual_section(
        doc, _CLAUSES, ["1.1", "1.2", "1.3", "1.4"],
        heading_en="1. Parties",
        heading_nl="1. Partijen",
        placeholder_subs={"provider": provider},
    )

    # §2 Background
    en, nl = _clause_2_background()
    render_bilingual_clause(doc, en, nl,
                            heading="2. Background and Purpose",
                            heading_nl="2. Achtergrond en Doel")

    # §3 Project Overview — composed from sub-clauses
    en_all: List[str] = []
    nl_all: List[str] = []
    for en_part, nl_part in [_clause_3_1_dec_development()]:
        en_all.extend(en_part)
        nl_all.extend(nl_part)
    if (deal.get("addons") or {}).get("bess_co_development"):
        en_part, nl_part = _clause_3_2_bess(_find_bess_details(site_partners))
        en_all.extend(en_part)
        nl_all.extend(nl_part)
    if _any_partner_returns(site_partners, "energy_heat"):
        en_part, nl_part = _clause_3_3_heat_supply()
        en_all.extend(en_part)
        nl_all.extend(nl_part)
    if _any_partner_contributes(site_partners, "grid_interconnection"):
        en_part, nl_part = _clause_3_4_grid()
        en_all.extend(en_part)
        nl_all.extend(nl_part)
    if _any_partner_contributes(site_partners, "land"):
        en_part, nl_part = _clause_3_5_land()
        en_all.extend(en_part)
        nl_all.extend(nl_part)
    en_part, nl_part = _clause_3_6_3_9()
    en_all.extend(en_part)
    nl_all.extend(nl_part)
    render_bilingual_clause(doc, en_all, nl_all,
                            heading="3. Project Overview",
                            heading_nl="3. Projectoverzicht")

    # §4 Pre-Feasibility + HoT
    en, nl = _clause_4_prefeas_hot()
    render_bilingual_clause(doc, en, nl,
                            heading="4. Pre-Feasibility and HoT",
                            heading_nl="4. Pre-haalbaarheidsonderzoek en HoT")

    # §5 Term
    en, nl = _clause_5_term()
    render_bilingual_clause(doc, en, nl,
                            heading="5. Term of this LOI",
                            heading_nl="5. Looptijd van deze LOI")

    # §6 Binding Provisions
    en, nl = _clause_6_binding()
    render_bilingual_clause(doc, en, nl,
                            heading="6. Binding Provisions",
                            heading_nl="6. Bindende Bepalingen")

    # §7 Execution
    en, nl = _clause_7_execution()
    render_bilingual_clause(doc, en, nl,
                            heading="7. Execution",
                            heading_nl="7. Ondertekening")

    # Section L — Locations (bilingual two-column table)
    locations = deal.get("locations") or []
    if locations:
        en_rows: List[str] = []
        nl_rows: List[str] = []
        for i, loc in enumerate(locations, 1):
            parcel = _sanitise(loc.get("parcel_id"))
            addr = _sanitise(loc.get("address"))
            dso = _sanitise(loc.get("dso"))
            muni = _sanitise(loc.get("municipality"))
            postcode = _sanitise(loc.get("postcode"), fallback="")
            zoning = _sanitise(loc.get("bestemmingsplan_designation"), fallback="")
            en_rows.append(
                f"Location {i}: parcel {parcel}; address {addr}"
                + (f", {postcode}" if postcode else "")
                + f"; municipality {muni}; DSO {dso}"
                + (f"; zoning {zoning}" if zoning else "")
                + "."
            )
            nl_rows.append(
                f"Locatie {i}: kadaster {parcel}; adres {addr}"
                + (f", {postcode}" if postcode else "")
                + f"; gemeente {muni}; netbeheerder {dso}"
                + (f"; bestemming {zoning}" if zoning else "")
                + "."
            )
        render_bilingual_clause(
            doc, en_rows, nl_rows,
            heading="Section L — Locations",
            heading_nl="Bijlage L — Locaties",
        )

    # Section R — Roles + Parties (bilingual two-column table)
    en_rows: List[str] = []
    nl_rows: List[str] = []
    for sp in site_partners:
        en_lbls, nl_lbls = sdb_derive_role_labels(sp)
        sp["_role_labels_en"] = en_lbls
        sp["_role_labels_nl"] = nl_lbls

        legal_name = _sanitise(sp.get("legal_name"))
        kvk = _sanitise(sp.get("kvk"), fallback="")
        kvk_en = f" (KvK {kvk})" if kvk else ""
        kvk_nl = f" (KvK {kvk})" if kvk else ""
        sig = sp.get("signatory") or {}
        sig_name = _sanitise(sig.get("name"))
        sig_title = _sanitise(sig.get("title"))

        en_roles = ", ".join(en_lbls) if en_lbls else _TBC_TOKEN
        nl_roles = ", ".join(nl_lbls) if nl_lbls else _TBC_TOKEN

        contribs_en = []
        contribs_nl = []
        for contrib in sp.get("contributions") or []:
            asset = contrib.get("asset", _TBC_TOKEN)
            instrument = contrib.get("instrument", _TBC_TOKEN)
            details = contrib.get("details") or {}
            kvs = []
            for k, v in details.items():
                if k.endswith("_doc") or k.endswith("_doc_ref") or k.endswith("_hash"):
                    continue
                vs = _sanitise(v, fallback="")
                if vs:
                    kvs.append(f"{k}={vs}")
            detail_str = ("; " + ", ".join(kvs)) if kvs else ""
            contribs_en.append(f"contributes {asset} via {instrument}{detail_str}")
            contribs_nl.append(f"levert {asset} via {instrument}{detail_str}")

        returns_en = []
        returns_nl = []
        for ret in sp.get("returns") or []:
            value = ret.get("value", _TBC_TOKEN)
            instrument = ret.get("instrument", _TBC_TOKEN)
            details = ret.get("details") or {}
            kvs = []
            for k, v in details.items():
                vs = _sanitise(v, fallback="")
                if vs and not k.endswith("_hash"):
                    kvs.append(f"{k}={vs}")
            detail_str = ("; " + ", ".join(kvs)) if kvs else ""
            returns_en.append(f"receives {value} via {instrument}{detail_str}")
            returns_nl.append(f"ontvangt {value} via {instrument}{detail_str}")

        en_rows.append(
            f"{legal_name}{kvk_en}. Roles: {en_roles}. Signatory: {sig_name} ({sig_title}). "
            + ("Contributions: " + "; ".join(contribs_en) + ". " if contribs_en else "")
            + ("Returns: " + "; ".join(returns_en) + "." if returns_en else "")
        )
        nl_rows.append(
            f"{legal_name}{kvk_nl}. Rollen: {nl_roles}. Ondertekenaar: {sig_name} ({sig_title}). "
            + ("Bijdragen: " + "; ".join(contribs_nl) + ". " if contribs_nl else "")
            + ("Vergoedingen: " + "; ".join(returns_nl) + "." if returns_nl else "")
        )
    render_bilingual_clause(
        doc, en_rows, nl_rows,
        heading="Section R — Roles and Parties",
        heading_nl="Bijlage R — Rollen en Partijen",
    )

    # Signature page
    doc.add_page_break()
    sig_parties: List[SigParty] = []
    for sp in site_partners:
        sig = sp.get("signatory") or {}
        sig_parties.append(
            SigParty(
                legal_name=_sanitise(sp.get("legal_name")),
                role_labels_en=sp.get("_role_labels_en", []),
                role_labels_nl=sp.get("_role_labels_nl", []),
                signatory_name=_sanitise(sig.get("name"), fallback=""),
                signatory_title=_sanitise(sig.get("title"), fallback=""),
                kvk=None,  # LOI omits KvK (per Van Gog non-binding pattern)
            )
        )
    render_signature_page(doc, provider_party, sig_parties, formality="non_binding")

    return doc


def write_qa_report(deal: dict, doc: Document, gate_verdicts: list[dict],
                    out_path: Path) -> Path:
    issues = run_format_validators(doc)
    lines = []
    lines.append(f"QA report for Site LOI draft: {deal.get('slug')}")
    lines.append(f"Generated: {date.today().isoformat()}")
    lines.append("")
    lines.append(f"Format-validator issues: {len(issues)}")
    for i in issues:
        lines.append(f"  - {i}")
    lines.append("")
    lines.append(f"Cross-doc gate verdicts: {len(gate_verdicts)}")
    for v in gate_verdicts:
        lines.append(f"  - {v}")
    lines.append("")
    lines.append("Integration status:")
    lines.append("  - Phase B6 role-label derivation: WIRED (site_doc_base.derive_role_labels)")
    lines.append("  - Phase B8 cross-doc gate invocation:   WIRED (see verdicts above)")
    lines.append("  - Phase B7 hubspot_sync round-trip:     WIRED (passthrough if no client supplied)")
    lines.append("  - Phase B5 document_parsers enrichment: WIRED (passthrough if no docs uploaded)")
    lines.append("  - Section L + R rendering:              WIRED as bilingual two-column tables (was bullet list)")
    lines.append("  - Cover page:                           WIRED via document-factory.add_cover (was inline paragraphs)")
    lines.append("  - Page margins:                         20mm L/R (narrow, for bilingual-table fit)")
    out_path.write_text("\n".join(lines), encoding="utf-8")
    return out_path


# ---------------------------------------------------------------------------
# SiteLOI — subclass of the shared chassis (rc3.1)
# ---------------------------------------------------------------------------


class SiteLOI(sdb.SiteDocBase):
    """LOI-stage Site document engine.

    Inherits the full ``PARSER_MAP`` (LOI stage is the broadest document-
    set stage), the real ``hydrate_from_hubspot`` + ``parse_documents`` +
    ``run_cross_doc_gate`` implementations, and ``normalise_placeholder``
    via the module-level ``_sanitise`` alias. Adds nothing beyond
    ``render_document`` (which wraps the module-level ``build_document``
    so we don't duplicate its body).

    Module-level wrappers (``hydrate_from_hubspot``, ``parse_documents``,
    ``run_cross_doc_gate``, ``build_document``) delegate to this class via
    the ``_engine()`` singleton. That preserves the rc1/rc2 CLI contract
    and existing test API.
    """

    stage = "loi"

    def render_document(self, deal: dict) -> Document:
        return build_document(deal)


def main(argv: Optional[list[str]] = None) -> int:
    parser = argparse.ArgumentParser(
        description="Render a Site LOI .docx from a deal.yaml file."
    )
    parser.add_argument("deal_yaml", type=Path, help="Path to deal.yaml")
    parser.add_argument(
        "--out-dir", type=Path, default=Path("/tmp"),
        help="Output directory (defaults to /tmp). Output_router hand-off "
             "via sites/_shared/output_router.route() is a Phase E wiring.",
    )
    parser.add_argument(
        "--dry-run", action="store_true",
        help="Build the Document in memory but do not write files.",
    )
    args = parser.parse_args(argv)

    deal = load_deal(args.deal_yaml)
    deal = hydrate_from_hubspot(deal)
    deal = parse_documents(deal, args.deal_yaml.parent / "documents")
    gate_verdicts = run_cross_doc_gate(deal)

    doc = build_document(deal)

    if args.dry_run:
        print(f"[dry-run] would have written LOI for {deal.get('slug')}")
        return 0

    slug = deal.get("slug", "unknown")
    today_str = date.today().strftime("%Y%m%d")
    out_docx = args.out_dir / f"{today_str}_DE_LOI_Site_{slug}_v1_(DRAFT).docx"
    out_qa = args.out_dir / f"{today_str}_DE_LOI_Site_{slug}_v1_qa.txt"

    args.out_dir.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))
    write_qa_report(deal, doc, gate_verdicts, out_qa)

    print(f"LOI: {out_docx}")
    print(f"QA:  {out_qa}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
