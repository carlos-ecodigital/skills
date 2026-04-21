"""Tests for generate_site_loi.py v0.1 — Phase D smoke suite.

End-to-end smoke tests plus unit tests for the asset-gated clause
composition logic. The engine itself is an integration-heavy surface;
these tests focus on the correctness-critical pieces: schema validation,
role-label derivation, asset-gate predicates, and a Van Gog full-pipeline
smoke that produces a .docx opening cleanly in python-docx.
"""

from __future__ import annotations

import sys
from pathlib import Path

import pytest
import yaml
from docx import Document

# Make the LOI module importable
_LOI_DIR = Path(__file__).resolve().parents[1]
if str(_LOI_DIR) not in sys.path:
    sys.path.insert(0, str(_LOI_DIR))

import generate_site_loi as engine  # noqa: E402

_VAN_GOG_FIXTURE = Path(__file__).resolve().parents[1] / "examples" / "deal_van-gog.yaml"


# ---------------------------------------------------------------------------
# Schema validation
# ---------------------------------------------------------------------------

def test_load_deal_reads_valid_yaml(tmp_path):
    p = tmp_path / "d.yaml"
    p.write_text('deal_yaml_schema_version: "1.0"\nslug: x\nsite_partners: []\n')
    deal = engine.load_deal(p)
    assert deal["slug"] == "x"


def test_load_deal_rejects_wrong_schema_version(tmp_path):
    p = tmp_path / "d.yaml"
    p.write_text('deal_yaml_schema_version: "0.9"\nslug: x\n')
    with pytest.raises(ValueError, match="schema version mismatch"):
        engine.load_deal(p)


def test_load_deal_rejects_missing_schema_version(tmp_path):
    p = tmp_path / "d.yaml"
    p.write_text('slug: x\n')
    with pytest.raises(ValueError):
        engine.load_deal(p)


# ---------------------------------------------------------------------------
# Asset-gate predicates
# ---------------------------------------------------------------------------

def test_any_partner_contributes_detects_grid():
    sps = [{"contributions": [{"asset": "land"}]},
           {"contributions": [{"asset": "grid_interconnection"}]}]
    assert engine._any_partner_contributes(sps, "grid_interconnection") is True


def test_any_partner_contributes_absent():
    sps = [{"contributions": [{"asset": "land"}]}]
    assert engine._any_partner_contributes(sps, "grid_interconnection") is False


def test_any_partner_returns_detects_heat():
    sps = [{"returns": [{"value": "equity"}]},
           {"returns": [{"value": "energy_heat"}]}]
    assert engine._any_partner_returns(sps, "energy_heat") is True


def test_find_bess_details_returns_configured_dict():
    sps = [{"contributions": [{"asset": "equipment_bess",
                               "details": {"mw": 25.5, "mwh": 51, "chemistry": "LFP"}}]}]
    d = engine._find_bess_details(sps)
    assert d["mw"] == 25.5
    assert d["chemistry"] == "LFP"


def test_find_bess_details_empty_when_no_bess():
    assert engine._find_bess_details([{"contributions": [{"asset": "land"}]}]) == {}


# ---------------------------------------------------------------------------
# Role-label derivation
# ---------------------------------------------------------------------------

def test_derive_labels_grid_contributor_only():
    sp = {"contributions": [{"asset": "grid_interconnection"}], "returns": []}
    en, nl = engine._derive_role_labels(sp)
    assert en == ["Grid Contributor"]
    assert nl == ["Netbijdrager"]


def test_derive_labels_heat_offtaker_only():
    sp = {"contributions": [], "returns": [{"value": "energy_heat"}]}
    en, nl = engine._derive_role_labels(sp)
    assert en == ["Heat Offtaker"]
    assert nl == ["Warmteafnemer"]


def test_derive_labels_multiple_roles_no_duplicates():
    sp = {
        "contributions": [
            {"asset": "grid_interconnection"},
            {"asset": "land"},
            {"asset": "grid_interconnection"},  # dup asset
        ],
        "returns": [],
    }
    en, nl = engine._derive_role_labels(sp)
    assert en == ["Grid Contributor", "Landowner"]
    assert nl == ["Netbijdrager", "Grondeigenaar"]


def test_derive_labels_empty_partner():
    assert engine._derive_role_labels({"contributions": [], "returns": []}) == ([], [])


# ---------------------------------------------------------------------------
# End-to-end: Van Gog
# ---------------------------------------------------------------------------

def test_van_gog_fixture_exists():
    assert _VAN_GOG_FIXTURE.exists(), f"missing: {_VAN_GOG_FIXTURE}"


def test_van_gog_pipeline_runs_without_error(tmp_path):
    deal = engine.load_deal(_VAN_GOG_FIXTURE)
    doc = engine.build_document(deal)
    out = tmp_path / "loi.docx"
    doc.save(str(out))
    # Re-open the saved doc to prove it's well-formed Word.
    reopened = Document(str(out))
    assert len(reopened.tables) >= 7  # §1..§7 + at least


def test_van_gog_output_activates_bess_block():
    """Van Gog has addons.bess_co_development == True. The §3 table should
    contain the BESS-specific text."""
    deal = engine.load_deal(_VAN_GOG_FIXTURE)
    doc = engine.build_document(deal)
    # Find §3 table
    section3_text = ""
    for t in doc.tables:
        first_en = t.rows[0].cells[0].text
        if first_en.startswith("3."):
            for r in t.rows:
                for c in r.cells:
                    section3_text += c.text + "\n"
            break
    assert "BESS" in section3_text
    assert "Battery Energy Storage" in section3_text


def test_van_gog_output_activates_heat_supply_block():
    deal = engine.load_deal(_VAN_GOG_FIXTURE)
    doc = engine.build_document(deal)
    section3_text = ""
    for t in doc.tables:
        first_en = t.rows[0].cells[0].text
        if first_en.startswith("3."):
            for r in t.rows:
                for c in r.cells:
                    section3_text += c.text + "\n"
            break
    assert "Heat Supply" in section3_text
    assert "Warmtelevering" in section3_text


def test_van_gog_output_contains_self_superseding_confidentiality():
    """Critical: §6.1.6 self-supersession must appear verbatim so template
    natively handles LOI→HoT confidentiality redundancy."""
    deal = engine.load_deal(_VAN_GOG_FIXTURE)
    doc = engine.build_document(deal)
    section6_text = ""
    for t in doc.tables:
        first_en = t.rows[0].cells[0].text
        if first_en.startswith("6."):
            for r in t.rows:
                for c in r.cells:
                    section6_text += c.text + "\n"
            break
    # EN: "is superseded by the confidentiality provisions therein"
    # NL: "wordt deze vervangen door de daarin opgenomen vertrouwelijkheidsbepalingen"
    assert "superseded by the confidentiality provisions therein" in section6_text
    assert "vervangen door de daarin opgenomen vertrouwelijkheidsbepalingen" in section6_text


def test_van_gog_output_has_3_site_partners_rendered_in_signature_page():
    deal = engine.load_deal(_VAN_GOG_FIXTURE)
    doc = engine.build_document(deal)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Van Gog Grubbenvorst B.V." in text
    assert "Van Gog Grubbenvorst Vastgoed B.V." in text
    assert "Van Gog kwekerijen Grubbenvorst B.V." in text


def test_van_gog_output_has_no_kvk_in_signatures_non_binding():
    """Per Van Gog pattern: LOI signature blocks omit KvK."""
    deal = engine.load_deal(_VAN_GOG_FIXTURE)
    doc = engine.build_document(deal)
    text = "\n".join(p.text for p in doc.paragraphs)
    # Provider KvK (98580086) must NOT appear in signature block when
    # formality=="non_binding". It will still appear in §1.1 recital where
    # the registration number is disclosed in full — that's separate.
    # Check that the literal "KvK:" label line is absent from the document.
    assert "KvK: " not in text


def test_van_gog_output_passes_format_validators():
    deal = engine.load_deal(_VAN_GOG_FIXTURE)
    doc = engine.build_document(deal)
    issues = engine.run_format_validators(doc)
    assert issues == [], f"validator issues: {issues}"


def test_main_cli_writes_two_files(tmp_path):
    import subprocess
    # Invoke main() directly rather than via subprocess so we can inspect
    # side-effects cleanly.
    rc = engine.main([str(_VAN_GOG_FIXTURE), "--out-dir", str(tmp_path)])
    assert rc == 0
    # Exactly two files: .docx and _qa.txt
    docx_files = list(tmp_path.glob("*_DE_LOI_Site_van-gog-grubbenvorst_v1_*.docx"))
    qa_files = list(tmp_path.glob("*_DE_LOI_Site_van-gog-grubbenvorst_v1_qa.txt"))
    assert len(docx_files) == 1
    assert len(qa_files) == 1


def test_dry_run_produces_no_output(tmp_path):
    rc = engine.main([str(_VAN_GOG_FIXTURE), "--out-dir", str(tmp_path), "--dry-run"])
    assert rc == 0
    assert list(tmp_path.glob("*.docx")) == []


def test_hydrate_from_hubspot_is_passthrough():
    deal = {"slug": "x", "site_partners": []}
    assert engine.hydrate_from_hubspot(deal) is deal


def test_parse_documents_is_passthrough(tmp_path):
    deal = {"slug": "x", "site_partners": []}
    assert engine.parse_documents(deal, tmp_path) is deal


def test_cross_doc_gate_returns_empty_list():
    deal = {"slug": "x", "site_partners": []}
    assert engine.run_cross_doc_gate(deal) == []
