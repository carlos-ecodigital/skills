"""End-to-end integration test: generate a real Van Gog LOI .docx and
assert structural + content invariants on the result."""

from __future__ import annotations

import sys
from pathlib import Path

import pytest
from docx import Document

# conftest.py handles sys.path setup

import generate_site_loi as engine  # noqa: E402


@pytest.fixture
def van_gog_yaml_path():
    return Path(__file__).resolve().parents[1] / "loi" / "examples" / "deal_van-gog.yaml"


def _text_of(doc):
    lines = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                lines.append(c.text)
    return "\n".join(lines)


def _generate(tmp_path, van_gog_yaml_path):
    rc = engine.main([str(van_gog_yaml_path), "--out-dir", str(tmp_path)])
    assert rc == 0
    docx_files = list(tmp_path.glob("*_DE_LOI_Site_van-gog-grubbenvorst_v1_*.docx"))
    qa_files = list(tmp_path.glob("*_DE_LOI_Site_van-gog-grubbenvorst_v1_qa.txt"))
    assert len(docx_files) == 1, docx_files
    assert len(qa_files) == 1, qa_files
    return docx_files[0], qa_files[0]


def test_van_gog_loi_file_produced(tmp_path, van_gog_yaml_path):
    docx, qa = _generate(tmp_path, van_gog_yaml_path)
    assert docx.stat().st_size > 20000          # real .docx, not stub
    assert qa.stat().st_size > 100               # non-empty QA


def test_van_gog_loi_has_eleven_bilingual_tables(tmp_path, van_gog_yaml_path):
    """Van Gog LOI renders eleven tables (rc3.2):

    - §1..§7 → 7 clause-body tables (2 cols, EN / NL)
    - Section L heading band → 1 bilingual clause table (2 cols)
    - Section L schedule → 1 schedule table (5 cols)
    - Section R heading band → 1 bilingual clause table (2 cols)
    - Section R schedule → 1 schedule table (5 cols)

    rc2: 9 tables (Section L + R were 2-column prose tables).
    rc3.2: Section L + R promoted to true N-column schedule tables;
    each gains a bilingual heading band so the schedule has a label.
    """
    docx, _ = _generate(tmp_path, van_gog_yaml_path)
    d = Document(str(docx))
    assert len(d.tables) == 11
    # 9 bilingual clause tables (2 cols) + 2 schedule tables (5 cols)
    two_col = [t for t in d.tables if len(t.columns) == 2]
    five_col = [t for t in d.tables if len(t.columns) == 5]
    assert len(two_col) == 9
    assert len(five_col) == 2


def test_van_gog_loi_activates_bess_block(tmp_path, van_gog_yaml_path):
    docx, _ = _generate(tmp_path, van_gog_yaml_path)
    text = _text_of(Document(str(docx)))
    assert "BESS" in text
    assert "Battery Energy Storage" in text


def test_van_gog_loi_activates_heat_supply_block(tmp_path, van_gog_yaml_path):
    docx, _ = _generate(tmp_path, van_gog_yaml_path)
    text = _text_of(Document(str(docx)))
    assert "Heat Supply" in text
    assert "Warmtelevering" in text


def test_van_gog_loi_activates_grid_connection_block(tmp_path, van_gog_yaml_path):
    docx, _ = _generate(tmp_path, van_gog_yaml_path)
    text = _text_of(Document(str(docx)))
    assert "Grid Connection" in text
    assert "Netaansluiting" in text


def test_van_gog_loi_activates_land_block(tmp_path, van_gog_yaml_path):
    docx, _ = _generate(tmp_path, van_gog_yaml_path)
    text = _text_of(Document(str(docx)))
    assert "recht van opstal" in text.lower()
    assert "300" in text       # "approximately 300 m² per MW"


def test_van_gog_loi_contains_self_superseding_confidentiality(
    tmp_path, van_gog_yaml_path
):
    """Critical: §6.1.6 must appear verbatim EN + NL so template natively
    handles LOI→HoT confidentiality redundancy."""
    docx, _ = _generate(tmp_path, van_gog_yaml_path)
    text = _text_of(Document(str(docx)))
    assert "superseded by the confidentiality provisions therein" in text
    assert "vervangen door de daarin opgenomen vertrouwelijkheidsbepalingen" in text


def test_van_gog_loi_signature_page_has_3_site_partners(
    tmp_path, van_gog_yaml_path
):
    docx, _ = _generate(tmp_path, van_gog_yaml_path)
    text = _text_of(Document(str(docx)))
    assert "Van Gog Grubbenvorst B.V." in text
    assert "Van Gog Grubbenvorst Vastgoed B.V." in text
    assert "Van Gog kwekerijen Grubbenvorst B.V." in text


def test_van_gog_loi_signature_page_omits_kvk_non_binding(
    tmp_path, van_gog_yaml_path
):
    """Per Van Gog pattern: LOI sig blocks omit KvK."""
    docx, _ = _generate(tmp_path, van_gog_yaml_path)
    text = _text_of(Document(str(docx)))
    # "KvK: <digits>" label must not appear (registration_number is in
    # §1.1 recital, not in signature block).
    assert "KvK: " not in text


def test_van_gog_loi_gate_surfaces_expected_verdicts(
    tmp_path, van_gog_yaml_path
):
    docx, qa = _generate(tmp_path, van_gog_yaml_path)
    qa_text = qa.read_text()
    # Expect at least 3× DataAcc-1 (one per site_partner with null entity_id)
    assert qa_text.count("DataAcc-1") >= 3
    # Expect Esc-2 (co-investment) for BESS equity return on partner 0
    assert "Esc-2" in qa_text


def test_van_gog_loi_format_validators_pass(tmp_path, van_gog_yaml_path):
    docx, qa = _generate(tmp_path, van_gog_yaml_path)
    qa_text = qa.read_text()
    assert "Format-validator issues: 0" in qa_text
