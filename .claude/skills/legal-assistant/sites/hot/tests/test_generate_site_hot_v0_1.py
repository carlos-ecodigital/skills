"""Tests for generate_site_hot.py v0.1 — Phase C smoke suite.

End-to-end smoke tests plus unit tests for SHA-guarded body copy, XML
form-fill correctness, pipeline integration, and CLI side effects.

Fixture: the Moerman-shaped deal.yaml v1.0 is embedded as a Python string
in ``fixtures_embedded.py`` and materialised into tmp_path per test. The
harness blocks checked-in .yaml file creation in this subtree, so this
on-the-fly write is the supported workaround.
"""

from __future__ import annotations

import hashlib
import json
import sys
from pathlib import Path

import pytest
import yaml
from docx import Document

_HOT_DIR = Path(__file__).resolve().parents[1]
if str(_HOT_DIR) not in sys.path:
    sys.path.insert(0, str(_HOT_DIR))

import generate_site_hot as engine  # noqa: E402

_TESTS_DIR = Path(__file__).resolve().parent
if str(_TESTS_DIR) not in sys.path:
    sys.path.insert(0, str(_TESTS_DIR))

import fixtures_embedded  # noqa: E402


@pytest.fixture
def deal_yaml_path(tmp_path):
    """Materialise the embedded Moerman fixture and return its .yaml path."""
    return fixtures_embedded.write_fixture(tmp_path)


# ---------------------------------------------------------------------------
# load_deal
# ---------------------------------------------------------------------------


def test_load_deal_accepts_v1_0(tmp_path):
    p = tmp_path / "d.yaml"
    p.write_text(
        'deal_yaml_schema_version: "1.0"\n'
        'slug: probe\n'
        'site_partners: []\n'
    )
    deal = engine.load_deal(p)
    assert deal["slug"] == "probe"


def test_load_deal_rejects_wrong_schema(tmp_path):
    p = tmp_path / "d.yaml"
    p.write_text('deal_yaml_schema_version: "0.9"\nslug: probe\n')
    with pytest.raises(ValueError, match="schema version mismatch"):
        engine.load_deal(p)


# ---------------------------------------------------------------------------
# copy_body
# ---------------------------------------------------------------------------


def test_copy_body_sha_matches_source(tmp_path):
    src = engine.BODY_TEMPLATE
    dst = tmp_path / "out_body.docx"
    out_path, sha = engine.copy_body(dst, source=src)
    assert out_path == dst
    assert dst.exists()
    h = hashlib.sha256()
    h.update(dst.read_bytes())
    assert h.hexdigest() == sha
    h_src = hashlib.sha256()
    h_src.update(src.read_bytes())
    assert h_src.hexdigest() == sha


# ---------------------------------------------------------------------------
# populate_annex_a
# ---------------------------------------------------------------------------


def test_populate_annex_a_opens_cleanly_in_python_docx(tmp_path):
    dst = tmp_path / "annex.docx"
    registry = engine.sdb.load_registry()
    stats = engine.populate_annex_a(
        engine.ANNEX_A_TEMPLATE,
        dst,
        values={"A.1": "Test B.V.", "A.2": "12345678"},
        registry=registry,
    )
    Document(str(dst))
    assert "A.1" in stats.fields_written
    assert "A.2" in stats.fields_written


def test_populate_annex_a_moerman_smoke(tmp_path):
    """Populate with Moerman-reference values; verify key substitutions."""
    dst = tmp_path / "annex_moerman.docx"
    registry = engine.sdb.load_registry()
    values = {
        "A.1": "Moerman Paprika B.V.",
        "A.2": "27178957",
        "A.3": "Harteveldlaan 16, 2671 VH Naaldwijk",
        "A.7": "Harteveldlaan 16, 2671 VH Naaldwijk",
        "B.1": "Westland Infra",
        "B.4": "4.8",
        "D.1": "Naaldwijk, Sectie B, Nr 4521",
        "E.1": "50 : 50",
        "G.Grower_email": "jan@moermanpaprika.nl",
    }
    stats = engine.populate_annex_a(engine.ANNEX_A_TEMPLATE, dst, values, registry)
    assert stats.warnings == [], stats.warnings
    doc = Document(str(dst))
    flat_text = "\n".join(
        c.text for tbl in doc.tables for r in tbl.rows for c in r.cells
    )
    assert "Moerman Paprika B.V." in flat_text
    assert "27178957" in flat_text
    assert "jan@moermanpaprika.nl" in flat_text


def test_populate_annex_a_preserves_non_shaded_cells(tmp_path):
    """Fixed-term cells (E.4-E.12 row) have no fill — must be untouched."""
    dst = tmp_path / "annex_preserve.docx"
    registry = engine.sdb.load_registry()
    engine.populate_annex_a(
        engine.ANNEX_A_TEMPLATE, dst,
        values={"A.1": "X B.V."}, registry=registry,
    )
    doc = Document(str(dst))
    flat_text = "\n".join(
        c.text for tbl in doc.tables for r in tbl.rows for c in r.cells
    )
    assert "30 years (fixed)" in flat_text
    assert "5 years each (fixed)" in flat_text
    assert "60 days (fixed)" in flat_text


# ---------------------------------------------------------------------------
# Pipeline passthroughs
# ---------------------------------------------------------------------------


def test_hydrate_from_hubspot_is_passthrough():
    deal = {"slug": "x", "site_partners": []}
    assert engine.hydrate_from_hubspot(deal) is deal


def test_parse_documents_is_passthrough(tmp_path):
    deal = {"slug": "x", "site_partners": []}
    assert engine.parse_documents(deal, tmp_path) is deal


def test_run_cross_doc_gate_returns_list(deal_yaml_path):
    """Gate returns a list of verdict dicts; may be non-empty since the
    cross_doc_gate fires several rules even on a partially populated deal."""
    deal = engine.load_deal(deal_yaml_path)
    verdicts = engine.run_cross_doc_gate(deal)
    assert isinstance(verdicts, list)
    for v in verdicts:
        assert isinstance(v, dict)


# ---------------------------------------------------------------------------
# CLI end-to-end
# ---------------------------------------------------------------------------


def test_main_cli_writes_four_files(deal_yaml_path, tmp_path):
    out_dir = tmp_path / "out"
    rc = engine.main([str(deal_yaml_path), "--out-dir", str(out_dir)])
    assert rc == 0
    body = list(out_dir.glob("*_DE_HoT_Site_*_v1_(DRAFT)_body.docx"))
    annex = list(out_dir.glob("*_DE_HoT_Site_*_v1_(DRAFT)_annex-a.docx"))
    qa = list(out_dir.glob("*_DE_HoT_Site_*_v1_qa.txt"))
    gate = list(out_dir.glob("*_DE_HoT_Site_*_v1_gate-report.json"))
    assert len(body) == 1, list(out_dir.iterdir())
    assert len(annex) == 1
    assert len(qa) == 1
    assert len(gate) == 1
    payload = json.loads(gate[0].read_text())
    assert payload["stage"] == "hot"
    assert "verdicts" in payload


def test_main_cli_dry_run_writes_no_files(deal_yaml_path, tmp_path):
    out_dir = tmp_path / "out"
    rc = engine.main([str(deal_yaml_path), "--out-dir", str(out_dir), "--dry-run"])
    assert rc == 0
    assert not out_dir.exists() or list(out_dir.glob("*.docx")) == []
    if out_dir.exists():
        assert list(out_dir.glob("*.json")) == []
        assert list(out_dir.glob("*.txt")) == []


def test_build_field_values_from_fixture(deal_yaml_path):
    """build_field_values must resolve the minimal fixture without errors
    and produce at least the A.* identity + G.* notices fields."""
    deal = engine.load_deal(deal_yaml_path)
    values = engine.build_field_values(deal)
    assert values.get("A.1") == "Moerman Paprika B.V."
    assert values.get("A.2") == "27178957"
    assert "G.DE_email" in values
    assert "project_name" in values
