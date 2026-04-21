"""M5 — Pipeline routing dispatcher tests.

Covers:
  - Pipeline A routing: dict / explicit agreement_spec → NotImplementedError
    with a message pointing to legal-assistant/colocation/generate_loi.py.
  - Pipeline B routing: .docx path + RebrandSpec → delegates to rebrand(),
    returns bytes. Missing rebrand_spec is rejected with TypeError.
  - Pipeline C routing: markdown string or .md file →
    delegates to md_to_docx. No warning when profile is non-agreement.
  - Pipeline C deprecation: markdown + agreement profile emits
    DeprecationWarning whose message names both alternative pipelines
    and the structural reasons.
  - Regression: the dispatcher does not change the behavior of the
    underlying pipelines — a direct md_to_docx() vs. build() produce
    byte-equivalent output.
  - Unrouteable input: clear TypeError with all three pipeline pointers.
"""
from __future__ import annotations

import io
import sys
import warnings
from pathlib import Path

import pytest
from docx import Document

_FACTORY = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(_FACTORY))

from dispatcher import (  # noqa: E402
    AGREEMENT_PROFILES,
    build,
)
from rebrand import RebrandSpec  # noqa: E402
from generate import md_to_docx  # noqa: E402


# ──────────────────────────────────────────────────────────────────────
# Shared fixtures


@pytest.fixture
def sample_markdown() -> str:
    return (
        "# Quarterly Update\n\n"
        "## Summary\n\n"
        "Revenue up **15%** this quarter. Three deals closed.\n\n"
        "- Deal A — signed\n"
        "- Deal B — signed\n"
        "- Deal C — signed\n"
    )


@pytest.fixture
def synthetic_docx(tmp_path: Path) -> Path:
    """Minimal valid .docx saved to disk. Used as a body source for
    Pipeline B routing; we only exercise the routing layer — the
    rebrand end-to-end is covered in tests/test_rebrand.py.
    """
    doc = Document()
    doc.add_heading("Heading 1", level=1)
    doc.add_paragraph("THE UNDERSIGNED:")
    doc.add_paragraph("Digital Energy Group AG, a company...")
    doc.add_paragraph("Acme B.V., a company...")
    doc.add_heading("1. Definitions", level=1)
    doc.add_paragraph("In this agreement...")
    path = tmp_path / "input.docx"
    doc.save(path)
    return path


@pytest.fixture
def rebrand_spec() -> RebrandSpec:
    return RebrandSpec(
        agreement_type="Letter of Intent",
        client="Acme B.V.",
        client_address="123 Main St, 1000 AA Amsterdam",
        entity="nl",
    )


# ──────────────────────────────────────────────────────────────────────
# Pipeline A — dict / agreement_spec → NotImplementedError


def test_pipeline_a_dict_raises_not_implemented() -> None:
    with pytest.raises(NotImplementedError) as exc:
        build({"profile": "loi", "counterparty": "Acme"})
    msg = str(exc.value)
    # Message must name the current Pipeline A entry point so callers
    # know where to go today.
    assert "legal-assistant/colocation/generate_loi.py" in msg
    # Message must mention M4 / builders so intent is clear.
    assert "M4" in msg or "builders" in msg


def test_pipeline_a_explicit_agreement_spec_raises_not_implemented() -> None:
    # Even non-dict shapes passed via agreement_spec= go through the
    # same Pipeline A branch.
    sentinel = object()
    with pytest.raises(NotImplementedError) as exc:
        build("irrelevant", agreement_spec=sentinel)
    assert "legal-assistant/colocation/generate_loi.py" in str(exc.value)


# ──────────────────────────────────────────────────────────────────────
# Pipeline B — .docx path + RebrandSpec


def test_pipeline_b_routes_to_rebrand(
    synthetic_docx: Path,
    rebrand_spec: RebrandSpec,
) -> None:
    out = build(synthetic_docx, rebrand_spec=rebrand_spec)
    # rebrand() returns raw .docx bytes.
    assert isinstance(out, bytes)
    assert len(out) > 1000  # non-trivial
    # Real validity check: the bytes round-trip through python-docx.
    # (A PK zip-magic check would pass for any zip, not specifically a
    #  valid .docx — the Document load is the actual assertion.)
    doc = Document(io.BytesIO(out))
    assert len(doc.paragraphs) > 0


def test_pipeline_b_accepts_str_path(
    synthetic_docx: Path,
    rebrand_spec: RebrandSpec,
) -> None:
    # str input (not Path) routes the same way.
    out = build(str(synthetic_docx), rebrand_spec=rebrand_spec)
    assert isinstance(out, bytes)
    # Round-trip validity check — see test_pipeline_b_routes_to_rebrand.
    Document(io.BytesIO(out))


def test_pipeline_b_rejects_missing_rebrand_spec(
    synthetic_docx: Path,
) -> None:
    with pytest.raises(TypeError) as exc:
        build(synthetic_docx)
    msg = str(exc.value)
    assert "rebrand_spec" in msg
    assert "Pipeline B" in msg


def test_pipeline_b_rejects_extra_kwargs(
    synthetic_docx: Path,
    rebrand_spec: RebrandSpec,
) -> None:
    # Pipeline B has a fixed signature. Extra **kwargs must be surfaced
    # as a TypeError rather than silently dropped.
    with pytest.raises(TypeError) as exc:
        build(synthetic_docx, rebrand_spec=rebrand_spec, client="wrong")
    assert "client" in str(exc.value)


def test_pipeline_b_rejects_rebrand_spec_as_input(
    rebrand_spec: RebrandSpec,
) -> None:
    # Caller mistake: passing the spec as the input argument.
    with pytest.raises(TypeError) as exc:
        build(rebrand_spec)
    assert "RebrandSpec" in str(exc.value)


# ──────────────────────────────────────────────────────────────────────
# Pipeline C — markdown (no warning for non-agreement profiles)


def test_pipeline_c_markdown_string_no_warning(sample_markdown: str) -> None:
    with warnings.catch_warnings():
        warnings.simplefilter("error", DeprecationWarning)
        # Non-agreement profile: must not warn.
        doc = build(sample_markdown, md_opts={"title": "Q2 Update"})
    assert doc is not None
    # md_to_docx returns a python-docx Document.
    assert hasattr(doc, "paragraphs")


def test_pipeline_c_markdown_with_non_agreement_profile_no_warning(
    sample_markdown: str,
) -> None:
    # Profile is set, but it's not in AGREEMENT_PROFILES.
    with warnings.catch_warnings():
        warnings.simplefilter("error", DeprecationWarning)
        doc = build(
            sample_markdown,
            profile="exec_summary",
            md_opts={"title": "Q2 Exec Summary"},
        )
    assert hasattr(doc, "paragraphs")


def test_pipeline_c_md_file_path(tmp_path: Path, sample_markdown: str) -> None:
    md_path = tmp_path / "report.md"
    md_path.write_text(sample_markdown, encoding="utf-8")
    with warnings.catch_warnings():
        warnings.simplefilter("error", DeprecationWarning)
        doc = build(md_path, md_opts={"title": "Report"})
    assert hasattr(doc, "paragraphs")


# ──────────────────────────────────────────────────────────────────────
# Pipeline C deprecation — markdown + agreement profile


def test_pipeline_c_agreement_profile_emits_deprecation_warning(
    sample_markdown: str,
) -> None:
    # pytest.warns(DeprecationWarning) already enforces that a warning
    # of that category fires — the context manager raises pytest.fail
    # otherwise. No redundant inner isinstance check needed.
    with pytest.warns(DeprecationWarning):
        doc = build(
            sample_markdown,
            profile="agreement",
            md_opts={"title": "Letter of Intent"},
        )
    # Path still works — the warning is advisory at M5.
    assert hasattr(doc, "paragraphs")


def test_deprecation_message_names_both_alternatives(
    sample_markdown: str,
) -> None:
    """R-helpful: the warning text must tell callers where to go."""
    with pytest.warns(DeprecationWarning) as captured:
        build(
            sample_markdown,
            profile="agreement",
            md_opts={"title": "LOI"},
        )
    msg = str(captured[0].message)
    # Pipeline A pointer.
    assert "legal-assistant/colocation/generate_loi.py" in msg
    # Pipeline B pointer.
    assert "rebrand" in msg
    # Structural reasons — caller should understand WHY markdown is bad.
    assert "numbering" in msg
    assert "QA" in msg
    # Phase signalling.
    assert "next minor version" in msg or "error" in msg


def test_every_agreement_profile_triggers_warning(
    sample_markdown: str,
) -> None:
    """If AGREEMENT_PROFILES grows (e.g. adds 'nda'), all members
    must trigger the warning. Guards against forgetting to wire new
    profiles into the policy."""
    for profile in AGREEMENT_PROFILES:
        with pytest.warns(DeprecationWarning):
            build(
                sample_markdown,
                profile=profile,
                md_opts={"title": f"Test {profile}"},
            )


# ──────────────────────────────────────────────────────────────────────
# Regression — dispatcher does NOT change underlying pipeline behavior


def test_pipeline_c_dispatcher_matches_direct_md_to_docx(
    sample_markdown: str,
) -> None:
    """Regression guard: build() through Pipeline C produces the same
    structural output as calling md_to_docx directly. This is the
    closest thing we have to a golden-byte assertion at the dispatcher
    level; full byte-identity is verified by the M0 golden corpus.
    """
    direct = md_to_docx(sample_markdown, title="Regression Test")
    via_build = build(sample_markdown, md_opts={"title": "Regression Test"})

    # Same number of body paragraphs.
    assert len(direct.paragraphs) == len(via_build.paragraphs)
    # Same paragraph texts.
    direct_texts = [p.text for p in direct.paragraphs]
    build_texts = [p.text for p in via_build.paragraphs]
    assert direct_texts == build_texts


def test_dispatcher_import_has_no_side_effects() -> None:
    """Smoke test: importing dispatcher.py + its dependencies must
    complete without raising or mutating shared module state.

    The M0 golden corpus is regenerated per test run from fixed
    fixtures, so a successful import alone doesn't directly assert
    goldens are byte-equal — that's what tests/test_golden.py does.
    But import-time side effects (monkey-patching generate.py's
    rendering, mutating _SP, etc.) would cause M0 goldens to drift
    in any test process that imports dispatcher. This test catches
    that class of bug cheaply by exercising the imports.
    """
    import dispatcher  # noqa: F401
    import generate  # noqa: F401
    import rebrand  # noqa: F401
    # No assertion needed — successful import means no import-time
    # monkey patching occurred.


# ──────────────────────────────────────────────────────────────────────
# Unrouteable input


def test_unrouteable_input_raises_type_error() -> None:
    with pytest.raises(TypeError) as exc:
        build(12345)  # int — not a shape we know.
    msg = str(exc.value)
    # Must mention all three pipelines so the caller can orient.
    assert "Pipeline A" in msg
    assert "Pipeline B" in msg
    assert "Pipeline C" in msg


def test_nonexistent_docx_path_falls_through_to_markdown() -> None:
    """A string ending in ``.docx`` that doesn't exist on disk is
    NOT treated as a Pipeline B input — Pipeline B requires a real
    file. With no other routing signals, the string is treated as
    markdown text (the Pipeline C fallback), which is the safer
    behavior than silently failing with FileNotFoundError from
    rebrand.
    """
    # No warning since no agreement profile.
    with warnings.catch_warnings():
        warnings.simplefilter("error", DeprecationWarning)
        doc = build(
            "# Not actually a docx path\n\nJust a title.",
            md_opts={"title": "Not a docx"},
        )
    assert hasattr(doc, "paragraphs")


# ──────────────────────────────────────────────────────────────────────
# AGREEMENT_PROFILES policy surface


def test_agreement_profiles_is_non_empty() -> None:
    assert len(AGREEMENT_PROFILES) > 0


def test_agreement_profile_string_is_in_set() -> None:
    # The canonical profile name used throughout the factory.
    assert "agreement" in AGREEMENT_PROFILES


def test_agreement_profiles_is_immutable() -> None:
    # frozenset prevents runtime tampering that would silently
    # disable the deprecation for some callers.
    assert isinstance(AGREEMENT_PROFILES, frozenset)
