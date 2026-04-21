"""M3 — Pipeline B rebrand tests.

Covers:
  - RebrandSpec validation (reuses M2 validators).
  - UNDERSIGNED block detection + removal.
  - Shell placeholder stripping.
  - Old-cover stripping.
  - numId remap collision-freeness.
  - End-to-end rebrand on a synthetic minimal input.

The end-to-end test builds a synthetic input docx in-memory rather than
checking in a large fixture. A full MDCS-style verification lives in M6.
"""
from __future__ import annotations

import os
import sys
from pathlib import Path

import pytest
from docx import Document

_FACTORY = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(_FACTORY))

from rebrand import (  # noqa: E402
    RebrandSpec,
    rebrand,
    _UNDERSIGNED_RE,
    _RECITALS_LABEL_RE,
    _AGREED_RE,
    _CLAUSE_START_RE,
    _merge_numbering,
)
from validators import AgreementValidationError  # noqa: E402
from audit_profiles import audit_agreement  # noqa: E402


# ───────────────────────────── regex sanity


def test_undersigned_re_matches_clean_line() -> None:
    assert _UNDERSIGNED_RE.search("THE UNDERSIGNED:")


def test_undersigned_re_matches_date_collapsed() -> None:
    # Key MDCS edge case: cover date bled into body paragraph.
    assert _UNDERSIGNED_RE.search("19 March 2026THE UNDERSIGNED:")


def test_undersigned_re_matches_lowercase() -> None:
    assert _UNDERSIGNED_RE.search("The Undersigned:")


def test_undersigned_re_matches_mid_sentence_known_boundary() -> None:
    """Documents a known false-positive boundary.

    The regex is intentionally unanchored (MDCS needed to match the
    date-prefix-collapsed case `19 March 2026THE UNDERSIGNED:`). As a
    consequence, it also matches the phrase mid-sentence in prose like
    `"signed by the undersigned party"`. Today this is acceptable
    because:

      - `audit_agreement`'s R-24 check uses `.match()` (start-of-string,
        with optional whitespace), so mid-paragraph prose mentioning
        "undersigned" does NOT trip the rule.
      - `_detect_body_start` uses `.search()` but runs on the INPUT
        document before any body content has been authored against the
        new cover; real inputs don't contain prose before the parties
        block.

    If a real input ever surfaces where prose mentioning "undersigned"
    appears before the parties block, tighten the regex (line-start or
    explicit colon suffix) — this test will start failing, which is
    the desired signal that the fix is needed."""
    assert _UNDERSIGNED_RE.search("signed by the undersigned party on the date")


def test_recitals_label_re() -> None:
    assert _RECITALS_LABEL_RE.match("RECITALS:")
    assert _RECITALS_LABEL_RE.match("RECITALS")
    assert _RECITALS_LABEL_RE.match("  Recitals  ")
    assert not _RECITALS_LABEL_RE.match("RECITALS: some text")


def test_agreed_re() -> None:
    assert _AGREED_RE.match("HAVE AGREED AS FOLLOWS:")
    assert _AGREED_RE.match("Have agreed as follows.")


def test_clause_start_re() -> None:
    assert _CLAUSE_START_RE.match("1. Definitions")
    assert _CLAUSE_START_RE.match("2.1 Scope")
    assert _CLAUSE_START_RE.match("10.9 Waiver")
    assert not _CLAUSE_START_RE.match("This is a paragraph about clause 1.")


# ───────────────────────────── spec validation (M2 reuse)


def test_rebrand_spec_rejects_missing_address() -> None:
    with pytest.raises(AgreementValidationError) as ei:
        RebrandSpec(
            agreement_type="Letter of Intent",
            client="Acme B.V.",
            client_address="",  # M2 rejects
        )
    assert "client_address" in str(ei.value)


def test_rebrand_spec_happy_path() -> None:
    # Does not raise.
    RebrandSpec(
        agreement_type="Letter of Intent",
        client="Acme B.V.",
        client_address="1 Main St, 1000 AA Amsterdam",
    )


# ───────────────────────────── numbering merge


def test_merge_numbering_on_empty_returns_empty(tmp_path: Path) -> None:
    shell = tmp_path / "shell_numbering.xml"
    inp = tmp_path / "input_numbering.xml"
    # No input numbering → returns empty map, no crash.
    shell.write_text(
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    )
    assert _merge_numbering(shell, inp) == {}


def test_merge_numbering_range_based_no_collisions(tmp_path: Path) -> None:
    """Regression guard: the M-1 spike's first attempt produced duplicate
    IDs. Fix: remap all input IDs into a range above shell's max."""
    NS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
    shell = tmp_path / "shell_numbering.xml"
    inp = tmp_path / "input_numbering.xml"

    shell.write_text(
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        f'<w:numbering {NS}>'
        f'  <w:abstractNum w:abstractNumId="0"/>'
        f'  <w:abstractNum w:abstractNumId="1"/>'
        f'  <w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>'
        f'  <w:num w:numId="2"><w:abstractNumId w:val="1"/></w:num>'
        f'</w:numbering>'
    )
    inp.write_text(
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<w:numbering {NS}>'
        f'  <w:abstractNum w:abstractNumId="0"/>'
        f'  <w:abstractNum w:abstractNumId="1"/>'
        f'  <w:abstractNum w:abstractNumId="2"/>'
        f'  <w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>'
        f'  <w:num w:numId="2"><w:abstractNumId w:val="1"/></w:num>'
        f'  <w:num w:numId="3"><w:abstractNumId w:val="2"/></w:num>'
        f'</w:numbering>'
    )

    num_map = _merge_numbering(shell, inp)

    # Input num IDs {1, 2, 3} must remap to values >= 3 (max shell + 1).
    assert num_map == {1: 3, 2: 4, 3: 5}

    # Reload merged, assert every ID appears exactly once.
    from lxml import etree
    W_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    QN = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    t = etree.parse(str(shell))
    abs_ids = [
        int(e.get(QN + "abstractNumId"))
        for e in t.getroot().findall("w:abstractNum", W_NS)
    ]
    num_ids = [
        int(e.get(QN + "numId"))
        for e in t.getroot().findall("w:num", W_NS)
    ]
    assert len(abs_ids) == len(set(abs_ids)), f"abstract ID collision: {abs_ids}"
    assert len(num_ids) == len(set(num_ids)), f"num ID collision: {num_ids}"

    # w:num elements must all have valid abstract refs (within merged set).
    abs_id_set = set(abs_ids)
    for num_el in t.getroot().findall("w:num", W_NS):
        ref = int(num_el.find("w:abstractNumId", W_NS).get(QN + "val"))
        assert ref in abs_id_set, f"dangling abstract ref: {ref}"


# ───────────────────────────── end-to-end on synthetic input


def _build_synthetic_loi_input(path: Path) -> Path:
    """Produce a minimal LOI-shaped .docx for rebrand testing.

    Structure:
      - Old cover (title + parties) — must be stripped.
      - THE UNDERSIGNED: paragraph with a collapsed date prefix — must be
        stripped together with the 2 party lines and the "together the
        Parties" line.
      - RECITALS: label — starts body content (must be preserved).
      - A recital paragraph.
      - HAVE AGREED AS FOLLOWS: label.
      - Clause 1 heading + body.
      - Clause 2 heading + body.
      - Signature block.
    """
    doc = Document()
    # Old cover
    doc.add_paragraph("LETTER OF INTENT")
    doc.add_paragraph("between Old DE Entity N.V. and Acme Test B.V.")
    doc.add_paragraph("1 January 2026")
    # Body
    doc.add_paragraph("1 January 2026THE UNDERSIGNED:")
    doc.add_paragraph("(1) Old DE Entity N.V., etc.")
    doc.add_paragraph("(2) Acme Test B.V., etc.")
    doc.add_paragraph('(together the "Parties"),')
    doc.add_paragraph("RECITALS:")
    doc.add_paragraph(
        "The Provider is part of the Digital Energy group and develops "
        "colocation facilities."
    )
    doc.add_paragraph("HAVE AGREED AS FOLLOWS:")
    h1 = doc.add_paragraph("1. Definitions")
    h1.style = "Heading 1"
    doc.add_paragraph('1.1 "Services" means the services described in Clause 2.')
    h2 = doc.add_paragraph("2. Scope")
    h2.style = "Heading 1"
    doc.add_paragraph("2.1 The Provider will deliver the Services.")
    doc.add_paragraph("")
    doc.add_paragraph("For and on behalf of Old DE Entity N.V.")
    doc.add_paragraph("Signature: _____________")
    doc.add_paragraph("For and on behalf of Acme Test B.V.")
    doc.add_paragraph("Signature: _____________")
    doc.save(str(path))
    return path


def test_rebrand_end_to_end_on_synthetic_input(tmp_path: Path) -> None:
    """End-to-end guarantee: rebrand a synthetic LOI, assert that:
      - Output renders (python-docx opens it).
      - Old cover is stripped (no "Old DE Entity N.V." on cover).
      - UNDERSIGNED block is stripped (no UNDERSIGNED anywhere in output).
      - DE cover has been inserted (new counterparty name + address).
      - Recitals preserved (RECITALS: label survives).
      - Clauses preserved (Heading 1 style on clause paragraphs).
      - Signature content preserved.
      - `audit_agreement` reports no R-24 and no placeholder violations.
    """
    input_path = _build_synthetic_loi_input(tmp_path / "input.docx")
    spec = RebrandSpec(
        agreement_type="Letter of Intent",
        subject="for Pipeline B test",
        client="Acme Test B.V.",
        client_address="1 Test Straat, 1000 AA Amsterdam",
        entity="nl",
        date_str="17 April 2026",
    )
    stats: dict = {}
    data = rebrand(input_path, spec, stats_out=stats)

    out_path = tmp_path / "out.docx"
    out_path.write_bytes(data)

    # Open — must not raise.
    doc = Document(str(out_path))
    texts = [p.text for p in doc.paragraphs]

    # UNDERSIGNED gone.
    assert not any("UNDERSIGNED" in t.upper() for t in texts), (
        "UNDERSIGNED leaked into output:\n" + "\n".join(t for t in texts if "UNDERSIGNED" in t.upper())
    )

    # Old cover gone (old DE entity must not appear on cover).
    # "Old DE Entity N.V." survives in the body signature block (that's
    # fine — we preserve body as-is); but it must not be on cover (first
    # ~20 paragraphs are cover territory).
    cover_texts = "\n".join(texts[:20])
    assert "Old DE Entity" not in cover_texts, "Old cover party leaked into new cover"

    # New DE cover inserted.
    assert any("Acme Test B.V." in t for t in texts[:20])
    assert any("1 Test Straat" in t for t in texts[:20])
    assert any("Letter of Intent" in t for t in texts[:20])
    assert any("17 April 2026" in t for t in texts[:20])

    # Recitals preserved.
    assert any(t.strip().upper().startswith("RECITALS") for t in texts)

    # Clauses preserved with heading style.
    heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert any("Definitions" in t for t in heading_texts)
    assert any("Scope" in t for t in heading_texts)

    # Signature content preserved (body verbatim).
    assert any("Signature:" in t for t in texts)

    # Semantic audit (M2): no R-24, no placeholder.
    vios = audit_agreement(doc)
    r24 = [v for v in vios if "R-24" in v]
    placeholder = [v for v in vios if "placeholder" in v]
    assert not r24, f"R-24 fired: {r24}"
    assert not placeholder, f"Placeholder fired: {placeholder}"

    # Rebrand statistics sanity.
    assert stats["stripped_shell_placeholder"] == 1, stats
    assert stats["stripped_undersigned"] >= 1, stats  # at least the header
    assert stats["copied_elements"] >= 5, stats


def test_rebrand_end_to_end_rejects_bad_spec(tmp_path: Path) -> None:
    input_path = _build_synthetic_loi_input(tmp_path / "input.docx")
    with pytest.raises(AgreementValidationError):
        rebrand(
            input_path,
            RebrandSpec(
                agreement_type="Letter of Intent",
                client="Acme Test B.V.",
                client_address="",  # rejected
            ),
        )


# ─────────────────────────────────────────────────────────────────────
# M3.1 — additional golden pairs: counsel MSA stub, counterparty NDA,
# tracked changes preservation.
#
# Per the plan, Pipeline B must also handle (a) documents we receive
# from outside counsel as drafts, (b) documents supplied by the
# counterparty, and (c) documents containing tracked changes that must
# survive the rebrand intact. All three cases are exercised below with
# in-memory synthetic inputs to keep fixtures out of git.


def _build_synthetic_msa_stub(path: Path) -> Path:
    """Counsel-style MSA stub.

    Distinguishing features vs. the LOI helper:
      - Explicit `MASTER SERVICE AGREEMENT` title + "by and between" (binding).
      - `RECITALS` heading (not `RECITALS:` — word only).
      - Parenthetical letter recitals `(A)/(B)/(C)`.
      - `AGREED TERMS` heading instead of `HAVE AGREED AS FOLLOWS:`.
      - `EXECUTED AS A DEED` signature block.
      - Pipe-delimited table-of-definitions (rendered as paragraphs, not
        a python-docx table — matches many counsel-drafted stubs).
    """
    doc = Document()
    # Cover (counsel style)
    doc.add_paragraph("MASTER SERVICE AGREEMENT")
    doc.add_paragraph("by and between")
    doc.add_paragraph("Party A Limited (company number 111111)")
    doc.add_paragraph("and")
    doc.add_paragraph("Party B Limited (company number 222222)")
    doc.add_paragraph("Dated: 1 January 2026")

    # Recitals
    doc.add_paragraph("RECITALS")
    doc.add_paragraph("(A) Party A is in the business of providing services.")
    doc.add_paragraph("(B) Party B wishes to procure those services.")
    doc.add_paragraph("(C) The parties wish to record the terms.")

    # Agreed terms
    doc.add_paragraph("AGREED TERMS")

    # Definitions table (simulated as paragraphs)
    doc.add_paragraph("Term | Definition")
    doc.add_paragraph("Services | means the services described in Schedule 1")
    doc.add_paragraph("Fees | means the fees payable by the Customer")
    doc.add_paragraph("Term | means the period from the Effective Date")

    # Clauses
    h1 = doc.add_paragraph("1. Definitions and Interpretation")
    h1.style = "Heading 1"
    doc.add_paragraph("1.1 In this Agreement, unless the context requires otherwise:")
    doc.add_paragraph("(a) capitalised terms have the meanings given above;")
    doc.add_paragraph("(b) references to Clauses are references to clauses of this Agreement.")

    h2 = doc.add_paragraph("2. Services")
    h2.style = "Heading 1"
    doc.add_paragraph("2.1 The Supplier shall provide the Services in accordance with Schedule 1.")

    h3 = doc.add_paragraph("3. Fees")
    h3.style = "Heading 1"
    doc.add_paragraph("3.1 The Customer shall pay the Fees in accordance with Schedule 2.")

    # Execution block
    doc.add_paragraph("")
    doc.add_paragraph("EXECUTED AS A DEED")
    doc.add_paragraph("For and on behalf of Party A Limited")
    doc.add_paragraph("Signature: ____________________")
    doc.add_paragraph("For and on behalf of Party B Limited")
    doc.add_paragraph("Signature: ____________________")
    doc.save(str(path))
    return path


def test_rebrand_msa_stub_from_counsel(tmp_path: Path) -> None:
    """Counsel-supplied binding MSA stub rebrands onto DE AG shell.

    Assertions:
      - Old counsel cover stripped (no "Party A Limited" or
        "Party B Limited" in cover region — they survive in the
        execution block, that's fine).
      - DE cover shows binding party labels ("By and between:") and
        both parties' registration numbers (required by M2 for
        formality=binding).
      - Recitals preserved: `(A)`, `(B)`, `(C)` lettering and content.
      - AGREED TERMS heading preserved.
      - Clauses preserved with Heading 1 style.
      - Definitions rows preserved (paragraph form).
      - EXECUTED AS A DEED block preserved.
      - audit_agreement clean on R-24 + placeholder checks.
    """
    input_path = _build_synthetic_msa_stub(tmp_path / "msa_stub.docx")
    spec = RebrandSpec(
        agreement_type="Master Service Agreement",
        subject="for Pipeline B counsel-stub test",
        client="External Counterparty Ltd",
        client_address="10 Old Bailey, London EC4M 7AA, United Kingdom",
        entity="ag",
        date_str="17 April 2026",
        # Binding requires registration (M2 validator).
        client_reg_type="CRN",
        client_reg_number="99999999",
    )
    stats: dict = {}
    data = rebrand(input_path, spec, stats_out=stats)
    out_path = tmp_path / "msa_out.docx"
    out_path.write_bytes(data)

    doc = Document(str(out_path))
    texts = [p.text for p in doc.paragraphs]

    # The DE cover is the first N paragraphs. "Party A/B Limited" must
    # not appear in the cover region (counsel cover stripped). They DO
    # legitimately appear in the execution block at the end.
    cover_region = "\n".join(texts[:25])
    assert "Party A Limited" not in cover_region, (
        f"Old counsel cover leaked into new cover:\n{cover_region}"
    )
    assert "Party B Limited" not in cover_region

    # DE cover inserted (binding → "By and between:").
    assert any("External Counterparty Ltd" in t for t in texts[:25])
    assert any("By and between" in t for t in texts[:25])
    # CRN number must appear somewhere in the cover (binding formality).
    assert any("99999999" in t for t in texts[:25])

    # Body content preserved.
    assert any("RECITALS" == t.strip() for t in texts)
    assert any(t.strip().startswith("(A)") for t in texts)
    assert any(t.strip().startswith("(B)") for t in texts)
    assert any(t.strip().startswith("(C)") for t in texts)
    assert any("AGREED TERMS" == t.strip() for t in texts)

    # Clauses preserved with Heading style.
    heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert any("Definitions" in t for t in heading_texts)
    assert any("Services" in t for t in heading_texts)
    assert any("Fees" in t for t in heading_texts)

    # Definitions rows preserved.
    assert any("means the services described in Schedule 1" in t for t in texts)

    # Execution block preserved.
    assert any("EXECUTED AS A DEED" in t for t in texts)

    # R-24 and placeholder checks pass.
    vios = audit_agreement(doc)
    r24 = [v for v in vios if "R-24" in v]
    placeholder = [v for v in vios if "placeholder" in v]
    assert not r24, f"R-24: {r24}"
    assert not placeholder, f"placeholder: {placeholder}"


def _build_synthetic_nda(path: Path) -> Path:
    """Counterparty-supplied NDA with a plain-text parties block (no
    UNDERSIGNED heading) and a short clause set."""
    doc = Document()
    # Counterparty-drafted cover
    doc.add_paragraph("NON-DISCLOSURE AGREEMENT")
    doc.add_paragraph("between ACME Corporation and DE Netherlands B.V.")
    doc.add_paragraph("Effective date: 1 March 2026")

    # Plain-text parties block (no UNDERSIGNED heading — different shape)
    doc.add_paragraph("THE PARTIES")
    doc.add_paragraph("ACME Corporation, a Delaware corporation (the 'Disclosing Party'); and")
    doc.add_paragraph("DE Netherlands B.V., a Dutch private limited company (the 'Receiving Party').")
    doc.add_paragraph("IN CONSIDERATION of the mutual covenants below, the parties agree as follows:")

    # Clauses
    for n, title, body in [
        ("1", "Confidential Information", "Means any non-public information disclosed by the Disclosing Party."),
        ("2", "Obligations of the Receiving Party", "The Receiving Party shall keep Confidential Information secret."),
        ("3", "Term", "This Agreement shall remain in force for three (3) years from the Effective Date."),
        ("4", "Return of Information", "Upon termination, the Receiving Party shall return all materials."),
    ]:
        h = doc.add_paragraph(f"{n}. {title}")
        h.style = "Heading 1"
        doc.add_paragraph(f"{n}.1 {body}")

    doc.add_paragraph("")
    doc.add_paragraph("Signed: ___________________")
    doc.add_paragraph("For ACME Corporation")
    doc.add_paragraph("Signed: ___________________")
    doc.add_paragraph("For DE Netherlands B.V.")
    doc.save(str(path))
    return path


def test_rebrand_nda_from_counterparty(tmp_path: Path) -> None:
    """Counterparty-supplied NDA with non-standard parties block.

    Exercises the fallback body-start detection: the input has no
    UNDERSIGNED marker, no RECITALS label, no AGREED TERMS, no
    parenthetical recital lettering. _detect_body_start falls back to
    the first Heading-styled paragraph (clause 1).

    As a consequence, "THE PARTIES" + the two party lines + the
    "IN CONSIDERATION" lead-in are all stripped as old-cover prelude.
    That's the intended behaviour for this shape: an NDA with a plain
    parties block BEFORE clauses is semantically the same duplication
    that UNDERSIGNED blocks are — both restate parties the new cover
    already carries.

    Known limitation: NDAs that include LEGITIMATE body content before
    the first heading (e.g. jurisdictional qualifiers, implicit
    recitals without a label) will see that content stripped too. If a
    concrete input case surfaces, extend _detect_body_start with a new
    marker — do NOT widen the UNDERSIGNED regex to cover PARTIES, the
    R-24 rule intentionally scopes to the literal UNDERSIGNED heading.
    """
    input_path = _build_synthetic_nda(tmp_path / "nda.docx")
    spec = RebrandSpec(
        agreement_type="Non-Disclosure Agreement",
        subject="for Pipeline B NDA test",
        client="ACME Corporation",
        client_address="1 Main St, Wilmington DE 19801, USA",
        entity="nl",
        date_str="17 April 2026",
        # _detect_formality may classify NDAs as binding by default;
        # the plan treats a supplier NDA as non-binding for this test.
        formality="non_binding",
    )
    data = rebrand(input_path, spec)
    out_path = tmp_path / "nda_out.docx"
    out_path.write_bytes(data)

    doc = Document(str(out_path))
    texts = [p.text for p in doc.paragraphs]

    # Old cover stripped: DE cover should carry the title-case DE version
    # ("Non-Disclosure Agreement") and the input's UPPERCASE version must
    # not appear anywhere in the output. Explicit positive + negative
    # instead of the previous split()-based hack that was meaningless if
    # "Non-Disclosure" happened to be absent from the cover.
    assert any("Non-Disclosure Agreement" in t for t in texts[:25]), (
        "DE cover missing the title-case agreement type"
    )
    assert not any("NON-DISCLOSURE AGREEMENT" in t for t in texts), (
        "Counterparty's uppercase NDA title leaked into the output"
    )
    # DE cover inserted with non-binding party labels.
    assert any("Between:" in t for t in texts[:25])
    assert any("ACME Corporation" in t for t in texts[:25])

    # THE PARTIES block + the "IN CONSIDERATION" lead-in are stripped
    # as cover prelude (fallback-to-first-heading behaviour). They
    # appear nowhere in the output.
    assert not any("THE PARTIES" == t.strip() for t in texts), (
        "THE PARTIES should have been stripped as cover prelude"
    )
    assert not any("IN CONSIDERATION" in t.upper() for t in texts), (
        "IN CONSIDERATION lead-in should have been stripped as cover prelude"
    )

    # Clauses preserved.
    heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert any("Confidential Information" in t for t in heading_texts)
    assert any("Obligations" in t for t in heading_texts)
    assert any("Term" in t for t in heading_texts)
    assert any("Return of Information" in t for t in heading_texts)

    # Signature block still preserved (it's after the last heading,
    # never part of any "cover prelude" region).
    assert any("For ACME Corporation" in t for t in texts)

    # Audit clean: no R-24 (no UNDERSIGNED heading anywhere) and no
    # placeholders leaked.
    vios = audit_agreement(doc)
    r24 = [v for v in vios if "R-24" in v]
    placeholder = [v for v in vios if "placeholder" in v]
    assert not r24, f"R-24: {r24}"
    assert not placeholder, f"placeholder: {placeholder}"


def _build_synthetic_input_with_tracked_changes(path: Path) -> Path:
    """Synthetic LOI with one <w:ins> and one <w:del> element.

    Direct lxml manipulation of python-docx's document tree after save
    (python-docx has no public tracked-changes API). Insert an <w:ins>
    around a run in one paragraph and a <w:del> in another.
    """
    # Start with the minimal LOI helper, save, then re-open via zip.
    _build_synthetic_loi_input(path)

    import zipfile
    import shutil
    from lxml import etree

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    NS = {"w": W}
    QN = lambda t: f"{{{W}}}{t}"

    # Unzip → edit document.xml → rezip.
    tmp_dir = path.parent / "_tracked_unpacked"
    if tmp_dir.exists():
        shutil.rmtree(tmp_dir)
    tmp_dir.mkdir()
    with zipfile.ZipFile(path, "r") as zf:
        zf.extractall(tmp_dir)

    doc_xml = tmp_dir / "word" / "document.xml"
    tree = etree.parse(str(doc_xml))
    root = tree.getroot()
    body = root.find("w:body", NS)

    paragraphs = body.findall("w:p", NS)
    # Wrap one run of the 2nd paragraph in <w:ins>.
    target_p = None
    for p in paragraphs:
        text = "".join(t.text or "" for t in p.iter(QN("t")))
        if "RECITALS" in text.upper():
            # Use the paragraph right AFTER recitals (a recital body).
            idx = list(body).index(p)
            if idx + 1 < len(list(body)):
                target_p = list(body)[idx + 1]
            break
    if target_p is None:
        target_p = paragraphs[5] if len(paragraphs) > 5 else paragraphs[-1]

    # Wrap first run in <w:ins>.
    first_run = target_p.find("w:r", NS)
    if first_run is not None:
        ins = etree.SubElement(target_p, QN("ins"))
        ins.set(QN("id"), "1")
        ins.set(QN("author"), "Carlos")
        ins.set(QN("date"), "2026-04-17T00:00:00Z")
        # Move the run into the <w:ins>
        target_p.remove(first_run)
        ins.append(first_run)
        # Re-insert the ins element at the original run position.
        target_p.insert(0, ins)

    # Add a <w:del> wrapper with a <w:delText> in another paragraph.
    del_p = paragraphs[-3] if len(paragraphs) >= 3 else paragraphs[-1]
    del_el = etree.SubElement(del_p, QN("del"))
    del_el.set(QN("id"), "2")
    del_el.set(QN("author"), "Carlos")
    del_el.set(QN("date"), "2026-04-17T00:00:00Z")
    del_r = etree.SubElement(del_el, QN("r"))
    del_t = etree.SubElement(del_r, QN("delText"))
    del_t.text = " DELETED TEXT "

    tree.write(
        str(doc_xml),
        xml_declaration=True,
        encoding="UTF-8",
        standalone=True,
    )

    # Rezip.
    out_path = path.parent / (path.stem + "_tracked.docx")
    if out_path.exists():
        out_path.unlink()
    with zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for root_dir, _dirs, files in os.walk(tmp_dir):
            for f in files:
                full = Path(root_dir) / f
                arcname = full.relative_to(tmp_dir)
                zf.write(full, str(arcname))
    # Replace the original.
    shutil.copy(out_path, path)
    shutil.rmtree(tmp_dir)
    out_path.unlink()
    return path


def test_rebrand_preserves_tracked_changes(tmp_path: Path) -> None:
    """Rebrand MUST NOT silently strip <w:ins> or <w:del> elements.

    Build a synthetic LOI, patch it to include one insertion and one
    deletion via lxml, rebrand it, unzip the output, and assert both
    revision elements survived with their author attributes intact.
    """
    input_path = _build_synthetic_input_with_tracked_changes(tmp_path / "input.docx")

    # Sanity: the patching worked — ins and del in the input.
    import zipfile
    from lxml import etree

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    NS = {"w": W}

    with zipfile.ZipFile(input_path, "r") as zf:
        input_tree = etree.fromstring(zf.read("word/document.xml"))
    assert input_tree.findall(".//w:ins", NS), "Input must contain <w:ins>"
    assert input_tree.findall(".//w:del", NS), "Input must contain <w:del>"

    spec = RebrandSpec(
        agreement_type="Letter of Intent",
        subject="for tracked-changes preservation test",
        client="Tracked B.V.",
        client_address="1 Track Straat, 1000 AA Amsterdam",
        entity="nl",
        date_str="17 April 2026",
    )
    data = rebrand(input_path, spec)
    out_path = tmp_path / "tracked_out.docx"
    out_path.write_bytes(data)

    # The output must still contain both revision elements.
    with zipfile.ZipFile(out_path, "r") as zf:
        output_tree = etree.fromstring(zf.read("word/document.xml"))

    ins_elements = output_tree.findall(".//w:ins", NS)
    del_elements = output_tree.findall(".//w:del", NS)
    W_AUTHOR = f"{{{W}}}author"
    assert ins_elements, (
        "Tracked-change <w:ins> element lost during rebrand. "
        "rebrand.py must NEVER strip revision elements."
    )
    assert any(ins.get(W_AUTHOR) == "Carlos" for ins in ins_elements), (
        "w:author attribute on <w:ins> lost during rebrand"
    )
    assert del_elements, (
        "Tracked-change <w:del> element lost during rebrand."
    )
    assert any(d.get(W_AUTHOR) == "Carlos" for d in del_elements), (
        "w:author attribute on <w:del> lost during rebrand"
    )
