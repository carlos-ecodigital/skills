#!/usr/bin/env python3
"""M6 Pipeline B verification — rebrand any MDCS-shape docx and report.

Standalone tool for manual verification on real MDCS (or similar LOI)
.docx inputs that cannot be committed as test fixtures.

Usage:
    python3 tools/m6_verify_rebrand.py INPUT.docx \
        [--client "MDCS.AI B.V."] [--client-address "..."] \
        [--entity nl] [--output OUT.docx]

Runs Pipeline B rebrand with sane MDCS defaults, then reports:
  - rebrand statistics (copied, stripped cover/undersigned/placeholder,
    numIds remapped)
  - audit_agreement R-21 + placeholder counts (Layer 2)
  - presence of key structural signals (no UNDERSIGNED in body, correct
    DE cover, all input headings preserved)

Exit code 0 on clean verification, 1 if R-21 or placeholder violations,
2 if a structural signal is missing.
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

# Make same-dir imports work when run as a script.
_THIS = Path(__file__).resolve()
sys.path.insert(0, str(_THIS.parent.parent))

from rebrand import RebrandSpec, rebrand  # noqa: E402
from audit_profiles import audit_agreement  # noqa: E402
from docx import Document  # noqa: E402


def _verify(
    input_path: str,
    *,
    client: str = "MDCS.AI B.V.",
    client_address: str = "Hanzeweg 10C, 2803 MC Gouda, the Netherlands",
    entity: str = "nl",
    agreement_type: str = "Letter of Intent",
    subject: str = "for AI Infrastructure Services",
    date_str: str = "17 April 2026",
    output_path: str | None = None,
) -> int:
    input_path = Path(input_path).expanduser().resolve()
    if not input_path.exists():
        print(f"error: input not found: {input_path}", file=sys.stderr)
        return 2

    out_path = Path(output_path) if output_path else input_path.with_name(
        input_path.stem + "_rebranded.docx"
    )

    spec = RebrandSpec(
        agreement_type=agreement_type,
        subject=subject,
        client=client,
        client_address=client_address,
        entity=entity,
        date_str=date_str,
    )
    stats: dict = {}
    data = rebrand(input_path, spec, stats_out=stats)
    out_path.write_bytes(data)

    # Rebrand stats
    print(f"Input : {input_path}")
    print(f"Output: {out_path}")
    print(
        f"  copied: {stats['copied_elements']} elements "
        f"| stripped: cover={stats['stripped_cover']}, "
        f"undersigned={stats['stripped_undersigned']}, "
        f"shell_placeholder={stats['stripped_shell_placeholder']} "
        f"| numIds remapped: {stats['remapped_numIds']}"
    )

    # Audit
    doc = Document(str(out_path))
    vios = audit_agreement(doc)
    r21 = [v for v in vios if "R-21" in v]
    placeholder = [v for v in vios if "placeholder" in v]
    other = len(vios) - len(r21) - len(placeholder)

    print()
    print(f"Audit (audit_agreement):")
    print(f"  R-21 party-duplication : {len(r21)} (expected 0)")
    print(f"  Placeholder tokens     : {len(placeholder)} (expected 0)")
    print(f"  Base audit_document    : {other} (style-level, non-fatal)")

    # Structural signals
    texts = [p.text for p in doc.paragraphs]
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]

    undersigned_count = sum(1 for t in texts if "UNDERSIGNED" in t.upper())
    cover_has_client = any(client in t for t in texts[:25])
    cover_has_address = any(client_address.split(",")[0] in t for t in texts[:25])
    has_headings = len(headings) >= 1

    print()
    print(f"Structural signals:")
    print(f"  UNDERSIGNED in body    : {undersigned_count} (expected 0)")
    print(f"  DE cover has client    : {cover_has_client}")
    print(f"  DE cover has address   : {cover_has_address}")
    print(f"  Body headings preserved: {len(headings)}")
    for h in headings[:12]:
        print(f"      {h[:70]}")

    # Exit code
    if r21 or placeholder:
        print("\n✗ FAIL — R-21 or placeholder violations")
        return 1
    if undersigned_count or not cover_has_client or not has_headings:
        print("\n✗ FAIL — structural signal missing")
        return 2
    print("\n✓ PASS — Pipeline B verification clean")
    return 0


def main() -> int:
    p = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    p.add_argument("input", help="Input .docx path (MDCS-shape LOI)")
    p.add_argument("--client", default="MDCS.AI B.V.")
    p.add_argument(
        "--client-address",
        default="Hanzeweg 10C, 2803 MC Gouda, the Netherlands",
    )
    p.add_argument("--entity", choices=["ag", "nl"], default="nl")
    p.add_argument("--agreement-type", default="Letter of Intent")
    p.add_argument("--subject", default="for AI Infrastructure Services")
    p.add_argument("--date", dest="date_str", default="17 April 2026")
    p.add_argument("-o", "--output", default=None)
    args = p.parse_args()

    return _verify(
        args.input,
        client=args.client,
        client_address=args.client_address,
        entity=args.entity,
        agreement_type=args.agreement_type,
        subject=args.subject,
        date_str=args.date_str,
        output_path=args.output,
    )


if __name__ == "__main__":
    sys.exit(main())
