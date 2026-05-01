"""
Post-assembly structural-health validators for generated ``.docx`` files.

These checks run AFTER a document has been assembled and are intended to
catch structural defects that are hard to catch at render time: table
width overflow, cells that overflow on long words, excessive list nesting,
unexpected diacritics, and font drift.

Additive to ``document-factory/generate.py``; no existing signatures
changed.

Usage:

    from format_validators import run_all
    issues = run_all(doc)
    if issues:
        for i in issues: print(i)

Each validator returns a ``list[str]`` of issue messages. Empty list =
clean. ``run_all`` concatenates the results of all validators in a stable
order.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import List

from docx.oxml.ns import qn

from generate import FONT

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

#: Maximum usable width in mm on A4 with 25/20 mm L/R margins.
DEFAULT_MAX_WIDTH_MM = 165.0

TWIPS_PER_INCH = 1440
MM_PER_INCH = 25.4

#: Default allowed diacritic range. Covers:
#:   - Latin-1 Supplement ``À-ÿ`` (é, ë, ï, ç, ñ, ö, ü, ß etc.)
#:   - Currency & typography: ``€ – — ° § ¶ ¢ £ ¥``
#:   - Superscripts used in legal (``m²``, ``m³``): ``² ³ ¹``
#:   - Smart quotes used for defined terms and quotations:
#:     left/right double ``" "`` and left/right single ``' '``, plus
#:     prime ``′ ″`` and horizontal ellipsis ``…``.
#:   - Bullet char used by bilingual_body list rendering: ``•``
#:   - Various dashes: en/em ``– —``, figure dash ``‒``, horizontal bar ``―``.
DEFAULT_ALLOWED_RANGE = (
    r"À-ÿ"                              # Latin-1 Supplement
    r"€°§¶¢£¥"                          # currency + common symbols
    r"¹²³"                              # superscripts
    r"\u2018\u2019\u201C\u201D"         # smart single/double quotes
    r"\u2022"                           # bullet •
    r"\u2013\u2014\u2012\u2015"         # dashes – — ‒ ―
    r"\u2032\u2033"                     # prime ′ ″
    r"\u2026"                           # ellipsis …
)

#: Default list-nesting ceiling.
DEFAULT_MAX_LIST_DEPTH = 3

#: Default expected font family.
DEFAULT_FONT_FAMILY = FONT  # "Inter"

#: Bundled Inter font path. Used by ``validate_cell_overflow`` to compute
#: an evidence-backed average advance width. SIL OFL 1.1 licence
#: bundled alongside.
_INTER_FONT_PATH = (
    Path(__file__).parent / "assets" / "fonts" / "Inter-Regular.ttf"
)

#: Specimen used to compute the average advance width — Latin alpha + digits
#: + the NL diacritics that appear in real Sites-stream content.
_INTER_SPECIMEN = (
    "abcdefghijklmnopqrstuvwxyz"
    "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
    "0123456789"
    "éëïöüç"
)


def _measure_inter_avg_advance_mm(point_size: float = 10.0) -> float:
    """Compute the average advance width of ``_INTER_SPECIMEN`` at
    ``point_size`` pt, in millimetres, against the bundled Inter Regular
    .ttf.

    Falls back to the legacy ``2.1 mm`` constant if ``fontTools`` is not
    installed or the font file is missing — the validator stays
    operational on minimal environments, with a known imprecise
    estimate. CI installs ``fonttools`` and bundles the font, so the
    fallback only matters for ad-hoc local runs without the dev deps.
    """
    try:
        from fontTools.ttLib import TTFont  # type: ignore
    except Exception:
        return 2.1
    if not _INTER_FONT_PATH.exists():
        return 2.1
    font = TTFont(str(_INTER_FONT_PATH))
    cmap = font.getBestCmap()
    hmtx = font["hmtx"]
    upm = font["head"].unitsPerEm
    total = 0
    n = 0
    for ch in _INTER_SPECIMEN:
        cp = ord(ch)
        if cp not in cmap:
            continue
        glyph_name = cmap[cp]
        adv, _lsb = hmtx[glyph_name]
        total += adv
        n += 1
    if n == 0:
        return 2.1
    avg_units = total / n
    # advance(units) / UPM gives ems; * pt-size = pt; / 72 = inch; * 25.4 = mm
    return avg_units / upm * point_size / 72.0 * MM_PER_INCH


#: Pre-computed at import time so each ``validate_cell_overflow`` call is
#: a constant lookup rather than an O(n) font walk. Re-measured on every
#: import — measured value is authoritative; the legacy 5 % tolerance is
#: removed since the constant is now empirical, not a guess.
_INTER_10PT_AVG_MM = _measure_inter_avg_advance_mm(10.0)


# ---------------------------------------------------------------------------
# Individual validators
# ---------------------------------------------------------------------------

def validate_table_widths(doc, max_mm: float = DEFAULT_MAX_WIDTH_MM) -> List[str]:
    """Flag any table whose ``tblW`` exceeds ``max_mm``."""
    issues: List[str] = []
    max_twips = int(max_mm / MM_PER_INCH * TWIPS_PER_INCH)

    for idx, table in enumerate(doc.tables):
        tblPr = table._tbl.find(qn("w:tblPr"))
        if tblPr is None:
            continue
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            continue
        wt = tblW.get(qn("w:type"))
        wv = tblW.get(qn("w:w"))
        if wt == "dxa" and wv:
            try:
                w = int(wv)
            except ValueError:
                continue
            if w > max_twips:
                w_mm = w / TWIPS_PER_INCH * MM_PER_INCH
                issues.append(
                    f"validate_table_widths: table {idx} width {w_mm:.1f} mm "
                    f"> max {max_mm:.1f} mm"
                )
    return issues


def validate_cell_overflow(doc, min_col_width_mm: float = 30.0) -> List[str]:
    """Detect cells whose longest single word is wider than the column.

    Uses ``_INTER_10PT_AVG_MM`` — the average advance width measured at
    import time against the bundled Inter Regular.ttf at 10 pt over a
    Latin + NL-diacritic specimen. The legacy 5 % tolerance is removed:
    the constant is empirical, not a guess.

    Catches words Word cannot break without hyphenation (e.g. legal
    compound nouns, slugged identifiers). Word wrapping handles the
    lighter edge cases on its own.
    """
    issues: List[str] = []
    mm_per_char = _INTER_10PT_AVG_MM

    for t_idx, table in enumerate(doc.tables):
        for c_idx, col in enumerate(table.columns):
            if col.width is None:
                continue
            width_mm = col.width.mm if hasattr(col.width, "mm") else col.width / 36000
            if width_mm < min_col_width_mm:
                continue  # very narrow columns often hold short labels — skip

            # Find longest word in any cell of this column
            longest = 0
            for cell in col.cells:
                for p in cell.paragraphs:
                    for word in p.text.split():
                        if len(word) > longest:
                            longest = len(word)
            predicted_mm = longest * mm_per_char
            if predicted_mm > width_mm:
                issues.append(
                    f"validate_cell_overflow: table {t_idx} col {c_idx} "
                    f"longest word {longest} chars ≈ {predicted_mm:.1f} mm "
                    f"> col width {width_mm:.1f} mm"
                )
    return issues


def validate_list_nesting(
    doc,
    max_depth: int = DEFAULT_MAX_LIST_DEPTH,
) -> List[str]:
    """Flag paragraphs whose ``w:ilvl`` exceeds ``max_depth``."""
    issues: List[str] = []
    for p_idx, p in enumerate(doc.paragraphs):
        numPr = p._element.find(qn("w:pPr") + "/" + qn("w:numPr"))
        if numPr is None:
            continue
        ilvl = numPr.find(qn("w:ilvl"))
        if ilvl is None:
            continue
        try:
            lvl = int(ilvl.get(qn("w:val")) or 0)
        except ValueError:
            continue
        if lvl > max_depth - 1:  # ilvl is 0-indexed (ilvl=0 is depth 1)
            issues.append(
                f"validate_list_nesting: paragraph {p_idx} list depth "
                f"{lvl + 1} > max {max_depth}"
            )
    return issues


def validate_diacritics(
    doc,
    allowed_range: str = DEFAULT_ALLOWED_RANGE,
) -> List[str]:
    """Flag runs containing any non-ASCII character NOT covered by
    ``allowed_range``.

    ``allowed_range`` is a bare character class body (no surrounding []);
    e.g. ``À-ÿ€–—``.
    """
    # Build regex: allow ASCII (\x00-\x7f) + user-supplied range
    pattern = re.compile(rf"[^\x00-\x7f{allowed_range}]")
    issues: List[str] = []

    # Walk paragraphs (body + tables via iter_inner_content when available;
    # simpler: paragraphs for body, then each table cell's paragraphs)
    def _check_paragraph(para, location):
        for r_idx, run in enumerate(para.runs):
            for ch in run.text:
                if pattern.search(ch):
                    issues.append(
                        f"validate_diacritics: {location} run {r_idx} contains "
                        f"disallowed character {ch!r} (U+{ord(ch):04X})"
                    )
                    break

    for p_idx, p in enumerate(doc.paragraphs):
        _check_paragraph(p, f"paragraph {p_idx}")
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for cp_idx, cp in enumerate(cell.paragraphs):
                    _check_paragraph(
                        cp,
                        f"table {t_idx} row {r_idx} col {c_idx} para {cp_idx}",
                    )
    return issues


def validate_font_consistency(
    doc,
    primary_family: str = DEFAULT_FONT_FAMILY,
) -> List[str]:
    """Flag runs whose ``rFonts/@w:ascii`` is neither empty nor the
    primary family."""
    issues: List[str] = []

    def _check_paragraph(para, location):
        for r_idx, run in enumerate(para.runs):
            rPr = run._r.find(qn("w:rPr"))
            if rPr is None:
                continue
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                continue
            ascii_val = rFonts.get(qn("w:ascii"))
            if ascii_val and ascii_val != primary_family:
                issues.append(
                    f"validate_font_consistency: {location} run {r_idx} font "
                    f"{ascii_val!r} != {primary_family!r}"
                )

    for p_idx, p in enumerate(doc.paragraphs):
        _check_paragraph(p, f"paragraph {p_idx}")
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for cp_idx, cp in enumerate(cell.paragraphs):
                    _check_paragraph(
                        cp,
                        f"table {t_idx} row {r_idx} col {c_idx} para {cp_idx}",
                    )
    return issues


# ---------------------------------------------------------------------------
# Run-all convenience
# ---------------------------------------------------------------------------

_VALIDATORS = (
    validate_table_widths,
    validate_cell_overflow,
    validate_list_nesting,
    validate_diacritics,
    validate_font_consistency,
)


def run_all(doc) -> List[str]:
    """Run every validator in a stable order. Concatenates issues."""
    out: List[str] = []
    for v in _VALIDATORS:
        out.extend(v(doc))
    return out
