"""Tests for bilingual_body.py — Phase B2.

Structural / deterministic tests; no golden .docx comparison (those are
fragile across python-docx versions). Instead we assert on the produced
OOXML structure via python-docx / lxml primitives.
"""

from __future__ import annotations

import pytest
from docx import Document
from docx.oxml.ns import qn

from bilingual_body import (
    BALANCE_RATIO_MAX,
    BALANCE_RATIO_MIN,
    BILINGUAL_MARKER,
    COL_WIDTH_MM,
    USABLE_WIDTH_MM,
    _list_depth,
    _strip_list_prefix,
    render_bilingual_clause,
    validate_pair_balance,
)


# ---------------------------------------------------------------------------
# render_bilingual_clause — structural tests
# ---------------------------------------------------------------------------

def test_renders_one_heading_plus_two_paragraph_rows():
    doc = Document()
    render_bilingual_clause(
        doc,
        en_paragraphs=["First EN paragraph.", "Second EN paragraph."],
        nl_paragraphs=["Eerste NL paragraaf.", "Tweede NL paragraaf."],
        heading="1. Parties",
        heading_nl="1. Partijen",
    )
    assert len(doc.tables) == 1
    table = doc.tables[0]
    # 1 heading row + 2 paragraph rows = 3 rows
    assert len(table.rows) == 3
    assert len(table.columns) == 2


def test_renders_only_paragraph_rows_when_no_heading():
    doc = Document()
    render_bilingual_clause(
        doc,
        en_paragraphs=["A", "B", "C"],
        nl_paragraphs=["A", "B", "C"],
        heading=None,
    )
    table = doc.tables[0]
    assert len(table.rows) == 3


def test_bilingual_marker_set_on_tblDescription():
    doc = Document()
    render_bilingual_clause(doc, ["a" * 50], ["b" * 50])
    tbl = doc.tables[0]._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    desc = tblPr.find(qn("w:tblDescription"))
    assert desc is not None
    assert desc.get(qn("w:val")) == BILINGUAL_MARKER


def test_table_has_fixed_layout():
    doc = Document()
    render_bilingual_clause(doc, ["a" * 50], ["b" * 50])
    tblPr = doc.tables[0]._tbl.find(qn("w:tblPr"))
    tblLayout = tblPr.find(qn("w:tblLayout"))
    assert tblLayout is not None
    assert tblLayout.get(qn("w:type")) == "fixed"


def test_table_borders_are_all_nil():
    doc = Document()
    render_bilingual_clause(doc, ["a" * 50], ["b" * 50])
    tblPr = doc.tables[0]._tbl.find(qn("w:tblPr"))
    borders = tblPr.find(qn("w:tblBorders"))
    assert borders is not None
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = borders.find(qn(f"w:{edge}"))
        assert b is not None, f"edge {edge} missing"
        assert b.get(qn("w:val")) == "nil", f"edge {edge} not nil"


def test_heading_cells_shaded_cobalt():
    from bilingual_body import COBALT_HEX as _imported
    doc = Document()
    render_bilingual_clause(
        doc,
        en_paragraphs=["X"],
        nl_paragraphs=["Y"],
        heading="H",
        heading_nl="H-NL",
    )
    for cell in doc.tables[0].rows[0].cells:
        tcPr = cell._tc.find(qn("w:tcPr"))
        shd = tcPr.find(qn("w:shd"))
        assert shd is not None
        fill = (shd.get(qn("w:fill")) or "").upper()
        from document_factory import COBALT_HEX
        assert fill == COBALT_HEX.upper()


def test_paragraph_count_mismatch_raises():
    doc = Document()
    with pytest.raises(ValueError, match="paragraph-count mismatch"):
        render_bilingual_clause(doc, ["a", "b"], ["one"])


def test_empty_input_is_no_op():
    doc = Document()
    render_bilingual_clause(doc, [], [])
    assert len(doc.tables) == 0


def test_preserves_diacritics_and_euro():
    doc = Document()
    en = ["Price: 12 €/MWh — connection fee applies."]
    nl = ["Prijs: 12 €/MWh — aansluitkosten van toepassing."]
    render_bilingual_clause(doc, en, nl)
    cells = doc.tables[0].rows[0].cells
    assert "€" in cells[0].paragraphs[0].text
    assert "€" in cells[1].paragraphs[0].text
    assert "—" in cells[0].paragraphs[0].text


def test_multiline_paragraph_produces_multiple_paragraphs_in_cell():
    doc = Document()
    en = ["Line 1.\nLine 2.\nLine 3."]
    nl = ["Regel 1.\nRegel 2.\nRegel 3."]
    render_bilingual_clause(doc, en, nl)
    cell = doc.tables[0].rows[0].cells[0]
    assert len(cell.paragraphs) == 3
    assert cell.paragraphs[0].text == "Line 1."
    assert cell.paragraphs[2].text == "Line 3."


def test_list_item_rendered_with_bullet_glyph_and_indent():
    doc = Document()
    en = ["- First bullet"]
    nl = ["- Eerste bolletje"]
    render_bilingual_clause(doc, en, nl)
    cell = doc.tables[0].rows[0].cells[0]
    p = cell.paragraphs[0]
    assert "\u2022" in p.text
    assert p.paragraph_format.left_indent is not None


def test_nested_list_has_deeper_indent_than_top_level():
    doc = Document()
    en = ["- top", "  - nested"]
    nl = ["- top", "  - genest"]
    render_bilingual_clause(doc, en, nl)
    rows = doc.tables[0].rows
    top_indent = rows[0].cells[0].paragraphs[0].paragraph_format.left_indent
    nested_indent = rows[1].cells[0].paragraphs[0].paragraph_format.left_indent
    assert nested_indent > top_indent


def test_tblgrid_has_two_columns_of_equal_width():
    doc = Document()
    render_bilingual_clause(doc, ["a" * 50], ["b" * 50])
    tbl = doc.tables[0]._tbl
    grid = tbl.find(qn("w:tblGrid"))
    assert grid is not None
    cols = grid.findall(qn("w:gridCol"))
    assert len(cols) == 2
    w0 = int(cols[0].get(qn("w:w")))
    w1 = int(cols[1].get(qn("w:w")))
    # Equal within 5 twips rounding tolerance
    assert abs(w0 - w1) <= 5


def test_heading_nl_falls_back_to_heading_when_omitted():
    doc = Document()
    render_bilingual_clause(
        doc,
        en_paragraphs=["x" * 50],
        nl_paragraphs=["y" * 50],
        heading="Only EN",
    )
    cells = doc.tables[0].rows[0].cells
    assert cells[0].paragraphs[0].text == "Only EN"
    assert cells[1].paragraphs[0].text == "Only EN"


# ---------------------------------------------------------------------------
# validate_pair_balance
# ---------------------------------------------------------------------------

def test_balance_clean_on_well_matched_pair():
    en = ["This is a perfectly ordinary English paragraph that is quite long."]
    nl = ["Dit is een perfect gewone Nederlandse paragraaf die behoorlijk lang is."]
    assert validate_pair_balance(en, nl) == []


def test_balance_flags_mismatched_paragraph_count():
    issues = validate_pair_balance(["a", "b"], ["one"])
    assert any("paragraph-count mismatch" in i for i in issues)


def test_balance_flags_empty_en_cell():
    en = ["", "b" * 50]
    nl = ["nl1" + "." * 50, "b" * 50]
    issues = validate_pair_balance(en, nl)
    assert any("EN cell empty" in i for i in issues)


def test_balance_flags_empty_nl_cell():
    en = ["a" * 50, "b" * 50]
    nl = ["", "b" * 50]
    issues = validate_pair_balance(en, nl)
    assert any("NL cell empty" in i for i in issues)


def test_balance_flags_ratio_outside_bounds():
    # Both totals ≥ MIN_BALANCE_CHARS (40), but ratio is 2.0 → out of bounds.
    en = ["x" * 100]
    nl = ["y" * 200]
    issues = validate_pair_balance(en, nl)
    assert any("length ratio" in i for i in issues)


def test_balance_skips_ratio_check_for_short_pair():
    # Both totals < MIN_BALANCE_CHARS → no ratio check
    en = ["x"]
    nl = ["yyyyyyyyyy"]  # way out of ratio if checked, but too short
    issues = validate_pair_balance(en, nl)
    assert not any("length ratio" in i for i in issues)


def test_balance_flags_nested_list_depth_mismatch():
    en = ["- a", "  - nested"]
    nl = ["- a", "- not-nested"]  # NL depth 1, EN depth 2
    issues = validate_pair_balance(en, nl)
    assert any("list-nesting mismatch" in i for i in issues)


# ---------------------------------------------------------------------------
# _list_depth / _strip_list_prefix
# ---------------------------------------------------------------------------

def test_list_depth_plain_prose_is_zero():
    assert _list_depth("Just prose.") == 0


def test_list_depth_top_level_bullet_is_one():
    assert _list_depth("- foo") == 1
    assert _list_depth("* foo") == 1
    assert _list_depth("+ foo") == 1


def test_list_depth_numbered_is_one():
    assert _list_depth("1. foo") == 1
    assert _list_depth("12. foo") == 1
    assert _list_depth("1) foo") == 1


def test_list_depth_two_space_indent_is_two():
    assert _list_depth("  - foo") == 2


def test_strip_list_prefix_removes_bullet():
    assert _strip_list_prefix("- foo bar") == "foo bar"
    assert _strip_list_prefix("  - foo bar") == "foo bar"
    assert _strip_list_prefix("1. foo bar") == "foo bar"


def test_strip_list_prefix_leaves_plain_prose_alone():
    assert _strip_list_prefix("Just prose.") == "Just prose."
