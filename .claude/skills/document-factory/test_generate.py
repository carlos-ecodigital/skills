#!/usr/bin/env python3
"""Document Factory test suite.

Two test classes:
  TestColumnWidths — pure function tests for _compute_col_widths()
  TestAudit — every generated document passes audit_document() with zero violations

Run: pytest test_generate.py -v
"""

import re
import sys
import os

import pytest
from docx.shared import Mm, Pt, Emu
from docx.oxml.ns import qn

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from generate import (
    _compute_col_widths, _display_len, _longest_word_len, _SP,
    _CHAR_WIDTH_EMU, _CELL_PAD_EMU, _MIN_COL_EMU,
    _ALPHA_ABSTRACT_ID, _ROMAN_ABSTRACT_ID, _DECIMAL_ABSTRACT_ID, _BULLET_ABSTRACT_ID,
    md_to_docx, new_doc, add_table, add_section, _run, _fix_zoom,
    audit_document, AuditFailedError, save_doc, profile_agreement, profile_letter,
    profile_seed_memo, profile_investor_memo, profile_exec_summary,
    Party, DE_ENTITIES, COBALT_HEX, SLATE_50_HEX,
)


def _make_doc(md_text, **kwargs):
    defaults = dict(title="Test Doc", cover=True, entity="nl")
    defaults.update(kwargs)
    doc = md_to_docx(md_text, **defaults)
    _fix_zoom(doc)  # save_doc() applies this before audit
    return doc


# ---------------------------------------------------------------------------
# TestColumnWidths — pure function tests (no audit overlap)
# ---------------------------------------------------------------------------

class TestColumnWidths:
    def test_2col_kv_heuristic(self):
        headers = ["Label", "Value"]
        rows = [["Short", "A" * 80]]
        avail = int(Mm(165))
        widths = _compute_col_widths(headers, rows, avail)
        assert abs(widths[0] - int(avail * 0.25)) < avail * 0.02

    def test_starved_columns_milestone_table(self):
        headers = ["Milestone", "Units", "% Pool", "Trigger", "Sunset", "Certified"]
        rows = [["Build-out phase 1", "5000", "2.5%", "A" * 300, "2027-Q4", "Board resolution"]]
        widths = _compute_col_widths(headers, rows, int(Mm(165)))
        for w in widths:
            assert w >= _MIN_COL_EMU, f"Column width {w} < minimum {_MIN_COL_EMU}"

    def test_all_compact_surplus_distributed(self):
        headers = ["A", "B", "C", "D", "E"]
        rows = [["12345", "12345", "12345", "12345", "12345"]]
        widths = _compute_col_widths(headers, rows, int(Mm(165)))
        assert all(w > _MIN_COL_EMU for w in widths)

    def test_display_len_strips_bold(self):
        assert _display_len("**bold**") == 4
        assert _display_len("normal") == 6
        assert _display_len("`code`") == 4

    def test_longest_word(self):
        assert _longest_word_len("Time-Based Tranche") == len("Time-Based")

    def test_total_equals_available(self):
        headers = ["A", "B", "C"]
        rows = [["Short", "Medium text here", "A" * 100]]
        avail = int(Mm(165))
        widths = _compute_col_widths(headers, rows, avail)
        assert sum(widths) == avail

    def test_overflow_guard(self):
        headers = [f"H{i}" for i in range(8)]
        rows = [[f"D{i}" for i in range(8)]]
        avail = int(Mm(165))
        widths = _compute_col_widths(headers, rows, avail)
        for w in widths:
            assert w >= _MIN_COL_EMU


# ---------------------------------------------------------------------------
# TestAudit — audit_document() on every document pattern
# ---------------------------------------------------------------------------

class TestAudit:
    """Every generated document must pass audit_document() with zero violations."""

    LEGAL_MD = """\
## 1. Definitions

"**Agreement**" means this Advisory Agreement.

## 2. Scope

The Advisor shall provide:

(a) Strategic guidance on positioning

(b) Introductions to potential partners

(c) Review of business plans

## 3. Deliverables

**Project Team:**
- Alice van der Berg
- Bob Smith

1. Phase one deliverables
2. Phase two deliverables

| Deliverable | Due Date | Owner |
|-------------|----------|-------|
| Report A    | Q1 2026  | Alice |
| Report B    | Q2 2026  | Bob   |

## 4. General

> This agreement is governed by Dutch law.

## Signatures

**DIGITAL ENERGY NETHERLANDS B.V.**

By: ___________________________
Name: Jelmer ten Wolde
Title: Bestuurder


**ADVISOR**

By: ___________________________
Name: Jane Doe
Title: Advisor
"""

    def test_simple_md(self):
        doc = _make_doc("## Heading\n\nBody text paragraph.\n\n- Bullet A\n- Bullet B\n\n1. Number one\n2. Number two\n")
        violations = audit_document(doc)
        assert violations == [], violations

    def test_legal_pattern(self):
        doc = _make_doc(self.LEGAL_MD)
        violations = audit_document(doc)
        assert violations == [], violations

    def test_bold_label_before_list(self):
        """Regression: signature handler must not swallow non-signature bold lines."""
        doc = _make_doc("**Management Board:**\n- Person A\n- Person B\n")
        violations = audit_document(doc)
        assert violations == [], violations

    def test_numbered_list_restart(self):
        """Two numbered blocks must get different numIds."""
        doc = _make_doc("1. A\n2. B\n\nSome paragraph.\n\n1. C\n2. D\n")
        violations = audit_document(doc)
        assert violations == [], violations
        # Verify restart: extract numIds
        num_ids = []
        for p in doc.paragraphs:
            numPr = p._element.find(f'.//{qn("w:numPr")}')
            if numPr is not None:
                nid = numPr.find(qn('w:numId'))
                if nid is not None:
                    num_ids.append(nid.get(qn('w:val')))
        assert len(num_ids) == 4
        assert num_ids[0] == num_ids[1], "First block should share numId"
        assert num_ids[2] == num_ids[3], "Second block should share numId"
        assert num_ids[0] != num_ids[2], "Different blocks must restart"

    def test_alpha_list_blank_lines(self):
        doc = _make_doc("## Section\n\n(a) First\n\n(b) Second\n\n(c) Third\n")
        violations = audit_document(doc)
        assert violations == [], violations
        # All 3 share one numId
        num_ids = set()
        for p in doc.paragraphs:
            numPr = p._element.find(f'.//{qn("w:numPr")}')
            if numPr is not None:
                nid = numPr.find(qn('w:numId'))
                if nid is not None:
                    num_ids.add(nid.get(qn('w:val')))
        assert len(num_ids) == 1, f"Expected 1 numId for continuous block, got {num_ids}"

    def test_table_formatting(self):
        """Table with enough rows to trigger header repeat."""
        rows_md = "\n".join(f"| Data {i} | Value {i} |" for i in range(10))
        md = f"## Table Test\n\n| Header A | Header B |\n|----------|----------|\n{rows_md}\n"
        doc = _make_doc(md)
        violations = audit_document(doc)
        assert violations == [], violations

    def test_list_after_table(self):
        md = "## Section\n\n| A | B |\n|---|---|\n| 1 | 2 |\n\n- Bullet after table\n"
        doc = _make_doc(md)
        violations = audit_document(doc)
        assert violations == [], violations

    def test_blockquote(self):
        doc = _make_doc("## Section\n\n> This is a blockquote with important text.\n")
        violations = audit_document(doc)
        assert violations == [], violations

    def test_md_without_cover(self):
        doc = md_to_docx("## Title\n\nBody.\n", title="Test", cover=False)
        _fix_zoom(doc)
        violations = audit_document(doc)
        assert violations == [], violations

    def test_md_with_cover(self):
        doc = _make_doc("## Title\n\nBody.\n")
        violations = audit_document(doc)
        assert violations == [], violations
        # Verify cover page break exists
        body_el = doc.element.body
        found_break = False
        for p_el in body_el.findall(qn('w:p')):
            if p_el.find(f'.//{qn("w:br")}[@{qn("w:type")}="page"]') is not None:
                found_break = True
                break
        assert found_break, "Cover page should have a page break"

    def _audit_profile(self, doc):
        """Apply save_doc fixes then audit."""
        _fix_zoom(doc)
        return audit_document(doc)

    def test_profile_agreement(self):
        doc = profile_agreement("Advisory Agreement", entity="nl")
        violations = self._audit_profile(doc)
        assert violations == [], violations

    def test_profile_letter(self):
        doc = profile_letter(title="Test Letter", date_str="16 April 2026", entity="ag")
        violations = self._audit_profile(doc)
        assert violations == [], violations

    def test_profile_seed_memo(self):
        doc = profile_seed_memo()
        violations = self._audit_profile(doc)
        assert violations == [], violations

    def test_profile_investor_memo(self):
        doc = profile_investor_memo()
        violations = self._audit_profile(doc)
        assert violations == [], violations

    def test_profile_exec_summary(self):
        doc = profile_exec_summary(title="Q1 Update")
        violations = self._audit_profile(doc)
        assert violations == [], violations

    # --- Step 0 regression tests: disclaimer lines must not be styled as sig parties ---

    def test_draft_line_at_top_not_styled_as_party(self):
        """Regression: **Grant Date:** / **DRAFT — ...** at doc top must NOT trigger sig handler."""
        md = """# COMPANY NAME

## Grant Agreement

**Grant Date:** 1 March 2026
**Issue Version:** 1.0

**DRAFT — Subject to review**

## 1. Parties

This agreement is entered into.
"""
        doc = md_to_docx(md, cover=False)
        _fix_zoom(doc)
        assert audit_document(doc) == []
        # DRAFT line should be 11pt body bold, NOT 12pt party bold
        draft_p = next((p for p in doc.paragraphs if 'DRAFT' in p.text), None)
        assert draft_p is not None
        assert draft_p.runs[0].font.size == Pt(11), \
            f"DRAFT line should be 11pt body, got {draft_p.runs[0].font.size}"

    def test_end_of_agreement_not_styled_as_party(self):
        """Regression: **End of Agreement.** after signatures must NOT be styled as party."""
        md = """## Signatures

**ACME CORP**

By: _________
Name: Jane Doe
Title: CEO

---

**End of Agreement.**

---

## Drafting Notes

Some note.
"""
        doc = md_to_docx(md, cover=False)
        _fix_zoom(doc)
        assert audit_document(doc) == []
        end_p = next((p for p in doc.paragraphs if 'End of Agreement' in p.text), None)
        assert end_p is not None
        assert end_p.runs[0].font.size == Pt(11), \
            f"End of Agreement should be 11pt body, got {end_p.runs[0].font.size}"

    # --- Strict audit mode tests ---

    def test_strict_mode_raises_on_violations(self, tmp_path):
        """strict=True + dirty doc → raises AuditFailedError, file NOT saved."""
        doc = md_to_docx("## Test\n\nBody.\n", cover=False)
        # Introduce a violation: strip widow_control + break heading
        for p in doc.paragraphs:
            p.paragraph_format.widow_control = False
            if p.runs and p.runs[0].font.bold:
                p.paragraph_format.keep_with_next = False
        out = tmp_path / "strict.docx"
        with pytest.raises(AuditFailedError) as exc_info:
            save_doc(doc, str(out), strict=True)
        assert exc_info.value.violations, "should have violations"
        assert not out.exists(), "file must NOT be saved when strict fails"

    def test_strict_mode_passes_clean_doc(self, tmp_path):
        """strict=True + clean doc → saves normally."""
        doc = md_to_docx("## Heading\n\nBody.\n", title="Test", cover=False)
        out = tmp_path / "clean.docx"
        save_doc(doc, str(out), strict=True)  # must not raise
        assert out.exists() and out.stat().st_size > 0

    def test_default_mode_warns_only(self, tmp_path, capsys):
        """strict=None + dirty doc → warns on stderr, file saved."""
        doc = md_to_docx("## Test\n\nBody.\n", cover=False)
        for p in doc.paragraphs:
            p.paragraph_format.widow_control = False
            if p.runs and p.runs[0].font.bold:
                p.paragraph_format.keep_with_next = False
        out = tmp_path / "warn.docx"
        save_doc(doc, str(out))  # no strict param
        assert out.exists(), "default mode saves even with violations"
        captured = capsys.readouterr()
        assert "[audit]" in captured.err

    def test_env_var_triggers_strict(self, tmp_path, monkeypatch):
        """DOCFACTORY_STRICT=1 makes default behavior strict."""
        monkeypatch.setenv("DOCFACTORY_STRICT", "1")
        doc = md_to_docx("## Test\n\nBody.\n", cover=False)
        for p in doc.paragraphs:
            p.paragraph_format.widow_control = False
            if p.runs and p.runs[0].font.bold:
                p.paragraph_format.keep_with_next = False
        out = tmp_path / "env.docx"
        with pytest.raises(AuditFailedError):
            save_doc(doc, str(out))  # no strict param; env var forces strict
        assert not out.exists()

    def test_audit_failed_error_carries_violations(self):
        """AuditFailedError exposes .violations list."""
        err = AuditFailedError(["v1", "v2", "v3"])
        assert err.violations == ["v1", "v2", "v3"]
        assert "3 audit violation" in str(err)
        assert "v1" in str(err)

    def test_real_party_names_still_detected(self):
        """Ensure legitimate party names still get 12pt styling."""
        md = """## Signatures

**For the Company — Digital Energy Netherlands B.V.**

By: _________
Name: Jelmer

**Recipient — Arco Vreugdenhil**

Signature: _________
Name: Arco
"""
        doc = md_to_docx(md, cover=False)
        _fix_zoom(doc)
        assert audit_document(doc) == []
        party_names = [p for p in doc.paragraphs
                       if 'For the Company' in p.text or 'Recipient' in p.text]
        assert len(party_names) == 2
        for p in party_names:
            assert p.runs[0].font.size == Pt(12), \
                f"Party '{p.text[:40]}' should be 12pt, got {p.runs[0].font.size}"
            assert p.runs[0].font.bold is True


# ---------------------------------------------------------------------------
# TestVisualQaUtils — unit tests for visual_qa.py utilities (no Word needed)
# ---------------------------------------------------------------------------

class TestVisualQaUtils:
    """Test visual_qa building blocks in isolation. No Word/PDF required."""

    @pytest.fixture(autouse=True)
    def _import_visual_qa(self):
        tools_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "tools")
        if tools_path not in sys.path:
            sys.path.insert(0, tools_path)

    def test_phash_is_hex_string(self, tmp_path):
        """phash() returns 16-char hex string (64 bits)."""
        from PIL import Image
        from visual_qa import phash
        img_path = tmp_path / "solid.png"
        Image.new("RGB", (100, 100), (128, 128, 128)).save(img_path)
        h = phash(img_path)
        assert isinstance(h, str)
        assert len(h) == 16
        assert all(c in "0123456789abcdef" for c in h)

    def test_hash_distance_identical(self, tmp_path):
        """Identical images → hash distance 0."""
        from PIL import Image
        from visual_qa import phash, hash_distance
        p1, p2 = tmp_path / "a.png", tmp_path / "b.png"
        img = Image.new("RGB", (200, 200), (100, 150, 200))
        img.save(p1)
        img.save(p2)
        assert hash_distance(phash(p1), phash(p2)) == 0

    def test_hash_distance_different(self, tmp_path):
        """Visually different images → non-zero distance."""
        from PIL import Image, ImageDraw
        from visual_qa import phash, hash_distance
        p1, p2 = tmp_path / "a.png", tmp_path / "b.png"
        # Create two clearly different images
        img1 = Image.new("RGB", (200, 200), "white")
        ImageDraw.Draw(img1).rectangle((20, 20, 80, 80), fill="black")
        img1.save(p1)
        img2 = Image.new("RGB", (200, 200), "white")
        ImageDraw.Draw(img2).rectangle((120, 120, 180, 180), fill="black")
        img2.save(p2)
        d = hash_distance(phash(p1), phash(p2))
        assert d > 0, f"Different images should have non-zero hash distance, got {d}"

    def test_pdf_to_pngs_produces_one_per_page(self, tmp_path):
        """Given a 2-page PDF, pdf_to_pngs produces 2 PNGs."""
        import fitz
        from visual_qa import pdf_to_pngs
        pdf_path = tmp_path / "test.pdf"
        doc = fitz.open()
        for _ in range(2):
            doc.new_page(width=200, height=300)
        doc.save(str(pdf_path))
        doc.close()
        page_dir = tmp_path / "pages"
        pngs = pdf_to_pngs(pdf_path, page_dir)
        assert len(pngs) == 2
        for p in pngs:
            assert p.exists() and p.stat().st_size > 0
            assert p.name.endswith(".png")

    def test_pdf_to_pngs_clears_stale_files(self, tmp_path):
        """Re-running on same dir removes old page_*.png files."""
        import fitz
        from visual_qa import pdf_to_pngs
        pdf_path = tmp_path / "test.pdf"
        doc = fitz.open()
        doc.new_page(width=200, height=300)
        doc.save(str(pdf_path))
        doc.close()
        page_dir = tmp_path / "pages"
        page_dir.mkdir()
        # Plant a stale page
        (page_dir / "page_099.png").write_bytes(b"stale")
        pdf_to_pngs(pdf_path, page_dir)
        assert not (page_dir / "page_099.png").exists(), "stale PNG must be removed"

    def test_build_report_valid_html(self, tmp_path):
        """HTML report renders without exceptions and contains expected elements."""
        from visual_qa import DocResult, build_report
        results = [
            DocResult(slug="test1", title="Test Doc One",
                      md_path="a.md", docx_path="a.docx", pdf_path="a.pdf",
                      page_pngs=[], page_hashes=[], audit_violations=[]),
            DocResult(slug="test2", title="Doc <with> & special chars",
                      md_path="b.md", docx_path="b.docx", pdf_path="b.pdf",
                      audit_violations=["v1", "v2"]),
        ]
        out = tmp_path / "report.html"
        build_report(results, out)
        html = out.read_text()
        assert "<!DOCTYPE html>" in html
        assert "Test Doc One" in html
        # Special chars must be escaped, not interpreted as HTML
        assert "&lt;with&gt; &amp; special" in html
        assert "v1" in html
        assert "2 audit violations" in html

    def test_save_load_results_roundtrip(self, tmp_path):
        """DocResult list serializes and deserializes correctly."""
        from visual_qa import DocResult, save_results, load_results
        original = [
            DocResult(slug="a", title="T1", md_path="a.md",
                      docx_path="a.docx", pdf_path="a.pdf",
                      page_pngs=["a/1.png"], page_hashes=["deadbeef"],
                      audit_violations=[], diff_distances={"1": 3}),
        ]
        save_results(original, tmp_path / "results.json")
        loaded = load_results(tmp_path / "results.json")
        assert len(loaded) == 1
        assert loaded[0].slug == "a"
        assert loaded[0].page_hashes == ["deadbeef"]
        assert loaded[0].diff_distances == {"1": 3}

    def test_compare_to_golden_flags_changes(self, tmp_path):
        """compare_to_golden flags pages with hash distance > threshold."""
        from visual_qa import DocResult, compare_to_golden
        results = [
            DocResult(slug="a", title="T", md_path="", docx_path="", pdf_path="",
                      page_hashes=["0000000000000000", "ffffffffffffffff"]),
        ]
        golden = {"a": {"1": "0000000000000000", "2": "0000000000000000"}}
        changes = compare_to_golden(results, golden, threshold=5)
        # Page 1: identical hashes → distance 0 → no flag
        # Page 2: all-1s vs all-0s → max distance 64 → flagged
        assert len(changes) == 1
        assert changes[0][0] == "a"
        assert changes[0][1] == 2

    def test_compare_to_golden_missing_page(self, tmp_path):
        """Current run missing a page that's in golden → flagged."""
        from visual_qa import DocResult, compare_to_golden
        results = [
            DocResult(slug="a", title="T", md_path="", docx_path="", pdf_path="",
                      page_hashes=["deadbeef00000000"]),  # only 1 page
        ]
        golden = {"a": {"1": "deadbeef00000000", "2": "deadbeef00000000"}}
        changes = compare_to_golden(results, golden, threshold=5)
        # Page 2 in golden but missing in current → flagged
        assert (("a", 2, 64) in changes)

    def test_resolve_corpus_default(self):
        """resolve_corpus(None) returns DEFAULT_CORPUS."""
        from visual_qa import resolve_corpus, DEFAULT_CORPUS
        assert resolve_corpus(None) == DEFAULT_CORPUS

    def test_resolve_corpus_from_json(self, tmp_path):
        """resolve_corpus(path) loads list from JSON file."""
        import json
        from visual_qa import resolve_corpus
        corpus_file = tmp_path / "corpus.json"
        corpus_file.write_text(json.dumps([["/tmp/a.md", "a"], ["/tmp/b.md", "b"]]))
        result = resolve_corpus(str(corpus_file))
        assert result == [("/tmp/a.md", "a"), ("/tmp/b.md", "b")]

    def test_resolve_corpus_from_json_dict(self, tmp_path):
        """resolve_corpus accepts dict-format too."""
        import json
        from visual_qa import resolve_corpus
        corpus_file = tmp_path / "corpus.json"
        corpus_file.write_text(json.dumps([{"path": "/tmp/x.md", "slug": "x"}]))
        assert resolve_corpus(str(corpus_file)) == [("/tmp/x.md", "x")]

    def test_build_golden_json_format(self):
        """build_golden_json converts DocResult list to {slug: {page_idx: hash}}."""
        from visual_qa import DocResult, build_golden_json
        results = [
            DocResult(slug="doc1", title="", md_path="", docx_path="", pdf_path="",
                      page_hashes=["aaaa", "bbbb", "cccc"]),
            DocResult(slug="doc2", title="", md_path="", docx_path="", pdf_path="",
                      page_hashes=[]),  # no pages → excluded
        ]
        golden = build_golden_json(results)
        assert golden == {
            "doc1": {"1": "aaaa", "2": "bbbb", "3": "cccc"},
        }


# ---------------------------------------------------------------------------
# TestVisualRegression — full corpus pipeline, requires --visual flag + LibreOffice
# ---------------------------------------------------------------------------

@pytest.fixture
def visual_enabled(request):
    return request.config.getoption("--visual")


class TestVisualRegression:
    """Full corpus run: markdown → docx → pdf → hashes. Opt-in via pytest --visual.

    Slow tests (~1 min full corpus). Require LibreOffice + markdown source files.
    """

    def _load_visual_qa(self):
        """Import visual_qa tool module, adding its path first."""
        tools_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "tools")
        if tools_path not in sys.path:
            sys.path.insert(0, tools_path)
        import visual_qa
        return visual_qa

    def test_doctor_passes(self, visual_enabled):
        """Pre-flight check must pass before any visual test is meaningful."""
        if not visual_enabled:
            pytest.skip("--visual not set")
        vq = self._load_visual_qa()
        import argparse
        rc = vq.cmd_doctor(argparse.Namespace(corpus=None))
        assert rc == 0, "pre-flight must pass (check LibreOffice, deps, corpus)"

    def test_corpus_clean_audit(self, visual_enabled, tmp_path):
        """Every corpus document must pass audit_document with zero violations."""
        if not visual_enabled:
            pytest.skip("--visual not set")
        vq = self._load_visual_qa()
        results = vq.run_pipeline(vq.DEFAULT_CORPUS, tmp_path, skip_pdf=True)
        bad = [(r.slug, r.audit_violations) for r in results if r.audit_violations or r.error]
        assert not bad, f"Corpus has violations/errors: {bad}"

    def test_corpus_renders_pdfs_and_hashes(self, visual_enabled, tmp_path):
        """End-to-end: each corpus doc produces PDFs + page PNGs + hashes."""
        if not visual_enabled:
            pytest.skip("--visual not set")
        vq = self._load_visual_qa()
        results = vq.run_pipeline(vq.DEFAULT_CORPUS, tmp_path, skip_pdf=False)
        assert len(results) == len(vq.DEFAULT_CORPUS)
        for r in results:
            assert r.error is None, f"{r.slug}: {r.error}"
            assert r.page_hashes, f"{r.slug}: no hashes produced"
            assert len(r.page_hashes) == len(r.page_pngs)

    def test_no_regression_vs_golden(self, visual_enabled, tmp_path):
        """If a golden baseline exists in git, current run must match within threshold."""
        if not visual_enabled:
            pytest.skip("--visual not set")
        vq = self._load_visual_qa()
        if not vq.GOLDEN_JSON.exists():
            pytest.skip("No golden baseline yet; run visual_qa.py approve first")
        results = vq.run_pipeline(vq.DEFAULT_CORPUS, tmp_path, skip_pdf=False)
        import json
        golden = json.loads(vq.GOLDEN_JSON.read_text())
        changes = vq.compare_to_golden(results, golden, threshold=5)
        assert not changes, f"Visual regressions vs golden: {changes}"
