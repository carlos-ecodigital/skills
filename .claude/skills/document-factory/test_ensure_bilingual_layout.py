"""Tests for ``bilingual_body.ensure_bilingual_layout`` (rc3.2)."""

from __future__ import annotations

from docx import Document
from docx.shared import Mm

import bilingual_body
from bilingual_body import ensure_bilingual_layout, render_bilingual_clause


def _mm(value) -> float:
    """python-docx margins are EMU; convert to mm with 1mm tolerance."""
    return value.mm if value is not None else 0.0


def test_sets_correct_margin_values():
    doc = Document()
    ensure_bilingual_layout(doc)
    section = doc.sections[0]
    assert abs(_mm(section.left_margin) - 20.0) < 0.5
    assert abs(_mm(section.right_margin) - 20.0) < 0.5
    assert abs(_mm(section.top_margin) - 25.0) < 0.5
    assert abs(_mm(section.bottom_margin) - 20.0) < 0.5


def test_idempotent_on_second_call():
    """Re-applying must be a no-op (and not raise)."""
    doc = Document()
    ensure_bilingual_layout(doc)
    # Tweak left margin to detect re-application
    doc.sections[0].left_margin = Mm(30)
    ensure_bilingual_layout(doc)  # should NOT reset to 20mm — idempotent guard
    assert abs(_mm(doc.sections[0].left_margin) - 30.0) < 0.5


def test_first_call_hook_in_render_bilingual_clause():
    """``render_bilingual_clause`` invokes ensure_bilingual_layout
    transparently on its first call per doc."""
    doc = Document()
    # Default margins should be ~31.75 mm — verify pre-condition
    pre_left = _mm(doc.sections[0].left_margin)
    assert pre_left > 25.0, f"expected default >25mm, got {pre_left}"

    render_bilingual_clause(
        doc,
        en_paragraphs=["Test."],
        nl_paragraphs=["Test."],
    )

    # After the first render call, margins should be 20 mm
    assert abs(_mm(doc.sections[0].left_margin) - 20.0) < 0.5
    assert abs(_mm(doc.sections[0].right_margin) - 20.0) < 0.5


def test_sites_loi_path_unbroken():
    """A doc that mixes ``add_cover()`` then ``render_bilingual_clause()``
    should land at 20 mm L/R for the bilingual content. This mirrors the
    Sites LOI engine call sequence."""
    from generate import add_cover

    doc = Document()
    add_cover(
        doc,
        agreement_type="Letter of Intent",
        bilingual=True,
        title_nl="Intentieverklaring",
    )
    render_bilingual_clause(
        doc,
        en_paragraphs=["First paragraph."],
        nl_paragraphs=["Eerste paragraaf."],
        heading="1. Parties",
        heading_nl="1. Partijen",
    )
    section = doc.sections[0]
    assert abs(_mm(section.left_margin) - 20.0) < 0.5
    assert abs(_mm(section.right_margin) - 20.0) < 0.5


def test_mia_style_doc_unaffected_outside_bilingual_render():
    """A doc that NEVER calls render_bilingual_clause / ensure_bilingual_layout
    keeps its python-docx default margins (i.e. existing MIA + colocation
    paths are not silently mutated)."""
    doc = Document()
    # No bilingual call; nothing should touch the margins.
    pre = _mm(doc.sections[0].left_margin)
    assert pre > 25.0, f"baseline expected >25mm, got {pre}"
    # Confirm id() is NOT in the applied set
    assert id(doc) not in bilingual_body._LAYOUT_APPLIED
