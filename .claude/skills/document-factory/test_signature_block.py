"""Tests for signature_block.py — Phase B2."""

from __future__ import annotations

import pytest
from docx import Document

from generate import DE_ENTITIES
from signature_block import (
    SIGNATURE_LINE,
    SigParty,
    _format_role_label_line,
    render_signature_page,
)


def _all_text(doc):
    return "\n".join(p.text for p in doc.paragraphs)


def test_renders_header_line():
    doc = Document()
    render_signature_page(doc, DE_ENTITIES["nl"], site_partners=[])
    assert "Signatures Page / Handtekeningen" in _all_text(doc)


def test_renders_provider_block():
    doc = Document()
    render_signature_page(
        doc,
        DE_ENTITIES["nl"],
        site_partners=[],
        provider_signatory_name="Carlos Reuven Mattis Glender",
        provider_signatory_title="Director",
    )
    txt = _all_text(doc)
    assert DE_ENTITIES["nl"].legal_name in txt
    assert "Name: Carlos Reuven Mattis Glender" in txt
    assert "Title: Director" in txt
    assert SIGNATURE_LINE in txt


def test_renders_single_site_partner_with_one_role_label():
    doc = Document()
    sp = SigParty(
        legal_name="Van Gog Grubbenvorst B.V.",
        role_labels_en=["Grid Contributor"],
        role_labels_nl=["Netbijdrager"],
        signatory_name="Marion van Gog",
        signatory_title="Directeur",
    )
    render_signature_page(doc, DE_ENTITIES["nl"], [sp])
    txt = _all_text(doc)
    assert "GRID CONTRIBUTOR / NETBIJDRAGER:" in txt
    assert "Van Gog Grubbenvorst B.V." in txt
    assert "Marion van Gog" in txt


def test_renders_partner_with_multiple_role_labels():
    doc = Document()
    sp = SigParty(
        legal_name="Combined Entity B.V.",
        role_labels_en=["Grid Contributor", "Landowner"],
        role_labels_nl=["Netbijdrager", "Grondeigenaar"],
        signatory_name="Name",
        signatory_title="Title",
    )
    render_signature_page(doc, DE_ENTITIES["nl"], [sp])
    txt = _all_text(doc)
    assert "GRID CONTRIBUTOR, LANDOWNER / NETBIJDRAGER, GRONDEIGENAAR:" in txt


def test_renders_three_site_partners_in_order():
    doc = Document()
    sps = [
        SigParty(legal_name="First B.V.",
                 role_labels_en=["Grid Contributor"],
                 role_labels_nl=["Netbijdrager"],
                 signatory_name="A", signatory_title="T"),
        SigParty(legal_name="Second B.V.",
                 role_labels_en=["Landowner"],
                 role_labels_nl=["Grondeigenaar"],
                 signatory_name="B", signatory_title="T"),
        SigParty(legal_name="Third B.V.",
                 role_labels_en=["Heat Offtaker"],
                 role_labels_nl=["Warmteafnemer"],
                 signatory_name="C", signatory_title="T"),
    ]
    render_signature_page(doc, DE_ENTITIES["nl"], sps)
    txt = _all_text(doc)
    assert txt.find("First B.V.") < txt.find("Second B.V.") < txt.find("Third B.V.")


def test_kvk_omitted_in_non_binding_formality():
    doc = Document()
    sp = SigParty(
        legal_name="Party",
        role_labels_en=["Grid Contributor"],
        role_labels_nl=["Netbijdrager"],
        signatory_name="n",
        signatory_title="t",
        kvk="12345678",
    )
    render_signature_page(doc, DE_ENTITIES["nl"], [sp], formality="non_binding")
    txt = _all_text(doc)
    assert "KvK: 12345678" not in txt
    # Provider KvK also absent
    assert "KvK: 98580086" not in txt


def test_kvk_rendered_in_binding_formality():
    doc = Document()
    sp = SigParty(
        legal_name="Party",
        role_labels_en=["Grid Contributor"],
        role_labels_nl=["Netbijdrager"],
        signatory_name="n",
        signatory_title="t",
        kvk="12345678",
    )
    render_signature_page(doc, DE_ENTITIES["nl"], [sp], formality="binding")
    txt = _all_text(doc)
    assert "KvK: 12345678" in txt
    # Provider KvK from DE_ENTITIES["nl"].registration_number
    assert "KvK: 98580086" in txt


def test_invalid_formality_raises():
    doc = Document()
    with pytest.raises(ValueError, match="formality"):
        render_signature_page(doc, DE_ENTITIES["nl"], [], formality="other")


def test_role_label_line_formatter_en_only():
    sp = SigParty(legal_name="X", role_labels_en=["Grid Contributor"], role_labels_nl=[])
    assert _format_role_label_line(sp) == "GRID CONTRIBUTOR:"


def test_role_label_line_formatter_nl_only():
    sp = SigParty(legal_name="X", role_labels_en=[], role_labels_nl=["Netbijdrager"])
    assert _format_role_label_line(sp) == "NETBIJDRAGER:"


def test_role_label_line_formatter_empty_returns_none():
    sp = SigParty(legal_name="X")
    assert _format_role_label_line(sp) is None


def test_every_signatory_has_name_title_date_rows():
    doc = Document()
    sp = SigParty(
        legal_name="Party",
        role_labels_en=["Grid Contributor"],
        role_labels_nl=["Netbijdrager"],
        signatory_name="Alice",
        signatory_title="CEO",
    )
    render_signature_page(doc, DE_ENTITIES["nl"], [sp])
    txt = _all_text(doc)
    # Provider has Name/Title/Date rows (empty but present)
    assert txt.count("Name: ") == 2  # provider + 1 partner
    assert txt.count("Title: ") == 2
    assert txt.count("Date: ") == 2
