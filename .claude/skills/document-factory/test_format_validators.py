"""Tests for format_validators.py — Phase B2."""

from __future__ import annotations

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from format_validators import (
    _INTER_10PT_AVG_MM,
    run_all,
    validate_cell_overflow,
    validate_diacritics,
    validate_font_consistency,
    validate_list_nesting,
    validate_table_widths,
)


# ---------------------------------------------------------------------------
# validate_table_widths
# ---------------------------------------------------------------------------

def test_table_widths_clean_on_fresh_document():
    doc = Document()
    assert validate_table_widths(doc) == []


def test_table_widths_flags_overwide_table():
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    # Remove any default tblW that python-docx may have already set so the
    # validator sees only our deliberately-overwide value.
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblW = OxmlElement("w:tblW")
    # 200 mm = ~11339 twips, way over 165 mm
    tblW.set(qn("w:w"), "11339")
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    issues = validate_table_widths(doc)
    assert any("width" in i and "table 0" in i for i in issues)


def test_table_widths_passes_on_exact_165mm():
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "9354")  # 165 mm exactly
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    assert validate_table_widths(doc) == []


# ---------------------------------------------------------------------------
# Inter font calibration constant
# ---------------------------------------------------------------------------

def test_inter_calibration_in_expected_range():
    """Sanity-check that the measured Inter 10 pt avg advance is in a
    plausible band (1.9–2.3 mm). Catches accidental UPM / point-size
    regressions and guards against the fontTools-missing fallback being
    silently shipped in CI.
    """
    assert 1.9 <= _INTER_10PT_AVG_MM <= 2.3, (
        f"Inter 10pt calibration {_INTER_10PT_AVG_MM:.4f} mm outside "
        f"expected 1.9–2.3 mm band — fontTools missing or font corrupt?"
    )


# ---------------------------------------------------------------------------
# validate_cell_overflow
# ---------------------------------------------------------------------------

def test_cell_overflow_clean_on_normal_words():
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Mm(80)
    table.columns[1].width = Mm(80)
    table.rows[0].cells[0].text = "normal prose here"
    table.rows[0].cells[1].text = "wat normale tekst"
    issues = validate_cell_overflow(doc)
    assert issues == []


def test_cell_overflow_flags_excessively_long_word():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.columns[0].width = Mm(40)
    # 24-char word at ~2.115 mm/char (measured Inter 10pt) ≈ 50.8 mm,
    # overflows 40 mm column.
    table.rows[0].cells[0].text = "Supercalifragilisticexpi"
    issues = validate_cell_overflow(doc)
    assert any("longest word" in i for i in issues)


# ---------------------------------------------------------------------------
# validate_list_nesting
# ---------------------------------------------------------------------------

def test_list_nesting_clean_when_no_numbered_lists():
    doc = Document()
    doc.add_paragraph("No lists here.")
    assert validate_list_nesting(doc) == []


def test_list_nesting_flags_over_max_depth():
    doc = Document()
    p = doc.add_paragraph()
    # Manually build w:pPr/w:numPr/w:ilvl
    pPr = OxmlElement("w:pPr")
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "4")  # depth 5 (0-indexed)
    numPr.append(ilvl)
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), "1")
    numPr.append(numId)
    pPr.append(numPr)
    p._element.insert(0, pPr)
    issues = validate_list_nesting(doc, max_depth=3)
    assert any("list depth" in i for i in issues)


# ---------------------------------------------------------------------------
# validate_diacritics
# ---------------------------------------------------------------------------

def test_diacritics_clean_on_ascii_plus_euro():
    doc = Document()
    doc.add_paragraph("Normal English text with a € sign.")
    assert validate_diacritics(doc) == []


def test_diacritics_clean_on_nl_extended_latin():
    doc = Document()
    doc.add_paragraph("Gëlukkig — de prijs is 12 €/MWh over het Nederlandse net.")
    assert validate_diacritics(doc) == []


def test_diacritics_flags_unexpected_character():
    doc = Document()
    # Chinese ideograph is outside default allowed range
    doc.add_paragraph("Contains unexpected ideograph 中 here.")
    issues = validate_diacritics(doc)
    assert any("disallowed character" in i for i in issues)


def test_diacritics_walks_table_cells():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "Hidden 中 in a table cell"
    issues = validate_diacritics(doc)
    assert any("table 0" in i for i in issues)


# ---------------------------------------------------------------------------
# validate_font_consistency
# ---------------------------------------------------------------------------

def test_font_consistency_clean_on_empty_runs():
    # Fresh python-docx paragraphs have no rFonts unless explicitly set.
    doc = Document()
    doc.add_paragraph("plain text, no font override")
    assert validate_font_consistency(doc) == []


def test_font_consistency_passes_when_font_is_inter():
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("Text")
    r.font.name = "Inter"
    assert validate_font_consistency(doc) == []


def test_font_consistency_flags_non_inter_font():
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("Text")
    r.font.name = "Comic Sans MS"
    issues = validate_font_consistency(doc)
    assert any("Comic Sans" in i for i in issues)


# ---------------------------------------------------------------------------
# run_all
# ---------------------------------------------------------------------------

def test_run_all_returns_empty_for_clean_doc():
    doc = Document()
    doc.add_paragraph("A clean document.")
    assert run_all(doc) == []


def test_run_all_aggregates_multiple_violations():
    doc = Document()
    # Bad font + bad char
    p = doc.add_paragraph()
    r = p.add_run("中")
    r.font.name = "Comic Sans MS"
    issues = run_all(doc)
    # Should have both a diacritic issue and a font issue
    diacritic_issues = [i for i in issues if "disallowed character" in i]
    font_issues = [i for i in issues if "Comic Sans" in i]
    assert diacritic_issues
    assert font_issues
