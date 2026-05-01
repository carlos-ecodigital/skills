"""
Bilingual EN/NL two-column clause rendering for Sites-stream documents.

Additive to ``document-factory/generate.py``; does NOT modify existing
function signatures. Existing Sales/colocation and MIA callers are
unaffected.

Design choices (see Phase B2 plan):
- python-docx ``Table`` for structural layout; lxml used only for
  (a) cell borders (invisible layout grid), (b) tblW + tblLayout=fixed,
  (c) Cobalt heading-bar shading, (d) tblDescription marker.
- 50/50 column split at 165 mm usable width (A4 210 − left 25 − right 20).
- List nesting via explicit indent + bullet glyph (avoids polluting
  numbering.xml with table-scoped abstractNums that reset across page
  breaks).
- Bilingual tables carry ``tblDescription="__bilingual_body__"`` so that
  ``generate.audit_document`` can skip T1..T10 checks (which assume
  data-table conventions that don't apply to bilingual clause bodies).
- Balance validation: length-ratio check applies only when both EN and
  NL totals ≥ 40 chars (short clauses give noisy ratios).

Usage:

    from bilingual_body import render_bilingual_clause, validate_pair_balance

    render_bilingual_clause(
        doc,
        heading="1. Parties",
        heading_nl="1. Partijen",
        en_paragraphs=[
            "1.1 Digital Energy Netherlands B.V., ...",
            "1.2 The counterparty or counterparties ...",
        ],
        nl_paragraphs=[
            "1.1 Digital Energy Netherlands B.V., ...",
            "1.2 De wederpartij of wederpartijen ...",
        ],
    )
"""

from __future__ import annotations

import re
from typing import List, Optional

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from generate import (
    COBALT_HEX,
    FONT,
    SLATE,
    SLATE_300_HEX,
    SLATE_800,
    SLATE_900,
    WHITE,
)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

#: Marker placed in ``w:tblDescription`` to signal that a table is a
#: bilingual clause body (so ``audit_document`` can skip data-table rules).
BILINGUAL_MARKER = "__bilingual_body__"

#: Usable width in mm on A4 with 25/20 mm L/R margins.
USABLE_WIDTH_MM = 165.0

#: 50/50 split. Two columns of equal width for the EN/NL pair.
COL_WIDTH_MM = USABLE_WIDTH_MM / 2

BODY_FONT_SIZE = Pt(10)
HEADING_FONT_SIZE = Pt(11)

#: Below this total character count, skip the NL/EN length-ratio check.
MIN_BALANCE_CHARS = 40

#: NL is typically 10–15 % longer than EN; allow ±25 % initially (calibrate
#: empirically to ≤ 15 % against ``hot-grower-body-v1.docx`` in Phase C).
BALANCE_RATIO_MIN = 0.75
BALANCE_RATIO_MAX = 1.33

#: Twips per mm (approx): 1 mm = 56.7 twips; we use 25.4 mm = 1440 twips.
TWIPS_PER_INCH = 1440
MM_PER_INCH = 25.4


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

#: Tracks ``id(doc)`` values for which ``ensure_bilingual_layout`` has run.
#: Module-level so re-application within the same process is a no-op.
_LAYOUT_APPLIED: set = set()


def ensure_bilingual_layout(doc) -> None:
    """Idempotently set narrow A4 margins on every section in ``doc``.

    Bilingual two-column tables are sized at ``USABLE_WIDTH_MM = 165 mm``,
    which only fits inside A4 (210 mm) when L/R margins are 20 mm each.
    The default python-docx margins (31.75 mm) leave only 146.5 mm
    usable — too narrow.

    Margins set per section:
      - Left:   20 mm
      - Right:  20 mm
      - Top:    25 mm
      - Bottom: 20 mm

    Idempotent: re-application on the same doc is a no-op (tracked by
    ``id(doc)``). Safe to call from multiple call-sites within a single
    render pass.

    This helper is invoked transparently from the first call to
    ``render_bilingual_clause()`` per doc, so engines that exclusively
    use the bilingual API never need to call it directly.
    """
    if id(doc) in _LAYOUT_APPLIED:
        return
    for section in doc.sections:
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)
        section.top_margin = Mm(25)
        section.bottom_margin = Mm(20)
    _LAYOUT_APPLIED.add(id(doc))


def render_bilingual_clause(
    doc,
    en_paragraphs: List[str],
    nl_paragraphs: List[str],
    heading: Optional[str] = None,
    heading_nl: Optional[str] = None,
    level: int = 2,
    style: str = "body",
) -> None:
    """Render a bilingual EN/NL two-column clause.

    A new table is appended to ``doc``. Column 0 holds the EN paragraphs,
    column 1 holds the NL paragraphs. If ``heading`` is given, a heading
    row is rendered above the paragraph rows with Cobalt fill and WHITE
    bold text.

    Args:
        doc: python-docx ``Document``.
        en_paragraphs: English paragraphs. Each item is one paragraph.
            Leading ``- ``, ``* ``, ``+ ``, or ``N. `` is interpreted as a
            list item; 2-space indent per nesting level.
        nl_paragraphs: Dutch paragraphs. Must be the same length and carry
            matching list-nesting structure to ``en_paragraphs``.
        heading: Optional clause heading (EN).
        heading_nl: Dutch heading. Defaults to ``heading`` if omitted.
        level: Heading level (2 or 3; affects font size).
        style: ``"body"``, ``"recital"``, or ``"signature_intro"``. Currently
            reserved — all three map to the body font / size.

    Raises:
        ValueError: if ``len(en_paragraphs) != len(nl_paragraphs)``.
    """
    if len(en_paragraphs) != len(nl_paragraphs):
        raise ValueError(
            "bilingual clause paragraph-count mismatch: "
            f"EN={len(en_paragraphs)} NL={len(nl_paragraphs)}"
        )

    # First-call hook — guarantees the doc has the narrow margins required
    # for the 165 mm USABLE_WIDTH bilingual table to fit on A4. Idempotent.
    ensure_bilingual_layout(doc)

    n_rows = (1 if heading is not None else 0) + len(en_paragraphs)
    if n_rows == 0:
        return  # empty clause; no-op

    table = doc.add_table(rows=n_rows, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    _apply_table_chrome(table._tbl)
    _set_col_widths(table, COL_WIDTH_MM, COL_WIDTH_MM)

    row_idx = 0
    if heading is not None:
        hrow = table.rows[0].cells
        _render_heading(hrow[0], heading)
        _render_heading(hrow[1], heading_nl or heading)
        row_idx = 1

    for en, nl in zip(en_paragraphs, nl_paragraphs):
        row = table.rows[row_idx].cells
        _render_paragraph(row[0], en)
        _render_paragraph(row[1], nl)
        row_idx += 1


def validate_pair_balance(
    en_paragraphs: List[str],
    nl_paragraphs: List[str],
) -> List[str]:
    """Validate bilingual-pair structural balance. Returns a list of issue
    strings (empty list = clean)."""
    issues: List[str] = []

    if len(en_paragraphs) != len(nl_paragraphs):
        issues.append(
            "paragraph-count mismatch: "
            f"EN={len(en_paragraphs)} NL={len(nl_paragraphs)}"
        )
        return issues  # downstream checks would be meaningless

    # Empty-cell detection
    for i, (en, nl) in enumerate(zip(en_paragraphs, nl_paragraphs)):
        if not en.strip():
            issues.append(f"paragraph {i}: EN cell empty")
        if not nl.strip():
            issues.append(f"paragraph {i}: NL cell empty")

    # Length-ratio check (only if both totals ≥ MIN_BALANCE_CHARS)
    en_total = sum(len(p) for p in en_paragraphs)
    nl_total = sum(len(p) for p in nl_paragraphs)
    if en_total >= MIN_BALANCE_CHARS and nl_total >= MIN_BALANCE_CHARS:
        ratio = nl_total / en_total
        if ratio < BALANCE_RATIO_MIN or ratio > BALANCE_RATIO_MAX:
            issues.append(
                f"NL/EN length ratio {ratio:.2f} outside "
                f"[{BALANCE_RATIO_MIN}, {BALANCE_RATIO_MAX}]"
            )

    # Nested-list depth parity (per-paragraph)
    for i, (en, nl) in enumerate(zip(en_paragraphs, nl_paragraphs)):
        en_d = _list_depth(en)
        nl_d = _list_depth(nl)
        if en_d != nl_d:
            issues.append(
                f"paragraph {i}: list-nesting mismatch "
                f"(EN depth {en_d}, NL depth {nl_d})"
            )

    return issues


# ---------------------------------------------------------------------------
# Private helpers
# ---------------------------------------------------------------------------

def _mm_to_twips(mm: float) -> int:
    """Convert millimetres to twentieths of a point (twips)."""
    return int(round(mm / MM_PER_INCH * TWIPS_PER_INCH))


def _get_or_create_tblPr(tbl) -> OxmlElement:
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    return tblPr


def _apply_table_chrome(tbl) -> None:
    """Fixed-layout, full-width, marker, invisible borders."""
    tblPr = _get_or_create_tblPr(tbl)

    # Width: full usable width in twips (dxa)
    tblW_existing = tblPr.find(qn("w:tblW"))
    if tblW_existing is not None:
        tblPr.remove(tblW_existing)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(_mm_to_twips(USABLE_WIDTH_MM)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    # Fixed layout (no autofit — widths are honoured)
    tblLayout_existing = tblPr.find(qn("w:tblLayout"))
    if tblLayout_existing is not None:
        tblPr.remove(tblLayout_existing)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)

    # Description marker — identifies this as a bilingual body table for
    # audit_document's T1..T10 exemption path.
    desc_existing = tblPr.find(qn("w:tblDescription"))
    if desc_existing is not None:
        tblPr.remove(desc_existing)
    tblDesc = OxmlElement("w:tblDescription")
    tblDesc.set(qn("w:val"), BILINGUAL_MARKER)
    tblPr.append(tblDesc)

    # Invisible borders (intentionally — the two columns are a layout grid,
    # not a data table).
    borders_existing = tblPr.find(qn("w:tblBorders"))
    if borders_existing is not None:
        tblPr.remove(borders_existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    tblPr.append(borders)


def _set_col_widths(table, col0_mm: float, col1_mm: float) -> None:
    """Set explicit column widths. python-docx alone is unreliable for this,
    so we also write tblGrid directly."""
    # Python-docx layer
    table.columns[0].width = Mm(col0_mm)
    table.columns[1].width = Mm(col1_mm)

    # XML layer (tblGrid) — python-docx auto-generates this but we overwrite
    # to guarantee exact widths.
    tbl = table._tbl
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for w_mm in (col0_mm, col1_mm):
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(_mm_to_twips(w_mm)))
        grid.append(col)
    tbl.insert(list(tbl).index(tbl.find(qn("w:tblPr"))) + 1, grid)


def _shade_cell(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    # Remove any prior shd
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _clear_cell_paragraphs(cell) -> None:
    """Delete all paragraphs in the cell except the first (python-docx
    requires at least one paragraph per cell), then clear its runs."""
    paragraphs = list(cell.paragraphs)
    for p in paragraphs[1:]:
        p._element.getparent().remove(p._element)
    # Clear runs of the first paragraph
    first = paragraphs[0]
    for r in list(first.runs):
        r._element.getparent().remove(r._element)


def _render_heading(cell, text: str) -> None:
    _shade_cell(cell, COBALT_HEX)
    _clear_cell_paragraphs(cell)
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = True
    r.font.name = FONT
    r.font.size = HEADING_FONT_SIZE
    r.font.color.rgb = WHITE


def _render_paragraph(cell, text: str) -> None:
    _clear_cell_paragraphs(cell)
    # Multi-line text: split on literal \n
    lines = text.split("\n") if "\n" in text else [text]
    first = True
    for line in lines:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()

        depth = _list_depth(line)
        clean = _strip_list_prefix(line)

        if depth > 0:
            p.paragraph_format.left_indent = Mm(5 * depth)
            p.paragraph_format.first_line_indent = Mm(-3)
            r = p.add_run(f"\u2022  {clean}")
        else:
            r = p.add_run(clean)

        r.font.name = FONT
        r.font.size = BODY_FONT_SIZE
        r.font.color.rgb = SLATE_800


_LIST_MARK_RE = re.compile(r"^(?:[-*+]|\d+[.)])\s+")


def _list_depth(line: str) -> int:
    """Return list-nesting depth: 0 = plain prose; 1 = top-level bullet/
    number; 2 = 2-space indented bullet/number; etc."""
    stripped = line.lstrip(" ")
    if not _LIST_MARK_RE.match(stripped):
        return 0
    indent = len(line) - len(stripped)
    return (indent // 2) + 1


def _strip_list_prefix(line: str) -> str:
    """Strip the leading list marker (``- ``, ``* ``, ``+ ``, or ``N. ``)
    along with any leading indent spaces."""
    stripped = line.lstrip(" ")
    m = _LIST_MARK_RE.match(stripped)
    if m:
        return stripped[m.end():]
    return stripped
