"""Tests for ``add_cover()`` bilingual extension (rc3.2).

Covers the new kwargs: ``bilingual``, ``title_nl``, ``subject_nl``,
``party_labels="bilingual"|"en"|"nl"``, and ``render_classification``.

Each test asserts on the python-docx paragraph stream produced by
``add_cover`` so we don't depend on golden .docx fixtures that drift
across python-docx versions.
"""

from __future__ import annotations

from docx import Document

from generate import Party, add_cover


def _para_texts(doc):
    """Concatenated text of every paragraph in the document, in order."""
    return [p.text for p in doc.paragraphs]


def _all_text(doc) -> str:
    return "\n".join(_para_texts(doc))


# ---------------------------------------------------------------------------
# Bilingual title + subject stacking
# ---------------------------------------------------------------------------

def test_bilingual_title_renders_stacked():
    doc = Document()
    add_cover(
        doc,
        agreement_type="Letter of Intent",
        subject="for a Digital Energy Center Project",
        bilingual=True,
        title_nl="Intentieverklaring",
        subject_nl="voor een Digital Energy Center-project",
    )
    texts = _para_texts(doc)
    en_idx = next(i for i, t in enumerate(texts) if t == "Letter of Intent")
    nl_idx = next(i for i, t in enumerate(texts) if t == "Intentieverklaring")
    assert nl_idx == en_idx + 1, "NL title must immediately follow EN title"


def test_bilingual_nl_subtitle_present():
    doc = Document()
    add_cover(
        doc,
        agreement_type="Letter of Intent",
        subject="for a Digital Energy Center Project",
        bilingual=True,
        title_nl="Intentieverklaring",
        subject_nl="voor een Digital Energy Center-project",
    )
    body = _all_text(doc)
    assert "voor een Digital Energy Center-project" in body
    assert "for a Digital Energy Center Project" in body


def test_bilingual_nl_title_italic_16pt():
    doc = Document()
    add_cover(
        doc,
        agreement_type="Letter of Intent",
        bilingual=True,
        title_nl="Intentieverklaring",
    )
    # Find the run with the NL title
    found = False
    for p in doc.paragraphs:
        for r in p.runs:
            if r.text == "Intentieverklaring":
                found = True
                assert r.italic is True
                # python-docx font size is in EMU; Pt(16) == 203200
                assert r.font.size is not None
                assert r.font.size.pt == 16
    assert found, "NL title run not found"


# ---------------------------------------------------------------------------
# party_labels — string spec ("en" | "nl" | "bilingual") + legacy list
# ---------------------------------------------------------------------------

_PROVIDER = Party(
    legal_name="Digital Energy Netherlands B.V.",
    address="Amsterdam, NL",
    registration_type="KvK",
    registration_number="12345678",
)
_COUNTERPARTY = Party(
    legal_name="Counterparty B.V.",
    address="Rotterdam, NL",
)


def test_party_labels_bilingual():
    doc = Document()
    add_cover(
        doc,
        agreement_type="Letter of Intent",
        parties=[_PROVIDER, _COUNTERPARTY],
        party_labels="bilingual",
    )
    body = _all_text(doc)
    assert "Between / Tussen:" in body
    assert "And / En:" in body
    assert "Between:" not in body.replace("Between / Tussen:", "")


def test_party_labels_en_explicit():
    doc = Document()
    add_cover(
        doc,
        agreement_type="Letter of Intent",
        parties=[_PROVIDER, _COUNTERPARTY],
        party_labels="en",
    )
    body = _all_text(doc)
    assert "Between:" in body
    assert "And:" in body
    assert "Tussen:" not in body


def test_party_labels_nl_explicit():
    doc = Document()
    add_cover(
        doc,
        agreement_type="Letter of Intent",
        parties=[_PROVIDER, _COUNTERPARTY],
        party_labels="nl",
    )
    body = _all_text(doc)
    assert "Tussen:" in body
    assert "En:" in body


# ---------------------------------------------------------------------------
# Backwards compatibility — monolingual default unchanged
# ---------------------------------------------------------------------------

def test_monolingual_default_unchanged():
    """Existing colocation/MIA callers must keep current behaviour."""
    doc = Document()
    add_cover(
        doc,
        agreement_type="Letter of Intent",
        subject="for a Test Project",
        parties=[_PROVIDER, _COUNTERPARTY],
        # No bilingual=, no party_labels=, no render_classification=
    )
    body = _all_text(doc)
    assert "Letter of Intent" in body
    assert "for a Test Project" in body
    assert "Between:" in body
    assert "And:" in body
    # Default classification still rendered
    assert "Confidential" in body
    # No NL stacking by default
    assert "Intentieverklaring" not in body


def test_page_break_preserved():
    """add_cover ends with a page break — must hold even on bilingual path."""
    doc = Document()
    add_cover(
        doc,
        agreement_type="Letter of Intent",
        bilingual=True,
        title_nl="Intentieverklaring",
    )
    # The final paragraph should contain a w:br with type=page in its XML.
    from docx.oxml.ns import qn
    last_p = doc.paragraphs[-1]
    breaks_in_last = last_p._p.findall(".//" + qn("w:br"))
    if not breaks_in_last:
        # The break may live in an empty trailing paragraph; check second-to-last.
        any_break = any(
            p._p.findall(".//" + qn("w:br"))
            for p in doc.paragraphs
        )
        assert any_break, "expected a page-break run somewhere in the cover"


# ---------------------------------------------------------------------------
# Classification rendering toggle
# ---------------------------------------------------------------------------

def test_classification_renders_by_default():
    doc = Document()
    add_cover(
        doc,
        agreement_type="Letter of Intent",
        classification="Confidential",
    )
    body = _all_text(doc)
    assert "Classification" in body
    assert "Confidential" in body


def test_classification_omitted_when_render_false():
    doc = Document()
    add_cover(
        doc,
        agreement_type="Letter of Intent",
        classification="Confidential",
        render_classification=False,
    )
    body = _all_text(doc)
    assert "Classification" not in body
    # Classification value (which doubles as the visible footer) also gone
    assert "Confidential" not in body


def test_render_classification_false_keeps_reference_and_version():
    """render_classification=False omits ONLY the Classification row;
    Reference + Version metadata still render if supplied."""
    doc = Document()
    add_cover(
        doc,
        agreement_type="Letter of Intent",
        reference="DE-LOI-VG-001",
        version="1.0",
        classification="Confidential",
        render_classification=False,
    )
    body = _all_text(doc)
    assert "Reference" in body
    assert "DE-LOI-VG-001" in body
    assert "Version" in body
    assert "v1.0" in body
    assert "Classification" not in body


# ---------------------------------------------------------------------------
# Legacy list-form party_labels still works (regression guard)
# ---------------------------------------------------------------------------

def test_party_labels_legacy_list_still_works():
    doc = Document()
    add_cover(
        doc,
        agreement_type="Term Sheet",
        parties=[_PROVIDER, _COUNTERPARTY],
        party_labels=["Prepared by:", "Prepared for:"],
    )
    body = _all_text(doc)
    assert "Prepared by:" in body
    assert "Prepared for:" in body
