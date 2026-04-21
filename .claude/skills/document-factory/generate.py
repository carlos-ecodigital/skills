#!/usr/bin/env python3
"""
Digital Energy Document Generator
===================================
Single-file generator. No XML manipulation. Only native python-docx API.

Usage:
    python3 generate.py --profile letter
    python3 generate.py --profile agreement --title "Colocation Agreement" --client "Younggrow BV"
    python3 generate.py --profile seed_memo --client "Acme Fund"
    python3 generate.py --profile investor_memo --client "Infrastructure Partners"
    python3 generate.py --profile exec_summary --title "PowerGrow Project Update"

Add --dotx to also save a Word template file.
"""

import argparse
import os
import sys
import re
from dataclasses import dataclass
from datetime import datetime, date
from typing import Optional

try:
    from docx import Document
    from docx.shared import Mm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml.ns import qn
    from lxml import etree
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)

# ---------------------------------------------------------------------------
# CONFIG
# ---------------------------------------------------------------------------

DE_WEBSITE = "digital-energy.group"

ENTITY_FOOTERS = {
    "ag": {
        "legal_name": "Digital Energy Group AG",
        "address": "Baarerstrasse 43, 6300 Zug, Switzerland",
        "registration": "CHE-408.639.320",
        "website": DE_WEBSITE,
        "return_address": "Digital Energy Group AG  \u2022  Baarerstrasse 43  \u2022  6300 Zug",
    },
    "nl": {
        "legal_name": "Digital Energy Netherlands B.V.",
        "address": "Mijnsherenweg 33 A, 1433 AP Kudelstaart, The Netherlands",
        "registration": "KvK 98580086",
        "website": DE_WEBSITE,
        "return_address": "Digital Energy Netherlands B.V.  \u2022  Mijnsherenweg 33 A  \u2022  1433 AP Kudelstaart",
    },
}
ENTITY = ENTITY_FOOTERS["ag"]  # Backward-compat default


def _footer_text(entity_key="ag"):
    e = ENTITY_FOOTERS.get(entity_key, ENTITY_FOOTERS["ag"])
    return "  \u00b7  ".join([e["legal_name"], e["address"], e["registration"], e["website"]])

FONT = "Inter"

# DE Approved Palette (2026-03-31) — see memory/project_de_brand_palette.md
# Intended use by domain:
#   COBALT  → primary CTA / headers / links (default accent)
#   VOLT    → success / growth accents (finance reports, KPI tables)
#   NEON    → technology / AI themes (product, pitch)
#   MINERAL → compute / premium (investor memos)
#   FORGE   → heat / construction (permits, engineering)
#   SAFFRON → value / returns (IRR tables, unit economics)
#   PATINA  → sustainability (ESG, emissions)
#   SLATE   → secondary / captions (default body secondary)
# Constants are exported for import by any skill that needs them.
COBALT = RGBColor(0x00, 0x34, 0xAF)      # Primary CTA, headers
VOLT = RGBColor(0x63, 0xE2, 0x34)        # Success, growth
NEON = RGBColor(0x16, 0xD3, 0xF2)        # Info, technology
MINERAL = RGBColor(0x58, 0x1C, 0x87)     # AI/compute, premium
FORGE = RGBColor(0xC0, 0x56, 0x21)       # Heat, construction
SAFFRON = RGBColor(0xE5, 0xB2, 0x1D)     # Value, returns
PATINA = RGBColor(0x0F, 0x76, 0x6E)      # Sustainability
SLATE = RGBColor(0x64, 0x74, 0x8B)       # Slate 500 — captions, secondary

# Neutral ramp (Slate)
SLATE_50_HEX = "F8FAFC"                   # Table alt-row shading (hex for XML attrs)
SLATE_300_HEX = "CBD5E1"                   # Table borders (hex for XML attrs)
COBALT_HEX = "0034AF"                      # Header fill (hex for XML attrs)
SLATE_800 = RGBColor(0x1E, 0x29, 0x3B)   # Body text
SLATE_900 = RGBColor(0x0F, 0x17, 0x2A)   # Headings
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ---------------------------------------------------------------------------
# SPACING CONSTANTS — single source of truth for all formatting
# ---------------------------------------------------------------------------
_SP = {
    'h1_before': Pt(18),   'h2_before': Pt(18),
    'h3_before': Pt(12),   'h4_before': Pt(12),
    'heading_after': Pt(6),
    'body_after': Pt(6),
    'body_line': Pt(16.5),
    'list_first_before': Pt(6),
    'list_gap': Pt(2),
    'list_last_after': Pt(6),
    'table_before': Pt(12),
    'table_after': Pt(12),
    'sig_section_before': Pt(36),
    'sig_party_gap': Pt(18),
    'sig_line_gap': Pt(2),
}

# Column width constants for table formatting
_CHAR_WIDTH_EMU = 63_500    # ~5pt avg char width at 10pt Inter
_CELL_PAD_EMU = int(Mm(4))  # 2mm left + 2mm right padding
_MIN_COL_EMU = int(Mm(15))  # absolute minimum column width


def _display_len(text):
    """Character count after removing markdown formatting markers."""
    t = re.sub(r'\*\*([^*]+)\*\*', r'\1', str(text))
    t = re.sub(r'\*([^*]+)\*', r'\1', t)
    t = re.sub(r'`([^`]+)`', r'\1', t)
    return len(t)


def _longest_word_len(text):
    """Length of the longest single word after stripping markdown markers."""
    t = re.sub(r'\*\*([^*]+)\*\*', r'\1', str(text))
    t = re.sub(r'\*([^*]+)\*', r'\1', t)
    t = re.sub(r'`([^`]+)`', r'\1', t)
    return max((len(w) for w in t.split()), default=0)


SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO = os.path.join(SCRIPT_DIR, "assets", "DE_Logo_Black.png")
if not os.path.exists(LOGO):
    # Fallback: check script directory root
    LOGO = os.path.join(SCRIPT_DIR, "DE_Logo_Black.png")
if not os.path.exists(LOGO):
    print("Warning: DE_Logo_Black.png not found — documents will be generated without logo",
          file=sys.stderr)
OUTPUT = os.path.join(SCRIPT_DIR, "output")

PROFILE_CODES = {
    "letter": "Letter",
    "agreement": "Agreement",
    "seed_memo": "Seed_Memo",
    "investor_memo": "Investor_Memo",
    "exec_summary": "Exec_Summary",
}



# ---------------------------------------------------------------------------
# PARTY DATA MODEL + ENTITY SYSTEM
# ---------------------------------------------------------------------------

@dataclass
class Party:
    legal_name: str = ""
    address: str = ""
    registration_type: Optional[str] = None    # "KvK", "CHE", "EIN"
    registration_number: Optional[str] = None
    parent: Optional[str] = None               # Parent company name + reg


DE_ENTITIES = {
    "ag": Party(
        legal_name="Digital Energy Group AG",
        address="Baarerstrasse 43, 6300 Zug, Switzerland",
        registration_type="CHE",
        registration_number="408.639.320",
    ),
    "nl": Party(
        legal_name="Digital Energy Netherlands B.V.",
        address="Mijnsherenweg 33 A, 1433 AP Kudelstaart, The Netherlands",
        registration_type="KvK",
        registration_number="98580086",
    ),  # No parent — NL presented as standalone entity on all documents
}


AGREEMENT_FORMALITY = {
    # Non-binding: name + address. Labels: "Between:" / "And:"
    "Letter of Intent": "non_binding",
    "Letter of Intent and NCNDA": "non_binding",
    "LOI": "non_binding",
    "NCNDA": "non_binding",
    "NDA": "non_binding",
    "HoT": "non_binding",
    "Heads of Terms": "non_binding",
    "MoU": "non_binding",
    "Memorandum of Understanding": "non_binding",
    "Term Sheet": "non_binding",
    "Non-Disclosure Agreement": "non_binding",
    "Non-Circumvention Non-Disclosure Agreement": "non_binding",
    # Binding: name + address + registration. Labels: "By and between:" / "And:"
    "Master Service Agreement": "binding",
    "MSA": "binding",
    "SPA": "binding",
    "Sales and Purchase Agreement": "binding",
    "JVA": "binding",
    "Joint Venture Agreement": "binding",
    "SHA": "binding",
    "Shareholders Agreement": "binding",
    "License Agreement": "binding",
    "Services Agreement": "binding",
    "Colocation Agreement": "binding",
    "Advisory Agreement": "binding",
    "SAR Grant Agreement": "binding",
    "Board Resolution": "binding",
    "Program Policy": "non_binding",
    "Terms of Reference": "non_binding",
    "Stock Appreciation Rights — Term Sheet": "non_binding",
}


def _detect_formality(title):
    """Look up formality by exact match, then by prefix match."""
    if not title:
        return "non_binding"
    if title in AGREEMENT_FORMALITY:
        return AGREEMENT_FORMALITY[title]
    # Prefix match: "Board Resolution — SAR Program" matches "Board Resolution"
    for key, val in AGREEMENT_FORMALITY.items():
        if title.startswith(key):
            return val
    return "non_binding"


_TYPE_ABBREVS = {
    "Letter of Intent": "LOI",
    "Letter of Intent and NCNDA": "LOI_NCNDA",
    "NCNDA": "NCNDA",
    "NDA": "NDA",
    "Heads of Terms": "HoT",
    "HoT": "HoT",
    "MoU": "MoU",
    "Term Sheet": "TS",
    "Master Service Agreement": "MSA",
    "MSA": "MSA",
    "SPA": "SPA",
    "Sales and Purchase Agreement": "SPA",
    "JVA": "JVA",
    "Joint Venture Agreement": "JVA",
    "SHA": "SHA",
    "Shareholders Agreement": "SHA",
    "License Agreement": "License",
    "Services Agreement": "Services",
    "Colocation Agreement": "Colo",
}


def _type_code(agreement_type):
    if not agreement_type:
        return None
    if agreement_type in _TYPE_ABBREVS:
        return _TYPE_ABBREVS[agreement_type]
    # Fallback: alphanumeric with underscores, word-boundary truncation at 20
    cleaned = re.sub(r"[^A-Za-z0-9]+", "_", agreement_type).strip("_")
    if not cleaned:
        return None
    max_len = 25
    if len(cleaned) <= max_len:
        return cleaned
    # Find the last underscore at-or-before max_len; if none, hard-trim at 20
    cut = cleaned.rfind("_", 0, max_len + 1)
    if cut <= 0:
        return cleaned[:20]
    return cleaned[:cut]


def auto_name(profile, client=None, version=1, dt=None, agreement_type=None):
    if dt is None:
        dt = date.today()
    if profile == "agreement" and agreement_type:
        type_part = _type_code(agreement_type) or PROFILE_CODES["agreement"]
    else:
        type_part = PROFILE_CODES.get(profile, profile)
    parts = [dt.strftime("%Y%m%d"), "DE", type_part]
    if client:
        parts.append(client.replace(" ", "_")[:30])
    parts.append(f"v{version}")
    return "_".join(parts) + ".docx"


# ---------------------------------------------------------------------------
# BUILDING BLOCKS (no OxmlElement, no raw XML)
# ---------------------------------------------------------------------------

def _fix_zoom(doc):
    """Ensure w:zoom has w:percent attribute (python-docx omits it, fails OOXML validation)."""
    settings = doc.settings.element
    zoom = settings.find(qn('w:zoom'))
    if zoom is not None and zoom.get(qn('w:percent')) is None:
        zoom.set(qn('w:percent'), '100')


def audit_document(doc):
    """Comprehensive formatting audit. Returns list of violation strings.

    Empty list = clean. Checks every invariant from formatting-standards.md:
    D1-D10 (document), P1-P6 (paragraph), T1-T10 (table), X1-X3 (OOXML).
    Cover page paragraphs (before first page break) are exempt from body checks.
    """
    v = []  # violations

    # --- Locate cover page boundary ---
    body_el = doc.element.body
    cover_end_idx = -1  # paragraph index where cover ends
    for p_idx, p_el in enumerate(body_el.findall(qn('w:p'))):
        if p_el.find(f'.//{qn("w:br")}[@{qn("w:type")}="page"]') is not None:
            cover_end_idx = p_idx
            break

    # --- Regexes ---
    _list_num_re = re.compile(r'^\d+[.)]\s+')
    _list_bul_re = re.compile(r'^[-*+\u2022]\s+')
    _list_alp_re = re.compile(r'^\([a-z]\)\s+')
    _list_rom_re = re.compile(r'^\((i{1,3}|iv|v|vi{0,3}|ix|x)\)\s+')
    # Signature line: starts with keyword or underscores (not buried mid-sentence)
    _sig_re = re.compile(r'^(_{3,}|(Name|Title|Signature|Date|Signed|By)\s*:)', re.IGNORECASE)

    def _is_heading(para):
        if not para.runs:
            return False
        r = para.runs[0]
        return bool(r.font.bold and r.font.size and r.font.size >= Pt(12))

    def _is_spacer(para):
        return not para.text.strip() or (para.runs and para.runs[0].font.size and para.runs[0].font.size <= Pt(2))

    def _is_italic(para):
        return para.runs and all(r.font.italic for r in para.runs if r.text.strip())

    # ===== D: Document checks =====
    sec = doc.sections[0]
    if not doc.core_properties.title:
        v.append("D1: core_properties.title is empty")
    if not doc.core_properties.author:
        v.append("D2: core_properties.author is empty")

    settings = doc.settings.element
    ah = settings.find(qn('w:autoHyphenation'))
    if ah is None or ah.get(qn('w:val')) != '0':
        v.append("D3: auto-hyphenation not disabled")

    zoom = settings.find(qn('w:zoom'))
    if zoom is not None and zoom.get(qn('w:percent')) is None:
        v.append("D4: w:zoom missing w:percent")

    # D5: font fallback
    fallback_ok = False
    for para in doc.paragraphs[:30]:
        for run in para.runs:
            rPr = run._element.find(qn('w:rPr'))
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None and rFonts.get(qn('w:cs')) and rFonts.get(qn('w:eastAsia')):
                    fallback_ok = True
                    break
        if fallback_ok:
            break
    if not fallback_ok:
        v.append("D5: no font fallback (w:cs/w:eastAsia) on any run")

    # D6-D8: page layout
    if abs(sec.page_width - Mm(210)) > Mm(1):
        v.append(f"D6: page_width {sec.page_width} != Mm(210)")
    if abs(sec.page_height - Mm(297)) > Mm(1):
        v.append(f"D6: page_height {sec.page_height} != Mm(297)")
    if abs(sec.left_margin - Mm(25)) > Mm(1):
        v.append(f"D7: left_margin {sec.left_margin} != Mm(25)")
    if abs(sec.right_margin - Mm(20)) > Mm(1):
        v.append(f"D7: right_margin {sec.right_margin} != Mm(20)")
    valid_top = {int(Mm(15)), int(Mm(20)), int(Mm(25))}
    if not any(abs(sec.top_margin - t) < Mm(1) for t in valid_top):
        v.append(f"D8: top_margin {sec.top_margin} not in {{Mm(15), Mm(20)}}")
    valid_bot = {int(Mm(25)), int(Mm(35))}
    if not any(abs(sec.bottom_margin - b) < Mm(1) for b in valid_bot):
        v.append(f"D8: bottom_margin {sec.bottom_margin} not in {{Mm(25), Mm(35)}}")

    # D9: abstractNum before num
    try:
        numbering_el = doc.part.numbering_part.element
        last_abstract_idx = -1
        first_num_idx = len(numbering_el) + 1
        for idx, child in enumerate(numbering_el):
            if child.tag == qn('w:abstractNum'):
                last_abstract_idx = idx
            elif child.tag == qn('w:num'):
                first_num_idx = min(first_num_idx, idx)
        if last_abstract_idx > first_num_idx:
            v.append("D9: w:abstractNum after w:num in numbering.xml")
    except Exception:
        pass  # no numbering part = no lists = OK

    # D10: custom abstractNums exist if lists present
    has_lists = any(
        p._element.find(f'.//{qn("w:numPr")}') is not None
        for p in doc.paragraphs
    )
    if has_lists:
        try:
            numbering_el = doc.part.numbering_part.element
            abs_ids = {
                e.get(qn('w:abstractNumId'))
                for e in numbering_el.findall(qn('w:abstractNum'))
            }
            for required in ('100', '101', '102', '103'):
                if required not in abs_ids:
                    v.append(f"D10: abstractNum {required} missing")
        except Exception:
            v.append("D10: no numbering part but lists exist")

    # ===== P: Paragraph checks (body only — skip cover) =====
    all_p_elements = body_el.findall(qn('w:p'))
    for p_idx, para in enumerate(doc.paragraphs):
        # Skip cover paragraphs
        if cover_end_idx >= 0 and p_idx <= cover_end_idx:
            continue

        text = para.text.strip()
        if not text:
            continue
        if _is_spacer(para):
            continue

        pf = para.paragraph_format
        is_h = _is_heading(para)

        # P1: heading keep_with_next
        if is_h and not pf.keep_with_next:
            v.append(f"P1: heading missing keep_with_next: '{text[:40]}'")

        # P2: body widow_control (skip headings, italic/guidance, blockquotes)
        if not is_h and not _is_italic(para) and pf.widow_control is not True:
            # Check it's actual body content (has runs, not just whitespace)
            if para.runs and any(r.text.strip() for r in para.runs):
                v.append(f"P2: missing widow_control: '{text[:40]}'")

        # P3: list text without native numbering
        if not is_h:
            is_list_text = (_list_num_re.match(text) or _list_bul_re.match(text)
                           or _list_alp_re.match(text) or _list_rom_re.match(text))
            has_numPr = para._element.find(f'.//{qn("w:numPr")}') is not None
            if is_list_text and not has_numPr:
                v.append(f"P3: plain-text list without w:numPr: '{text[:50]}'")

        # P4: signature keep_together
        if _sig_re.match(text) and not pf.keep_together:
            v.append(f"P4: sig line missing keep_together: '{text[:40]}'")

        # P5: heading color
        if is_h and para.runs[0].font.color and para.runs[0].font.color.rgb:
            if para.runs[0].font.color.rgb != SLATE_900:
                v.append(f"P5: heading color {para.runs[0].font.color.rgb} != SLATE_900: '{text[:30]}'")

        # P6: body color (skip italic/guidance, skip headings)
        # Allow SLATE_800 (primary body) and SLATE (secondary/closing text)
        if not is_h and not _is_italic(para) and para.runs:
            r = para.runs[0]
            if r.font.size and abs(r.font.size - Pt(11)) < Pt(1) and not r.font.bold:
                if r.font.color and r.font.color.rgb:
                    if r.font.color.rgb not in (SLATE_800, SLATE):
                        v.append(f"P6: body color {r.font.color.rgb} not SLATE_800/SLATE: '{text[:30]}'")

        # P7: numbered heading must have pPr/rPr with bold (for numbering character)
        if is_h and para._element.find(f'.//{qn("w:numPr")}') is not None:
            pPr = para._element.find(qn('w:pPr'))
            pRpr = pPr.find(qn('w:rPr')) if pPr is not None else None
            if pRpr is None or pRpr.find(qn('w:b')) is None:
                v.append(f"P7: numbered heading missing pPr/rPr bold: '{text[:40]}'")

    # ===== T: Table checks =====
    for t_idx, table in enumerate(doc.tables):
        tblPr = table._tbl.find(qn('w:tblPr'))

        # T0: exempt bilingual-body tables from data-table rules.
        # Two-column EN/NL clause bodies are intentional layout tables, not
        # data tables — the Cobalt-header / CENTER-align / WHITE-bold
        # conventions don't apply. Rendered by bilingual_body.py.
        if tblPr is not None:
            _desc = tblPr.find(qn('w:tblDescription'))
            if _desc is not None and _desc.get(qn('w:val')) == '__bilingual_body__':
                continue

        # T1: width
        if tblPr is not None:
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is not None:
                wt = tblW.get(qn('w:type'))
                wv = tblW.get(qn('w:w'))
                if wt == 'dxa' and wv and int(wv) > 9360:
                    v.append(f"T1: table {t_idx} width {wv} > 9360 twips")

            # T2: fixed layout
            tblLayout = tblPr.find(qn('w:tblLayout'))
            if tblLayout is None or tblLayout.get(qn('w:type')) != 'fixed':
                v.append(f"T2: table {t_idx} not fixed layout")

        # T3: header repeat
        if len(table.rows) > 5:
            tr = table.rows[0]._tr
            trPr = tr.find(qn('w:trPr'))
            if trPr is None or trPr.find(qn('w:tblHeader')) is None:
                v.append(f"T3: table {t_idx} ({len(table.rows)} rows) missing header repeat")

        # T4: column minimum width (500 EMU tolerance for rounding)
        for c_idx, col in enumerate(table.columns):
            if col.width and col.width < _MIN_COL_EMU - 500:
                v.append(f"T4: table {t_idx} col {c_idx} width {col.width} < {_MIN_COL_EMU}")

        if not table.rows:
            continue

        # T5-T7: header row
        for c_idx, cell in enumerate(table.rows[0].cells):
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                shd = tcPr.find(qn('w:shd'))
                if shd is not None:
                    fill = (shd.get(qn('w:fill')) or '').upper()
                    if fill != COBALT_HEX.upper():
                        v.append(f"T5: table {t_idx} header col {c_idx} fill {fill} != {COBALT_HEX}")
            # T7: vertical alignment
            if cell.vertical_alignment != WD_CELL_VERTICAL_ALIGNMENT.CENTER:
                v.append(f"T7: table {t_idx} header col {c_idx} not CENTER aligned")

        # T6: header text
        for c_idx, cell in enumerate(table.rows[0].cells):
            for run in cell.paragraphs[0].runs if cell.paragraphs else []:
                if run.text.strip():
                    if run.font.color and run.font.color.rgb and run.font.color.rgb != WHITE:
                        v.append(f"T6: table {t_idx} header col {c_idx} text color != WHITE")
                    if not run.font.bold:
                        v.append(f"T6: table {t_idx} header col {c_idx} text not bold")
                    break

        # T8: data cell vertical alignment
        for r_idx in range(1, min(len(table.rows), 4)):  # sample first 3 data rows
            for c_idx, cell in enumerate(table.rows[r_idx].cells):
                if cell.vertical_alignment != WD_CELL_VERTICAL_ALIGNMENT.TOP:
                    v.append(f"T8: table {t_idx} row {r_idx} col {c_idx} not TOP aligned")

        # T9: borders (sample first data cell)
        if len(table.rows) > 1:
            tc = table.rows[1].cells[0]._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                tcBorders = tcPr.find(qn('w:tcBorders'))
                if tcBorders is None:
                    v.append(f"T9: table {t_idx} missing cell borders")
                else:
                    for edge in ('w:top', 'w:left', 'w:bottom', 'w:right'):
                        b = tcBorders.find(qn(edge))
                        if b is None:
                            v.append(f"T9: table {t_idx} missing {edge} border")
                        elif b.get(qn('w:val')) != 'single' or b.get(qn('w:sz')) != '4':
                            v.append(f"T9: table {t_idx} {edge} border wrong style/size")

        # T10: alt-row shading (sample)
        for r_idx in range(1, min(len(table.rows), 6)):
            cell = table.rows[r_idx].cells[0]
            tcPr = cell._tc.find(qn('w:tcPr'))
            if tcPr is not None:
                shd = tcPr.find(qn('w:shd'))
                if shd is not None:
                    fill = (shd.get(qn('w:fill')) or '').upper()
                    expected = SLATE_50_HEX.upper() if r_idx % 2 == 0 else 'FFFFFF'
                    if fill != expected:
                        v.append(f"T10: table {t_idx} row {r_idx} fill {fill} != {expected}")

    # ===== X: OOXML structural checks (sampled) =====
    # X1: numPr before spacing in pPr
    for para in doc.paragraphs:
        pPr = para._element.find(qn('w:pPr'))
        if pPr is None:
            continue
        numPr_idx = spacing_idx = None
        for idx, child in enumerate(pPr):
            if child.tag == qn('w:numPr'):
                numPr_idx = idx
            elif child.tag == qn('w:spacing'):
                spacing_idx = idx
        if numPr_idx is not None and spacing_idx is not None and numPr_idx > spacing_idx:
            v.append(f"X1: w:numPr after w:spacing in pPr")
            break  # one is enough

    # X2: tcBorders before shd in tcPr (sample first table)
    if doc.tables:
        for row in doc.tables[0].rows[:2]:
            for cell in row.cells:
                tcPr = cell._tc.find(qn('w:tcPr'))
                if tcPr is None:
                    continue
                b_idx = s_idx = None
                for idx, child in enumerate(tcPr):
                    if child.tag == qn('w:tcBorders'):
                        b_idx = idx
                    elif child.tag == qn('w:shd'):
                        s_idx = idx
                if b_idx is not None and s_idx is not None and b_idx > s_idx:
                    v.append("X2: w:tcBorders after w:shd in tcPr")
                    break
            else:
                continue
            break

    # X3: tblW before tblLayout in tblPr
    for table in doc.tables:
        tblPr = table._tbl.find(qn('w:tblPr'))
        if tblPr is None:
            continue
        w_idx = l_idx = None
        for idx, child in enumerate(tblPr):
            if child.tag == qn('w:tblW'):
                w_idx = idx
            elif child.tag == qn('w:tblLayout'):
                l_idx = idx
        if w_idx is not None and l_idx is not None and w_idx > l_idx:
            v.append("X3: w:tblW after w:tblLayout in tblPr")
            break

    return v


class AuditFailedError(Exception):
    """Raised in strict mode when audit_document() finds violations.

    Attributes:
        violations: list of violation strings from audit_document()
    """
    def __init__(self, violations):
        self.violations = violations
        msg = f"{len(violations)} audit violation(s)"
        if violations:
            msg += f": {violations[0]}"
            if len(violations) > 1:
                msg += f" (+{len(violations) - 1} more)"
        super().__init__(msg)


def save_doc(doc, path, strict=None):
    """Save document with post-processing fixes applied.

    Args:
        doc: python-docx Document
        path: output .docx path
        strict:
            None (default) — warn on audit violations via stderr, save anyway.
                             If env var DOCFACTORY_STRICT=1, behaves as True.
            True           — raise AuditFailedError on violations, do NOT save.
            False          — silent, save unconditionally (skip audit print).
    """
    _fix_zoom(doc)
    # Resolve the effective mode: "strict" | "warn" | "silent"
    if strict is True:
        mode = "strict"
    elif strict is False:
        mode = "silent"
    else:  # None — consult env var
        mode = "strict" if os.environ.get("DOCFACTORY_STRICT") == "1" else "warn"

    violations = audit_document(doc)

    if violations and mode == "strict":
        raise AuditFailedError(violations)

    if violations and mode == "warn":
        print(f"[audit] {len(violations)} issue(s):", file=sys.stderr)
        for w in violations[:10]:
            print(f"  - {w}", file=sys.stderr)

    doc.save(path)


def _run(para, text, size=Pt(11), color=SLATE_800, bold=False, italic=False, font=FONT):
    """Add a styled run with full font fallback chain."""
    r = para.add_run(text)
    r.font.name = font
    r.font.size = size
    r.font.color.rgb = color
    r.font.bold = bold
    r.font.italic = italic
    # Set w:cs and w:eastAsia fallback — without Inter, Word falls back to
    # Calibri (wider metrics → layout shift). Arial has similar metrics to Inter.
    rFonts = r._element.get_or_add_rPr().find(qn('w:rFonts'))
    if rFonts is not None:
        fallback = 'Arial' if font == FONT else font
        rFonts.set(qn('w:cs'), fallback)
        rFonts.set(qn('w:eastAsia'), fallback)
    return r


def new_doc(top=Mm(25), bottom=Mm(35), diff_first=False):
    """Create A4 document with professional margins.

    top=25mm gives ~10mm clearance below header (header occupies ~15mm).
    """
    doc = Document()
    s = doc.sections[0]
    s.page_width = Mm(210)
    s.page_height = Mm(297)
    s.left_margin = Mm(25)
    s.right_margin = Mm(20)
    s.top_margin = top
    s.bottom_margin = bottom
    s.different_first_page_header_footer = diff_first
    # Suppress auto-hyphenation — legal documents must not hyphenate
    settings = doc.settings.element
    autoHyph = settings.find(qn('w:autoHyphenation'))
    if autoHyph is None:
        autoHyph = etree.SubElement(settings, qn('w:autoHyphenation'))
    autoHyph.set(qn('w:val'), '0')
    return doc


def setup_first_page_header(section, logo_height=Mm(12)):
    """First page: logo only."""
    h = section.first_page_header
    h.is_linked_to_previous = False
    if os.path.exists(LOGO):
        h.paragraphs[0].add_run().add_picture(LOGO, height=logo_height)


def setup_cont_header(section, title=""):
    """Continuation pages: small logo left, title right."""
    h = section.header
    h.is_linked_to_previous = False
    p = h.paragraphs[0]
    p.paragraph_format.tab_stops.add_tab_stop(Mm(165), WD_TAB_ALIGNMENT.RIGHT)
    if os.path.exists(LOGO):
        p.add_run().add_picture(LOGO, height=Mm(8))
    if title:
        p.add_run("\t")
        r = p.add_run(title)
        r.font.size = Pt(9)
        r.font.name = FONT
        r.font.color.rgb = SLATE


def setup_footer(footer, classification=None, entity="ag"):
    """Footer: entity details, centre-aligned.

    Classification is intentionally NOT prefixed into the footer — it appears
    once on the cover page's metadata block. Repeating "Confidential" on every
    page footer is redundant noise. The `classification` arg is retained for
    signature compatibility but ignored here.

    v3.5 scope A'''' (LOI): footer now centre-aligned (was left). Entity
    selection (`entity="ag"` or `"nl"`) must match the signing Provider —
    callers pass `entity="nl"` for Digital Energy Netherlands B.V.
    instruments and `entity="ag"` for Digital Energy Group AG instruments.
    """
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(_footer_text(entity))
    r.font.size = Pt(7)
    r.font.name = FONT
    r.font.color.rgb = SLATE


def setup_first_footer(section, classification=None, entity="ag"):
    f = section.first_page_footer
    f.is_linked_to_previous = False
    setup_footer(f, classification, entity)


def setup_cont_footer(section, classification=None, entity="ag"):
    f = section.footer
    f.is_linked_to_previous = False
    setup_footer(f, classification, entity)


def _cover_title(agreement_type):
    """Derive the concise cover-page title from the full agreement type.

    Compound agreement names (e.g. "Letter of Intent and Non-Circumvention
    Non-Disclosure Agreement") are shortened to the primary type on the cover.
    The full legal name continues to appear in the body/recitals.

    Examples:
        "Letter of Intent and Non-Circumvention Non-Disclosure Agreement" -> "Letter of Intent"
        "Letter of Intent and NCNDA" -> "Letter of Intent"
        "Master Service Agreement" -> "Master Service Agreement"
    """
    if not agreement_type:
        return agreement_type
    # Split on " and " — primary type is the first segment
    primary = agreement_type.split(" and ")[0].strip()
    return primary or agreement_type


def add_cover(doc, agreement_type, subject=None, date_str=None,
              parties=None, party_labels=None, formality="non_binding",
              reference=None, version=None, classification="Confidential",
              cover_title=None):
    """IB-standard cover page with structured hierarchy.

    Rendering order:
      1. Agreement type (28pt bold) — e.g. "Letter of Intent"
         (auto-shortened from compound names; override via cover_title)
      2. Subject / deal description (14pt slate) — e.g. "for AI Infrastructure Distribution"
      3. Date (11pt, document-level)
      4. Party blocks — legal name, address, registration (binding only)
      5. Metadata — reference, version, classification
    """
    # Push title down from logo
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(200)
    spacer.paragraph_format.space_after = Pt(0)

    # 1. Agreement type
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(8)
    tp.paragraph_format.space_after = Pt(4)
    displayed_title = cover_title if cover_title else _cover_title(agreement_type)
    _run(tp, displayed_title, size=Pt(28), color=SLATE_900, bold=True)

    # 2. Subject / deal description
    if subject:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(2)
        sp.paragraph_format.space_after = Pt(0)
        _run(sp, subject, size=Pt(14), color=SLATE_800)

    # 3. Date (document-level, below title block)
    if date_str:
        dp = doc.add_paragraph()
        dp.paragraph_format.space_before = Pt(16)
        dp.paragraph_format.space_after = Pt(0)
        _run(dp, date_str, size=Pt(11), color=SLATE_900)

    # 4. Party blocks
    if parties:
        # Determine labels
        if party_labels is None:
            if formality == "binding":
                party_labels = ["By and between:"] + ["And:"] * (len(parties) - 1)
            else:
                party_labels = ["Between:"] + ["And:"] * (len(parties) - 1)
        # Pad labels if fewer than parties
        while len(party_labels) < len(parties):
            party_labels.append("And:")

        ps = doc.add_paragraph()
        ps.paragraph_format.space_before = Pt(40)
        ps.paragraph_format.space_after = Pt(0)

        for i, party in enumerate(parties):
            # Party label
            lp = doc.add_paragraph()
            lp.paragraph_format.space_before = Pt(12) if i > 0 else Pt(0)
            lp.paragraph_format.space_after = Pt(2)
            _run(lp, party_labels[i], size=Pt(10), color=COBALT)

            # Legal name
            np = doc.add_paragraph()
            np.paragraph_format.space_before = Pt(0)
            np.paragraph_format.space_after = Pt(0)
            _run(np, party.legal_name, size=Pt(11), color=SLATE_900, bold=True)

            # Address (skip empty — prevents blank paragraph for investor-style covers)
            if party.address and party.address.strip():
                ap = doc.add_paragraph()
                ap.paragraph_format.space_before = Pt(0)
                ap.paragraph_format.space_after = Pt(0)
                _run(ap, party.address, size=Pt(10), color=SLATE_900)

            # Registration (binding only)
            if formality == "binding" and party.registration_type and party.registration_number:
                rp = doc.add_paragraph()
                rp.paragraph_format.space_before = Pt(0)
                rp.paragraph_format.space_after = Pt(0)
                _run(rp, f"{party.registration_type}: ", size=Pt(10), color=SLATE_900, bold=True)
                _run(rp, party.registration_number, size=Pt(10), color=SLATE_900)

    # 5. Metadata
    meta_items = []
    if reference:
        meta_items.append(("Reference", reference))
    if version is not None:
        meta_items.append(("Version", f"v{version}"))
    if classification:
        meta_items.append(("Classification", classification))

    if meta_items:
        ms = doc.add_paragraph()
        ms.paragraph_format.space_before = Pt(40)
        ms.paragraph_format.space_after = Pt(0)
        for key, val in meta_items:
            mp = doc.add_paragraph()
            mp.paragraph_format.space_before = Pt(2)
            mp.paragraph_format.space_after = Pt(2)
            _run(mp, f"{key}:  ", size=Pt(10), color=COBALT, bold=True)
            _run(mp, str(val), size=Pt(10), color=SLATE_900)

    doc.add_page_break()


def add_section(doc, title, guidance, level=2):
    """Section heading + italic guidance text."""
    sizes = {1: Pt(24), 2: Pt(18), 3: Pt(14), 4: Pt(12)}
    hp = doc.add_paragraph()
    hp.paragraph_format.space_before = _SP['h1_before'] if level <= 2 else _SP['h3_before']
    hp.paragraph_format.space_after = _SP['heading_after']
    hp.paragraph_format.keep_with_next = True
    _run(hp, title, size=sizes.get(level, Pt(12)), color=SLATE_900, bold=True)

    gp = doc.add_paragraph()
    gp.paragraph_format.space_before = Pt(0)
    gp.paragraph_format.space_after = _SP['body_after']
    gp.paragraph_format.widow_control = True
    _run(gp, guidance, size=Pt(11), color=SLATE, italic=True)


def _compute_col_widths(headers, rows, avail_emu):
    """Compute column widths: header-first, then data-weighted surplus.

    Priority order:
    1. Every column fits its header text on one line (non-negotiable)
    2. Remaining space distributed proportionally by data volume
    3. Columns with trivial data (≤ header length) don't get surplus
    4. 2-column key-value layout uses 25/75 split
    """
    ncols = len(headers)

    # Measure display lengths (strip markdown markers)
    header_lens = [_display_len(h) for h in headers]
    max_data_lens = []
    for j in range(ncols):
        max_len = 0
        for row in rows:
            if j < len(row):
                max_len = max(max_len, _display_len(str(row[j])))
        max_data_lens.append(max_len)

    # 2-column key-value heuristic
    if ncols == 2 and max_data_lens[0] < 30 and max_data_lens[1] > 60:
        return [int(avail_emu * 0.25), int(avail_emu * 0.75)]

    # Step 1: base width = enough to fit header on one line + padding
    base_widths = []
    for j in range(ncols):
        header_w = (header_lens[j] + 2) * _CHAR_WIDTH_EMU + _CELL_PAD_EMU
        base_widths.append(max(header_w, _MIN_COL_EMU))

    base_total = sum(base_widths)

    # If headers alone exceed available space, scale down proportionally
    if base_total > avail_emu:
        scale = avail_emu / base_total
        widths = [max(int(w * scale), _MIN_COL_EMU) for w in base_widths]
        drift = avail_emu - sum(widths)
        if drift != 0:
            widths[max(range(ncols), key=lambda j: widths[j])] += drift
        return widths

    # Step 2: distribute surplus by data volume
    surplus = avail_emu - base_total
    # Weight = how much MORE data each column has beyond its header width
    # Columns where data fits in header width get zero surplus
    data_needs = []
    for j in range(ncols):
        data_chars = min(max_data_lens[j], 80)  # cap at 80 to prevent domination
        header_chars = header_lens[j] + 2
        need = max(0, data_chars - header_chars)
        data_needs.append(need)

    total_need = sum(data_needs)
    if total_need > 0 and surplus > 0:
        widths = [base_widths[j] + int(surplus * data_needs[j] / total_need)
                  for j in range(ncols)]
    else:
        # All data fits in headers — distribute surplus evenly
        per_col = surplus // ncols
        widths = [base_widths[j] + per_col for j in range(ncols)]

    # Correct rounding drift
    drift = avail_emu - sum(widths)
    if drift != 0:
        widths[max(range(ncols), key=lambda j: widths[j])] += drift

    return widths


def _insert_element(parent, tag_name, after_tags=()):
    """Insert a new child element respecting OOXML schema ordering.

    Finds the position after the last occurrence of any element in after_tags,
    or appends at end if none found. Returns the new element.
    """
    new_elem = etree.Element(qn(tag_name))
    insert_idx = 0
    for idx, child in enumerate(parent):
        for at in after_tags:
            if child.tag == qn(at):
                insert_idx = idx + 1
    if insert_idx < len(parent):
        parent.insert(insert_idx, new_elem)
    else:
        parent.append(new_elem)
    return new_elem


# ---------------------------------------------------------------------------
# CUSTOM NUMBERING (alphabetic + roman lists)
# ---------------------------------------------------------------------------

_ALPHA_ABSTRACT_ID = 100
_ROMAN_ABSTRACT_ID = 101
_DECIMAL_ABSTRACT_ID = 102
_BULLET_ABSTRACT_ID = 103
_HEADING_ABSTRACT_ID = 104


def _setup_custom_numbering(doc):
    """Create custom abstractNum definitions for all list types.

    Called once at the start of md_to_docx(). Adds abstractNum definitions:
      100 = lowerLetter (a), (b), (c)
      101 = lowerRoman  (i), (ii), (iii)
      102 = decimal      1. 2. 3.
      103 = bullet        bullet character

    IDs start at 100 to avoid conflict with the default template's 0-9.
    Each list block gets its own w:num instance via _new_list_instance() so
    numbering restarts correctly.

    OOXML ordering: w:abstractNum elements must come BEFORE all w:num elements.
    """
    try:
        numbering_part = doc.part.numbering_part
    except Exception:
        # No numbering part — create one by adding a dummy list paragraph
        dummy = doc.add_paragraph(style='List Bullet')
        dummy._element.getparent().remove(dummy._element)
        try:
            numbering_part = doc.part.numbering_part
        except Exception:
            return  # Graceful fallback — lists will use manual text

    numbering_el = numbering_part.element

    _defs = [
        (_ALPHA_ABSTRACT_ID, 'lowerLetter', '(%1)', 360, None),
        (_ROMAN_ABSTRACT_ID, 'lowerRoman', '(%1)', 480, None),
        (_DECIMAL_ABSTRACT_ID, 'decimal', '%1.', 360, None),
        # Bullet uses Symbol font for the bullet character (same as template)
        (_BULLET_ABSTRACT_ID, 'bullet', '\uf0b7', 360, ('Symbol', 'Symbol')),
    ]

    for abs_id, fmt, lvl_text, hanging, bullet_font in _defs:
        # Skip if already exists
        existing = numbering_el.findall(qn('w:abstractNum'))
        if any(e.get(qn('w:abstractNumId')) == str(abs_id) for e in existing):
            continue

        abstractNum = etree.Element(qn('w:abstractNum'))
        abstractNum.set(qn('w:abstractNumId'), str(abs_id))

        # Required schema elements (Word ignores abstractNums without these)
        nsid = etree.SubElement(abstractNum, qn('w:nsid'))
        nsid.set(qn('w:val'), f'{0xDE000000 + abs_id:08X}')
        multiLevel = etree.SubElement(abstractNum, qn('w:multiLevelType'))
        multiLevel.set(qn('w:val'), 'singleLevel')
        tmpl = etree.SubElement(abstractNum, qn('w:tmpl'))
        tmpl.set(qn('w:val'), f'{0xDE100000 + abs_id:08X}')

        lvl = etree.SubElement(abstractNum, qn('w:lvl'))
        lvl.set(qn('w:ilvl'), '0')

        start = etree.SubElement(lvl, qn('w:start'))
        start.set(qn('w:val'), '1')

        numFmt = etree.SubElement(lvl, qn('w:numFmt'))
        numFmt.set(qn('w:val'), fmt)

        lvlText_el = etree.SubElement(lvl, qn('w:lvlText'))
        lvlText_el.set(qn('w:val'), lvl_text)

        lvlJc = etree.SubElement(lvl, qn('w:lvlJc'))
        lvlJc.set(qn('w:val'), 'left')

        pPr = etree.SubElement(lvl, qn('w:pPr'))
        tab_stops = etree.SubElement(pPr, qn('w:tabs'))
        tab = etree.SubElement(tab_stops, qn('w:tab'))
        tab.set(qn('w:val'), 'num')
        tab.set(qn('w:pos'), str(hanging))
        ind = etree.SubElement(pPr, qn('w:ind'))
        ind.set(qn('w:left'), str(hanging))
        ind.set(qn('w:hanging'), str(hanging))

        # Bullet character needs font specification
        if bullet_font:
            rPr = etree.SubElement(lvl, qn('w:rPr'))
            rFonts = etree.SubElement(rPr, qn('w:rFonts'))
            rFonts.set(qn('w:ascii'), bullet_font[0])
            rFonts.set(qn('w:hAnsi'), bullet_font[1])
            rFonts.set(qn('w:hint'), 'default')

        # Insert before first w:num (OOXML schema order)
        first_num = numbering_el.find(qn('w:num'))
        if first_num is not None:
            first_num.addprevious(abstractNum)
        else:
            numbering_el.append(abstractNum)

    # --- Multilevel heading numbering (104): 1. / 1.1 / 1.1.1 ---
    existing = numbering_el.findall(qn('w:abstractNum'))
    if not any(e.get(qn('w:abstractNumId')) == str(_HEADING_ABSTRACT_ID) for e in existing):
        abstractNum = etree.Element(qn('w:abstractNum'))
        abstractNum.set(qn('w:abstractNumId'), str(_HEADING_ABSTRACT_ID))

        nsid = etree.SubElement(abstractNum, qn('w:nsid'))
        nsid.set(qn('w:val'), f'{0xDE000000 + _HEADING_ABSTRACT_ID:08X}')
        multiLevel = etree.SubElement(abstractNum, qn('w:multiLevelType'))
        multiLevel.set(qn('w:val'), 'multilevel')
        tmpl = etree.SubElement(abstractNum, qn('w:tmpl'))
        tmpl.set(qn('w:val'), f'{0xDE100000 + _HEADING_ABSTRACT_ID:08X}')

        # Level 0: "1." / Level 1: "1.1" / Level 2: "1.1.1"
        # No indent — headings are flush left with number followed by tab to text
        heading_levels = [
            (0, '%1.', 432, 432),       # "1." tab to text
            (1, '%1.%2', 576, 576),     # "1.1" slightly wider for two-part
            (2, '%1.%2.%3', 720, 720),  # "1.1.1" wider for three-part
        ]
        for ilvl, lvl_text, left, hanging in heading_levels:
            lvl = etree.SubElement(abstractNum, qn('w:lvl'))
            lvl.set(qn('w:ilvl'), str(ilvl))
            s = etree.SubElement(lvl, qn('w:start'))
            s.set(qn('w:val'), '1')
            nf = etree.SubElement(lvl, qn('w:numFmt'))
            nf.set(qn('w:val'), 'decimal')
            lt = etree.SubElement(lvl, qn('w:lvlText'))
            lt.set(qn('w:val'), lvl_text)
            jc = etree.SubElement(lvl, qn('w:lvlJc'))
            jc.set(qn('w:val'), 'left')
            pp = etree.SubElement(lvl, qn('w:pPr'))
            tabs = etree.SubElement(pp, qn('w:tabs'))
            tb = etree.SubElement(tabs, qn('w:tab'))
            tb.set(qn('w:val'), 'num')
            tb.set(qn('w:pos'), str(left))
            ind = etree.SubElement(pp, qn('w:ind'))
            ind.set(qn('w:left'), str(left))
            ind.set(qn('w:hanging'), str(hanging))
            # No w:rPr — Word inherits font/size/bold from paragraph formatting

        first_num = numbering_el.find(qn('w:num'))
        if first_num is not None:
            first_num.addprevious(abstractNum)
        else:
            numbering_el.append(abstractNum)


def _new_list_instance(doc, abstract_id):
    """Create a new w:num instance pointing to abstract_id with startOverride=1.

    Each separate list block gets its own numId so numbering restarts.
    Returns the numId integer.
    """
    try:
        numbering_el = doc.part.numbering_part.element
    except Exception:
        return None

    # Find max existing numId
    max_id = 0
    for num_el in numbering_el.findall(qn('w:num')):
        nid = num_el.get(qn('w:numId'))
        if nid:
            max_id = max(max_id, int(nid))
    new_id = max_id + 1

    num = etree.SubElement(numbering_el, qn('w:num'))
    num.set(qn('w:numId'), str(new_id))

    abstractNumId = etree.SubElement(num, qn('w:abstractNumId'))
    abstractNumId.set(qn('w:val'), str(abstract_id))

    # Restart numbering from 1
    lvlOverride = etree.SubElement(num, qn('w:lvlOverride'))
    lvlOverride.set(qn('w:ilvl'), '0')
    startOverride = etree.SubElement(lvlOverride, qn('w:startOverride'))
    startOverride.set(qn('w:val'), '1')

    return new_id


def _apply_numbering(para, num_id, ilvl=0):
    """Set w:numPr on a paragraph in correct OOXML schema position.

    w:numPr must come AFTER widowControl and BEFORE suppressLineNumbers in w:pPr.
    Using SubElement would append at end — out of schema order if spacing/ind already set.
    """
    pPr = para._element.get_or_add_pPr()
    # Remove existing numPr if any
    existing = pPr.find(qn('w:numPr'))
    if existing is not None:
        pPr.remove(existing)

    numPr = _insert_element(pPr, 'w:numPr',
                            after_tags=('w:pStyle', 'w:keepNext', 'w:keepLines',
                                        'w:pageBreakBefore', 'w:framePr', 'w:widowControl'))
    ilvl_el = etree.SubElement(numPr, qn('w:ilvl'))
    ilvl_el.set(qn('w:val'), str(ilvl))
    numId_el = etree.SubElement(numPr, qn('w:numId'))
    numId_el.set(qn('w:val'), str(num_id))


def _format_table(table, headers=None, rows=None):
    """IB-grade table formatting: Cobalt header, alternating rows, Slate borders.

    Uses lxml.etree + docx.oxml.ns.qn to set cell properties. Respects OOXML
    element ordering (tcBorders before shd; tblW before tblLayout before tblBorders).
    This is NOT OxmlElement or raw XML string parsing — it works through the
    same lxml tree that python-docx builds internally and is safe on Word for Mac.
    """
    ncols = len(table.columns)
    avail_emu = int(Mm(165))     # page 210mm - 25mm left - 20mm right (in EMU for col widths)
    avail_twips = int(165 / 25.4 * 1440)  # same width in twips (dxa) for tblW

    # Set total table width via tblPr to prevent overflow
    # OOXML tblPr order: tblStyle, tblpPr, tblOverlap, bidiVisual,
    #   tblStyleRowBandSize, tblStyleColBandSize, tblW, jc, tblCellSpacing,
    #   tblInd, tblBorders, shd, tblLayout, tblCellMar, tblLook, ...
    tbl_elem = table._tbl
    tblPr = tbl_elem.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = etree.SubElement(tbl_elem, qn('w:tblPr'))
        tbl_elem.insert(0, tblPr)

    for tag in ('w:tblW', 'w:tblLayout'):
        existing = tblPr.find(qn(tag))
        if existing is not None:
            tblPr.remove(existing)

    tblW = _insert_element(tblPr, 'w:tblW',
                           after_tags=('w:tblStyle', 'w:tblpPr', 'w:tblOverlap',
                                       'w:bidiVisual', 'w:tblStyleRowBandSize',
                                       'w:tblStyleColBandSize'))
    tblW.set(qn('w:w'), str(avail_twips))
    tblW.set(qn('w:type'), 'dxa')

    tblLayout = _insert_element(tblPr, 'w:tblLayout',
                                after_tags=('w:tblStyle', 'w:tblpPr', 'w:tblOverlap',
                                            'w:bidiVisual', 'w:tblStyleRowBandSize',
                                            'w:tblStyleColBandSize', 'w:tblW', 'w:jc',
                                            'w:tblCellSpacing', 'w:tblInd', 'w:tblBorders',
                                            'w:shd'))
    tblLayout.set(qn('w:type'), 'fixed')

    # Proportional column widths (EMU for python-docx col.width)
    if headers and rows:
        col_widths = _compute_col_widths(headers, rows, avail_emu)
    else:
        col_width = avail_emu // ncols
        col_widths = [col_width] * ncols
    for j, col in enumerate(table.columns):
        col.width = col_widths[j] if j < len(col_widths) else col_widths[-1]

    # Table header row repeat — repeats header on page break for tables >5 rows
    if len(table.rows) > 5:
        tr = table.rows[0]._tr
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = etree.SubElement(tr, qn('w:trPr'))
            tr.insert(0, trPr)
        if trPr.find(qn('w:tblHeader')) is None:
            etree.SubElement(trPr, qn('w:tblHeader'))

    # OOXML tcPr order: cnfStyle, tcW, gridSpan, hMerge, vMerge,
    #   tcBorders, shd, noWrap, tcMar, textDirection, ...
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = etree.SubElement(tc, qn('w:tcPr'))
                tc.insert(0, tcPr)

            for tag in ('w:tcBorders', 'w:shd'):
                existing = tcPr.find(qn(tag))
                if existing is not None:
                    tcPr.remove(existing)

            # tcBorders FIRST (before shd in OOXML schema)
            tcBorders = _insert_element(tcPr, 'w:tcBorders',
                                        after_tags=('w:cnfStyle', 'w:tcW', 'w:gridSpan',
                                                    'w:hMerge', 'w:vMerge'))
            for edge in ('top', 'left', 'bottom', 'right'):
                b = etree.SubElement(tcBorders, qn('w:{}'.format(edge)))
                b.set(qn('w:val'), 'single')
                b.set(qn('w:sz'), '4')       # 4 eighth-points = 0.5pt
                b.set(qn('w:space'), '0')
                b.set(qn('w:color'), SLATE_300_HEX)

            # shd AFTER tcBorders
            shd = _insert_element(tcPr, 'w:shd',
                                  after_tags=('w:cnfStyle', 'w:tcW', 'w:gridSpan',
                                              'w:hMerge', 'w:vMerge', 'w:tcBorders'))
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            if i == 0:
                shd.set(qn('w:fill'), COBALT_HEX)
            elif i % 2 == 0:
                shd.set(qn('w:fill'), SLATE_50_HEX)
            else:
                shd.set(qn('w:fill'), 'FFFFFF')

            # Cell padding for breathing room
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)


def add_table(doc, headers, rows):
    """Branded table: Cobalt header, proportional columns, inline formatting."""
    # Pre-table spacing (invisible spacer — contributes gap via space_after)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = _SP['table_before']
    spacer.paragraph_format.line_spacing = Pt(1)
    r = spacer.add_run()
    r.font.size = Pt(1)

    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers), style="Table Grid")
    for j, h in enumerate(headers):
        cell = tbl.cell(0, j)
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _run(cell.paragraphs[0], h, size=Pt(10), color=WHITE, bold=True)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.cell(i + 1, j)
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            _add_inline(cell.paragraphs[0], str(val), size=Pt(10), color=SLATE_800)
    _format_table(tbl, headers, rows)

    # Post-table spacing (invisible spacer — contributes gap via space_before)
    spacer2 = doc.add_paragraph()
    spacer2.paragraph_format.space_before = _SP['table_after']
    spacer2.paragraph_format.space_after = Pt(0)
    spacer2.paragraph_format.line_spacing = Pt(1)
    r2 = spacer2.add_run()
    r2.font.size = Pt(1)


# ---------------------------------------------------------------------------
# PROFILES
# ---------------------------------------------------------------------------

def profile_letter(title="", date_str="", entity="ag", **kw):
    ent = ENTITY_FOOTERS.get(entity, ENTITY_FOOTERS["ag"])
    doc = new_doc(diff_first=True)
    setup_first_page_header(doc.sections[0])
    setup_cont_header(doc.sections[0], title=title or "Letter")
    setup_first_footer(doc.sections[0], entity=entity)
    setup_cont_footer(doc.sections[0], entity=entity)

    # Return address
    rp = doc.add_paragraph()
    rp.paragraph_format.space_before = Pt(24)
    rp.paragraph_format.space_after = Pt(2)
    rp.paragraph_format.widow_control = True
    r = _run(rp, ent["return_address"], size=Pt(7), color=SLATE)
    r.underline = True

    # Recipient
    for line in ["[Recipient Name]", "[Company Name]", "[Street Address]",
                 "[Postcode] [City]", "[Country]"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.widow_control = True
        _run(p, line, size=Pt(11))

    # Date
    dp = doc.add_paragraph()
    dp.paragraph_format.space_before = Pt(16)
    dp.paragraph_format.widow_control = True
    dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(dp, f"Zug, {date_str}", size=Pt(10))

    # Subject
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(16)
    sp.paragraph_format.space_after = Pt(12)
    sp.paragraph_format.widow_control = True
    _run(sp, "[Subject Line]", size=Pt(11), color=SLATE_900, bold=True)

    # Salutation + body
    sal = doc.add_paragraph()
    sal.paragraph_format.widow_control = True
    _run(sal, "Dear [Mr./Ms. Last Name],", size=Pt(11))

    bp = doc.add_paragraph()
    bp.paragraph_format.line_spacing = Pt(16.5)
    bp.paragraph_format.widow_control = True
    _run(bp, "[Letter body text. DIN 5008 Form B compliant for windowed envelopes.]", size=Pt(11))

    # Closing
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(18)
    cp.paragraph_format.widow_control = True
    _run(cp, "Kind regards,", size=Pt(11))

    sig = doc.add_paragraph()
    sig.paragraph_format.space_before = Pt(36)
    sig.paragraph_format.widow_control = True
    _run(sig, "[Name]", size=Pt(11), bold=True)
    p_title = doc.add_paragraph()
    p_title.paragraph_format.widow_control = True
    _run(p_title, "[Title]", size=Pt(11), color=SLATE)
    p_entity = doc.add_paragraph()
    p_entity.paragraph_format.widow_control = True
    _run(p_entity, ent["legal_name"], size=Pt(11), color=SLATE)

    doc.core_properties.title = title or "Letter"
    doc.core_properties.author = "Digital Energy"
    return doc


def profile_agreement(agreement_type="[Agreement Type]", subject=None,
                      client="[Counterparty]", client_address=None,
                      client_reg_type=None, client_reg_number=None,
                      date_str="", version=1, formality=None,
                      entity="ag", title=None, cover_title=None, reference=None, **kw):
    # Backward compat: --title maps to agreement_type
    if title and agreement_type == "[Agreement Type]":
        agreement_type = title

    # M2: refuse to render an agreement with placeholder data.
    # Import lazily to avoid a circular import at module load.
    from validators import validate_agreement_inputs
    _resolved_formality = formality if formality is not None else _detect_formality(agreement_type)
    validate_agreement_inputs(
        agreement_type=agreement_type,
        client=client,
        client_address=client_address,
        formality=_resolved_formality,
        client_reg_type=client_reg_type,
        client_reg_number=client_reg_number,
    )

    doc = new_doc(diff_first=True)
    setup_first_page_header(doc.sections[0])
    setup_cont_header(doc.sections[0], title=_cover_title(agreement_type))
    setup_first_footer(doc.sections[0], classification="Confidential", entity=entity)
    setup_cont_footer(doc.sections[0], classification="Confidential", entity=entity)

    # Auto-detect formality from agreement type
    if formality is None:
        formality = _detect_formality(agreement_type)

    # Build party list
    de_party = DE_ENTITIES.get(entity, DE_ENTITIES["ag"])
    counterparty = Party(
        legal_name=client,
        address=client_address or "[Address]",
        registration_type=client_reg_type,
        registration_number=client_reg_number,
    )

    add_cover(doc,
              agreement_type=agreement_type,
              subject=subject,
              date_str=date_str,
              parties=[de_party, counterparty],
              formality=formality,
              reference=reference or "[DE-AGR-YYYY-NNN]",
              version=version,
              classification="Confidential",
              cover_title=cover_title)

    # Document properties
    doc.core_properties.title = agreement_type
    doc.core_properties.author = "Digital Energy"
    doc.core_properties.subject = subject or ""

    gp = doc.add_paragraph()
    gp.paragraph_format.line_spacing = Pt(16.5)
    _run(gp, "[Agreement content begins here. Replace with agreement body.]",
         size=Pt(11), color=SLATE, italic=True)

    return doc


SEED_SECTIONS = [
    ("1. Company Overview",
     "One paragraph: what DE does, where, for whom. Founding date, HQ, one-sentence positioning."),
    ("2. Problem",
     "Grid scarcity (60 GW queue vs 20 GW peak), stranded greenhouse heat (30-50% gas costs), no sovereign EU AI compute. Use 2-3 sourced data points."),
    ("3. Solution",
     "How DECs solve it. 'One Watt, Three Jobs': compute + heat + grid flexibility. Co-location on existing grid connections, 97% energy recycling."),
    ("4. Business Model",
     "Four revenue streams: colocation, heat sales, BESS/flexibility, development fees. DevCo recycling model. Include unit economics table."),
    ("5. Market Opportunity",
     "TAM/SAM/SOM. European AI compute demand, Dutch grid-constrained supply, greenhouse heat market (EUR 2.8B energy spend)."),
    ("6. Traction",
     "14 sites, 600 MW grid infrastructure, partnerships (Lenovo, Nokia, Ampower), Fonti 4 months from FID, PowerGrow 12 months."),
    ("7. Team",
     "Core team with infrastructure, energy, and technology credentials. Board/advisors if notable."),
    ("8. The Ask",
     "Round size, instrument (SAFE/note), valuation cap, use of proceeds breakdown, timeline to close."),
]


def profile_seed_memo(client="[Investor Name]", date_str="", version=1, entity="ag", **kw):
    de_party = DE_ENTITIES.get(entity, DE_ENTITIES["ag"])
    doc = new_doc(diff_first=True)
    setup_first_page_header(doc.sections[0])
    setup_cont_header(doc.sections[0], title="Seed Investment Memo")
    setup_first_footer(doc.sections[0], classification="Confidential", entity=entity)
    setup_cont_footer(doc.sections[0], classification="Confidential", entity=entity)

    add_cover(doc,
              agreement_type="Seed Investment Memo",
              subject="Sovereign AI Infrastructure for Europe",
              date_str=date_str,
              parties=[
                  de_party,
                  Party(legal_name=client, address=""),
              ],
              party_labels=["Prepared by:", "Prepared for:"],
              formality="non_binding",
              version=version,
              classification="Confidential")

    for sect_title, guidance in SEED_SECTIONS:
        add_section(doc, sect_title, guidance)

    doc.core_properties.title = "Seed Investment Memo"
    doc.core_properties.author = "Digital Energy"
    return doc


IM_SECTIONS = [
    ("1. Executive Overview",
     "One-page standalone summary: investment thesis, key metrics, round details, why now."),
    ("2. Asset Overview",
     "DEC architecture: modular 4.2 MW, liquid cooling (40-140 kW/rack, PUE 1.2), heat integration, BESS, Super-Factory concept."),
    ("3. Market Dynamics",
     "EU AI compute demand, Dutch grid congestion (60 GW queue), Wcw regulation, SDE++ subsidies, restwarmteplicht."),
    ("4. Business Model",
     "Four fee streams: colocation (EUR/kW/month), heat (EUR/MWh + CPI), BESS, development fees. DevCo/HoldCo recycling. AG structure."),
    ("5. Pipeline & Traction",
     "14 sites, 600 MW, 12M m\u00b2 greenhouse area. Pipeline table: site, location, MW, status, target FID. Lead projects: Fonti, PowerGrow."),
    ("6. Unit Economics",
     "Per-DEC: CAPEX/MW, revenue/MW, EBITDA margin, payback, IRR. Three scenarios (conservative/base/optimistic). Source: FM v3.51."),
    ("7. Risk Factors",
     "Grid/DSO, permitting, technology, market, regulatory, execution risks. Each with probability, impact, mitigant."),
    ("8. Financial Projections",
     "3-5 year projections: revenue, EBITDA, cash flow, CAPEX, net debt. Three scenarios. Key assumptions stated."),
    ("9. Capital Structure",
     "Current cap table, round terms, use of proceeds breakdown, pro forma post-raise."),
    ("10. Appendix",
     "Site profiles, FM v3.51 outputs, team bios, glossary (Wcw, SDE++, transportschaarste, omgevingsvergunning, recht van opstal)."),
]


def profile_investor_memo(client="[Investor Name]", date_str="", version=1, entity="ag", **kw):
    de_party = DE_ENTITIES.get(entity, DE_ENTITIES["ag"])
    doc = new_doc(diff_first=True)
    setup_first_page_header(doc.sections[0])
    setup_cont_header(doc.sections[0], title="Investment Memorandum")
    setup_first_footer(doc.sections[0], classification="Confidential", entity=entity)
    setup_cont_footer(doc.sections[0], classification="Confidential", entity=entity)

    add_cover(doc,
              agreement_type="Investment Memorandum",
              subject="European Sovereign AI Infrastructure",
              date_str=date_str,
              parties=[
                  de_party,
                  Party(legal_name=client, address=""),
              ],
              party_labels=["Prepared by:", "Prepared for:"],
              formality="non_binding",
              version=version,
              classification="Confidential")

    for sect_title, guidance in IM_SECTIONS:
        add_section(doc, sect_title, guidance)

    doc.core_properties.title = "Investment Memorandum"
    doc.core_properties.author = "Digital Energy"
    return doc


def profile_exec_summary(title="[Executive Summary]", date_str="", entity="ag", **kw):
    doc = new_doc(top=Mm(15), bottom=Mm(25))
    setup_cont_header(doc.sections[0], title=title)
    setup_cont_footer(doc.sections[0], entity=entity)

    # Title
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(4)
    tp.paragraph_format.space_after = Pt(2)
    _run(tp, title, size=Pt(24), color=SLATE_900, bold=True)

    _run(doc.add_paragraph(), date_str, size=Pt(10), color=SLATE)

    # Metrics placeholder
    add_table(doc,
              ["[Metric 1]", "[Metric 2]", "[Metric 3]"],
              [["[Value]", "[Value]", "[Value]"]])

    add_section(doc, "Context",
                "[Why this document exists. 2-3 sentences setting the scene.]")
    add_section(doc, "Key Findings",
                "[3-5 numbered findings, each 1-2 sentences. Lead with conclusion, then evidence.]")

    doc.add_page_break()

    add_section(doc, "Analysis",
                "[Supporting detail. Can include a summary table.]")
    add_section(doc, "Recommendation",
                "[What to do, by when, who owns it. 3-5 sentences max.]")
    add_section(doc, "Next Steps",
                "[3-5 numbered action items: action + owner + deadline.]")

    doc.core_properties.title = title or "Executive Summary"
    doc.core_properties.author = "Digital Energy"
    return doc


# ---------------------------------------------------------------------------
# MARKDOWN CONVERTER
# ---------------------------------------------------------------------------

def _detect_entity(title, md_text=""):
    """Auto-detect entity from title or document content.

    Returns "nl" if the document is clearly about Digital Energy Netherlands,
    otherwise "ag" (default).
    """
    combined = (title or "") + " " + (md_text[:500] if md_text else "")
    combined_upper = combined.upper()
    nl_signals = ("DENL", "NETHERLANDS B.V.", "NETHERLANDS BV",
                  "DIGITAL ENERGY NETHERLANDS", "KVK:", "KVKNUMMER",
                  "BESLOTEN VENNOOTSCHAP", "KUDELSTAART", "MIJNSHERENWEG")
    for signal in nl_signals:
        if signal in combined_upper:
            return "nl"
    return "ag"


def _extract_title_from_md(md_text):
    """Extract clean document title from markdown headings.

    Looks at H1 and H2. If H1 is an entity name (e.g. "DIGITAL ENERGY NETHERLANDS B.V."),
    uses H2 as the title. If H2 is also generic, falls back to H1.
    Also extracts H3 as subject if it looks like a subtitle.
    Returns (title, subject).
    """
    h1 = h2 = h3 = None
    for line in md_text.split('\n')[:30]:
        if line.startswith('# ') and not line.startswith('## '):
            h1 = line[2:].strip()
        elif line.startswith('## ') and not line.startswith('### ') and h2 is None:
            h2 = line[3:].strip()
            # Strip numbering prefix: "## 1. Parties" → "Parties" — but keep unnumbered
            h2_clean = re.sub(r'^\d+[A-Z]?\.\s+', '', h2)
            if h2_clean != h2:
                h2 = None  # First H2 was numbered section, not a title
                continue
        elif line.startswith('### ') and not line.startswith('#### ') and h3 is None and h2:
            h3_text = line[4:].strip()
            # Only use H3 as subject if it's not a numbered section
            if not re.match(r'^\d+', h3_text):
                h3 = h3_text

    entity_names = ('DIGITAL ENERGY', 'DE GROUP', 'DEG AG', 'DENL', 'DEC THERMAL')
    if h1 and any(en in h1.upper() for en in entity_names):
        # H1 is entity name — use H2 as title
        return (h2 or h1, h3)
    return (h1 or h2, h3)


def md_to_docx(md_text, title=None, client=None, date_str=None, cover=False,
               entity=None, subject=None, formality=None):
    """Convert markdown to branded docx.

    Args:
        title: Document title. If None, auto-extracted from markdown H1/H2.
        formality: "binding" or "non_binding". If None, auto-detects from
                   AGREEMENT_FORMALITY dict using title, defaulting to "non_binding".
        subject: Cover page subtitle (e.g. "Chairman, Horticulture Advisory Board").
    """
    # Auto-extract title and subject from markdown if not provided
    if title is None or (cover and subject is None):
        auto_title, auto_subject = _extract_title_from_md(md_text)
        if title is None:
            title = auto_title
        if subject is None and cover:
            subject = auto_subject

    if entity is None:
        entity = _detect_entity(title, md_text)
    has_cover = cover and title
    doc = new_doc(diff_first=has_cover)

    if has_cover:
        setup_first_page_header(doc.sections[0])
        setup_cont_header(doc.sections[0], title=_cover_title(title))
        setup_first_footer(doc.sections[0], entity=entity)
        setup_cont_footer(doc.sections[0], entity=entity)
        de_party = DE_ENTITIES.get(entity, DE_ENTITIES["ag"])
        parties = [de_party]
        if client:
            parties.append(Party(legal_name=client, address=""))
        if formality is None:
            formality = _detect_formality(title)
        add_cover(doc,
                  agreement_type=title,
                  subject=subject,
                  date_str=date_str,
                  parties=parties,
                  party_labels=["Prepared by:", "Prepared for:"] if client else None,
                  formality=formality)
    else:
        setup_cont_header(doc.sections[0], title=title or "")
        setup_cont_footer(doc.sections[0], entity=entity)

    # Document properties
    doc.core_properties.title = title or ""
    doc.core_properties.author = "Digital Energy"
    doc.core_properties.subject = subject or ""

    # Set up custom numbering definitions for lists + heading numbering
    _setup_custom_numbering(doc)
    heading_num_id = _new_list_instance(doc, _HEADING_ABSTRACT_ID)

    heading_sizes = {1: Pt(24), 2: Pt(18), 3: Pt(14), 4: Pt(12), 5: Pt(11), 6: Pt(11)}
    # Detect numbered headings: "## 1. Title", "### 1.1 Title", "#### 1.1.1 Title"
    # Detect numbered headings: "1. Title", "1.1 Title", "2A. Title", "3.2 Title"
    _numbered_heading_re = re.compile(r'^(\d+(?:[A-Z]|\.\d+)*)\.?\s+(.*)')
    # Map markdown heading level → ilvl for heading numbering
    # ## (level 2) → ilvl 0 (top-level "1."), ### (level 3) → ilvl 1 ("1.1" or "2A"), #### → ilvl 2
    _heading_level_to_ilvl = {2: 0, 3: 1, 4: 2}

    # Regex patterns for list detection
    _alpha_re = re.compile(r'^[\s]*\([a-z]\)\s+')
    _roman_re = re.compile(r'^[\s]*\((i{1,3}|iv|v|vi{0,3}|ix|x)\)\s+')

    # Signature section mode
    _SIG_HEADING_RE = re.compile(
        r'(signatures?|counter[- ]?signatures?|execution|signature\s*page)',
        re.IGNORECASE
    )
    _SIG_PARTY_RE = re.compile(r'^\*\*[^*]+\*\*\s*$')
    in_signature_section = False

    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        # Heading
        m = re.match(r'^(#{1,6})\s+(.+)$', line)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip()

            # Check if this heading enters/exits signature section
            if _SIG_HEADING_RE.search(heading_text):
                in_signature_section = True
            elif in_signature_section:
                in_signature_section = False  # Non-signature heading exits mode

            p = doc.add_paragraph()
            p.paragraph_format.space_before = _SP['h1_before'] if level <= 2 else _SP['h3_before']
            p.paragraph_format.space_after = _SP['heading_after']
            p.paragraph_format.keep_with_next = True

            # Native heading numbering: "## 1. Title" → strip number, apply w:numPr
            heading_size = heading_sizes.get(level, Pt(11))
            ilvl = _heading_level_to_ilvl.get(level)
            nm = _numbered_heading_re.match(heading_text)
            if nm and ilvl is not None and heading_num_id is not None:
                heading_text = nm.group(2).strip()  # strip "1." prefix
                _apply_numbering(p, heading_num_id, ilvl=ilvl)
                # Set paragraph-level rPr so numbering character inherits
                # font/size/bold (Word ignores run-level formatting for numbers)
                pPr = p._element.get_or_add_pPr()
                pRpr = etree.SubElement(pPr, qn('w:rPr'))
                etree.SubElement(pRpr, qn('w:b'))
                rFonts = etree.SubElement(pRpr, qn('w:rFonts'))
                rFonts.set(qn('w:ascii'), FONT)
                rFonts.set(qn('w:hAnsi'), FONT)
                rFonts.set(qn('w:cs'), 'Arial')
                sz = etree.SubElement(pRpr, qn('w:sz'))
                sz.set(qn('w:val'), str(int(heading_size / 6350)))  # EMU → half-points
                szCs = etree.SubElement(pRpr, qn('w:szCs'))
                szCs.set(qn('w:val'), str(int(heading_size / 6350)))
                # Color
                color_el = etree.SubElement(pRpr, qn('w:color'))
                color_el.set(qn('w:val'), '0F172A')  # SLATE_900

            _add_inline(p, heading_text, size=heading_size,
                        color=SLATE_900, bold=True)
            i += 1
            continue

        # HR
        if re.match(r'^[-*_]{3,}\s*$', line):
            i += 1
            continue

        # Unordered list — explicit w:numPr numbering (with blank-line look-ahead)
        if re.match(r'^[\s]*[-*+]\s+', line):
            bullet_num_id = _new_list_instance(doc, _BULLET_ABSTRACT_ID)
            first = True
            p = None
            while i < len(lines):
                if re.match(r'^[\s]*[-*+]\s+', lines[i]):
                    text = re.sub(r'^[\s]*[-*+]\s+', '', lines[i]).strip()
                    p = doc.add_paragraph()
                    if bullet_num_id is not None:
                        _apply_numbering(p, bullet_num_id)
                    else:
                        p.style = doc.styles['List Bullet']  # fallback
                    p.paragraph_format.space_before = _SP['list_first_before'] if first else _SP['list_gap']
                    p.paragraph_format.space_after = _SP['list_gap']
                    p.paragraph_format.widow_control = True
                    _add_inline(p, text)
                    first = False
                    i += 1
                elif not lines[i].strip():
                    # Blank line — look ahead for list continuation
                    j = i + 1
                    while j < len(lines) and not lines[j].strip():
                        j += 1
                    if j < len(lines) and re.match(r'^[\s]*[-*+]\s+', lines[j]):
                        i = j
                    else:
                        break
                else:
                    break
            if p:
                p.paragraph_format.space_after = _SP['list_last_after']
            continue

        # Ordered list — explicit w:numPr numbering (with blank-line look-ahead)
        if re.match(r'^[\s]*\d+[.)]\s+', line):
            decimal_num_id = _new_list_instance(doc, _DECIMAL_ABSTRACT_ID)
            first = True
            p = None
            while i < len(lines):
                if re.match(r'^[\s]*\d+[.)]\s+', lines[i]):
                    text = re.sub(r'^[\s]*\d+[.)]\s+', '', lines[i]).strip()
                    p = doc.add_paragraph()
                    if decimal_num_id is not None:
                        _apply_numbering(p, decimal_num_id)
                    else:
                        p.style = doc.styles['List Number']  # fallback
                    p.paragraph_format.space_before = _SP['list_first_before'] if first else _SP['list_gap']
                    p.paragraph_format.space_after = _SP['list_gap']
                    p.paragraph_format.widow_control = True
                    _add_inline(p, text)
                    first = False
                    i += 1
                elif not lines[i].strip():
                    j = i + 1
                    while j < len(lines) and not lines[j].strip():
                        j += 1
                    if j < len(lines) and re.match(r'^[\s]*\d+[.)]\s+', lines[j]):
                        i = j
                    else:
                        break
                else:
                    break
            if p:
                p.paragraph_format.space_after = _SP['list_last_after']
            continue

        # Table
        if '|' in line and i + 1 < len(lines) and re.match(r'^[\s]*\|?[\s]*[-:]', lines[i + 1]):
            headers = [c.strip() for c in line.strip().strip('|').split('|')]
            i += 2
            rows = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip():
                rows.append([c.strip() for c in lines[i].strip().strip('|').split('|')])
                i += 1
            add_table(doc, headers, rows)
            continue

        # Blockquote
        if line.startswith('>'):
            parts = []
            while i < len(lines) and lines[i].startswith('>'):
                parts.append(lines[i].lstrip('>').strip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Mm(10)
            p.paragraph_format.widow_control = True
            _run(p, ' '.join(parts), size=Pt(11), color=SLATE, italic=True)
            continue

        # Alphabetic list: (a), (b), (c) — native lowerLetter numbering
        if _alpha_re.match(line):
            # Disambiguate (i): if next item is (ii) → roman, if (j) → alpha
            abstract_id = _ALPHA_ABSTRACT_ID
            if re.match(r'^[\s]*\(i\)\s+', line):
                # Peek ahead for next list item
                j = i + 1
                while j < len(lines) and not lines[j].strip():
                    j += 1
                if j < len(lines) and re.match(r'^[\s]*\(ii\)\s+', lines[j]):
                    abstract_id = _ROMAN_ABSTRACT_ID

            num_id = _new_list_instance(doc, abstract_id)
            first = True
            p = None
            while i < len(lines):
                if _alpha_re.match(lines[i]) or _roman_re.match(lines[i]):
                    text = re.sub(r'^[\s]*\([a-z]+\)\s+', '', lines[i]).strip()
                    p = doc.add_paragraph()
                    if num_id is not None:
                        _apply_numbering(p, num_id)
                    else:
                        # Fallback: manual text prefix
                        m_item = re.match(r'^[\s]*(\([a-z]+\))\s+', lines[i])
                        if m_item:
                            text = m_item.group(1) + ' ' + text
                    p.paragraph_format.space_before = _SP['list_first_before'] if first else _SP['list_gap']
                    p.paragraph_format.space_after = _SP['list_gap']
                    p.paragraph_format.line_spacing = _SP['body_line']
                    p.paragraph_format.widow_control = True
                    _add_inline(p, text)
                    first = False
                    i += 1
                elif not lines[i].strip():
                    # Blank line — look ahead for list continuation
                    j = i + 1
                    while j < len(lines) and not lines[j].strip():
                        j += 1
                    if j < len(lines) and (_alpha_re.match(lines[j]) or _roman_re.match(lines[j])):
                        i = j  # skip blanks, continue same list
                    else:
                        break
                else:
                    break
            if p:
                p.paragraph_format.space_after = _SP['list_last_after']
            continue

        # Roman list: (i), (ii), (iii) — native lowerRoman numbering
        if _roman_re.match(line) and not _alpha_re.match(line):
            num_id = _new_list_instance(doc, _ROMAN_ABSTRACT_ID)
            first = True
            p = None
            while i < len(lines):
                if _roman_re.match(lines[i]):
                    text = re.sub(r'^[\s]*\([ivxlcdm]+\)\s+', '', lines[i]).strip()
                    p = doc.add_paragraph()
                    if num_id is not None:
                        _apply_numbering(p, num_id)
                    p.paragraph_format.space_before = _SP['list_first_before'] if first else _SP['list_gap']
                    p.paragraph_format.space_after = _SP['list_gap']
                    p.paragraph_format.line_spacing = _SP['body_line']
                    p.paragraph_format.widow_control = True
                    _add_inline(p, text)
                    first = False
                    i += 1
                elif not lines[i].strip():
                    j = i + 1
                    while j < len(lines) and not lines[j].strip():
                        j += 1
                    if j < len(lines) and _roman_re.match(lines[j]):
                        i = j
                    else:
                        break
                else:
                    break
            if p:
                p.paragraph_format.space_after = _SP['list_last_after']
            continue

        # Signature block detection
        # Mode 1: Section mode (in_signature_section flag set by heading)
        # Mode 2: Fallback line detection for documents without ## Signatures heading
        # Negative lookahead (?!\*) rejects **bold metadata** like **Grant Date:** from
        # falsely triggering the signature handler.
        sig_line_re = re.compile(
            r'^(?!\*)(_{3,}|(Name|Title|Signature|Date|Signed|By)\s*:)',
            re.IGNORECASE
        )

        # Party-name detector: bold-only line that is NOT a disclaimer/meta marker.
        # Prevents **DRAFT — ...** / **End of Agreement.** from being styled as party.
        _SIG_NON_PARTY_RE = re.compile(
            r'^\*\*(DRAFT|END|NOTE|WARNING|CONFIDENTIAL|PRIVILEGED|TBD|TODO)\b',
            re.IGNORECASE
        )

        def _is_sig_party_name(line_stripped):
            """True if line is a bold-only party name (not a disclaimer)."""
            if not _SIG_PARTY_RE.match(line_stripped):
                return False
            if _SIG_NON_PARTY_RE.match(line_stripped):
                return False
            return True

        # Signature fallback: only trigger on bold party names if NEXT non-blank
        # lines look like signature content (By:, Name:, Title:, ___, etc.)
        _sig_bold_with_context = False
        if _is_sig_party_name(line.strip()) and not in_signature_section:
            # Peek ahead for signature-like lines
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines) and sig_line_re.match(lines[j].strip()):
                _sig_bold_with_context = True

        if in_signature_section or sig_line_re.match(line) or _sig_bold_with_context:
            # Collect all lines in this signature region
            party_blocks = []
            current_block = []
            blank_count = 0
            is_first_party = True

            while i < len(lines):
                cl = lines[i]
                # Exit on heading (will be handled by heading handler which may exit sig mode)
                if cl.strip().startswith('#'):
                    break
                # Blank lines: track for party block separation
                if not cl.strip():
                    blank_count += 1
                    if blank_count >= 2 and current_block:
                        party_blocks.append(current_block)
                        current_block = []
                        blank_count = 0
                    i += 1
                    continue
                blank_count = 0
                current_block.append(cl)
                i += 1

            if current_block:
                party_blocks.append(current_block)

            # Split blocks at bold party names — separates intro text from first party
            split_blocks = []
            for block in party_blocks:
                current = []
                for sl in block:
                    if _is_sig_party_name(sl.strip()) and current:
                        split_blocks.append(current)
                        current = []
                    current.append(sl)
                if current:
                    split_blocks.append(current)
            party_blocks = split_blocks

            # Render party blocks — corporate-grade signature layout
            _SIG_FIELD_RE = re.compile(r'^(By|Name|Title|Signature|Date|Signed)\s*:', re.IGNORECASE)

            for block_idx, block in enumerate(party_blocks):
                # Skip HR lines (---) inside signature blocks
                block = [sl for sl in block if not re.match(r'^[-*_]{3,}\s*$', sl.strip())]
                if not block:
                    continue

                for line_idx, sl in enumerate(block):
                    sl_stripped = sl.strip()
                    p = doc.add_paragraph()
                    p.paragraph_format.keep_together = True
                    p.paragraph_format.widow_control = True
                    is_last_line = (line_idx == len(block) - 1)
                    is_last_block = (block_idx == len(party_blocks) - 1)
                    p.paragraph_format.keep_with_next = not (is_last_line and is_last_block)

                    is_party_name = _is_sig_party_name(sl_stripped)
                    is_field = _SIG_FIELD_RE.match(sl_stripped)

                    # Spacing: breathing room within block, clear break between parties
                    if line_idx == 0 and block_idx == 0:
                        p.paragraph_format.space_before = _SP['sig_section_before']
                    elif line_idx == 0:
                        p.paragraph_format.space_before = _SP['sig_party_gap']
                    else:
                        p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(2)

                    if is_party_name:
                        # Party name: 12pt bold, 48pt after = ~17mm for physical signature
                        p.paragraph_format.space_after = Pt(48)
                        _add_inline(p, sl_stripped, size=Pt(12), color=SLATE_900, bold=True)
                    elif is_field:
                        # Normalize signature lines to consistent width
                        fm = _SIG_FIELD_RE.match(sl_stripped)
                        field_label = fm.group(0)
                        field_value = sl_stripped[fm.end():].strip()
                        if field_value and re.match(r'^[_\s]+$', field_value):
                            field_value = '_' * 40
                        _run(p, field_label + ' ', size=Pt(11), color=SLATE_800)
                        if field_value:
                            _run(p, field_value, size=Pt(11), color=SLATE_800)
                    else:
                        _add_inline(p, sl_stripped)
            continue

        # Paragraph (with guards to avoid swallowing lists/tables)
        para_lines = []
        while (i < len(lines) and lines[i].strip()
               and not lines[i].startswith('#')
               and not re.match(r'^[-*_]{3,}\s*$', lines[i])
               and not re.match(r'^[\s]*[-*+]\s+', lines[i])
               and not re.match(r'^[\s]*\d+[.)]\s+', lines[i])
               and not re.match(r'^[\s]*\([a-z]\)\s+', lines[i])
               and not re.match(r'^[\s]*\((i{1,3}|iv|v|vi{0,3}|ix|x)\)\s+', lines[i])
               and not (lines[i].strip().startswith('|') and '|' in lines[i][1:])):
            para_lines.append(lines[i])
            i += 1
        if para_lines:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = _SP['body_line']
            p.paragraph_format.space_after = _SP['body_after']
            p.paragraph_format.widow_control = True
            _add_inline(p, ' '.join(para_lines))
            continue

        i += 1

    return doc


def _add_inline(para, text, size=Pt(11), color=SLATE_800, bold=False):
    """Add text with **bold**, *italic*, `code` inline formatting."""
    for part in re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            _run(para, part[2:-2], size=size, color=color, bold=True)
        elif part.startswith('*') and part.endswith('*'):
            _run(para, part[1:-1], size=size, color=color, italic=True)
        elif part.startswith('`') and part.endswith('`'):
            _run(para, part[1:-1], size=Pt(10), color=color, font="JetBrains Mono")
        else:
            _run(para, part, size=size, color=color, bold=bold)


# ---------------------------------------------------------------------------
# PDF CONVERSION (Microsoft Word via docx2pdf)
# ---------------------------------------------------------------------------

def docx_to_pdf(docx_path, pdf_path=None):
    """Convert .docx → .pdf via Microsoft Word (docx2pdf).

    Word is the canonical .docx renderer on the team's machines.
    On macOS this uses AppleScript automation; on Windows, COM.
    No fallback — if Word is unavailable, raises RuntimeError with
    install instructions.

    Returns the output PDF path on success.
    """
    docx_path = os.path.abspath(docx_path)
    if not os.path.exists(docx_path):
        raise FileNotFoundError(docx_path)
    if pdf_path is None:
        pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
    pdf_path = os.path.abspath(pdf_path)

    try:
        from docx2pdf import convert as _word_convert
    except ImportError:
        raise RuntimeError(
            "docx2pdf not installed. Run: pip install docx2pdf\n"
            "Also ensure Microsoft Word is installed (macOS: Office 365; Windows: Office)."
        )

    try:
        _word_convert(docx_path, pdf_path)
    except Exception as e:
        raise RuntimeError(
            f"Word conversion failed: {e}\n"
            "Verify Microsoft Word is installed and, on macOS, that AppleScript is permitted."
        )

    if not os.path.exists(pdf_path):
        raise RuntimeError("docx2pdf ran but produced no output file.")
    return pdf_path


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

PROFILES = {
    "letter": profile_letter,
    "agreement": profile_agreement,
    "seed_memo": profile_seed_memo,
    "investor_memo": profile_investor_memo,
    "exec_summary": profile_exec_summary,
}


def _run_validate(docx_path):
    """Run office_bridge validation if available."""
    try:
        from office_bridge import OfficeBridge
        bridge = OfficeBridge()
        result = bridge.validate(docx_path)
        print(f"Validate:  {result.strip()}")
    except ImportError:
        print("Validate:  skipped (office_bridge not available)", file=sys.stderr)
    except Exception as e:
        print(f"Validate:  FAILED — {e}", file=sys.stderr)


def main():
    parser = argparse.ArgumentParser(description="Generate branded DE documents")
    parser.add_argument("--profile", choices=list(PROFILES.keys()),
                       help="Document profile")
    parser.add_argument("--title", default=None,
                       help="DEPRECATED for agreement profile; use --agreement-type. Still valid for other profiles.")
    parser.add_argument("--agreement-type", default=None,
                       help="Agreement type name (e.g. 'Letter of Intent', 'Master Service Agreement')")
    parser.add_argument("--cover-title", default=None,
                       help="Override cover-page title (28pt). Defaults to auto-shortened agreement-type")
    parser.add_argument("--reference", default=None,
                       help="Document reference (e.g. 'DE-AGR-2026-001'). Defaults to placeholder")
    parser.add_argument("--subject", default=None,
                       help="Deal description sub-header (e.g. 'for AI Infrastructure Distribution')")
    parser.add_argument("--client", default=None)
    parser.add_argument("--client-address", default=None, help="Counterparty address")
    parser.add_argument("--client-reg-type", default=None, help="Registration type (KvK, CHE, EIN)")
    parser.add_argument("--client-reg-number", default=None, help="Registration number")
    parser.add_argument("--entity", choices=list(DE_ENTITIES.keys()), default="ag",
                       help="DE contracting entity (default: ag)")
    parser.add_argument("--formality", choices=["binding", "non_binding"], default=None,
                       help="Override formality auto-detection")
    parser.add_argument("--date", default=None, help="YYYY-MM-DD")
    parser.add_argument("--version", type=int, default=1)
    parser.add_argument("--output", default=None)
    parser.add_argument("--dotx", action="store_true")
    parser.add_argument("--pdf", action="store_true",
                       help="Also produce a PDF via Microsoft Word")
    parser.add_argument("--validate", action="store_true",
                       help="Validate .docx via office_bridge after generation")
    parser.add_argument("--strip-review", action="store_true",
                       help="Strip [REVIEW REQUIRED] markers from output (md mode)")
    # Markdown mode
    parser.add_argument("--md", default=None, help="Input markdown file")
    parser.add_argument("--cover", action="store_true", help="Add cover page to md output")

    args = parser.parse_args()

    if args.date:
        dt = datetime.strptime(args.date, "%Y-%m-%d").date()
    else:
        dt = date.today()
    date_str = dt.strftime("%d %B %Y").lstrip("0")  # Cross-platform (no %-d)

    os.makedirs(OUTPUT, exist_ok=True)

    # Markdown mode
    if args.md:
        if not os.path.exists(args.md):
            print(f"Error: {args.md} not found")
            sys.exit(1)
        with open(args.md, 'r') as f:
            md_text = f.read()
        if args.strip_review:
            md_text, n_stripped = re.subn(r'\[REVIEW REQUIRED[^\]]*\]', '', md_text)
            if n_stripped:
                print(f"Stripped:  {n_stripped} [REVIEW REQUIRED] marker(s)", file=sys.stderr)
        doc = md_to_docx(md_text, title=args.title, client=args.client,
                         date_str=date_str, cover=args.cover, entity=args.entity,
                         subject=args.subject, formality=args.formality)
        out = args.output or os.path.join(OUTPUT, "md_output.docx")
        save_doc(doc, out)
        print(f"Saved: {out}")
        if args.validate:
            _run_validate(out)
        if args.pdf:
            try:
                pdf = docx_to_pdf(out)
                print(f"PDF:   {pdf}")
            except RuntimeError as e:
                print(f"PDF conversion failed: {e}", file=sys.stderr)
                sys.exit(1)
        return

    # Profile mode
    if not args.profile:
        parser.error("--profile or --md required")

    fn = PROFILES[args.profile]
    kwargs = {"date_str": date_str, "version": args.version}

    # Agreement profile uses new structured flags
    if args.profile == "agreement":
        agreement_type = args.agreement_type or args.title or "[Agreement Type]"
        kwargs["agreement_type"] = agreement_type
        kwargs["subject"] = args.subject
        kwargs["entity"] = args.entity
        if getattr(args, "cover_title", None):
            kwargs["cover_title"] = args.cover_title
        if getattr(args, "reference", None):
            kwargs["reference"] = args.reference
        if args.client:
            kwargs["client"] = args.client
        if getattr(args, "client_address", None):
            kwargs["client_address"] = args.client_address
        if getattr(args, "client_reg_type", None):
            kwargs["client_reg_type"] = args.client_reg_type
        if getattr(args, "client_reg_number", None):
            kwargs["client_reg_number"] = args.client_reg_number
        if args.formality:
            kwargs["formality"] = args.formality
    else:
        if args.title:
            kwargs["title"] = args.title
        if args.client:
            kwargs["client"] = args.client
        if args.profile in ("letter", "seed_memo", "investor_memo", "exec_summary"):
            kwargs["entity"] = args.entity

    doc = fn(**kwargs)

    if args.output:
        out = args.output
    else:
        out = os.path.join(OUTPUT, auto_name(
            args.profile, args.client, args.version, dt,
            agreement_type=kwargs.get("agreement_type")))

    save_doc(doc, out)
    print(f"Generated: {out}")

    if args.validate:
        _run_validate(out)

    if args.dotx:
        dotx = out.replace(".docx", ".dotx")
        save_doc(doc, dotx)
        print(f"Template:  {dotx}")

    if args.pdf:
        try:
            pdf = docx_to_pdf(out)
            print(f"PDF:       {pdf}")
        except RuntimeError as e:
            print(f"PDF conversion failed: {e}", file=sys.stderr)
            sys.exit(1)


if __name__ == "__main__":
    main()
