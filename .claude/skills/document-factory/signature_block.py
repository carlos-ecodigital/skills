"""
Bilingual multi-party signature-block rendering for Sites-stream documents.

Additive to ``document-factory/document_factory.py``. Renders Van Gog-style
signature pages with:

- A single "Signatures Page / Handtekeningen:" bilingual header
- The DE provider block (legal name, signature line, Name/Title/Date)
- One block per ``SigParty``, prefixed by a role-label header in the form
  ``GRID CONTRIBUTOR / NETBIJDRAGER:`` (or comma-joined when a party holds
  multiple labels, e.g., ``GRID CONTRIBUTOR, LANDOWNER / NETBIJDRAGER, GRONDEIGENAAR:``)
- KvK line is included only when ``formality == "binding"`` (i.e., HoT).
  LOI sig blocks deliberately omit KvK per Van Gog pattern.

Import via:

    from signature_block import SigParty, render_signature_page

No existing ``document_factory.py`` signatures are modified.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from typing import List, Optional

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from document_factory import (
    COBALT,
    FONT,
    Party,
    SLATE,
    SLATE_800,
    SLATE_900,
)

SIGNATURE_LINE = "_" * 32


@dataclass
class SigParty:
    """A Site Partner signing the document with one or more role labels."""

    legal_name: str
    #: EN labels, e.g. ``["Grid Contributor"]`` or ``["Grid Contributor", "Landowner"]``.
    role_labels_en: List[str] = field(default_factory=list)
    #: NL equivalents, e.g. ``["Netbijdrager"]`` or ``["Netbijdrager", "Grondeigenaar"]``.
    role_labels_nl: List[str] = field(default_factory=list)
    signatory_name: str = ""
    signatory_title: str = ""
    #: KvK required only for ``formality == "binding"``. Absent from LOI.
    kvk: Optional[str] = None


def render_signature_page(
    doc,
    provider_party: Party,
    site_partners: List[SigParty],
    formality: str = "non_binding",
    provider_signatory_name: str = "",
    provider_signatory_title: str = "",
) -> None:
    """Append a signature page for DE + 1..N Site Partners.

    Args:
        doc: python-docx ``Document`` to append to.
        provider_party: DE's ``Party`` record (from ``generate.DE_ENTITIES``).
        site_partners: list of ``SigParty``. Order preserved as rendered.
        formality: ``"non_binding"`` (LOI) or ``"binding"`` (HoT). Controls
            whether KvK rows are rendered.
        provider_signatory_name: DE signatory name (e.g., "Carlos Reuven
            Mattis Glender"). Left blank if not yet assigned.
        provider_signatory_title: DE signatory title (e.g., "Director").
    """
    if formality not in ("binding", "non_binding"):
        raise ValueError(
            f"formality must be 'binding' or 'non_binding', got {formality!r}"
        )

    # Header
    p = doc.add_paragraph()
    r = p.add_run("Signatures Page / Handtekeningen:")
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(11)
    r.font.color.rgb = SLATE_900

    # DE provider block (no role-label header)
    _render_party_block(
        doc,
        role_label_line=None,
        legal_name=provider_party.legal_name,
        signatory_name=provider_signatory_name,
        signatory_title=provider_signatory_title,
        kvk=provider_party.registration_number if formality == "binding" else None,
    )

    # Site Partner blocks
    for sp in site_partners:
        label_line = _format_role_label_line(sp)
        _render_party_block(
            doc,
            role_label_line=label_line,
            legal_name=sp.legal_name,
            signatory_name=sp.signatory_name,
            signatory_title=sp.signatory_title,
            kvk=sp.kvk if formality == "binding" else None,
        )


def _format_role_label_line(sp: SigParty) -> Optional[str]:
    """Return ``"GRID CONTRIBUTOR, LANDOWNER / NETBIJDRAGER, GRONDEIGENAAR:"``
    or ``None`` if no labels."""
    if not sp.role_labels_en and not sp.role_labels_nl:
        return None
    en = ", ".join(lbl.upper() for lbl in sp.role_labels_en)
    nl = ", ".join(lbl.upper() for lbl in sp.role_labels_nl)
    if en and nl:
        return f"{en} / {nl}:"
    return f"{en or nl}:"


def _render_party_block(
    doc,
    role_label_line: Optional[str],
    legal_name: str,
    signatory_name: str,
    signatory_title: str,
    kvk: Optional[str],
) -> None:
    # Visual spacer
    doc.add_paragraph()

    # Role-label header (if any)
    if role_label_line:
        p = doc.add_paragraph()
        r = p.add_run(role_label_line)
        r.bold = True
        r.font.name = FONT
        r.font.size = Pt(10)
        r.font.color.rgb = COBALT

    # Legal name
    p = doc.add_paragraph()
    r = p.add_run(legal_name)
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(11)
    r.font.color.rgb = SLATE_900

    # KvK (binding only)
    if kvk:
        p = doc.add_paragraph()
        r = p.add_run(f"KvK: {kvk}")
        r.font.name = FONT
        r.font.size = Pt(9)
        r.font.color.rgb = SLATE

    # Signature line
    p = doc.add_paragraph()
    r = p.add_run(SIGNATURE_LINE)
    r.font.name = FONT
    r.font.size = Pt(11)
    r.font.color.rgb = SLATE_800

    # Name / Title / Date
    for label, value in (
        ("Name", signatory_name),
        ("Title", signatory_title),
        ("Date", ""),
    ):
        p = doc.add_paragraph()
        r = p.add_run(f"{label}: {value}")
        r.font.name = FONT
        r.font.size = Pt(10)
        r.font.color.rgb = SLATE_800
